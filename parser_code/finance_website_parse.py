from pymongo import MongoClient
import os
import re
import subprocess
from xlrd import open_workbook
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
import docx

from utility import request_site_page, format_date, get_content_text, table_to_list, remove_special_char
from pdf2html import pdf_to_text
from init import logger_init, config_init
from oss_utils import init_ali_oss, oss_add_file

logger = logger_init('财政厅 数据解析')
config = config_init()
if config['mongodb']['dev_mongo'] == '1':
    db = MongoClient(config['mongodb']['ali_mongodb_url'], username=config['mongodb']['ali_mongodb_username'],
                     password=config['mongodb']['ali_mongodb_password'],
                     port=int(config['mongodb']['ali_mongodb_port']))[config['mongodb']['ali_mongodb_name']]
else:
    db = MongoClient(
        host=config['mongodb']['mongodb_host'],
        port=int(config['mongodb']['mongodb_port']),
        username=None if config['mongodb']['mongodb_username'] == '' else config['mongodb']['mongodb_username'],
        password=None if config['mongodb']['mongodb_password'] == '' else config['mongodb']['mongodb_password'])[
        config['mongodb']['mongodb_db_name']]

ali_bucket = init_ali_oss()


# 新疆维吾尔自治区财政厅
def xinjiangczt_parse():
    for each_document in db.finance_data.find({'origin': '新疆维吾尔自治区财政厅', 'status': {'$nin': ['ignored']}}):
        announcement_url = each_document['url']
        announcement_title = each_document['title']

        if '财政厅监督检查局重点针对' in announcement_title:
            logger.warning('url has nothing to do with punishment ...')
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update finance data success')
            continue

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')
        announcement_title = content_soup.find(class_='article-title').text.strip()

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_class_name': 'article-content'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_text = get_content_text(content_soup.find(attrs={"class": "article-content"}))

        if re.search('监管关注函', announcement_title):
            announcement_type = '监管关注'
        elif re.search('行政处罚', announcement_title):
            announcement_type = '行政处罚决定'
        else:
            announcement_type = '责令整改通知'
        litigant = re.search(r'^([\s\S]*?)\n(根据财政部|查工作的通知》（财监〔2016〕4号）).*?(精神|要求)',
                             content_text).group(1).strip()
        facts = re.search(r'(现将查出的主要问题及拟作出的行政处罚决定告知如下|检查出的问题及处理决定如下|'
                          r'检查出的主要问题及做出的行政处罚决定如下|检查出的主要问题及处理决定如下|'
                          r'现将检查出的主要问题及处理决定通知如下|现将检查出的主要问题及拟作出的行政处罚决定告知如下).*?\n'
                          r'([\s\S]*?)\n.*?(拟作出的行政处罚|行政处罚决定|处理决定|处理意见)',
                          content_text).group(2).strip()
        punishment_decision = re.search(r'(拟作出的行政处罚|行政处罚决定|处理决定|处理意见)\n([\s\S]*)\n新疆维吾尔自治区财政厅',
                                        content_text).group(2).strip()
        publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', content_text)[-1].strip()
        real_publish_date = format_date(publish_date)

        punishment_basis_str_list = [
            r'([^\n。；]*?)(问题|行为|事项|情况|事实)([^\n。；\s]*?)(违反|构成).*?\n?.*?第.*?条?\n?.*?((的|之|等)(相关)?规定)?',
        ]
        punishment_basis_str = '|'.join(punishment_basis_str_list)
        punishment_basis_compiler = re.compile(r'[。\n；]' + '(' + punishment_basis_str + ')' +
                                               r'.\n', re.MULTILINE)
        punishment_basis_list = punishment_basis_compiler.findall(content_text)
        punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '新疆财政厅',
            'announcementDate': real_publish_date,
            'announcementCode': '',
            'facts': facts,
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': announcement_type,
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('新疆维吾尔自治区财政厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('新疆维吾尔自治区财政厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('新疆维吾尔自治区财政厅 数据解析 ' + ' -- 修改parsed完成')


# 甘肃省财政厅
def gansuczt_parse():
    for each_document in db.finance_data.find({'origin': '甘肃省财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_id_name': 'newscontent'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_text = get_content_text(content_soup.find(attrs={"id": 'newscontent'}))
        if '未作出行政处罚决定。' in content_text:
            logger.warning('url has nothing to do with punishment ...')
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update finance data success')
            continue
        # 处罚类型为：行政处罚决定
        litigant = re.search(r'(当事人[\s\S]*?)\n根据.*?有关规定和要求', content_text).group(1).strip()
        facts = re.search(r'检查发现的问题和行政处罚决定如下.*?\n([\s\S]*?)(以上事实.*?等相关证据予以证实。|上述行为违反了)',
                          content_text).group(1).strip()
        punishment_decision = re.search(r'(根据.*?第.*?条.*?规定.*?我厅决定[\s\S]*?)\n(如不服本处罚决定)',
                                        content_text).group(1).strip()
        publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', content_text)[-1].strip()
        real_publish_date = format_date(publish_date)
        punishment_basis_str_list = [
            r'([^\n。；]*?)(问题|行为|事项|情况|事实)([^\n。；\s]*?)(违反|构成).*?\n?.*?第.*?条?\n?.*?((的|之|等)(相关)?规定)?',
        ]
        punishment_basis_str = '|'.join(punishment_basis_str_list)
        punishment_basis_compiler = re.compile(r'[。\n；]' + '(' + punishment_basis_str + ')' +
                                               r'.\n', re.MULTILINE)
        punishment_basis_list = punishment_basis_compiler.findall(content_text)
        punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '甘肃财政厅',
            'announcementDate': real_publish_date,
            'announcementCode': '',
            'facts': facts,
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('甘肃省财政厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('甘肃省财政厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('甘肃省财政厅 数据解析 ' + ' -- 修改parsed完成')


# 云南省财政厅
def yunnanczt_parse():
    for each_document in db.finance_data.find({'origin': '云南省财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_id_name': 'Zoom'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_text = get_content_text(content_soup.find(id='Zoom'))

        if re.search(r'^(云南省财政厅文件\n.*?\n)?.*?决定\n([\s\S]*?)根据.*?向.*?下达', content_text):
            litigant = re.search(r'^(云南省财政厅文件\n.*?\n)?.*?决定\n([\s\S]*?)根据.*?向.*?下达',
                                 content_text).group(2).strip()
            facts = re.search(litigant + r'\n([\s\S]*?)\n根据.*?第.*?条.*?规定', content_text).group(1).strip()
            punishment_decision = re.search(r'(根据.*?第.*?条.*?规定.*?)\n.*?如对本决定有异议',
                                            content_text).group(1).strip()
        else:
            litigant = re.search(r'^(.*?)因未在限期整改期限内达到符合执业许可条', content_text).group(1).strip()
            facts = re.search(r'^([\s\S]*?)根据.*?第.*?条.*?规定', content_text).group(1).strip()
            punishment_decision = re.search(r'(根据.*?第.*?条.*?规定.*?)\n特此公告', content_text).group(1).strip()

        if re.search(r'(云财会〔\d{4}〕\d+号)', content_text):
            announcement_code = re.search(r'(云财会〔\d{4}〕\d+号)', content_text).group(1).strip()
        else:
            announcement_code = ''
        publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', content_text)[-1].strip()
        real_publish_date = format_date(publish_date)
        punishment_basis_str_list = [
            r'([^\n。；]*?)(问题|行为|事项|情况|事实)([^\n。；\s]*?)(违反|构成).*?\n?.*?第.*?条?\n?.*?((的|之|等)(相关)?规定)?',
        ]
        punishment_basis_str = '|'.join(punishment_basis_str_list)
        punishment_basis_compiler = re.compile(r'[。\n；]' + '(' + punishment_basis_str + ')' +
                                               r'.\n', re.MULTILINE)
        punishment_basis_list = punishment_basis_compiler.findall(content_text)
        punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '云南财政厅',
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': each_document['type'],
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('云南省财政厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('云南省财政厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('云南省财政厅 数据解析 ' + ' -- 修改parsed完成')


# 贵州省财政厅
def guizhouczt_parse():
    for each_document in db.finance_data.find({'origin': '贵州省财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_id_name': 'Zoom'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        if '处罚事务所名单' in announcement_title:
            table_list = table_to_list(content_soup.find(id='Zoom').find('table'))

            for each_tr in table_list:
                if '会计师事务所名称' in each_tr[0]:
                    continue
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': '贵州财政厅',
                    'announcementDate': format_date(each_tr[6].strip()),
                    'announcementCode': each_tr[8].strip(),
                    'facts': '',
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': '会计师事务所名称：' + each_tr[0].strip() + '\n' + '执业证书编号：' + each_tr[
                        2].strip() + '\n' + '事务所编号：' + each_tr[3].strip(),
                    'punishmentBasement': '',
                    'punishmentDecision': each_tr[5].strip(),
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id,
                                         'announcementCode': each_tr[8].strip()}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('贵州省财政厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('贵州省财政厅 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('贵州省财政厅 数据解析 ' + ' -- 修改parsed完成')


# 四川省财政厅
def sichuanczt_parse():
    for each_document in db.finance_data.find({'origin': '四川省财政厅', 'status': {'$nin': ['ignored']}}):
        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        # 解析excel
        if '行政处罚公告' in announcement_title:
            content_response = request_site_page(announcement_url)
            if content_response is None:
                logger.error('网页请求错误 %s' % announcement_url)
                continue
            content_response.encoding = content_response.apparent_encoding
            content_soup = BeautifulSoup(content_response.text, 'lxml')
            content_text = get_content_text(content_soup.find(class_='txt2-in'))
            if len(re.findall(r'.{4}年.{1,2}月.{1,3}日', content_text)) > 0:
                publish_date = re.findall(r'.{4}年.{1,2}月.{1,3}日', content_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            else:
                real_publish_date = format_date(each_document['publishDate'])

            announcement_title = re.search(r'相关下载：\n(.*?).xls',
                                           get_content_text(content_soup.find(class_='newstext04'))).group(1).strip()
            excel_download_id = str(
                re.findall(r'\d+', content_soup.find(class_='newstext04').find_all('td')[-1].attrs['onclick'])[0])
            xlsx_link = 'http://www.sccz.gov.cn/Site/DownAttach?id=' + str(excel_download_id) + '&TS=1548126893737'

            response = request_site_page(xlsx_link)
            with open('./test/tmp.xlsx', 'wb') as f:
                f.write(response.content)

            with open('./test/tmp.xlsx', 'rb') as xlsx_file:
                xlsx_content = xlsx_file.read()

            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': xlsx_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': xlsx_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'xlsx',
                    'oss_file_name': announcement_title,
                    'oss_file_content': xlsx_content,
                    'parsed': False
                }
                insert_response = db.parsed_data.insert_one(oss_file_map)
                file_id = insert_response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.xlsx', xlsx_content)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': xlsx_link})['_id']

            excel_data = open_workbook('./test/tmp.xlsx')
            logger.info('删除tmp文件')
            if os.path.exists('./test/tmp.xlsx'):
                os.remove('./test/tmp.xlsx')
            sheet = excel_data.sheets()[0]

            result_map_list = []
            for i in range(sheet.nrows):
                if '附件' in str(sheet.cell_value(i, 0)) or str(sheet.cell_value(i, 0)) == '' or \
                        '序号' in str(sheet.cell_value(i, 0)):
                    continue
                punishment_result = \
                    [kk for kk in re.split('(事务所|注册会计师)', sheet.cell_value(i, 3).strip()) if kk != '']
                final_punishment_result = []
                for index, each_punishment in enumerate(punishment_result):
                    if each_punishment not in ['事务所', '注册会计师']:
                        final_punishment_result.append((punishment_result[index - 1] + each_punishment).strip())
                for each_final_punishment in final_punishment_result:
                    announcement_code = re.search(r'(川财罚〔\d{4}〕\d+号(、川财罚〔\d{4}〕\d+号)?)',
                                                  each_final_punishment).group(1).strip()
                    facts = str(sheet.cell_value(i, 2)).strip()
                    if '事务所' in each_final_punishment:
                        litigant = str(sheet.cell_value(i, 1)).strip()
                        punishment_decision = re.search(r'事务所：([\s\S]*?)（' + announcement_code + '）',
                                                        each_final_punishment).group(1).strip()
                    else:
                        litigant = each_final_punishment.split('\n')[0].strip() + '\n' + '就职公司：' \
                                   + str(sheet.cell_value(i, 1)).strip()
                        punishment_decision = re.search(
                            r'注册会计师.*?\n?[(（]\d+[）)]([\s\S]*?)（' + announcement_code + ' ?）',
                            each_final_punishment).group(1).strip()

                    result_map = {
                        'announcementTitle': announcement_title + '（' + announcement_code + '）',
                        'announcementOrg': '四川财政厅',
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
                        'punishmentBasement': '',
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
                    logger.info(result_map)
                    if db.announcement.find(
                            {'announcementTitle': announcement_title, 'oss_file_id': file_id,
                             'announcementCode': announcement_code}).count() == 0:
                        db.announcement.insert_one(result_map)
                        logger.info('四川省财政厅 数据解析 ' + ' -- 数据导入完成')
                    else:
                        logger.info('四川省财政厅 数据解析 ' + ' -- 数据已经存在')
            if len(result_map_list) > 0:
                logger.info('四川省财政厅 数据解析 一共有%d条数据' % len(result_map_list))
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('四川省财政厅 数据解析 ' + ' -- 修改parsed完成')
            else:
                logger.info('四川省财政厅 数据解析 ' + ' -- 解析出现问题')
        # 解析纯文本
        else:
            content_response = request_site_page(announcement_url)
            if content_response is None:
                logger.error('网页请求错误 %s' % announcement_url)
                continue
            content_response.encoding = content_response.apparent_encoding
            content_soup = BeautifulSoup(content_response.text, 'lxml')

            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text,
                    'parsed': False,
                    'content_class_name': 'txt2-in'
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']

            content_text = get_content_text(content_soup.find(class_='txt2-in'))

            announcement_code = re.search(r'\n(川财罚〔\d{4}〕\d+号)\n', content_text).group(1).strip()
            litigant = re.search(r'\n(当事人[\s\S]*?)\n本厅调查查明', content_text).group(1).strip()
            facts = re.search(r'\n(本厅调查查明[\s\S]*?)\n以上事实.*?等予以证明', content_text).group(1).strip()
            punishment_decision = re.search(r'\n(根据.*第.*?条.*?规定，本厅对.*?处以[\s\S]*?)\n(请你公司自收到本决定书之日)',
                                            content_text).group(1).strip()
            publish_date = re.findall(r'.{4}年.{1,2}月.{1,3}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
            punishment_basis_str_list = [
                r'([^\n。；]*?)(问题|行为|事项|情况|事实)([^\n。；\s]*?)(违反|构成).*?\n?.*?第.*?条?\n?.*?((的|之|等)(相关)?规定)?',
            ]
            punishment_basis_str = '|'.join(punishment_basis_str_list)
            punishment_basis_compiler = re.compile(r'[。\n；]' + '(' + punishment_basis_str + ')' +
                                                   r'.\n', re.MULTILINE)
            punishment_basis_list = punishment_basis_compiler.findall(content_text)
            punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '四川财政厅',
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('四川省财政厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('四川省财政厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('四川省财政厅 数据解析 ' + ' -- 修改parsed完成')


# 重庆市财政局
def chongqingczj_parse():
    for each_document in db.finance_data.find({'origin': '重庆市财政局', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find(
                {'url': announcement_url, 'status': 'parsed'}).count() == 1 and \
                db.parsed_data.find({'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue

        # ignored
        if '自愿撤回执业许可' in announcement_title:
            logger.warning('url has nothing to do with punishment ...')
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update finance data success')
            continue

        logger.info('url to parse ' + announcement_url)

        # 行政处罚决定书
        if re.search('行政处罚决定书', announcement_title):
            content_response = request_site_page(announcement_url)
            if content_response is None:
                logger.error('网页请求错误 %s' % announcement_url)
                continue
            content_response.encoding = content_response.apparent_encoding
            content_soup = BeautifulSoup(content_response.text, 'lxml')

            content_text = get_content_text(content_soup.find(id='text'))
            link_list = content_soup.find(id='showcontent').find_all('a')

            # 多个附件问题
            if len(link_list) > 1:
                doc_list = [urljoin(announcement_url, kk.attrs['href'].strip()) for kk in link_list]
                for each_doc_link in doc_list:
                    response = request_site_page(each_doc_link)
                    if response is None:
                        logger.error('网页请求错误 %s' % each_doc_link)
                        continue
                    response.encoding = response.apparent_encoding

                    with open('./test/tmp.' + each_doc_link.split('.')[-1], 'wb') as f:
                        f.write(response.content)
                    if not os.path.exists('./test/tmp.docx'):
                        shell_str = config['soffice']['soffice_path'] + ' --headless --convert-to docx ' + \
                                    './test/tmp.doc' + ' --outdir ./test/'
                        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                                   shell=True, stdout=subprocess.PIPE)
                        process.wait()

                    with open('./test/tmp.docx', 'rb') as docx_file:
                        docx_content = docx_file.read()

                    if db.parsed_data.find(
                            {'origin_url': announcement_url, 'oss_file_origin_url': each_doc_link}).count() == 0:
                        oss_file_map = {
                            'origin_url': announcement_url,
                            'oss_file_origin_url': each_doc_link,
                            'origin_url_id': each_document['_id'],
                            'oss_file_type': 'docx',
                            'oss_file_name': announcement_title,
                            'oss_file_content': docx_content,
                            'parsed': False
                        }
                        insert_response = db.parsed_data.insert_one(oss_file_map)
                        file_id = insert_response.inserted_id
                        oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                        db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    else:
                        db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                        file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                           'oss_file_origin_url': each_doc_link})['_id']
                    doc = Document('./test/tmp.docx')
                    full_text = []
                    for para in doc.paragraphs:
                        if para.text.strip() != '':
                            full_text.append(para.text.strip())
                    content_text = '\n'.join(full_text)

                    if re.search(r'(渝财处罚〔\d{4}〕\d+号)', content_text):
                        announcement_code = re.search(r'(渝财处罚〔\d{4}〕\d+号)', content_text).group(1).strip()
                        announcement_title = '重庆市财政局行政处罚决定书' + '（' + announcement_code + '）'
                    else:
                        announcement_code = ''
                        announcement_title = '重庆市财政局行政处罚决定书'

                    litigant = re.search(r'(当事人[\s\S]*?)\n'
                                         r'(根据.*?等.*?规定|你公司于.*?参加了.*?项目|重庆市安全生产监督管理局.*?分别于.*?向我局递交|'
                                         r'.*?我局收到有关反映)',
                                         content_text).group(1).strip()
                    facts = re.search(r'(违法事实和证据\n|查明的违法事实\n|违法事实\n|会计信息质量进行了检查。)'
                                      r'([\s\S]*?)\n'
                                      r'(你.*?上述行为.*?违反了|(以上|上述)(违法)?事实.*?(证据在案佐证|予以证实)|.{2}行政处罚决定)',
                                      content_text).group(2).strip()
                    punishment_decision = re.search(r'(根据.*?第.*?条.*?规定，(我局|市财政局|本机关)决定[\s\S]*?)'
                                                    r'\n(.*?在接到本决定书之日|如不服本处罚决定)',
                                                    content_text).group(1).strip()
                    publish_date_text = re.search(punishment_decision + r'([\s\S]*?)$', content_text).group(1).strip()
                    if len(re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)) > 0:
                        publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)[-1].strip()
                        real_publish_date = format_date(publish_date)
                    else:
                        real_publish_date = format_date(each_document['publishDate'])
                    punishment_basis_str_list = [
                        r'([^\n。；]*?)(问题|行为|事项|情况|事实)([^\n。；\s]*?)(违反|构成).*?\n?.*?第.*?条?\n?.*?((的|之|等)(相关)?规定)?',
                    ]
                    punishment_basis_str = '|'.join(punishment_basis_str_list)
                    punishment_basis_compiler = re.compile(r'[。\n；]' + '(' + punishment_basis_str + ')' +
                                                           r'.(\n|根据)', re.MULTILINE)
                    punishment_basis_list = punishment_basis_compiler.findall(content_text)
                    punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': '重庆财政局',
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    logger.info(result_map)
                    if db.announcement.find(
                            {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                        db.announcement.insert_one(result_map)
                        logger.info('重庆市财政局 数据解析 ' + ' -- 数据导入完成')
                    else:
                        logger.info('重庆市财政局 数据解析 ' + ' -- 数据已经存在')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('重庆市财政局 数据解析 ' + ' -- 修改parsed完成')

                    logger.info('删除tmp文件')
                    if os.path.exists('./test/tmp.doc'):
                        os.remove('./test/tmp.doc')
                    if os.path.exists('./test/tmp.docx'):
                        os.remove('./test/tmp.docx')

            # 没有附件或者既有附件又有文字的格式
            else:
                if db.parsed_data.find(
                        {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': announcement_url,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'html',
                        'oss_file_name': announcement_title,
                        'oss_file_content': content_response.text,
                        'parsed': False,
                        'content_id_name': 'text'
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                    db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                       'oss_file_origin_url': announcement_url})['_id']

                if re.search(r'(渝财处罚〔\d{4}〕\d+号)', content_text):
                    announcement_code = re.search(r'(渝财处罚〔\d{4}〕\d+号)', content_text).group(1).strip()
                    announcement_title = '重庆市财政局行政处罚决定书' + '（' + announcement_code + '）'
                else:
                    announcement_code = ''
                    announcement_title = '重庆市财政局行政处罚决定书'

                litigant = re.search(r'(当事人[\s\S]*?)\n'
                                     r'(根据.*?等.*?规定|你公司于.*?参加了.*?项目|重庆市安全生产监督管理局.*?分别于.*?向我局递交|'
                                     r'.*?我局收到有关反映)',
                                     content_text).group(1).strip()
                facts = re.search(r'(违法事实和证据\n|查明的违法事实\n|违法事实\n|会计信息质量进行了检查。)'
                                  r'([\s\S]*?)\n'
                                  r'(你.*?上述行为.*?违反了|(以上|上述)(违法)?事实.*?(证据在案佐证|予以证实)|.{2}行政处罚决定)',
                                  content_text).group(2).strip()
                punishment_decision = re.search(r'(根据.*?第.*?条.*?规定，(我局|市财政局|本机关)决定[\s\S]*?)'
                                                r'\n(.*?在接到本决定书之日|如不服本处罚决定)',
                                                content_text).group(1).strip()
                publish_date_text = re.search(punishment_decision + r'([\s\S]*?)$', content_text).group(1).strip()
                if len(re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)) > 0:
                    publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)[-1].strip()
                    real_publish_date = format_date(publish_date)
                else:
                    real_publish_date = format_date(each_document['publishDate'])
                punishment_basis_str_list = [
                    r'([^\n。；]*?)(问题|行为|事项|情况|事实)([^\n。；\s]*?)(违反|构成).*?\n?.*?第.*?条?\n?.*?((的|之|等)(相关)?规定)?',
                ]
                punishment_basis_str = '|'.join(punishment_basis_str_list)
                punishment_basis_compiler = re.compile(r'[。\n；]' + '(' + punishment_basis_str + ')' +
                                                       r'.(\n|根据)', re.MULTILINE)
                punishment_basis_list = punishment_basis_compiler.findall(content_text)
                punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': '重庆财政局',
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('重庆市财政局 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('重庆市财政局 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('重庆市财政局 数据解析 ' + ' -- 修改parsed完成')
        # 撤回执业许可
        elif re.search('撤回执业许可', announcement_title):
            content_response = request_site_page(announcement_url)
            if content_response is None:
                logger.error('网页请求错误 %s' % announcement_url)
                continue
            content_response.encoding = content_response.apparent_encoding
            content_soup = BeautifulSoup(content_response.text, 'lxml')

            content_text = get_content_text(content_soup.find(id='text'))

            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text,
                    'parsed': False,
                    'content_id_name': 'text'
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']

            if re.search(r'(渝财会〔\d{4}〕\d+号)', content_text):
                announcement_code = re.search(r'(渝财会〔\d{4}〕\d+号)', content_text).group(1).strip()
                announcement_title = '重庆市财政局关于撤回执业许可的通知' + '（' + announcement_code + '）'
            else:
                announcement_code = ''
                announcement_title = '重庆市财政局关于撤回执业许可的通知'

            litigant = re.search(r'(.*?)\n根据.*?本局于.*?作出', content_text).group(1).strip()
            facts = re.search(litigant + r'([\s\S]*?)根据.*?第.*条.*?规定，本局现决定',
                              content_text).group(1).strip()
            punishment_decision = re.search(r'(根据.*?第.*?条.*?规定，本局现决定[\s\S]*?)'
                                            r'\n如不服本决定',
                                            content_text).group(1).strip()
            publish_date_text = re.search(punishment_decision + r'([\s\S]*?)$', content_text).group(1).strip()
            if len(re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)) > 0:
                publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            else:
                real_publish_date = format_date(each_document['publishDate'])
            punishment_basis_str_list = [
                r'([^\n。；]*?)(问题|行为|事项|情况|事实)([^\n。；\s]*?)(违反|构成).*?\n?.*?第.*?条?\n?.*?((的|之|等)(相关)?规定)?',
            ]
            punishment_basis_str = '|'.join(punishment_basis_str_list)
            punishment_basis_compiler = re.compile(r'[。\n；]' + '(' + punishment_basis_str + ')' +
                                                   r'.(\n|根据)', re.MULTILINE)
            punishment_basis_list = punishment_basis_compiler.findall(content_text)
            punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '重庆财政局',
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '撤回执业许可',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('重庆市财政局 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('重庆市财政局 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('重庆市财政局 数据解析 ' + ' -- 修改parsed完成')
        # 投诉处理决定书
        else:
            content_response = request_site_page(announcement_url)
            if content_response is None:
                logger.error('网页请求错误 %s' % announcement_url)
                continue
            content_response.encoding = content_response.apparent_encoding
            content_soup = BeautifulSoup(content_response.text, 'lxml')

            content_text = get_content_text(content_soup.find(id='text'))

            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text,
                    'parsed': False,
                    'content_id_name': 'text'
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']

            announcement_code = re.search(r'(渝财处理〔\d{4}〕\d+号)', announcement_title).group(1).strip()
            litigant = re.search(r'被投诉人\n([\s\S]*?)\n.{2}投诉项目', content_text).group(1).strip()
            facts = re.search(r'(投诉人\n[\s\S]*?)\n.{2}被投诉人', content_text).group(1).strip() + '\n' + re.search(
                r'(投诉项目[\n：][\s\S]*?)\n.{2}投诉受理', content_text).group(1).strip() + '\n' + re.search(
                r'(投诉内容\n[\s\S]*?)\n.{2}(调查核实情况|调查取证情况)', content_text).group(1).strip()
            punishment_basis = re.search(r'(调查核实情况|调查取证情况)\n([\s\S]*?)\n.{2}处理决定',
                                         content_text).group(2).strip()
            punishment_decision = re.search(r'处理决定\n([\s\S]*?)\n(.*?对本处理决定不服)',
                                            content_text).group(1).strip()
            publish_date_text = re.search(punishment_decision + r'([\s\S]*?)$', content_text).group(1).strip()
            if len(re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)) > 0:
                publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            else:
                real_publish_date = format_date(each_document['publishDate'])

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '重庆财政局',
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('重庆市财政局 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('重庆市财政局 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('重庆市财政局 数据解析 ' + ' -- 修改parsed完成')


# 广西壮族自治区财政厅
def guangxiczt_parse():
    for each_document in db.finance_data.find({'origin': '广西壮族自治区财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if content_soup.find(attrs={"class": "TRS_Editor"}):
            content_text_list = content_soup.find(attrs={"class": "TRS_Editor"}).findAll('p')
            content_text = '\n'.join([kk.text.strip() for kk in content_text_list])
            content_text = remove_special_char(content_text)
            content_class_name = 'TRS_Editor'
        else:
            content_text_list = content_soup.find(attrs={"class": "article-con"}).findAll('p')
            content_text = '\n'.join([kk.text.strip() for kk in content_text_list])
            content_text = remove_special_char(content_text)
            content_class_name = 'article-con'

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_class_name': content_class_name
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        if '行政处罚决定' in announcement_title:
            announcement_code = re.search(r'(桂财监〔\d{4}〕\d+号)', content_text).group(1).strip()
            litigant = re.search(r'(当事人[\s\S]*?)\n根据.*?(要求|规定)', content_text).group(1).strip()
            facts = re.search(r'((经查|根据广西第一工业学校提供的收支票据反映)[\s\S]*?)'
                              r'(。违反了.*?的规定|\n以上不符合.*?的规定)', content_text).group(1).strip()
            punishment_basis = re.search(r'(违反了.*?的规定。|\n以上不符合.*?的规定。)'
                                         r'\n(根据.*?第.*?条.*?的规定|鉴于.*?并结合你对“小金库”举报的表现)',
                                         content_text).group(1).strip()
            punishment_decision = re.search(r'(根据.*?第.*?条.*?的规定，(我厅拟)?对.*?处以[\s\S]*?)'
                                            r'\n(请你会通过一般缴款书将该笔罚款在收到本决定书之日起十五日内缴入自治区本级国库|'
                                            r'请你自收到本决定书之日)',
                                            content_text).group(1).strip()
            publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        else:
            announcement_code_compiler = re.compile(r'(桂财监〔\d{4}〕\d+号)')
            announcement_code = announcement_code_compiler.search(content_text).group(1).strip()
            litigant = re.search(r'\n(.*?)\n根据.*?要求', content_text).group(1).strip()
            facts_compiler = re.compile(r'\n([一二三四五六七八九十]{1,2}、.*?\n（[一二三四五六七八九十]{1,2}）.*?|'
                                        r'（[一二三四五六七八九十]{1,2}）.*?|[一二三四五六七八九十]{1,2}、.*?|'
                                        r'你公司.*?经查|检查发现|'
                                        r'\d{4}年\d+月\d+日.你公司收到|'
                                        r'经查，\d{4}年，你公司收到|'
                                        r'2014年12月29日和2015年5月29日，你公司分别获得|'
                                        r'经查，你公司2015年6月12日收到自治区体育局|'
                                        r'经查.\d{4}年\d+月\d+日.你公司收到自治区体育局)'
                                        r'([\s\S]*?)((以上)?[^。，\n]*?不符合.*?(的|有关)规定[。，]|'
                                        r'不符合《小企业会计准则—会计科目、主要账务处理和财务报表》规定的会计科目、科目代码设置。|'
                                        r'不符合.*?第.*?条.*?的要求。|'
                                        r'以上.*?不符合.*?第.*?条.*?的规定”。)'
                                        r'(\n?((针对|对于|对以上).*?问题.|今后.?|以上.*?点.|针对以上.*?点.)?'
                                        r'(你厅|你委|你会|你局|县财政局和扶贫办|县扶贫办|县财政局|你公司|'
                                        r'你单位)[^。，\n]*?[应要在]'
                                        r'(重视|督促|立即组织|调整|及时进行财务调整|主动履行职责|严格按照相关规定|'
                                        r'重新对固定资产进行盘点查清|根据有关法律法规的规定|补充完善原始凭证|尽快补充登记|'
                                        r'严格按照|当?进一步加强|立即进行账务调整|立即整改|认真履行监管职责|认真落实|'
                                        r'认真履行政府采购监管职责|将挤占的专项资金归还原资金渠道|按规定|高度重视|'
                                        r'将被挤占的土地整治专项资金按原渠道归还|按相关法律法规|补充完善|立即督促|'
                                        r'牵头加强项目库建设|立即进行核实整改|按.*?规定进行账务调整|自觉履行|将违规|'
                                        r'完善相关手续|调整会计账簿上的会计科目|按企业财务会计报告条例的规定|'
                                        r'针对上述会计事项进行会计处理|对出纳员进行岗位调整|立即向当地税务机关申报|'
                                        r'完善上述原始凭证|当?加强|当?建立|按《2015年农业技术推广项目申报指南》的规定，及时办理|'
                                        r'加强公务用车使用管理|按相关规定严格控制|当?规范|当?贯彻|当?严格|当?责成|'
                                        r'将.*?全部收回|当?补充和完善|当?按照|尽快将.*?全部收回|'
                                        r'当?及时调整|当?及时补充完善|当?进行账务调整|当?按有关规定要求|'
                                        r'当?尽快对目前已完工项目申请验收|申请项目验收).*|'
                                        r'\n?根据.*?第.*?条.*?的规定.*)?')
            facts_list = facts_compiler.findall(content_text)
            if str(facts_list[0][0].strip()).startswith(('一', '（')):
                facts = '\n'.join([kk[0].strip() + kk[1].strip() for kk in facts_list])
                punishment_decision = '\n'.join([kk[0].strip() + kk[5].strip() for kk in facts_list])
                punishment_decision = re.sub(r'\n+', r'\n', punishment_decision)
                punishment_basis = '\n'.join([kk[0].strip() + kk[2].strip() for kk in facts_list])
                punishment_basis = re.sub(r'\n+', r'\n', punishment_basis)
            else:
                facts = '\n'.join([kk[0].strip() + kk[1].strip() for kk in facts_list])
                punishment_decision = '\n'.join([kk[5].strip() for kk in facts_list])
                punishment_decision = re.sub(r'\n+', r'\n', punishment_decision)
                punishment_basis = '\n'.join([kk[2].strip() for kk in facts_list])
                punishment_basis = re.sub(r'\n+', r'\n', punishment_basis)
            try:
                publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.warning(e)
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '广西财政厅',
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('广西壮族自治区财政厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('广西壮族自治区财政厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('广西壮族自治区财政厅 数据解析 ' + ' -- 修改parsed完成')


# 广东省财政厅
def guangdongczt_parse():
    for each_document in db.finance_data.find({'origin': '广东省财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find(
                {'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
            {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        # ignored
        if '事项目录' in announcement_title or 'test' in announcement_title:
            logger.warning('url has nothing to do with punishment ...')
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update finance data success')
            continue

        announcement_id = re.search(r'http://czt.gd.gov.cn/xyxxsgs/xzcf/xzcfxl/\?i=(\d+)&type=xzcf',
                                    announcement_url).group(1).strip()
        announcement_request_url = \
            'http://app.southcn.com/ctsgs/ajax/act.ajax.php?callback=success_jsonpCallback' \
            + '&i=' + str(announcement_id) + '&t=xzcf&act=Gi&_=1548149623100'

        content_response = request_site_page(announcement_request_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_request_url)
            continue
        html_string = re.search(r'success_jsonpCallback\({\"code\":1,\"data\":\"(.*)\"}\)',
                                content_response.text.encode(content_response.encoding).decode('utf-8')) \
            .group(1).strip()
        content_soup = BeautifulSoup(
            html_string.replace('\\"', '"').replace('\\/', '/').encode('latin-1').decode('unicode_escape'), 'lxml')

        pdf_list = content_soup.find(class_='detail_block').find_all('a')
        if len(pdf_list) > 1:
            pdf_link = urljoin(announcement_url, pdf_list[-1].attrs['href'])
            response = request_site_page(pdf_link)
            if response is None:
                logger.error('网页请求错误 %s' % pdf_link)
                continue
            with open('./test/tmp.pdf', 'wb') as f:
                f.write(response.content)
            with open('./test/tmp.pdf', 'rb') as pdf_file:
                pdf_content = pdf_file.read()

            content_text = pdf_to_text('./test/tmp.pdf')
            logger.info(content_text)
            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.pdf'):
                os.remove('./test/tmp.pdf')

            announcement_code = re.search(r'(粤财罚〔\d{4}〕\d+号)', content_text).group(1).strip()
            announcement_title = '广东省财政厅行政处罚决定书（' + announcement_code + '）'

            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': pdf_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': pdf_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False
                }
                insert_response = db.parsed_data.insert_one(oss_file_map)
                file_id = insert_response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': pdf_link})['_id']
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'origin_url_id': each_document['_id']}})

            litigant = re.search(r'((当事人|企业名称)[\s\S]*?)(\n本(机关|单位)对.*?的违法行为)',
                                 content_text).group(1).strip()
            facts = re.search(r'违法事实\n([\s\S]*?)(以上事实有.*?等证据证实。|本机关认定|综上，本机关认为|本单位认定)',
                              content_text).group(1).strip()
            if re.search(r'((综上，本机关认为|本单位认定|本机关认定)[\s\S]*?)(\n.{2}行政处罚决定|以上事实有)',
                         content_text):
                punishment_basis = re.search(r'((综上，本机关认为|本单位认定|本机关认定)[\s\S]*?)(\n.{2}行政处罚决定|以上事实有)',
                                             content_text).group(1).strip()
            else:
                punishment_basis = ''
            punishment_decision = re.search(r'((根据.*?第.*?条.*?规定，(决定对你|对你单位作出)|'
                                            r'根据.*?按照.*?规定，你的违法行为为轻微|'
                                            r'根据.*?按照《广东省财政厅关于规范财政行政处罚自由裁量权的规定》，你的违法行为为轻微)'
                                            r'[\s\S]*?).{2}权利告知',
                                            content_text).group(1).strip()
            publish_date_text = re.search(punishment_decision + r'([\s\S]*?)$', content_text).group(1).strip()
            if len(re.findall('\n.{4}年.{1,2}月.{1,3}日\n', publish_date_text)) > 0:
                publish_date = re.findall('\n.{4}年.{1,2}月.{1,3}日\n', publish_date_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            else:
                real_publish_date = format_date(each_document['publishDate'])

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '广东财政厅',
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('广东省财政厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('广东省财政厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('广东省财政厅 数据解析 ' + ' -- 修改parsed完成')
        else:
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': str(content_soup),
                    'parsed': False,
                    'content_class_name': 'item_table'
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']

            item_table = content_soup.find(class_='item_table').find_all('tr')
            announcement_code = item_table[0].find_all('td')[1].text.strip()
            announcement_title = '广东省财政厅行政处罚决定书（' + announcement_code + '）'
            facts = item_table[4].find_all('td')[1].text.strip()
            punishment_decision = '依据' + item_table[5].find_all('td')[1].text.strip() + '，' + \
                                  item_table[3].find_all('td')[1].text.strip()
            punishment_basis = ''
            real_publish_date = format_date(item_table[21].find_all('td')[1].text.strip())
            litigant = ''
            for each_tr in item_table[6:17]:
                if each_tr.find_all('td')[1].text.strip() != '':
                    litigant += each_tr.find_all('td')[0].text.strip() + '：' + \
                                each_tr.find_all('td')[1].text.strip() + '\n'
            litigant = litigant.strip()
            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '广东财政厅',
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('广东省财政厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('广东省财政厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('广东省财政厅 数据解析 ' + ' -- 修改parsed完成')


# TODO 有点问题
# 湖南省财政厅
def hunanczt_parse():
    for each_document in db.finance_data.find({'origin': '湖南省财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find(
                {'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
            {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        # ignored
        if '事项目录' in announcement_title or 'test' in announcement_title:
            logger.warning('url has nothing to do with punishment ...')
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update finance data success')
            continue

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if len(content_soup.find(class_='tys-main-zt-show').find_all('a')) > 0:
            continue
        elif len(content_soup.find(class_='tys-main-zt-show').find_all('img')) > 0:
            continue
        elif len(content_soup.find(class_='tys-main-zt-show').find_all('table')) > 0:
            continue
        else:
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text,
                    'parsed': False,
                    'content_id_name': 'j-show-body'
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             content_response.text)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']

            content_text = get_content_text(content_soup.find(id='j-show-body'))
            announcement_code = re.search(r'(湘财行罚〔\d{4}〕\d+号)', announcement_title).group(1).strip()
            litigant = re.search(r'(当事人[\s\S]*?)'
                                 r'(根据.*?(精神|有关法律法规的规定|要求)|\d{4}年\d+月(\d+日)?.*?我厅收到)',
                                 content_text).group(1).strip()
            facts_text = re.search(r'.、查明的问题\n([\s\S]*?)'
                                   r'\n(以上(违法)?事实.*?有.*?作为证据(予以证明)?。|.、你公司申辩及事实认定|'
                                   r'以上(违法)?事实.*?有.*?等证据证明。|.、处罚(及处理)?决定)',
                                   content_text).group(1).strip()
            facts_list = re.findall(
                r'((^|\n)（[一二三四五六七八九十]{1,2}）.*\n)([\s\S]*?)(上述(问题|事项).*?(违反了|属于).*)', facts_text)
            if len(facts_list) > 0:
                facts = '\n'.join([(kk[0].strip() + '\n' + kk[2].strip()) for kk in facts_list])
                punishment_basis = '\n'.join([(kk[0].strip() + '\n' + kk[3].strip()) for kk in facts_list])
            else:
                facts = re.search(r'.、查明的问题\n([\s\S]*?)'
                                  r'\n(以上(违法)?事实，有.*?作为证据。|.、你公司申辩及事实认定|上述问题，违反了|'
                                  r'以上(违法)?事实.*?有.*?等证据(证明|材料)。)',
                                  content_text).group(1).strip()
                try:
                    punishment_basis = re.search(r'\n(上述(问题|行为)违反了.*?规定。)'
                                                 r'\n(以上(违法)?事实，有.*?作为证据(予以证明)?。|.、行政处罚决定|'
                                                 r'.、你公司申请陈述申辩、听证及我厅复核认定情况)',
                                                 content_text).group(1).strip()
                except Exception as e:
                    logger.warning(e)
                    punishment_basis = ''
            if '申辩' in content_text:
                try:
                    defense_opinion = re.search(
                        r'((（一）你公司在书面陈述申辩及听证会上提出申辩意见如下|（一）你公司提出的申辩理由|'
                        r'\(一）你公司提出的申辩理由|你公司于2017年12月20日向我厅提出陈述申辩，申辩的理由|'
                        r'2017年7月31日，你单位就此告知书进行了申辩，提出)'
                        r'[\s\S]*?)(（二）(经复核，?)?我厅认定情况|（二）我厅经复核后，认定情况如下|'
                        r'（二）经复核后我厅认定情况|经复核，我厅认为|我厅认为，该申辩)',
                        content_text).group(1).strip()
                    defense_response = re.search(
                        r'((（二）(经复核，?)?我厅认定情况如下|（二）我厅经复核后，认定情况如下|'
                        r'（二）经复核后我厅认定情况|经复核，我厅认为|我厅认为，该申辩)[\s\S]*?)'
                        r'\n(.、.*?处理处罚决定\n|根据《中华人民共和国政府采购法》第七十一条|'
                        r'根据《中华人民共和国政府采购法》第七十七条“供应商有下列情形之一的|'
                        r'根据《中华人民共和国注册会计师法》第三十九条“会计师事务所违反本法第二十条、第二十一条规定)',
                        content_text).group(1).strip()
                except Exception as e:
                    logger.warning(e)
                    defense_opinion = re.search(r'(你公司于2017年1月6日收到后，未在规定期限内提出陈述申辩及听证申请。|'
                                                r'你.*?在规定期限内未提出异议。)',
                                                content_text).group(1).strip()
                    defense_response = ''
            else:
                defense_opinion = defense_response = ''

            try:
                punishment_decision = re.search(r'.、.*?(处理|行政)?处罚(及处理)?决定\n([\s\S]*?)'
                                                r'\n(.*?应将整改落实情况在收到本决定书之日起|'
                                                r'.*?应在收到本行政处罚决定书30日内|'
                                                r'.*?应当在收到本处罚决定书之日|'
                                                r'如不服本决定)', content_text).group(3).strip()
            except Exception as e:
                logger.warning(e)
                punishment_decision = re.search(
                    r'(根据.*?第.*?条.*?规定.*?我厅对.*?给予[\s\S]*?)你公司应在收到本行政处罚决定书',
                    content_text).group(1).strip()

            try:
                publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.warning(e)
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '湖南财政厅',
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': defense_opinion,
                'defenseResponse': defense_response,
                'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('湖南省财政厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖南省财政厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖南省财政厅 数据解析 ' + ' -- 修改parsed完成')


# 青岛市财政局
def qingdaoczt_parse():
    for each_document in db.finance_data.find({'origin': '青岛市财政局', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_class_name': 'content_wz'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_text = get_content_text(content_soup.find(class_='content_wz'))
        litigant = re.search(r'(当事人[\s\S]*?)\n根据.*?市财政局于.*?监督检查', content_text).group(1).strip()
        facts = re.search(r'(经查[\s\S]*?)\n根据.*?第.*?条规定', content_text).group(1).strip()
        punishment_decision = re.search(r'(根据.*?第.*?条规定.*?决定.*?)'
                                        r'\n(你单位应当自收到本行政处罚决定书之日|请自收到本行政处罚决定书之日起15日内)',
                                        content_text).group(1).strip()
        try:
            publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.warning(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '青岛财政局',
            'announcementDate': real_publish_date,
            'announcementCode': '',
            'facts': facts,
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
            'punishmentBasement': '',
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('青岛市财政局 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('青岛市财政局 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('青岛市财政局 数据解析 ' + ' -- 修改parsed完成')


# 江西省财政厅
def jiangxiczt_parse():
    for each_document in db.finance_data.find({'origin': '江西省财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_class_name': 'showcontent'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_text = get_content_text(content_soup.find(class_='showcontent'))
        announcement_code = re.search(r'(赣财监处字.\d{4}.\d+号)', content_text).group(1).strip()
        real_publish_date = format_date(each_document['publishDate'])
        litigant = re.search(announcement_code.replace(r'[', r'\[').replace(r']', r'\]') +
                             r'\n([\s\S]*?)(\n自.*?我厅派出检查组)', content_text).group(1).strip()
        facts = re.search(r'((经查明|经检查发现)[\s\S]*?)(前述行为，违反了|你所另外一名合伙人江灿辉长期不在你所执业，你所未保持设立条件的行为，违反了)',
                          content_text).group(1).strip()
        punishment_basis = re.search(
            r'((前述行为，违反了|你所另外一名合伙人江灿辉长期不在你所执业，你所未保持设立条件的行为，违反了)[\s\S]*?)根据',
            content_text).group(1).strip()
        punishment_decision = re.search(r'(根据.*?第.*?条的规定，我厅决定[\s\S]*?)本决定自送达之日起生效',
                                        content_text).group(1).strip()

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '江西财政厅',
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('江西省财政厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('江西省财政厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('江西省财政厅 数据解析 ' + ' -- 修改parsed完成')


# 山东省财政厅
def shandongczt_parse():
    for each_document in db.finance_data.find({'origin': '山东省财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if announcement_url.endswith('pdf') or len(content_soup.find(id='level3_content').find_all('a')) > 0:
            if not announcement_url.endswith('pdf'):
                pdf_link = urljoin(announcement_url, content_soup.find(id='level3_content').find('a').attrs['href'])
            else:
                pdf_link = announcement_url
            response = request_site_page(pdf_link)
            if response is None:
                logger.error('网页请求错误')
                continue
            with open('./test/tmp.pdf', 'wb') as out_file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        out_file.write(chunk)
            with open('./test/tmp.pdf', 'rb') as pdf_file:
                pdf_content = pdf_file.read()

            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': pdf_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': pdf_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': pdf_link})['_id']

            content_text = pdf_to_text('./test/tmp.pdf')
            content_text = re.sub(r'—\d+—', '', content_text)
            content_text = re.sub(r'\n+', '\n', content_text)

            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.pdf'):
                os.remove('./test/tmp.pdf')
        else:
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text,
                    'parsed': False,
                    'content_id_name': 'level3_content'
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             content_response.text)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']

            for tag in content_soup("style"):
                tag.decompose()
            content_text = '\n'.join(
                [kk.text.strip() for kk in content_soup.find(id='level3_content').findChildren(recursive=False)])
            content_text = re.sub(r'\n+', '\n', content_text)

        content_text = re.sub(r' ', '', content_text)
        content_text = re.sub(r'\n+', '\n', content_text).strip()
        logger.info(content_text)

        if '会计师事务所整改的公告' in announcement_title or '投诉处理决定书' in announcement_title:
            if '会计师事务所整改的公告' in announcement_title:
                facts = re.search(r'(根据.*?报备情况，经我厅核实.*?(不符合.*?第.*?条.*?规定的(设立)?条件。|已被省注册会计师协会撤销（注销）注册。))',
                                  content_text).group(1).strip()
                punishment_decision = re.search(r'(根据.*?第.*?条.*?规定，责令[\s\S]*?)(逾期仍未达到设立条件的)',
                                                content_text).group(1).strip()
                try:
                    publish_date = re.findall(r'\n\d{4}年 *?\d{1,2} *?月 *?\d{1,3} *?日\n', content_text)[-1].strip()
                    real_publish_date = format_date(publish_date.replace(' ', ''))
                except Exception as e:
                    logger.warning(e)
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)

                if len(content_soup.find(id='level3_content').find_all('table')) > 0:
                    litigant_table = content_soup.find(id='level3_content').find('table')
                    litigant_table_list = table_to_list(litigant_table)
                    for index, each_litigant in enumerate(litigant_table_list):
                        if index < 2:
                            continue
                        result_map = {
                            'announcementTitle': '关于要求' + each_litigant[2].strip() + '整改的公告',
                            'announcementOrg': '山东财政厅',
                            'announcementDate': real_publish_date,
                            'announcementCode': '',
                            'facts': facts,
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': '事务所代码：' + each_litigant[1].strip() + '\n' + '事务所名称：' +
                                        each_litigant[2].strip(),
                            'punishmentBasement': '',
                            'punishmentDecision': punishment_decision,
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        logger.info(result_map)
                        if db.announcement.find(
                                {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                            db.announcement.insert_one(result_map)
                            logger.info('山东省财政厅 数据解析 ' + ' -- 数据导入完成')
                        else:
                            logger.info('山东省财政厅 数据解析 ' + ' -- 数据已经存在')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('山东省财政厅 数据解析 ' + ' -- 修改parsed完成')
                else:
                    litigant = re.search(r'^(.*?)\n', content_text).group(1).strip()
                    result_map = {
                        'announcementTitle': '关于要求' + (litigant[:-1] if litigant[-1] == '：' else litigant) + '整改的公告',
                        'announcementOrg': '山东财政厅',
                        'announcementDate': real_publish_date,
                        'announcementCode': '',
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                        'punishmentBasement': '',
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    logger.info(result_map)
                    if db.announcement.find(
                            {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                        db.announcement.insert_one(result_map)
                        logger.info('山东省财政厅 数据解析 ' + ' -- 数据导入完成')
                    else:
                        logger.info('山东省财政厅 数据解析 ' + ' -- 数据已经存在')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('山东省财政厅 数据解析 ' + ' -- 修改parsed完成')
            else:
                announcement_code = re.search(r'(鲁财采.\d{4}.\d+号) ?\n', content_text).group(1).strip()
                litigant = re.search(r'(被投诉人[\s\S]*?)(投诉人因不满意)', content_text).group(1).strip()
                facts = re.search(r'投诉处理决定书([\s\S]*?)(被投诉人)',
                                  content_text).group(1).strip() + '\n' + re.search(r'(投诉人因不满意[\s\S]*?)(根据投诉事项)',
                                                                                    content_text).group(1).strip()
                punishment_basis = re.search(r'(根据投诉事项，本机关调阅审查[\s\S]*?)(综上，本机关认定)',
                                             content_text).group(1).strip()
                punishment_decision = re.search(r'(综上，本机关认定[\s\S]*?)(投诉人对本投诉处理决定不服)',
                                                content_text).group(1).strip()
                try:
                    publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
                    real_publish_date = format_date(publish_date)
                except Exception as e:
                    logger.warning(e)
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)
                result_map = {
                    'announcementTitle': '投诉处理决定书（' + announcement_code + '）',
                    'announcementOrg': '山东财政厅',
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('山东省财政厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('山东省财政厅 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('山东省财政厅 数据解析 ' + ' -- 修改parsed完成')
        else:
            announcement_code = re.search(r'(鲁财罚决字.\d{4}.\d+号) ?\n', content_text).group(1).strip()
            if re.search(r'处罚决定书\n(.*?)\n', content_text):
                litigant = re.search(r'处罚决定书\n(.*?)\n', content_text).group(1).strip()
            else:
                litigant = re.search(r'鲁财罚决字.\d{4}.\d+号\n(.*?)\n', content_text).group(1).strip()
            facts = re.search(r'((经查|经审查|本机关在.*?(工作|投诉过程)中，发现|'
                              r'根据《财政部关于开展2015年度全国政府采购代理机构监督检查工作的通知》|'
                              r'根据.*?鲁财监.\d{4}.\d+号.*?要求)'
                              r'[\s\S]*?)'
                              r'(本机关认为|(上述|该)行为违反了|你作为上述审计报告的签字注册会计师，对上述问题负有审计责任|'
                              r'根据.*?第.*?条规定)',
                              content_text).group(1).strip()
            punishment_basis = re.search(r'((本机关认为.*?违反了|(上述|该)行为违反了|你作为上述审计报告的签字注册会计师，对上述问题负有审计责任|'
                                         r'该行为违反了.*?第.*?条规定。)'
                                         r'[\s\S]*?)'
                                         r'(本机关于.*?下达了|根据.*?第.*?条.*?规定|以上事实有.*?在案佐证。|'
                                         r'\d+年\d+月\d+日，我厅已将《行政处罚事先告知书》.*?送达)',
                                         content_text).group(1).strip()
            punishment_decision = re.search(r'((现依据|根据).*?(第.*?条|第三十一条、第三十\n九条).*?规定，'
                                            r'(决定对.*?作出以下行政处罚|本机关决定对.*?处以|现\n?决定对.*?处以|决\n?定对你处以|'
                                            r'决定对.*?给予|本机关对.*?给予|我厅决定对.*?处以|我厅拟对.*?处以|'
                                            r'责令.*?改正)'
                                            r'[\s\S]*?)\n(履行方式和期限|如不服本(处罚)?决定)',
                                            content_text).group(1).strip()
            if '申辩' in content_text:
                defense_opinion = re.search(r'(在规定的期限内.*?未提出陈述、申辩。)', content_text).group(1).strip()
                defense_response = ''
            else:
                defense_opinion = defense_response = ''
            try:
                publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.warning(e)
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '山东财政厅',
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': defense_opinion,
                'defenseResponse': defense_response,
                'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('山东省财政厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('山东省财政厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('山东省财政厅 数据解析 ' + ' -- 修改parsed完成')


# 河南省财政厅
def henanczt_parse():
    for each_document in db.finance_data.find({'origin': '河南省财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        if '行政处罚目录' in announcement_title or announcement_url in ['http://www.hncz.gov.cn/2016/0617/3392.html',
                                                                  'http://www.hncz.gov.cn/2016/0513/3390.html']:
            logger.warning('url has nothing to do with punishment ...')
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update finance data success')
            continue

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if len([kk.attrs.get('href', '') for kk in content_soup.find(class_='newstxt').find_all('a')
                if str(kk.attrs.get('href', '')).endswith('doc')]) > 0:
            doc_link = urljoin(announcement_url, content_soup.find(class_='newstxt').find('a').attrs['href'])
            link_type = doc_link.split('.')[-1]
            response = request_site_page(doc_link)
            if response is None:
                logger.error('网页请求错误')
                return '', ''
            with open('./test/tmp.' + link_type, 'wb') as out_file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        out_file.write(chunk)
            if link_type == 'doc' or link_type == 'wps':
                shell_str = config['soffice']['soffice_path'] + ' --headless --convert-to docx ' + \
                            './test/tmp.' + link_type + ' --outdir ./test/'
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.wait()
            with open('./test/tmp.docx', 'rb') as docx_file:
                docx_content = docx_file.read()
            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': doc_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': doc_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'docx',
                    'oss_file_name': announcement_title,
                    'oss_file_content': docx_content,
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': doc_link})['_id']
            doc = docx.Document('./test/tmp.docx')
            result_text_list = []
            for para in doc.paragraphs:
                result_text_list.append(para.text)
            content_text = '\n'.join(result_text_list)
            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.doc'):
                os.remove('./test/tmp.doc')
            if os.path.exists('./test/tmp.docx'):
                os.remove('./test/tmp.docx')
        else:
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text,
                    'parsed': False,
                    'content_class_name': 'newstxt'
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             content_response.text)
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']

            for tag in content_soup("style"):
                tag.decompose()
            content_text = '\n'.join(
                [kk.text.strip() for kk in content_soup.find(class_='newstxt').findChildren(recursive=False)])
            content_text = re.sub(r'\n+', '\n', content_text)

        announcement_code_compiler = re.compile(r'((豫财监|豫财会|豫财购罚).\d{4}.第?\d+ ?号)')
        if announcement_code_compiler.search(announcement_title):
            announcement_code = announcement_code_compiler.search(announcement_title).group(1).strip()
        elif announcement_code_compiler.search(content_text):
            announcement_code = announcement_code_compiler.search(content_text).group(1).strip()
        else:
            announcement_code = ''

        if re.search('(撤回.*?设立许可的决定|撤销.*?资产评估资格的通知)', announcement_title):
            announcement_type = '撤回设立许可'
            litigant = re.search(announcement_code + r'\n(.*?：)', content_text).group(1).strip()
            facts = re.search(litigant.replace(r'(', r'\(').replace(r')', r'\)') +
                              r'([\s\S]*?)((依据|根据).*?第.*?条.*?规定，我厅决定)', content_text).group(1).strip()
            defense_opinion = defense_response = ''
            punishment_basis = ''
            punishment_decision = re.search(r'((依据|根据).*?第.*?条.*?规定，我厅决定[\s\S]*?)(如对上述决定有异议|请接到《行政处罚决定书》10日内)',
                                            content_text).group(1).strip()
        else:
            announcement_type = '行政处罚决定'
            if re.search(r'((豫财监|豫财会|豫财购罚).\d{4}.第?\d+ ?号\n)'
                         r'([\s\S]*?)'
                         r'\n(.*?(本机关|本厅|我厅|我们|厅于).*?进行了(检查|调查核实|专项检查))', content_text):
                litigant = re.search(r'((豫财监|豫财会|豫财购罚).\d{4}.第?\d+ ?号\n)'
                                     r'([\s\S]*?)'
                                     r'\n(.*?(本机关|本厅|我厅|我们|厅于).*?进行了(检查|调查核实|专项检查))',
                                     content_text).group(3).strip()
            else:
                litigant = re.search(r'^([\s\S]*?)'
                                     r'(\n(.*?(本机关|本厅|我厅|我们|厅于).*?进行了(检查|调查核实|专项检查))|'
                                     r'根据《河南省财政厅关于开展2014年会计师事务所执业质量检查的通知》（豫财监〔2014〕175号）|'
                                     r'2014年12月消费日报网报道|'
                                     r'你公司在2014年6月参与河南省产品质量监督检验院检测设备采购项目|'
                                     r'你公司在2014年9月参与河南省医疗器械检验所仪器设备采购项目|'
                                     r'你公司在2015年6月19日洛阳师范学院伊滨校区二期工程)',
                                     content_text).group(1).strip()

            facts = re.search(r'(处罚决定通知如下：|在法定时间内未进行陈述和申辩。|联系电话：\d+-\d+)'
                              r'([\s\S]*?)((上述|以上)(行为|事项|行业)?违反了|在调查过程中，你公司拒不提供投标人河南轻工进出口公司的投标文件|'
                              r'以上事实有.*?(材料为证。|说明为证，事实清楚。))',
                              content_text).group(2).strip()
            punishment_basis = re.search(r'(((上述|以上|你公司的)(行为|事项|行业)?违反了|在调查过程中，你公司拒不提供投标人河南轻工进出口公司的投标文件)'
                                         r'[\s\S]*?(规定|隐瞒审计中发现的问题)[\s\S]*?)'
                                         r'(根据|(结合|按照).*?违法行为的事实|经查，你单位为省政府直属事业单位|但在调查过程中|'
                                         r'依法拟对你公司做出如下处罚决定|依法对你公司做出如下处罚决定)',
                                         content_text).group(1).strip()
            punishment_decision = re.search(r'((((结合|按照).*?违法行为的事实.*?|经查，你单位为.*?)?'
                                            r'根据.*?第.*?条.*?规定.*?'
                                            r'(我厅决定|本机关决定|责令.*?立即改正|此情节认定|我厅拟对|依法拟对.*?做出如下处罚决定)|'
                                            r'你公司有主动配合情节，依法拟对你公司|依法拟对你公司做出如下处罚决定|'
                                            r'依法对你公司做出如下处罚决定)'
                                            r'[\s\S]*?)'
                                            r'(本决定自送达之日起生效|.*?如不服本决定)',
                                            content_text).group(1).strip()
            if '申辩' in content_text:
                defense_opinion = re.search(r'(你(单位|所)?在法定时间内提出的?陈述和申辩|'
                                            r'你(单位|所)?在法定(时间|期限)内(未进行|进行了|未提出)陈述和申辩(，未要求听证。)?|'
                                            r'你(单位|所)?在法定时间内以书面形式进行了陈述和申辩，未要求听证。|'
                                            r'你(单位|所)?在法定期限内提出了减轻处罚的(陈述和)?申辩。|'
                                            r'你(单位|所)?未在法定时间内提出陈述和申辩。|'
                                            r'你(单位|所)?在法定时间内未提出的陈述和申辩意见|'
                                            r'你(单位|所)?提出的陈述和申辩意见)',
                                            content_text).group(1).strip()
                try:
                    defense_response = re.search(r'(我厅不予采纳。|'
                                                 r'我厅对陈述和申辩的合理部分予以采纳|'
                                                 r'我厅对你(单位|所)?的陈述和申辩不予采纳|'
                                                 r'我厅采纳了合理部分。)',
                                                 content_text).group(1).strip()
                except Exception as e:
                    logger.warning(e)
                    defense_response = ''
            else:
                defense_opinion = defense_response = ''

        try:
            publish_date = re.findall(r'.{4}年.{1,2}月.{1,3}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.warning(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '河南财政厅',
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': announcement_type,
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('河南省财政厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('河南省财政厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('河南省财政厅 数据解析 ' + ' -- 修改parsed完成')


# 深圳市财政委员会
def szczt_parse():
    for each_document in db.finance_data.find({'origin': '深圳市财政局', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        if content_response.status_code == 404:
            logger.error('网页请求404 %s' % announcement_url)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update finance data success')
            continue
        if 'GB' in content_response.apparent_encoding:
            content_response.encoding = 'GB18030'
        else:
            content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')
        announcement_title = content_soup.find(class_='tit').find('h1').text.strip() \
            .replace('�z', '﹝').replace('�{', '﹞')

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_class_name': 'news_cont_d_wrap'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_text = get_content_text(content_soup.find(class_='news_cont_d_wrap')) \
            .replace('�z', '﹝').replace('�{', '﹞').replace('�|', '磡')

        announcement_code_compiler = re.compile(r'((福财罚决字|罗财书|深宝财购|深龙财|深财书|深坪财书|'
                                                r'深财购函|深福罚决字|深龙华财书|深南财行决|深鹏发财|深坪发财罚字|罗财)'
                                                r'.\d{4}.第?\d+ ?号? ?)')
        if announcement_code_compiler.search(announcement_title):
            announcement_code = announcement_code_compiler.search(announcement_title).group(1).strip()
        elif announcement_code_compiler.search(content_text):
            announcement_code = announcement_code_compiler.search(content_text).group(1).strip()
        else:
            announcement_code = ''
        litigant_compiler = re.compile(r'(当事人[\s\S]*?)'
                                       r'\n(本(机关|单位)于.*?对.*?立案调查|经查|'
                                       r'.*?(参加|参与).*?(政府采购活动|招投标活动|招标投标|招标活动|投标|由深圳市华昊信息技术有限公司|采购过程|网上竞价项目))')
        if litigant_compiler.search(content_text):
            litigant = litigant_compiler.search(content_text).group(1).strip()
        else:
            litigant = re.search(r'^(.*?\d{4}.\d+号\n)?(.*?)\n', content_text).group(2).strip()

        facts = re.search(litigant.replace(r'(', r'\(').replace(r')', r'\)')
                          + r'\n' + r'([\s\S]*?)(以上事实有.*?等(证据证实|材料为证|为证)|本机关依法向.*?送达|'
                                    r'[^。\n]*?行为属于.*?第.*?条.*?情形|(我委|我局).*?向.*?(发出|送达)|'
                                    r'综上，根据《深圳经济特区政府采购条例》第五十七条第一款第（五）项的规定|'
                                    r'上述事实已构成.*?第.*?条.*?串通投标行为。|'
                                    r'综上.根据.*?第.*?条.*?规定|'
                                    r'根据.*?第五十七条|'
                                    r'我委已于2015年5月22日将《行政处罚告知书》（深财书〔2015〕105号）送达至地景公司，地景公司未在法定期限内提出申辩。|'
                                    r'根据《中华人民共和国政府采购法》第七十七条第一款、第一款第（一）项及第二款的规定|'
                                    r'据此，我委于.*?向.*?印发了《行政处罚告知书》|'
                                    r'你公司上述行为存在隐瞒真实情况，提供虚假资料的情形|'
                                    r'你公司.*?行为违反了|根据.*?第.*?条.*?规定|'
                                    r'你公司采取隐瞒真实情况，通过虚假承诺的手段，试图骗取政府采购项目中标的行为，违反了有关规定|'
                                    r'依照《中华人民共和国政府采购法》第七十七条第（一）款的规定)',
                          content_text).group(1).strip()

        if '申辩' in content_text:
            defense_opinion = re.search(r'([^。\n]*?(提出|提交|递交).*?(陈述[、和]?|书面)?申辩的?(意见|材料)|'
                                        r'[^。\n]*?(提交申辩意见、提出听证申请。|亦作出书面申辩。)|'
                                        r'你单位于2018年8月16日向本单位提起听证申请，并于2018年8月30日申请听证延期。|'
                                        r'深圳润和恒达公司在法定期间向本机关提交了《说明函》。|'
                                        r'你单位于2018年6月8日向本单位提起听证申请|市环保协会未在规定期限内提出陈述和申辩。|'
                                        r'你单位未在规定时间内向本单位提起陈述、申辩|'
                                        r'[^。\n]*?(在(规定期间|法定期间))?'
                                        r'向(本机关|我局|我委)提交了?(书面)?《(陈述意见|说明函|申述书|申辩函|申辩书|行政处罚申辩书|申诉函|陈述书)》|'
                                        r'澳芝川公司未提出申辩。|[^。\n]*?未在规定时间内向我局进行陈述、申辩。|'
                                        r'我局已向.*?发出.*?未提出申辩。|北京海鑫公司向我局提交《陈述申辩书》。|'
                                        r'我局已向.*?发出.*?于.*?签收，(未向我局提出书面申辩书。|未提出陈述和申辩。)|'
                                        r'我局已向.*?发出.*?于.*?签收.*?向我局(提交|提出)书面(陈述|申辩)[函书]|'
                                        r'本机关依法作出深财书\[2017\]14号《行政处罚告知》并送达盛世传奇公司，盛世传奇公司亦作出书面申辩|'
                                        r'[^。\n]*?未在规定期限内提出陈述和申辩，视为放弃陈述、申辩的权利。|'
                                        r'本机关已作出《行政处罚告知》（深财书\[2017\]58号）并送达至伟兴业公司，伟兴业公司作出书面申辩。|'
                                        r'[^。\n]*?在法定期间未向我局提交有关陈述和申辩的材料，视为放弃陈述和申辩的权利。|'
                                        r'欣邦公司在法定期间内向我局提交了《关于深圳市龙岗区财政局行政处罚公告函告知书的回复》|'
                                        r'江西瑞浩公司在法定期限内未提出陈述和申辩。|2017年2月24日,鹏运公司向我委提出书面陈述和申辩。|'
                                        r'我委已于1月5日向林邦公司发出了深财书〔2017〕4号行政处罚告知书，林邦公司于1月18日向我委提出申辩。|'
                                        r'[^。\n]*?于.*?向我委提出陈述和申辩。|'
                                        r'我委于2014年2月7日向你公司发出了《行政处罚告知书》（深财书\[2014\]37号），并收到了你公司《请求从轻处罚的申辩》|'
                                        r'碧韵达公司在申辩书中称涉案检验报告系生产厂商提供，但并没有提供该申辩理由的相关证据|'
                                        r'我局于.*?收到.*?书面申辩意见。|你公司在规定时间内提交了《申辩函》|'
                                        r'其后.*?未在规定时间内书面提出陈述和申辩，我委对该采购项目的调查已终结。|'
                                        r'我委已向.*?于.*?签收.*?(向我委提交申辩书。|但未在申辩期限内提出申辩，视为放弃陈述、申辩的权利。)|'
                                        r'我委于.*?发出.*?于.*?收到.*?提交的《行政处罚陈述申辩书》。|'
                                        r'[^。\n]*?于.*?向我委提交申辩函。|我委于.*?向你公司发出了.*?并收到了你公司的《申辩书》。|'
                                        r'我委.*?于.*?依法向.*?送达.*?于.*?收到.*?提交的.*?。|'
                                        r'我委已于2015年5月22日将《行政处罚告知书》（深财书〔2015〕105号）送达至地景公司，地景公司未在法定期限内提出申辩。|'
                                        r'2014年12月30日，长城交通公司向我委提交了《申辩及陈述说明》。在《申辩及陈述说明》中，长城交通公司没有提供新的证明材料。|'
                                        r'我委依法向你公司发出了?行政处罚告知书.*?于.*?收到.*?(《行政处罚陈述申辩书》|《行政处罚申辩与陈述》)|'
                                        r'2014年9月22日，鸿华通公司向我委提交了本项目申辩书。2014年10月11日，根据鸿华通公司申请，我委组织了行政处罚听证。|'
                                        r'我委于.*?向你公司发出.*?并收到了你公司的?(《申辩》|《请求从轻处罚的申辩》)。|'
                                        r'之后你公司申辩认为，你公司在投标文件中提交的虚假资料是厂商向你公司提供，你公司不知情。|'
                                        r'我局于2014年4月2日向你公司发出《行政处罚告知书》（深龙财〔2014〕57号），在规定时间内我局未收到你公司书面陈述和申辩意见。|'
                                        r'我[委局]于.*?向.*?发出了.*?并收到了.*?的(《律师意见函》|《事实陈述书》|《关于福田区财政局行政处罚陈述和申辩书》)。|'
                                        r'我委于.*?向你公司发出了.*?你公司并未申辩|你公司提出的不签约理由是：工程量清单不清晰，存在“安全措施”“施工图纸”等诸多问题。|'
                                        r'我委于.*?向你公司发出了.*?并收到了你公司的《申辩函》|[^。\n]*?在签收后.*?提出了书面申辩。|'
                                        r'我局于.*?向你单位发出了.*?你单位未提出陈述申辩。|你公司逾期未向本机关提出申辩、说明，本机关视为你公司放弃陈述、申辩的权利。|'
                                        r'你公司逾期未向本机关书面提出陈述和申辩,本机关视为你公司放弃陈述、申辩的权利。|'
                                        r'我委于2013年6月8日向你公司发出了《行政处罚告知书》（深财书\[2013\]32号），'
                                        r'并收到了你公司的书面答复。你公司承认相关证书虚假，但表示是代理办证公司所为，并未能提出有力的申辩理由和事实证据。|'
                                        r'我委于2015年3月11日依法向三绿园林公司发出《深圳市财政委员会行政处罚告知书》（深财书〔2015〕41号），'
                                        r'并于2015年3月24日收到三绿园林公司提交的《关于<深圳市财政委员会行政处罚告知书>（深财书〔2015〕41号）的申辩函》。)',
                                        content_text).group(1).strip()
            try:
                defense_response = re.search(r'((经审查|经听证)，本机关认为.*?申辩理由不成立。|'
                                             r'本机关(组织听证并)?经研究后,认为.*?申辩意见不成立。|'
                                             r'本机关依法举行了听证会.经审查.本机关认为.*?申辩(意见|理由)不成立。|'
                                             r'经研究.其?申辩(理由)?不成立。|经听证，你单位申辩、陈述理由不成立。|'
                                             r'经研究.*?申辩(理由)?不成立.(本机关|本局|我局)(决定)?不予采纳。|'
                                             r'经研究.*?申辩意见不成立。.*?未向本机关申请举行听证,视为放弃听证权利。|'
                                             r'经研究，本机关对广州康琼公司的说明不予支持。|'
                                             r'经研究，我局决定不予采纳。|经研究.*?申辩理由均不能成立。|'
                                             r'(我局|本机关)(依法)?组织听证并研究后.认为.*?申辩意见不成立。|'
                                             r'经研究，佳佳顺公司的申辩意见不成立。佳佳顺公司未向本机关申请举行听证，视为放弃听证权利。|'
                                             r'经研究.*?申辩(理由)?不成立(.*?决定不予采纳)?。|'
                                             r'经研究.*?我局决定不予采纳。|经研究，我局认为贵司的申辩意见不成立|'
                                             r'经调查研究，我委不支持.*?申辩理由|经研究，我委认定你公司的申辩情况不真实、理由不成立。|'
                                             r'经组织专家论证，并认真分析申辩理由，认为申辩理由不充分，不予采纳，现认定“串通投标”的违规行为属实。)',
                                             content_text).group(1).strip()
            except Exception as e:
                logger.warning(e)
                defense_response = ''
        else:
            defense_opinion = defense_response = ''

        punishment_basis_str_list = [
            r'([^\n。；]*?)(问题|行为|事项|情况|事实|属于|情形)([^\n。；\s]*?)(违反|不符合|构成|属于).*?(违法情形|行为|规定)',
        ]
        punishment_basis_str = '|'.join(punishment_basis_str_list)
        punishment_basis_compiler = re.compile(r'[。\n；]' + '(' + punishment_basis_str + ')' +
                                               '(.本机关.*?送达|.本单位已作出|.\n?我局已(依法)?向.*?(发出|送达))', re.MULTILINE)
        punishment_basis_list = punishment_basis_compiler.findall(content_text)
        punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

        punishment_basis_laws_compiler = re.compile(r'(附(相关)?法律条文[\s\S]*?)$')
        if punishment_basis_laws_compiler.search(content_text):
            punishment_basis = punishment_basis + '\n' + \
                               punishment_basis_laws_compiler.search(content_text).group(1).strip()
        else:
            pass

        punishment_decision = re.search(
            r'(\n(.*?第.*?条规定.*?(根据上述规定和本案调查的事实.*?本机关决定|结合本项目实际情况，我委决定对|结合本项目实际情况，我委认定)[\s\S]*?)|'
            r'根据.*?第.*?条.*?作出(以下|如下)(行政)?处罚[\s\S]*?|'
            r'根据.*?第.*?条.*?作出(行政)?处罚(以下|如下)[\s\S]*?|'
            r'根据.*?第.*?条.*?违规行为处罚如下[\s\S]*?|'
            r'根据.*?第.*?条.*?决定对.*?(处罚如下|处以如下处罚)[\s\S]*?|'
            r'根据.*?第.*?条.*?对.*?处罚如下[\s\S]*?|'
            r'根据.*?第.*?条.*?'
            r'(我局决定|我局拟对|本机关决定|我委决定|决定对.*?作出如下处罚|我委对.*?给予警告处理|我局对.*?作出如下行政处罚|'
            r'对.*?给予如下处罚|决定对你公司进行以下处罚|并决定对你公司处罚如下|决定如下|我委对.*?处以|现决定中止|我委决定取消)[\s\S]*?|'
            r'我局根据.*?第.*?条.*?规定，(决定将|决定对)[\s\S]*?|'
            r'经研究，现决定对你公司给予如下处罚[\s\S]*?|'
            r'依照《中华人民共和国政府采购法》第七十七条第（一）款的规定，本局决定[\s\S]*?|'
            r'在法定的暂停期限结束前或者我委发出恢复采购活动通知前，不得进行该项目的采购活动。)'
            r'\n(.*?如不服本(处罚)?决定|.*?应当自收到本决定书之日|.*?如对上述行政处罚决定不服|.*?如对上述决定不服|本决定书?自送达之日起生效|'
            r'.*?如果对上述决定不服|.*?如对本决定不服|此函|深圳市财政委员会|此书)',
            content_text).group(1).strip()

        try:
            publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.warning(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '深圳财政局',
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('深圳市财政局 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('深圳市财政局 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('深圳市财政局 数据解析 ' + ' -- 修改parsed完成')


# 厦门市财政局
def xiamenczt_parse():
    for each_document in db.finance_data.find({'origin': '厦门市财政局', 'status': {'$nin': ['ignored']}}):

        announcement_title = each_document['title']
        announcement_url = each_document['url']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)

        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue

        if content_response.status_code == 404:
            logger.error('网页请求404 %s' % announcement_url)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update finance data success')
            continue

        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_id_name': 'newsview'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_text_list = []
        for each_p in content_soup.find(id="newsview").children:
            if each_p.name is None:
                continue
            content_text_list.append(each_p.text.strip().replace('\n', '').replace(' ', '').replace('\xa0', ' '))
        content_text = '\n'.join(content_text_list)
        content_text = re.sub(r'\n+', r'\n', content_text).strip()
        logger.info(content_text)

        announcement_code = re.search(r'(厦财.*?\d{4}.\d+号)', content_text).group(1).strip()

        if '采购投诉处理决定书' in announcement_title:
            litigant = re.search(r'(被投诉人[\s\S]*?)(投诉人.*?参与[\s\S]*?过程中)', content_text).group(1).strip()
            facts = re.search(
                r'(投诉人[\s\S]*?)被投诉人', content_text).group(1).strip() + '\n' + re.search(
                r'(投诉人.*?参与[\s\S]*?过程中[\s\S]*?)'
                r'(\n.*?被投诉人.*?(认为|称)|'
                r'根据《政府采购质疑和投诉办法》（财政部令第94号）第三十条，决定如下)',
                content_text).group(1).strip()
            try:
                defense = re.search(r'\n(.*?被投诉人.*?(认为|称)[\s\S]*?)'
                                    r'((四、)?根据以上情况.*?我局认为|'
                                    r'综上，我局认为)',
                                    content_text).group(1).strip()
            except Exception as e:
                logger.warning(e)
                defense = ''
            try:
                punishment_basis = re.search(r'(((四、)?根据以上情况.*?我局认为|综上，我局认为)'
                                             r'[\s\S]*?)'
                                             r'\n((综上，)?(根据|依据).*?第.*?条)',
                                             content_text).group(1).strip()
            except Exception as e:
                logger.warning(e)
                punishment_basis = ''
            punishment_decision = re.search(r'((\n(综上，)?(根据|依据).*?第.*?条.*(决定如下)|'
                                            r'根据《政府采购质疑和投诉办法》（财政部令第94号）第三十条，决定如下)'
                                            r'[\s\S]*?)'
                                            r'(如不服本决定)',
                                            content_text).group(1).strip()
            announcement_type = '投诉处理决定'
        elif '采购监督意见书' in announcement_title:
            litigant = re.search(announcement_code + r'\n(.*?)\n', content_text).group(1).strip()
            facts = re.search(r'\n((经核查|经调查)[\s\S]*?)(二、监督意见|根据.*?等相关规定，现对你司提出以下监督意见)',
                              content_text).group(1).strip()
            defense = ''
            punishment_basis = ''
            try:
                punishment_decision = re.search(r'二、监督意见\n([\s\S]*?)\n三、整改要求',
                                                content_text).group(1).strip() + '\n' + re.search(
                    r'三、整改要求\n([\s\S]*?)\n厦门市财政局', content_text).group(1).strip()
            except Exception as e:
                logger.warning(e)
                punishment_decision = re.search(r'(根据.*?等相关规定，现对你司提出以下监督意见[\s\S]*?)\n厦门市财政局',
                                                content_text).group(1).strip()
            announcement_type = '监督意见'
        else:
            litigant = re.search(r'((当事人|违法行为人)[\s\S]*?)'
                                 r'(2009年5月15日\n?，你司通过市发改委以生物法制取|'
                                 r'我局于\n?2015年\n?7月17日\n?对当事人涉嫌提供虚假材料谋取中标的行为予以立案调查。|'
                                 r'当事人在参与“柴油发电机组”项目（采购编号：XM2010-DZ0089）政府采购投标活动中|'
                                 r'我局于2018年11月21日对违法行为人应当采用公开招标方式而擅自采用其他方式采购的行为予以立案调查)',
                                 content_text).group(1).strip()
            facts = re.search(litigant + r'([\s\S]*?)'
                                         r'(具体有专项审计报告、会计凭证复印件、谈话笔录等证据为凭。|'
                                         r'具体有.*?等证据为凭。|'
                                         r'以上事实有.*?为证。)',
                              content_text).group(1).strip()
            try:
                defense = re.search(r'(当事人在法定期限内未提出书面陈述、申辩意见。|'
                                    r'本局于2010年9月13日告知了当事人享有的陈述、申辩的权利，在规定的截止日期前当事人未向本局提交书面申诉意见。)',
                                    content_text).group(1).strip()
            except Exception as e:
                logger.warning(e)
                defense = ''
            punishment_basis = re.search(r'([^\n。]*?上述行为违反了.*?第.*?条.*?规定。|'
                                         r'[^\n。]*?上述行为属于.*?第.*?条.*?违法情形(，依法应给予行政处罚)?。)',
                                         content_text).group(1).strip()
            punishment_decision = re.search(r'([^\n。]*?(依据|根据).*?第.*?条.*?规定[\s\S]*?)'
                                            r'\n?([^\n。]*?如不服本处罚决定|[^\n。]*?如对本处罚决定不服)',
                                            content_text).group(1).strip()
            announcement_type = '行政处罚决定'

        publish_date = re.findall(r'\d{4} ?年\d{1,2}月\d{1,3}日', content_text)[-1].strip().replace(' ', '')
        real_publish_date = format_date(publish_date)

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '厦门财政局',
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense,
            'defenseResponse': '',
            'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': announcement_type,
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('厦门市财政局 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('厦门市财政局 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('厦门市财政局 数据解析 ' + ' -- 修改parsed完成')


# 福建省财政厅
def fujianczt_parse():
    for each_document in db.finance_data.find({'origin': '福建省财政厅', 'status': {'$nin': ['ignored']}}):

        announcement_title = each_document['title']
        announcement_url = each_document['url']

        # 判断是否解析过
        if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)

        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue

        if content_response.status_code == 404:
            logger.error('网页请求404 %s' % announcement_url)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update finance data success')
            continue

        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_class_name': 'czxw-ny'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text)
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']
        if content_soup.style:
            content_soup.style.decompose()
        if 'http://czt.fujian.gov.cn/zfxxgkzl/zfxxgkml/xzzf/xzcf/' in announcement_url:
            content_text_list = []
            if content_soup.find(class_="czxw-ny").find(class_='Section0'):
                for each_p in content_soup.find(class_="czxw-ny").find(class_='Section0').children:
                    if each_p.name is None:
                        continue
                    content_text_list.append(
                        each_p.text.strip().replace('\n', '').replace(' ', '').replace('\xa0', ' '))
                content_text = '\n'.join(content_text_list)
                content_text = re.sub(r'\n+', r'\n', content_text).strip()
            else:
                if content_soup.find(class_='czxw-ny').style:
                    content_soup.find(class_='czxw-ny').style.decompose()
                content_text = get_content_text(content_soup.find(class_='czxw-ny'))
        else:
            content_text = get_content_text(content_soup.find(class_='czxw-ny'))
            content_text = re.sub(r'\n+', r'\n', content_text).strip()

        announcement_code = each_document['code'].strip()
        litigant = re.search(r'(当事人[\s\S]*?|厦门德诚会计师事务所有限公司（以下简称“德诚所”）|'
                             r'双悦（福建）动力机械有限公司：|'
                             r'厦门呈祥源联合会计师事务所：|'
                             r'中国注册会计师韩萍、董如山：|'
                             r'中国注册会计师陈福尧：|'
                             r'中国注册会计师周冬华：)'
                             r'(我局收到|经调查核实|'
                             r'爱玛客服务产业（中国）有限公司对泉州市政府采购中心组织的关于|'
                             r'根据平潭综合实验区财政金融局财政检查通知书.*?进行检查|'
                             r'你.*?参与了|我局在投诉处理过程中发现|'
                             r'本机关依法调查|你公司在代理|'
                             r'2009年5月15日，你司通过市发改委以生物法制取30000吨木糖醇|'
                             r'\n.*?(根据|依据).*?有关规定|'
                             r'本机关在开展“1\+X”专项督查中|'
                             r'本机关根据举报材料反映的情况并经查实|'
                             r'你公司于\n?\d{4}年\d{1,2}月\d{1,3}日参加了|'
                             r'2016年1-6月执业质量情况进行监督检查。|'
                             r'本机关在依法对福建省红十字会帐篷|'
                             r'福建省审计厅2015年7月对政府采购业务进行审计时|'
                             r'\n我厅向你单位下达的|'
                             r'福建省森林资源管理总站委托福建省公共资源交易中心于2015年2月4日发布|'
                             r'\n因你所在2013年1-6月开展业务期间|'
                             r'\n因你对.*?负有责任|'
                             r'经举报查实|'
                             r'本机关根据举报查实)',
                             content_text).group(1).strip()
        facts = re.search(
            litigant.replace(r')', r'\)') + r'([\s\S]*?)'
                                            r'(上述(违法)?事实.有.*?予以(证实|证明)。|以上事实有下列证据证实|上述行为违反了|'
                                            r'根据《福建省财政厅关于盘活财政存量资金有关事项的通知》|'
                                            r'你单位存在提供虚假材料谋取招标、成交的情形|'
                                            r'(现)?根据.*?第.*?条.*?规定.*?(决定|行政处罚)|'
                                            r'上述违法事实，有本机关《福建省财政厅关于省贸促会“1\+X”|'
                                            r'上述事实，有福建工业学校的关于中标方自愿放弃中标资格的《说明函》|'
                                            r'二、行政处罚|'
                                            r'本机关认为|'
                                            r'根据《中华人民共和国行政强制法》第五十四条规定|'
                                            r'上述行为违反了《中华人民共和国注册会计师法》第二十一)',
            content_text).group(1).strip()
        if '申辩' in content_text:
            try:
                defense = re.search(r'(2018年6月19日，你公司提交《行政处罚陈述申辩书》，未要求举行听证。|'
                                    r'你公司未提出陈述、申辩意见。)',
                                    content_text).group(1).strip()
            except Exception as e:
                logger.warning(e)
                defense = ''
            defense_response = ''
        else:
            defense = defense_response = ''
        try:
            punishment_basis = re.search(r'((我局认为|本机关认为|[^。\n]*?上述行为违反了|'
                                         r'上述行为属于|你单位存在提供虚假材料谋取招标、成交的情形)'
                                         r'[\s\S]*?)'
                                         r'((现)?(根据|按照|依据).*?第.*?条.*?规定.*?(决定|行政处罚|处以如下处罚)|'
                                         r'以上违法事实，有.*?等证据为证。|'
                                         r'依据.*?第十八条有关规定|'
                                         r'上述违法事实，有.*?予以证实。|'
                                         r'上述违法事实，有检查报告、检查工作底稿和当事人笔录等)',
                                         content_text).group(1).strip()
        except Exception as e:
            logger.warning(e)
            punishment_basis = ''
        punishment_decision = re.search(r'(((现)?(根据|按照|依据).*?第.*?条.*?规定.*?(决定|行政处罚|处以如下处罚)|'
                                        r'根据《福建省财政厅关于盘活财政存量资金有关事项的通知》|'
                                        r'依据.*?第十八条有关规定|'
                                        r'根据《中华人民共和国政府采购法》第十三条、第七十一条|'
                                        r'二、行政处罚：|'
                                        r'根据《中华人民共和国行政强制法》第五十四条规定，我厅对你公司进行催告|'
                                        r'按照《中华人民共和国注册会计师法》第三)'
                                        r'[\s\S]*?)'
                                        r'(.*?如不服本(处罚|处理)?决定|.*?如对本决定不服|'
                                        r'\n泉州市财政局|.*?对本处罚决定如有异议|特此公告。|'
                                        r'\n福建省财政厅)',
                                        content_text).group(1).strip()

        publish_date = re.findall(r'\d{4} ?年\d{1,2}月\d{1,3}日', content_text)[-1].strip().replace(' ', '')
        real_publish_date = format_date(publish_date)

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '福建财政厅',
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense,
            'defenseResponse': defense_response,
            'litigant': litigant[:-1] if litigant[-1] in [':', '：'] else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('福建省财政厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('福建省财政厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('福建省财政厅 数据解析 ' + ' -- 修改parsed完成')


# 安徽省财政厅
def anhuiczt_parse():
    for each_document in db.finance_data.find(
            {'origin': '安徽省财政厅', 'status': {'$nin': ['ignored']}, 'url': {'$exists': True}}):
        if 'url' in each_document.keys():
            announcement_url = each_document['url']
            announcement_title = each_document['title']
            # 判断是否解析过
            if db.finance_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                    {'origin_url': announcement_url, 'parsed': True}).count() == 1:
                continue

            logger.info('url to parse ' + announcement_url)

            content_response = request_site_page(announcement_url)
            if content_response is None:
                logger.error('网页请求错误 %s' % announcement_url)
                continue
            content_response.encoding = content_response.apparent_encoding
            content_soup = BeautifulSoup(content_response.text, 'lxml')
            doc_link_list = [kk.attrs['href'] for kk in content_soup.find(id='BodyLabel').find_all('a') if
                             'href' in kk.attrs.keys() and str(kk.attrs['href']).endswith('doc')]
            if len(doc_link_list) > 0:
                doc_link = urljoin(announcement_url, doc_link_list[0])
                response = request_site_page(doc_link)
                link_type = doc_link.split('.')[-1]
                if response is None:
                    logger.error('网页请求错误')
                    return

                with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                if not os.path.exists('./test/tmp.docx'):
                    shell_str = '/usr/local/bin/soffice --headless --convert-to docx ' + \
                                './test/tmp.' + link_type + ' --outdir ./test'
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()

                with open('./test/tmp.docx', 'rb') as docx_file:
                    docx_content = docx_file.read()

                if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': doc_link}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': doc_link,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'docx',
                        'oss_file_name': announcement_title,
                        'oss_file_content': docx_content,
                        'parsed': False
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                    db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = db.parsed_data.find_one(
                        {'origin_url': announcement_url, 'oss_file_origin_url': doc_link})['_id']

                doc = docx.Document('./test/tmp.docx')
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip() != '':
                        full_text.append(para.text.strip())
                content_text = '\n'.join(full_text)

                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.doc'):
                    os.remove('./test/tmp.doc')
                if os.path.exists('./test/tmp.docx'):
                    os.remove('./test/tmp.docx')
                if os.path.exists('./test/tmp.pdf'):
                    os.remove('./test/tmp.pdf')
            else:
                if db.parsed_data.find(
                        {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': announcement_url,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'html',
                        'oss_file_name': announcement_title,
                        'oss_file_content': content_response.text,
                        'parsed': False,
                        'content_id_name': 'BodyLabel'
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                                 content_response.text)
                    db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                       'oss_file_origin_url': announcement_url})['_id']

                content_text = get_content_text(content_soup.find(id='BodyLabel'))
            if re.search(r'(财监.\d{4}.\d+号)', announcement_title):
                announcement_code = re.search(r'(财监.\d{4}.\d+号)', announcement_title).group(1).strip()
            else:
                announcement_code = ''

            litigant = re.search(r'(当事人[\s\S]*?)'
                                 r'(根据.*?要求|'
                                 r'我厅在对.*?执业质量检查时|'
                                 r'根据省住房和城乡建设厅转来的人民来信|'
                                 r'根据《中华人民共和国会计法》、《财政部门监督办法》、《安徽省财政监督暂行办法》|'
                                 r'根据《中华人民共和国注册会计师法》、《会计师事务所审批和监督暂行办法》|'
                                 r'国家外汇管理局安徽省分局《关于安徽一凡会计师事务所违规验资有关情况的函》|'
                                 r'财政部驻安徽省财政监察专员办事处（以下简称“安徽专员办”）)',
                                 content_text).group(1).strip()
            facts_compiler = re.compile(r'(([一二三四五六七八九十]{1,2}、.*?\n)[\s\S]*?)'
                                        r'(上述行为违反了.*?第.*?条.*?规定。)')
            facts_list = facts_compiler.findall(content_text)
            if len(facts_list) > 0:
                facts = '\n'.join([kk[0].strip() for kk in facts_list])
                punishment_basis = '\n'.join([kk[1].strip() + '\n' + kk[2].strip() for kk in facts_list])
            else:
                facts = re.search(r'(查出的主要问题[和及]作出的行政处罚如下：|检查发现：|检查发现存在以下问题：|'
                                  r'你单位存在以下问题：|检查发现的主要问题：|查实的主要问题和行政处罚决定如下：|'
                                  r'所附案件材料显示：)'
                                  r'([\s\S]*?)'
                                  r'(上述(行为|事项)违反了.*?第.*?条.*?规定(，且情节严重|，所出具的审计报告不具有证明效力)?。|'
                                  r'\n上述(违规|违法)(事实|事项).*?予以证实。|'
                                  r'安徽美术出版社虚构购销业务、虚增营业收入及成本的行为违反了)',
                                  content_text).group(2).strip()
                punishment_basis = re.search(r'(上述(行为|事项)违反了.*?第.*?条.*?规定'
                                             r'(，且情节严重|，所出具的审计报告不具有证明效力)?。|'
                                             r'安徽美术出版社虚构购销业务、虚增营业收入及成本的行为违反了.*?第.*?条.*?规定。)'
                                             r'(\n上述(违规|违法)(事实|事项).*?予以证实。|'
                                             r'\n根据.*?第.*?条.*?决定给予|'
                                             r'鉴此，我厅作出《行政处罚告知书》|'
                                             r'\n你作为.*?对上述.*?负有.*?责任)',
                                             content_text).group(1).strip()
            punishment_decision = re.search(r'((根据|依据).*?第.*?条.*?(决定|维持)给予[\s\S]*?)'
                                            r'(如不服本处罚决定|\n.*?应在接到本处罚决定之日)',
                                            content_text).group(1).strip()
            try:
                publish_date = re.search(r'\n(\d{4}年\d{1,3}月\d{1,2}日)($|\n)', content_text).group(1).strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.warning(e)
                real_publish_date = format_date(each_document['publishDate'])

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '安徽财政厅',
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('安徽省财政厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('安徽省财政厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('安徽省财政厅 数据解析 ' + ' -- 修改parsed完成')
    for each_document in db.finance_data.find(
            {'origin': '安徽省财政厅', 'status': {'$nin': ['ignored', 'parsed']}, 'url': {'$exists': False},
             'cfZt': {'$ne': '1'}}):
        if each_document['cfZt'] == '1':
            announcement_title = '行政处罚决定书' + '（' + each_document['cfWsh'].strip() + '）（撤销）'
        else:
            announcement_title = '行政处罚决定书' + '（' + each_document['cfWsh'].strip() + '）'

        if 'cfJdrq' in each_document.keys():
            real_publish_date = format_date(each_document['cfJdrq'])
        else:
            real_publish_date = '2019年'

        litigant = ''
        if each_document.get('cfXdrMc', '') != '':
            litigant += '行政相对人名称：' + each_document.get('cfXdrMc', '').strip()
        if each_document.get('cfXdrShxym', '') != '':
            litigant += '\n统一社会信用代码：' + each_document.get('cfXdrShxym', '').strip()
        if each_document.get('cfFr', '') != '':
            litigant += '\n法定代表人姓名：' + each_document.get('cfFr', '').strip()
        if each_document.get('cfXdrSfz', '') != '':
            litigant += '\n身份证：' + each_document.get('cfXdrSfz', '').strip()

        punishment_decision = '依据' + each_document.get('cfYj', '').strip() + '，我厅对当事人进行如下处罚：' + \
                              each_document.get('cfJg', '').strip()
        punishment_decision = punishment_decision.replace('依据根据', '依据').replace(
            '根据根据', '依据').replace('依据上述行为违反了', '上述行为违反了').replace('。，', '，')

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '安徽财政厅',
            'announcementDate': real_publish_date,
            'announcementCode': each_document['cfWsh'].strip(),
            'facts': each_document['cfSy'].strip(),
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant,
            'punishmentBasement': '',
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': '',
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementCode': each_document['cfWsh'].strip(),
                                 'announcementOrg': '安徽财政厅'}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('安徽省财政厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('安徽省财政厅 数据解析 ' + ' -- 数据已经存在')
        db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        logger.info('安徽省财政厅 数据解析 ' + ' -- 修改parsed完成')

    for each_document in db.finance_data.find(
            {'origin': '安徽省财政厅', 'status': {'$nin': ['ignored', 'parsed']}, 'url': {'$exists': False},
             'cfZt': '1'}):
        if each_document['cfZt'] == '1':
            announcement_title = '行政处罚决定书' + '（' + each_document['cfWsh'].strip() + '）（撤销）'
        else:
            announcement_title = '行政处罚决定书' + '（' + each_document['cfWsh'].strip() + '）'

        if 'cfJdrq' in each_document.keys():
            real_publish_date = format_date(each_document['cfJdrq'])
        else:
            real_publish_date = '2019年'

        litigant = ''
        if each_document.get('cfXdrMc', '') != '':
            litigant += '行政相对人名称：' + each_document.get('cfXdrMc', '').strip()
        if each_document.get('cfXdrShxym', '') != '':
            litigant += '\n统一社会信用代码：' + each_document.get('cfXdrShxym', '').strip()
        if each_document.get('cfFr', '') != '':
            litigant += '\n法定代表人姓名：' + each_document.get('cfFr', '').strip()
        if each_document.get('cfXdrSfz', '') != '':
            litigant += '\n身份证：' + each_document.get('cfXdrSfz', '').strip()

        punishment_decision = '依据' + each_document.get('cfYj', '').strip() + '，我厅对当事人进行如下处罚：' + \
                              each_document.get('cfJg', '').strip()
        punishment_decision = punishment_decision.replace('依据根据', '依据').replace(
            '根据根据', '依据').replace('依据上述行为违反了', '上述行为违反了').replace('。，', '，')

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': '安徽财政厅',
            'announcementDate': real_publish_date,
            'announcementCode': each_document['cfWsh'].strip(),
            'facts': each_document['cfSy'].strip(),
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant,
            'punishmentBasement': '',
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': '',
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementCode': each_document['cfWsh'].strip(),
                                 'announcementOrg': '安徽财政厅'}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('安徽省财政厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('安徽省财政厅 数据解析 ' + ' -- 数据已经存在')
        db.finance_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        logger.info('安徽省财政厅 数据解析 ' + ' -- 修改parsed完成')


def parse_all():
    xinjiangczt_parse()
    gansuczt_parse()
    yunnanczt_parse()
    guizhouczt_parse()
    sichuanczt_parse()
    chongqingczj_parse()
    guangxiczt_parse()
    guangdongczt_parse()
    # TODO 部分内容没有解析好
    # hunanczt_parse()
    qingdaoczt_parse()
    jiangxiczt_parse()
    shandongczt_parse()
    henanczt_parse()
    szczt_parse()
    xiamenczt_parse()
    fujianczt_parse()
    anhuiczt_parse()


if __name__ == "__main__":
    parse_all()
