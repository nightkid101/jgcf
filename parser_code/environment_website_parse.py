from pymongo import MongoClient
import re
import os
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
from xlrd import open_workbook, xldate_as_tuple
from init import logger_init, config_init
from utility import format_date, remove_special_char, request_site_page, get_content_text, table_to_list
from oss_utils import init_ali_oss, oss_add_file
import subprocess
from pdf2html import pdf_ocr_to_text
import xlrd
import docx2txt
import patoolib

logger = logger_init(' 数据解析')
config = config_init()
if config['mongodb']['dev_mongo'] == '1':
    db = MongoClient(config['mongodb']['ali_mongodb_url'], username=config['mongodb']['ali_mongodb_username'],
                     password=config['mongodb']['ali_mongodb_password'],
                     port=int(config['mongodb']['ali_mongodb_port']))[config['mongodb']['ali_mongodb_name']]
else:
    db = MongoClient(
        host=config['mongodb']['mongodb_host'],
        port=int(config['mongodb']['mongodb_port']),
        username=None if config['mongodb']['mongodb_username'] == '' else config['mongodb']['mongodb_username'],
        password=None if config['mongodb']['mongodb_password'] == '' else config['mongodb']['mongodb_password'])[
        config['mongodb']['mongodb_db_name']]
ali_bucket = init_ali_oss()


# 生态环境部
def mee_parse():
    for each_document in db.environment_data.find({'origin': '环保部门', 'status': {'$nin': ['ignored']}}):
        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        # ignored
        if '专项行动破解突出环境问题' in announcement_title:
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        result_map_list = []
        if content_soup.style:
            content_soup.style.decompose()

        if len(content_soup.find_all(class_='wzxq2')) > 0:
            content_class_name = 'wzxq2'
        elif content_soup.find(class_='content_box'):
            content_class_name = 'content_box'
        else:
            content_class_name = 'mainContainer'

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'htm',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'parsed': False,
                'content_class_name': content_class_name
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        if re.search('关于对.*?挂牌督办的通知', announcement_title):
            announcement_type = '挂牌督办通知'
            content_text = get_content_text(content_soup.find(id='ContentRegion'))
            logger.info(content_text)

            head_info_text = get_content_text(content_soup.find(class_='headInfo'))
            announcement_code = re.search(r'\n文号:(.*?)\n', head_info_text).group(1).strip()
            publish_date = re.search(r'(生态环境部办公厅|最高人民检察院办公厅|环境保护部办公厅)\n'
                                     r'(\d{4}年\d{1,2}月\d{1,2}日)\n抄送',
                                     content_text).group(2).strip()
            real_publish_date = format_date(publish_date)

            if re.search(r'\n附件.*?\n(严重超标|环境|环境违法)(问题|案件)基本情况及督办要求', content_text):
                punishment_decision_prefix = re.search(
                    r'(经研究，(我部)?决定对.*?挂牌督办。\n请.*?(落实|按照)督办要求.*?|'
                    r'按照《环境违法案件挂牌督办管理办法》的规定，现决定对.*?挂牌督办。\n请.*?(落实|按照)督办要求.*?|'
                    r'经研究，决定对上述3家企业环境违法问题挂牌督办。\n请.*?(落实|按照)督办要求.*?|'
                    r'(为督促其尽快解除环境安全隐患，|为督促你市尽快解决环境污染问题，|经商江西省人民政府办公厅同意，|为强化地方政府对环境质量负责的主体责任意识，|根据《环境违法案件挂牌督办管理办法》，)?'
                    r'我部[^。\n]*?决定.*?对.*?挂牌督办。\n请.*?(落实|按照)督办要求.*?)\n',
                    content_text).group(1).strip()
                punishment_decision_prefix = punishment_decision_prefix.replace('（见附件）', '').strip()

                if announcement_url in ['http://www.mee.gov.cn/gkml/hbb/bgt/201505/t20150521_302081.htm']:
                    punishment_decision = punishment_decision_prefix + '\n' + \
                                          re.search(r'(督办要求：.*?\n督办期限：.*?)\n', content_text).group(1).strip()
                    result_map_list = []
                    for each_table in content_soup.find_all(class_='MsoNormalTable'):
                        for each_tr in each_table.find_all('tr'):
                            each_td_list = [kk.text.strip() for kk in each_tr.find_all('td')]
                            if '序号' in each_td_list[0]:
                                continue
                            each_map = {
                                'announcementTitle': '关于对' + each_td_list[1] + '挂牌督办的通知',
                                'announcementOrg': '生态环境部',
                                'announcementDate': real_publish_date,
                                'announcementCode': announcement_code,
                                'facts': each_td_list[2],
                                'defenseOpinion': '',
                                'defenseResponse': '',
                                'litigant': each_td_list[1],
                                'punishmentBasement': '',
                                'punishmentDecision': punishment_decision,
                                'type': announcement_type,
                                'oss_file_id': file_id,
                                'status': 'not checked'
                            }
                            logger.info(each_map)
                            if db.announcement.count(
                                    {
                                        'oss_file_id': file_id,
                                        'facts': each_td_list[2],
                                        'litigant': each_td_list[1]
                                    }) == 0:
                                db.announcement.insert_one(each_map)
                                result_map_list.append(each_map)
                                logger.info('生态环境部 数据解析 ' + ' -- 数据导入完成')
                            else:
                                logger.info('生态环境部 数据解析 ' + ' -- 数据已经存在')
                    if len(result_map_list) > 0:
                        logger.info('生态环境部 数据解析 ' + ' -- 一共有%d条数据' % len(result_map_list))
                        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                        logger.info('生态环境部 数据解析 ' + ' -- 修改parsed完成')
                    else:
                        logger.info('生态环境部 数据解析 ' + ' -- 解析未能完成')
                else:
                    facts_text = re.search(r'\n(严重超标|环境|环境违法)(问题|案件)基本情况及督办要求\n([\s\S]*?)$',
                                           content_text).group(3).strip()
                    if facts_text.startswith('一、') and not facts_text.startswith('一、基本情况'):
                        litigant_info_list = []
                        if len(content_soup.find(id='ContentRegion').find_all('table')) > 1:
                            table_info = content_soup.find(id='ContentRegion').find_all('table')[0]
                            for each_tr in table_info.find_all('tr'):
                                td_list = [kk.text.strip() for kk in each_tr.find_all('td')]
                                if '序号' in td_list[0]:
                                    continue
                                litigant_info_list.append(td_list)
                        litigant_list = re.findall(r'([一二三四五六七八九十]{1,3})、(.*?)\n', facts_text)
                        for index, each_litigant in enumerate(litigant_list):
                            each_litigant_final = re.search(r'(.*?)'
                                                            r'(倾倒工业固体废物及有毒有害物质|'
                                                            r'非法外排废水|'
                                                            r'环境违法案|'
                                                            r'生活污水直排湘江|'
                                                            r'堆积大量非法盗采砂石|'
                                                            r'堆积大量砂石侵占岸线|'
                                                            r'将修船打磨的铁锈和油漆散落在江滩上|'
                                                            r'生产过程中产生巨大矿坑|$)',
                                                            each_litigant[1].strip()).group(1).strip()
                            each_facts = re.search(
                                each_litigant[1].strip().replace(r'(', r'\(').replace(r')', r'\)') +
                                r'\n基本情况：([\s\S]*?)\n督办要求：',
                                facts_text).group(1).strip()
                            each_punishment_decision = re.search(
                                each_litigant[1].strip().replace(r'(', r'\(').replace(r')', r'\)') +
                                r'\n基本情况：[\s\S]*?\n(督办要求：.*?(\n督办期限：.*?)?)\n',
                                facts_text).group(1).strip()
                            if len(litigant_info_list) != 0:
                                each_litigant_info = each_litigant_final + '\n' + '行政区划：' + \
                                                     litigant_info_list[index][1] + litigant_info_list[index][2]
                            else:
                                each_litigant_info = each_litigant_final
                            each_map = {
                                'announcementTitle': '关于对' + each_litigant_final + '挂牌督办的通知',
                                'announcementOrg': '生态环境部',
                                'announcementDate': real_publish_date,
                                'announcementCode': announcement_code,
                                'facts': each_facts,
                                'defenseOpinion': '',
                                'defenseResponse': '',
                                'litigant': each_litigant_info,
                                'punishmentBasement': '',
                                'punishmentDecision': punishment_decision_prefix + '\n' + each_punishment_decision,
                                'type': announcement_type,
                                'oss_file_id': file_id,
                                'status': 'not checked'
                            }
                            logger.info(each_map)
                            if db.announcement.count(
                                    {'oss_file_id': file_id, 'facts': each_facts, 'litigant': each_litigant_info}) == 0:
                                db.announcement.insert_one(each_map)
                                result_map_list.append(each_map)
                                logger.info('生态环境部 数据解析 ' + ' -- 数据导入完成')
                            else:
                                logger.info('生态环境部 数据解析 ' + ' -- 数据已经存在')
                        if len(result_map_list) > 0:
                            logger.info('生态环境部 数据解析 ' + ' -- 一共有%d条数据' % len(result_map_list))
                            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                            logger.info('生态环境部 数据解析 ' + ' -- 修改parsed完成')
                        else:
                            logger.info('生态环境部 数据解析 ' + ' -- 解析未能完成')
                    else:
                        litigant = re.search(r'关于对(.*?)(环境问题|环境违法案)?(联合)?挂牌督办的通知',
                                             announcement_title).group(1).strip()
                        facts = re.search(r'基本情况：([\s\S]*?)\n(二、)?督办要求', facts_text).group(1).strip()
                        punishment_decision = punishment_decision_prefix + '\n' + re.search(
                            r'(督办要求[\s\S]*?\n(三、)?督办(期限|时限)[\s\S]*?).bdsharebuttonbox',
                            facts_text).group(1).strip()

                        result_map = {
                            'announcementTitle': announcement_title,
                            'announcementOrg': '生态环境部',
                            'announcementDate': real_publish_date,
                            'announcementCode': announcement_code,
                            'facts': facts,
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': litigant,
                            'punishmentBasement': '',
                            'punishmentDecision': punishment_decision,
                            'type': announcement_type,
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        logger.info(result_map)
                        if db.announcement.count(
                                {'oss_file_id': file_id, 'facts': facts, 'litigant': litigant}) == 0:
                            db.announcement.insert_one(result_map)
                            logger.info('生态环境部 数据解析 ' + ' -- 数据导入完成')
                        else:
                            logger.info('生态环境部 数据解析 ' + ' -- 数据已经存在')
                        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                        logger.info('生态环境部 数据解析 ' + ' -- 修改parsed完成')
            elif re.search(r'\n附件\n涉案企业名单及督办要求', content_text):
                punishment_decision_prefix = re.search(
                    r'(为消除环境安全隐患，严厉打击环境违法犯罪，生态环境部、公安部、最高人民检察院决定对.*?联合挂牌督办。.*?)\n',
                    content_text).group(1).strip()
                punishment_decision_prefix = punishment_decision_prefix.replace('（见附件）', '').strip()
                facts = re.search(r'((2018年3月.*?进行了现场调查。经查)[\s\S]*?)\n(为消除环境安全隐患)',
                                  content_text).group(1).strip()
                litigant = re.search(r'涉案企业名单：\n(.*?)\n督办要求', content_text).group(1).strip()
                punishment_decision = re.search(r'\n(督办要求：\n[\s\S]*?).bdsharebuttonbox',
                                                content_text).group(1).strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': '生态环境部',
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': punishment_decision_prefix + '\n' + punishment_decision,
                    'type': announcement_type,
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.count(
                        {'oss_file_id': file_id, 'facts': facts, 'litigant': litigant}) == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('生态环境部 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('生态环境部 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('生态环境部 数据解析 ' + ' -- 修改parsed完成')
            elif re.search(r'\n附件\n督办要求', content_text):
                punishment_decision_prefix = re.search(
                    r'(为消除环境安全隐患，严厉打击环境违法犯罪[\s\S]*?)附件：督办要求',
                    content_text).group(1).strip()
                punishment_decision_prefix = punishment_decision_prefix.replace('（见附件）', '').strip()
                facts = re.search(r'((2016年以来.*?发生.*?环境违法犯罪案件)[\s\S]*?)\n(为消除环境安全隐患)',
                                  content_text).group(1).strip()
                litigant = re.search(r'关于对(.*?)联合挂牌督办的通知', announcement_title).group(1).strip()
                punishment_decision = re.search(r'\n督办要求\n([\s\S]*?).bdsharebuttonbox',
                                                content_text).group(1).strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': '生态环境部',
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': punishment_decision_prefix + '\n' + punishment_decision,
                    'type': announcement_type,
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.count(
                        {'oss_file_id': file_id, 'facts': facts, 'litigant': litigant}) == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('生态环境部 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('生态环境部 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('生态环境部 数据解析 ' + ' -- 修改parsed完成')
            else:
                litigant = re.search(
                    r'关于对(.*?)(联合)?挂牌督办的通知',
                    announcement_title).group(1).strip()
                facts = re.search(r'((2018年1月17日.*?发生.*?案件。经公安机关调查|'
                                  r'近期，我部接群众举报龙口市尾矿库存在尾矿水直排、防扬散措施不到位等环境问题|'
                                  r'2015年11月，我部在环境保护大检查督查中发现|'
                                  r'我部华北环境保护督查中心联合你厅对运城市综合督查整改情况进行后督察时发现|'
                                  r'2015年3月，我部西南环境保护督查中心对你省黔西南州开展环境执法稽查时发现|'
                                  r'2015年3月，我部华东环境保护督查中心对你省苏州市、无锡市开展大气污染防治专项督查时发现|'
                                  r'我部华东环境保护督查中心对你省临沂市大气污染物排放企业和区域环境空气状况专项督查中发现|'
                                  r'我部华北环境保护督查中心在冬季大气污染防治督查中发现)'
                                  r'[\s\S]*?)'
                                  r'(为严厉打击环境违法犯罪|为尽快解决|根据《环境违法案件挂牌督办管理办法》|'
                                  r'我部决定对运城市环境问题挂牌督办。|我部决定对临沂市大气污染问题挂牌督办。|'
                                  r'我部决定对承德市大气污染问题挂牌督办。)',
                                  content_text).group(1).strip()
                punishment_decision = re.search(r'((为严厉打击环境违法犯罪|为尽快解决.*?问题|根据《环境违法案件挂牌督办管理办法》|'
                                                r'我部决定对运城市环境问题挂牌督办。|我部决定对临沂市大气污染问题挂牌督办。|'
                                                r'我部决定对承德市大气污染问题挂牌督办。)'
                                                r'[\s\S]*?)\n(附件|联系人)',
                                                content_text).group(1).strip()
                punishment_decision = punishment_decision.replace('（见附件）', '').strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': '生态环境部',
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': punishment_decision,
                    'type': announcement_type,
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.count(
                        {'oss_file_id': file_id, 'facts': facts, 'litigant': litigant}) == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('生态环境部 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('生态环境部 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('生态环境部 数据解析 ' + ' -- 修改parsed完成')
        elif re.search('关于解除(对)?.*?挂牌督办的通知', announcement_title):
            if announcement_url in ['http://www.mee.gov.cn/xxgk2018/xxgk/xxgk06/201904/t20190401_698026.html']:
                continue
            announcement_type = '解除挂牌督办'
            content_text = get_content_text(content_soup.find(class_=content_class_name))
            logger.info(content_text)
            publish_date = re.search(r'(生态环境部办公厅|最高人民检察院办公厅|环境保护部办公厅|（环境保护部办公厅代章）)\n'
                                     r'(\d{4}年\d{1,2}月\d{1,2}日)\n抄送',
                                     content_text).group(2).strip()
            real_publish_date = format_date(publish_date)
            litigant = re.search(r'关于解除对?(.*?)挂牌督办的通知', announcement_title).group(1).strip()
            facts = re.search(
                r'(按照.*?要求.*?督促.*?积极(开展)?(整改|整治).*?并向我部提交了解除挂牌督办申请.*?'
                r'(经我部现场[核检]查|我部组织对案件的整改情况进行核查|'
                r'经我部联合江苏、山东两省环境保护厅现场核查|'
                r'我部.*?对.*?情况进行了?核查).*?)'
                r'\n经研究，现决定',
                content_text).group(1).strip()
            punishment_decision = re.search(r'(经研究，现决定.*?)\n(生态环境部办公厅|环境保护部办公厅)',
                                            content_text).group(1).strip()
            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '生态环境部',
                'announcementDate': real_publish_date,
                'announcementCode': '',
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant,
                'punishmentBasement': '',
                'punishmentDecision': punishment_decision,
                'type': announcement_type,
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.count(
                    {'oss_file_id': file_id, 'facts': facts, 'litigant': litigant}) == 0:
                db.announcement.insert_one(result_map)
                logger.info('生态环境部 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('生态环境部 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('生态环境部 数据解析 ' + ' -- 修改parsed完成')
        elif re.search(r'环境保护部对廊坊市大气环境问题进行挂牌督办', announcement_title):
            content_text = get_content_text(content_soup.find(class_=content_class_name))
            facts = re.search(
                r'(环境保护部近期督查发现，廊坊市虽然加大工作力度，狠抓污染治理，大气污染防治工作取得积极进展。[\s\S]*?)'
                r'(为督促廊坊市进一步落实环境保护责任)', content_text).group(1).strip()
            punishment_decision = re.search(r'(为督促廊坊市进一步落实环境保护责任[\s\S]*?).bdsharebuttonbox',
                                            content_text).group(1).strip()
            real_publish_date = format_date(each_document['publishDate'])
            announcement_type = '挂牌督办通知'
            litigant = re.search(r'环境保护部对(.*?)进行挂牌督办', announcement_title).group(1).strip()
            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '生态环境部',
                'announcementDate': real_publish_date,
                'announcementCode': '',
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant,
                'punishmentBasement': '',
                'punishmentDecision': punishment_decision,
                'type': announcement_type,
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.count(
                    {'oss_file_id': file_id, 'facts': facts, 'litigant': litigant}) == 0:
                db.announcement.insert_one(result_map)
                logger.info('生态环境部 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('生态环境部 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('生态环境部 数据解析 ' + ' -- 修改parsed完成')
        else:
            announcement_type = '挂牌督办通知'
            content_text = get_content_text(content_soup.find(class_=content_class_name))
            facts_list = re.findall(r'(（?[一二三四五六七八九十]{1,2}）?|[1234567890]{1,2})、?'
                                    r'(.*?)\n(基本情况|存在问题)：\n?(.*?)\n(督办要求|整改要求|联查联办要求)：\n?(.*?)\n',
                                    content_text)
            result_map_list = []
            for each_fact in facts_list:
                each_map = {
                    'announcementTitle': '关于对' + each_fact[1].strip() + '挂牌督办的通知',
                    'announcementOrg': '生态环境部',
                    'announcementDate': format_date(each_document['publishDate']),
                    'announcementCode': '',
                    'facts': each_fact[3].strip(),
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': each_fact[1].strip(),
                    'punishmentBasement': '',
                    'punishmentDecision': each_fact[5].strip(),
                    'type': announcement_type,
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(each_map)
                if db.announcement.count(
                        {'oss_file_id': file_id, 'facts': each_fact[3].strip(), 'litigant': each_fact[1].strip()}) == 0:
                    db.announcement.insert_one(each_map)
                    result_map_list.append(each_map)
                    logger.info('生态环境部 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('生态环境部 数据解析 ' + ' -- 数据已经存在')
            if len(result_map_list) > 0:
                logger.info('生态环境部 数据解析 ' + ' -- 一共有%d条数据' % len(result_map_list))
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('生态环境部 数据解析 ' + ' -- 修改parsed完成')
            else:
                logger.info('生态环境部 数据解析 ' + ' -- 解析未能完成')


# 北京生态环境局
def beijing():
    for each_document in db.environment_data.find({'origin': '北京市环境保护局', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        # 处罚公告标题
        announcement_title = content_soup.find(class_='h_dl_t').text.strip()

        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                'parsed': False,
                'content_id_name': 'h_dl_c'
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text.encode(content_response.encoding).decode('utf-8'))
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_text = get_content_text(content_soup.find(id='h_dl_c'))

        # 处罚机构
        announcement_org = '北京生态环境局'
        # 处罚日期
        try:
            publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.warning(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)
        # 文号
        announcement_code_compiler = re.compile(r'(\n|^)(京环.*?\d{4}.*?\d+号)\n')
        if announcement_code_compiler.search(content_text):
            announcement_code = announcement_code_compiler.search(content_text).group(2).strip()
        else:
            announcement_code = re.search(r'(.\d{4}.*?\d+号?)', announcement_title).group(1).strip()
        # 当事人
        litigant = re.search(r'((当事人名称|法定代表人)[\s\S]*?)我局', content_text).group(1).strip()
        # 违规事实
        facts = re.search(r'((我局|我局于)[\s\S]*?)(以上事实有现场检查(记录|笔录)|(上述事实))', content_text).group(1).strip()

        # 认定意见
        try:
            punishment_basis = re.search(r'((你单位.*?行为已?违反)[\s\S]*?规定)',
                                         content_text).group(1).strip()
        except Exception as e:
            logger.warning(e)
            punishment_basis = ''

        # 处罚决定
        if re.search(r'((依据|责令你单位停止违法行为)'
                     r'[\s\S]*?'
                     r'我局将(依法|申请)(实施行政处罚|强制执行|进行处罚|.*?法院强制执行))', content_text):
            punishment_decision = re.search(r'((依据|责令你单位停止违法行为)'
                                            r'[\s\S]*?'
                                            r'我局将(依法|申请)(实施行政处罚|强制执行|进行处罚|.*?法院强制执行))',
                                            content_text).group(1).strip()
        elif re.search(r'(《北京市大气污染防治条例》[\s\S]*?我局将依法实施行政处罚)', content_text):
            punishment_decision = re.search(r'(《北京市大气污染防治条例》[\s\S]*?我局将依法实施行政处罚)',
                                            content_text).group(1).strip()
        else:
            punishment_decision = ''

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('北京生态环境局 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('北京生态环境局 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('北京生态环境局 数据解析 ' + ' -- 修改parsed完成')


# 天津生态环境局
def tianjin():
    for each_document in db.environment_data.find({'origin': '天津市环境保护局', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        # ignored
        if '2016年11月我市未作出环境违法行为红牌处罚' in announcement_title:
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')
        if len(content_soup.find_all(class_='wzxq_neirong2')) > 0:
            content_class_name = 'wzxq_neirong2'
        else:
            content_class_name = 'pages_content'

        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                'parsed': False,
                'content_class_name': content_class_name
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text.encode(content_response.encoding).decode('utf-8'))
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_text = content_soup.find(class_=content_class_name).text.strip()
        content_text = '\n'.join([kk.strip() for kk in content_text.split('\n')])
        content_text = re.sub('\n+', '\n', content_text)
        content_text = re.sub(r'^P.MsoNormal[\s\S]*?(WordSection1|line-through|COLOR: red)\n\}\n', '', content_text)

        if '执法情况通报' in announcement_title or '最高法首次公布环境保护行政案件十大案例' in announcement_title:
            continue

        # 处罚机构
        announcement_org = '天津生态环境局'
        # 处罚日期
        if re.search(r'\n(\d{4}年\d{1,2}月\d{1,3}日)$', content_text):
            publish_date = re.search(r'\n(\d{4}年\d{1,2}月\d{1,3}日)$', content_text).group(1).strip()
        elif len(content_soup.find_all(class_='pages_print')) > 0:
            publish_date = content_soup.find(class_='pages_print').text.strip()
        else:
            publish_date = each_document['publishDate']
        real_publish_date = format_date(publish_date)

        # 文号
        announcement_code_compiler = re.compile(r'(津市环.*?\d{4}.\d+号)')
        announcement_code = announcement_code_compiler.search(content_text).group(1).strip()

        # 当事人
        try:
            litigant = re.search(announcement_code.replace(r'[', r'\[').replace(r']', r'\]') +
                                 r'([\s\S]*?)(一、环境违法事实和证据|你公司环境违法一案|你单位|我局|一、调查情况及发现的环境违法事实)',
                                 content_text).group(1).strip()
        except Exception as e:
            litigant = ''
            logger.warning(e)

        # 违规事实
        facts = re.search(
            r'(一、调查情况及发现的环境违法事实、证据和陈述申辩（听证）及采纳情况|我局于|一、环境违法事实和证据|环保执法人员)'
            r'([\s\S]*?)(你单位的上述行为违反|《环境行政处罚办法》|上述行为违反了|以上事实，?有.*?等证据为凭。|现依据)',
            content_text).group(2).strip()

        # 认定意见
        try:
            punishment_basis = re.search(r'((你单位的上述行为[\s\S]*?予以处罚)|(上述行为违反[\s\S]*?的规定))',
                                         content_text).group(1).strip()
        except Exception as e:
            logger.warning(e)
            punishment_basis = ''

        # 申辩意见
        if re.search(r'(应当予以处罚。|违反了《中华人民共和国大气污染防治法》第二十条第二款的规定。|'
                     r'违反了《放射性同位素与射线装置安全和防护条例》第二十一条的规定。)'
                     r'([\s\S]*?)'
                     r'(经复核|二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类|二、责令改正的依据和种类|'
                     r'以上事实，有.*?为凭。)', content_text):
            defense_opinion = re.search(r'(应当予以处罚。|违反了《中华人民共和国大气污染防治法》第二十条第二款的规定。|'
                                        r'违反了《放射性同位素与射线装置安全和防护条例》第二十一条的规定。)'
                                        r'([\s\S]*?)'
                                        r'(经复核|二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类|二、责令改正的依据和种类|'
                                        r'以上事实，有.*?为凭。)',
                                        content_text).group(2).strip()
        else:
            defense_opinion = ''
        # 申辩意见反馈
        if re.search(r'(经复核[\s\S]*?)(二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类|'
                     r'二、责令停止使用和行政处罚的依据、种类|以上事实，有.*?为凭。)', content_text):
            defense_response = re.search(r'(经复核[\s\S]*?)'
                                         r'(二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类|二、责令改正的依据和种类|'
                                         r'以上事实，有.*?为凭。|二、责令停止使用和行政处罚的依据、种类)',
                                         content_text).group(1).strip()
        else:
            defense_response = ''
        # 处罚决定
        if re.search(
                r'(二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类|二、责令改正的依据和种类|二、责令停止使用和行政处罚的依据、种类|'
                r'（津市环罚字\[2017\]50号），|二、责令停止生产和行政处罚的依据、种类)'
                r'([\s\S]*?)(.、申请行政复议或者提起行政诉讼的途径和期限|$)',
                content_text):
            punishment_decision = re.search(
                r'(二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类|二、责令改正的依据和种类|二、责令停止使用和行政处罚的依据、种类|'
                r'（津市环罚字\[2017\]50号），|二、责令停止生产和行政处罚的依据、种类)'
                r'([\s\S]*?)(.、申请行政复议或者提起行政诉讼的途径和期限|$)',
                content_text).group(2).replace(
                '三、责令改正和处罚决定的履行方式和期限', '').replace(
                '三、责令改正的履行', '').replace(
                '三、处罚决定的履行方式和期限', '').replace(
                '三、责令停止使用和处罚决定的履行方式和期限', '').replace(
                '三、责令停止生产和处罚决定的履行方式和期限', '').strip()
        elif re.search(r'(依据[\s\S]*?)(如你单位对查封决定不服的|根据)', content_text):
            punishment_decision = re.search(r'(依据[\s\S]*?)(如你单位对查封决定不服的|根据)', content_text).group(1).strip()
        else:
            punishment_decision = ''
        punishment_decision = re.sub('\n+', '\n', punishment_decision)

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('天津生态环境局 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('天津生态环境局 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('天津生态环境局 数据解析 ' + ' -- 修改parsed完成')


# 河北生态环境厅
def hebei():
    for each_document in db.environment_data.find({'origin': '河北省环境保护厅',
                                                   # 'title': {'$regex': '.*督[查察]通知$'},
                                                   'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title'].replace('\u3000', '')

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)
        logger.info(announcement_title)

        # ignored
        if re.search('(河北省污染源日常监管“双随机”抽查工作实施方案|电视电话会议上的讲话|'
                     '关于表彰河北省秸秆禁烧和综合利用工作先进单位和先进个人名单的通知|'
                     '关于做好全国“两会”期间我省环境信访工作的通知|'
                     '关于进一步做好“环境保护省长电话”办理工作的通知|'
                     '河北省环境执法稽查大队开展环保严查专项行动情况总结|'
                     '河北省2003年秸秆禁烧和综合利用工作情况|'
                     '关于2003年全省夏季秸秆禁烧和综合利用工作情况的通报)', announcement_title):
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        content_response = request_site_page(announcement_url, trust_environment=True)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = 'GB18030'
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if str(announcement_url).endswith('.jpg'):
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                with open('./test/tmp.jpg', 'wb') as tmp_file:
                    for chunk in content_response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                if not os.path.exists('./test/tmp.pdf'):
                    # 加入图片尺寸参数，百度ocr有图片尺寸限制
                    shell_str = 'img2pdf ./test/tmp.jpg --imgsize 20cmx30cm  -o ./test/tmp.pdf'
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()
                result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')

                with open('./test/tmp.pdf', 'rb') as pdf_file:
                    pdf_content = pdf_file.read()

                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False,
                    'if_ocr': True,
                    'ocr_result': result_text
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']

            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.pdf'):
                os.remove('./test/tmp.pdf')
            if os.path.exists('./test/tmp.jpg'):
                os.remove('./test/tmp.jpg')
            continue
        elif str(announcement_url).endswith('.pdf'):
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                with open('./test/tmp.pdf', 'wb') as tmp_file:
                    for chunk in content_response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')

                with open('./test/tmp.pdf', 'rb') as pdf_file:
                    pdf_content = pdf_file.read()

                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False,
                    'if_ocr': True,
                    'ocr_result': result_text
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                content_text = result_text
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']
                content_text = db.parsed_data.find_one({'origin_url': announcement_url,
                                                        'oss_file_origin_url': announcement_url})['ocr_result']

            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.pdf'):
                os.remove('./test/tmp.pdf')
            continue
        elif str(announcement_url).endswith('.doc'):
            with open('./test/tmp.doc', 'wb') as tmp_file:
                for chunk in content_response.iter_content(chunk_size=1024):
                    if chunk:
                        tmp_file.write(chunk)

            if not os.path.exists('./test/tmp.docx'):
                shell_str = '/usr/local/bin/soffice --headless --convert-to docx ' + \
                            './test/tmp.doc' + ' --outdir ./test'
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.communicate()

            with open('./test/tmp.docx', 'rb') as docx_file:
                docx_content = docx_file.read()
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'docx',
                    'oss_file_name': announcement_title,
                    'oss_file_content': docx_content,
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url})['_id']

            document = Document('./test/tmp.docx')
            content_text = '\n'.join([each_paragraph.text.strip() for each_paragraph in document.paragraphs])

            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.doc'):
                os.remove('./test/tmp.doc')
            if os.path.exists('./test/tmp.docx'):
                os.remove('./test/tmp.docx')
        elif content_soup.find('div', class_='TRS_Editor'):
            if len(content_soup.find('div', class_='TRS_Editor').find_all('img')) >= 1 and \
                    announcement_url not in \
                    ['http://www.hebhb.gov.cn/lishilanmu/zwhd/zywj/tbfw/201305/t20130528_37248.html']:
                img_link_list = [urljoin(announcement_url, kk.attrs['src'])
                                 for kk in content_soup.find('div', class_='TRS_Editor').find_all('img')]
                if db.parsed_data.find(
                        {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:

                    for index, each_img_link in enumerate(img_link_list):
                        img_response = request_site_page(each_img_link, trust_environment=True)
                        with open('./test/' + str(index) + '.jpg', 'wb') as tmp_file:
                            for chunk in img_response.iter_content(chunk_size=1024):
                                if chunk:
                                    tmp_file.write(chunk)

                    if not os.path.exists('./test/tmp.pdf'):
                        shell_str = 'img2pdf '
                        for index in range(len(img_link_list)):
                            shell_str += './test/' + str(index) + '.jpg '
                        shell_str += '--imgsize 20cmx30cm  -o ./test/tmp.pdf'  # 加入图片尺寸参数，百度ocr有图片尺寸限制
                        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                                   shell=True, stdout=subprocess.PIPE)
                        process.communicate()

                    result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')

                    with open('./test/tmp.pdf', 'rb') as pdf_file:
                        pdf_content = pdf_file.read()
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': announcement_url,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'pdf',
                        'oss_file_name': announcement_title,
                        'oss_file_content': pdf_content,
                        'parsed': False,
                        'if_ocr': True,
                        'ocr_result': result_text
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    content_text = result_text
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                       'oss_file_origin_url': announcement_url})['_id']
                    content_text = db.parsed_data.find_one({'origin_url': announcement_url,
                                                            'oss_file_origin_url': announcement_url})['ocr_result']
                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.pdf'):
                    os.remove('./test/tmp.pdf')
                for index in range(len(img_link_list)):
                    if os.path.exists('./test/' + str(index) + '.jpg'):
                        os.remove('./test/' + str(index) + '.jpg')
                continue
            else:
                real_content_soup = content_soup.find_all('div', class_='TRS_Editor')[-1]
                if real_content_soup.style:
                    real_content_soup.style.decompose()
                content_text = get_content_text(real_content_soup)
                if db.parsed_data.find(
                        {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': announcement_url,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'html',
                        'oss_file_name': announcement_title,
                        'oss_file_content': content_response.text,
                        'parsed': False,
                        'content_class_name': 'TRS_Editor'
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                       'oss_file_origin_url': announcement_url})['_id']
        elif content_soup.find('div', class_='congczt'):
            content_text = get_content_text(content_soup.find(class_='congczt'))
            if re.search(r'^/upfiles/.*?htm$', content_text):
                final_url = urljoin(announcement_url, content_text)
                final_content_response = request_site_page(final_url, trust_environment=True)
                final_content_response.encoding = 'GB18030'
                final_content_soup = BeautifulSoup(final_content_response.text, 'lxml')
                content_text = get_content_text(final_content_soup.find('table'))
                if db.parsed_data.find(
                        {'origin_url': announcement_url, 'oss_file_origin_url': final_url}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': final_url,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'htm',
                        'oss_file_name': announcement_title,
                        'oss_file_content': final_content_response.text,
                        'parsed': False,
                        'content_id_name': 'zoom'
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                       'oss_file_origin_url': final_url})['_id']
            else:
                if db.parsed_data.find(
                        {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': announcement_url,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'html',
                        'oss_file_name': announcement_title,
                        'oss_file_content': content_response.text,
                        'parsed': False,
                        'content_class_name': 'congczt'
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                       'oss_file_origin_url': announcement_url})['_id']
        else:
            logger.info('Others ~~~~~~~')
            continue

        content_text = content_text.replace('査', '查').replace('単', '单')
        content_text = re.sub('\n+', '\n', content_text)
        logger.info(content_text)

        # 处罚机构
        announcement_org = '河北生态环境厅'
        # 处罚日期
        real_publish_date = format_date(each_document['publishDate'].split(' ')[0])

        if '决定书' in announcement_title:  # 如果是照片形式的责令改正违法行为的决定书
            try:
                # 文号
                announcement_code = re.search(r'\n((冀环|环罚).*?\d{4}.\d+号?)\n', content_text).group(1).strip()
                # 当事人
                if re.search(announcement_code.replace(r'[', r'\[').replace(r']', r'\]').replace(
                        r'(', r'\(').replace(r')', r'\)') +
                             r'\n([\s\S]*?)\n我[厅斤]于.*?进行了调\n?查', content_text):
                    litigant = re.search(announcement_code.replace(r'[', r'\[').replace(r']', r'\]')
                                         .replace(r'(', r'\(').replace(r')', r'\)') +
                                         r'\n([\s\S]*?)\n我[厅斤]于.*?进行了调\n?查', content_text).group(1).strip()
                else:
                    litigant = ''
                # 违规事实
                if re.search(r'\n((我厅|我斤|现你单位|代表人)[\s\S]*?)\n(以上事实|以上事|等证据为凭|等证据为|上事实)', content_text):
                    facts = re.search(r'\n((我厅|我斤|现你单位|代表人)[\s\S]*?)\n(以上事实|以上事|等证据为凭|等证据为|上事实)',
                                      content_text).group(1).strip()
                else:
                    facts = ''
                facts = facts.replace('\n', '')

                # 认定意见
                try:
                    punishment_basis = re.search(r'\n((本机关认为|上述|述行为|.*?上述行为)[\s\S]*?的规定)', content_text).group(
                        1).strip()
                except Exception as e:
                    logger.warning(e)
                    punishment_basis = ''
                punishment_basis = punishment_basis.replace('\n', '')
                if '申辩' in content_text or '中辩和申请听证' in content_text or '陈述、中和申请听证' in content_text or \
                        '陈述、申和申请听证' in content_text or '陈述、申\n辩和申请听证' in content_text:
                    # 申辩意见
                    defense_opinion = re.search(r'(我[厅斤]斤?已告知[\s\S]*?你单位[\s\S]*?'
                                                r'(未行使相关权利|提出了?陈述、?申[辯辩]|并进行\n?了听证。?|提出了?听证申请。?|'
                                                r'提出了陈述、申辦申请|提出了陈述、申))',
                                                content_text).group(1).strip()
                    # 申辩意见反馈
                    try:
                        defense_response = re.search(r'(我\n?[厅斤](部分|未)\n?采纳.*?(申辩理由|听证意见|听证\n?理由))。?\n',
                                                     content_text).group(1).strip()
                    except Exception as e:
                        logger.warning(e)
                        defense_response = ''
                else:
                    defense_opinion = defense_response = ''
                defense_opinion = defense_opinion.replace('\n', '')
                defense_response = defense_response.replace('\n', '')
                # 处罚决定
                try:
                    punishment_decision = re.search(r'(依据[\s\S]*?)\n(限你单位|我厅|我斤)', content_text).group(1).strip()
                except Exception as e:
                    logger.warning(e)
                    punishment_decision = ''
                punishment_decision = punishment_decision.replace('\n', '')
            except Exception as e:
                logger.warning(e)
                continue
        elif re.search(r'.*?督[查察]通知$', announcement_title):
            try:
                announcement_code = re.search(
                    r'\n((冀环办函|冀环督字).\d{4}.\d+号|冀环督字\n?.\d{4}.\d+\n?号)\n', content_text).group(1).strip()
            except Exception as e:
                logger.warning(e)
                announcement_code = ''
            try:
                litigant = re.search(r'关于'
                                     r'(严肃查处|查处|调查处理|严肃查处和依法取缔|督办|责令|排查违法排污企业解决|限期解决|排查|解决围场县群众反映)'
                                     r'(.*?)'
                                     r'(非法加工|停产整改|立即停止生产|污水超标问题|噪声扰民问题|环境污染|污染问题|环境问题|'
                                     r'未批先建问题|环境违法问题|污染环境问题|违法处置危险废物|扬尘污染问题|非法排污|擅自停运|'
                                     r'限期搬迁|停止生产|环境违法|存在问题|被沿岸企业污染|违法排污|违法建设生产|违法生产)'
                                     r'.*?督[查察]通知',
                                     announcement_title).group(2).strip()
            except Exception as e:
                logger.warning(e)
                if announcement_title != '藁城市三座砖窑督查通知':
                    litigant = re.search(r'(关于|加强)(.*?)的?督[查察]通知', announcement_title).group(2).strip()
                else:
                    litigant = '藁城市三座砖窑'
            org = re.search(r'督[查察]通知\n((冀环办函|冀环督字).\d{4}.\d+号\n)?([^\d\n]*?)\n', content_text).group(3).strip()
            facts = re.search(
                org + r'([\s\S]*?)'
                      r'(现责成你局|'
                      r'现(请|要求|责成)你局(督导|责令|会同|抓紧|组织|协调|督促|进一步对.*?进行查处|'
                      r'商饶阳县人民政府|依法|按.*?要求逐项落实|对.*?采取有效监管措施|加强|对.*?进行现场核查|责成)|'
                      r'为此，请你局对.*?加强监管|'
                      r'\n针对上述问题，现责成你局会同|'
                      r'为防止信访问题的发生，确保群众的环境质量，现要求你局商正定县人民政府|'
                      r'为严肃环保法律法规，维护群众环境权益，现请你局责成|'
                      r'为(严格执行|严厉打击).*?现责成你局|'
                      r'2008年5月，你局对该厂下达了限期治理通知，现责成你局对|'
                      r'\n2008北京奥运会即将来临，为进一步做好奥运环保保障工作，现责成你局|'
                      r'请你局责成|'
                      r'为此，责成你局立即督促|'
                      r'现要求你局督.*?彻底整改|现要求你局督|'
                      r'现责成你市局督|现要求你局：|'
                      r'\n为做好国家、省级挂牌督办案件的监管工作，请你们督促|'
                      r'请你局会同|'
                      r'责成你市按照省政府办公厅《关于开展清理整顿不法排污企业，保障群众健康环保专项行动》的通知要求|'
                      r'请你局立即组织执法人员|'
                      r'请你市政府立即采取果断措施，对拒不执行政府决定的要严肃查处。|'
                      r'\n请你市组织|'
                      r'为此，你局要严格执法，迅速开展调查|'
                      r'责成你们对辖区内的污染问题|'
                      r'[因为]此，责成你局按照|'
                      r'当前正值全国全省开展清理整顿不法排污企业|'
                      r'因此，责成你局进一步调查|'
                      r'经局领导批示，现责成你|'
                      r'此信访件属群众重复举报，省环保局领导多次批示)',
                content_text).group(1).strip()
            punishment_decision = re.search(
                facts.replace(r'[', r'\[').replace(r']', r'\]').replace(
                    r'(', r'\(').replace(r')', r'\)') + r'([\s\S]*?)'
                                                        r'(\n联系电话及传真|\n联系人|\n联系单位|\n二〇〇八年四月十七日)',
                content_text).group(1).strip()
            punishment_decision = punishment_decision.replace('你局', org.replace('：', ''))
            if re.search(r'\n(.{4}年.{1,2}月.{1,3}日)(\n|$)', content_text):
                real_publish_date = format_date(
                    re.search(r'\n(.{4}年.{1,2}月.{1,3}日)(\n|$)', content_text).group(1).strip())
            else:
                real_publish_date = format_date(each_document['publishDate'])
            defense_opinion = defense_response = ''
            punishment_basis = ''
        else:
            continue

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('河北生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('河北生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('河北生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# xlsx/xls文件
def xlsx_to_json(file_link, announcement_url, origin_data_id, origin_data_title, publishdate):
    file_response = request_site_page(file_link)
    if file_response is None:
        logger.error('xlsx文件下载失败 %s' % file_link)
        return
    with open('./test/tmp.xlsx', 'wb') as tmp_file:
        for chunk in file_response.iter_content(chunk_size=1024):
            if chunk:
                tmp_file.write(chunk)
    with open('./test/tmp.xlsx', 'rb') as xlsx_file:
        xlsx_content = xlsx_file.read()
    if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
        oss_file_map = {
            'origin_url': announcement_url,
            'oss_file_origin_url': announcement_url,
            'origin_url_id': origin_data_id,
            'oss_file_type': 'xlsx',
            'oss_file_name': origin_data_title,
            'oss_file_content': xlsx_content,
            'parsed': False
        }
        response = db.parsed_data.insert_one(oss_file_map)
        file_id = response.inserted_id
        oss_add_file(ali_bucket, str(file_id) + '/' + origin_data_title + '.xlsx', xlsx_content)
        db.environment_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
    else:
        db.environment_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
        file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                           'oss_file_origin_url': announcement_url})['_id']
    logger.info('存入parsed_data')

    excel_data = xlrd.open_workbook('./test/tmp.xlsx')
    sheet = excel_data.sheets()[0]
    result_map_list = []
    for i in range(2, sheet.nrows):
        if sheet.ncols >= 6:
            description = sheet.cell(i, 2).value

            # 文号
            announcement_code_compiler = re.compile(r'(\S环(罚|法|责|停)[\s\S]*?号)')
            try:
                announcement_code = announcement_code_compiler.search(description).group(1).strip()
            except Exception as e:
                logger.info(e)
                announcement_code = ''

            # 当事人
            litigant = sheet.cell(i, 1).value

            # 违规事实
            facts = ''

            # 认定意见
            punishment_basis = ''

            # 申辩意见
            defenseOpinion = ''

            # 申辩意见反馈
            defenseResponse = ''

            # 处罚决定
            punishment_decision = '罚款 ' + str(sheet.cell(i, 3).value) + ' 万元'

            result_map = {
                'announcementTitle': origin_data_title,
                'announcementOrg': '山西生态环境厅',
                'announcementDate': publishdate,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': defenseOpinion,
                'defenseResponse': defenseResponse,
                'litigant': litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            result_map_list.append(result_map)
    logger.info(result_map_list)
    if db.announcement.find({'announcementTitle': origin_data_title, 'oss_file_id': file_id}).count() == 0:
        db.announcement.insert_many(result_map_list)
        logger.info('山西生态环境厅 数据解析 ' + ' -- 数据导入完成')
    else:
        logger.info('山西生态环境厅 数据解析 ' + ' -- 数据已经存在')

    logger.info('删除xlsx文件')
    if os.path.exists('./test/tmp.xlsx'):
        os.remove('./test/tmp.xlsx')

    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
    logger.info('山西生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 山西生态环境厅
def shanxi():
    for each_document in db.environment_data.find({'origin': '山西省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        # 如果是xlsx或xls文件
        if ('.xlsx' in content_soup.find('div', class_='td-con').text) or (
                '.xls' in content_soup.find('div', class_='td-con').text):
            tmp = content_soup.find('div', class_='td-con').find_all('a')
            for each_a in tmp:
                if '.xlsx' or '.xls' in each_a['href']:
                    xlsx_url = str(urljoin('http://sthjt.shanxi.gov.cn', each_a['href']))
                    xlsx_to_json(xlsx_url, announcement_url, each_document['_id'], announcement_title,
                                 format_date(each_document['publishDate'].split(' ')[0]))
            continue

        # 如果是html公告
        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            # 如果是图片
            if content_soup.find('div', class_='td-con').find('img'):
                img_link_list = [urljoin(announcement_url, kk.attrs['src'])
                                 for kk in content_soup.find('div', class_='td-con').find_all('img')]

                for index, each_img_link in enumerate(img_link_list):
                    img_response = request_site_page(each_img_link)
                    with open('./test/' + str(index) + '.jpg', 'wb') as tmp_file:
                        for chunk in img_response.iter_content(chunk_size=1024):
                            if chunk:
                                tmp_file.write(chunk)

                if not os.path.exists('./test/tmp.pdf'):
                    shell_str = 'img2pdf '
                    for index in range(len(img_link_list)):
                        shell_str += './test/' + str(index) + '.jpg '
                    shell_str += '--imgsize 20cmx30cm  -o ./test/tmp.pdf'  # 加入图片尺寸参数，百度ocr有图片尺寸限制
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()

                result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')

                with open('./test/tmp.pdf', 'rb') as pdf_file:
                    pdf_content = pdf_file.read()

                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False,
                    'if_ocr': True,
                    'ocr_result': result_text
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                content_text = result_text

                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.pdf'):
                    os.remove('./test/tmp.pdf')
                for index in range(len(img_link_list)):
                    if os.path.exists('./test/' + str(index) + '.jpg'):
                        os.remove('./test/' + str(index) + '.jpg')
            else:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             content_response.text.encode(content_response.encoding).decode('utf-8'))
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                if content_soup.find('div', class_='td-con'):
                    content_text = content_soup.find('div', class_='td-con').text

        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']
            content_text = db.parsed_data.find_one({'origin_url': announcement_url,
                                                    'oss_file_origin_url': announcement_url})['oss_file_content']

        logger.info('存入parsed_data')

        # 如果无相关信息公开
        if '无相关信息公开' in announcement_title:
            logger.info('山西生态环境厅 数据解析 ' + ' -- 无相关信息公开')
            continue

        if announcement_url == 'http://sthjt.shanxi.gov.cn/html/xzcfjd/20170614/54089.html' \
                or announcement_url == 'http://sthjt.shanxi.gov.cn/html/xzcfjd/20170227/54071.html':
            content_text_list = content_soup.find('div', class_='td-con').find_all('tr')[2:]
            i = 0
            while i < len(content_text_list):
                if ('季度' in content_text_list[i].text) or ('企业名称' in content_text_list[i].text):
                    del (content_text_list[i])
                else:
                    i = i + 1
            result_map_list = []
            for content_text in content_text_list:
                context = content_text.find_all('td')
                # 处罚机构
                announcement_org = context[3].text
                # 处罚日期
                real_publish_date = format_date(each_document['publishDate'].split(' ')[0])
                # 文号
                announcement_code = context[4].text
                # 当事人
                litigant = context[1].text
                # 违规事实
                facts = '超标率： ' + context[2].text
                # 认定意见
                punishment_basis = ''
                # 申辩意见
                defenseOpinion = ''

                # 申辩意见反馈
                defenseResponse = ''

                # 处罚决定
                punishment_decision = context[5].text

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defenseOpinion,
                    'defenseResponse': defenseResponse,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('山西生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('山西生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('山西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif announcement_url == 'http://sthjt.shanxi.gov.cn/html/xzcfjd/20161116/54018.html' or announcement_url == 'http://sthjt.shanxi.gov.cn/html/xzcfjd/20160817/54010.html' or announcement_url == 'http://sthjt.shanxi.gov.cn/html/xzcfjd/20160630/53991.html':

            content_text_list = content_soup.find('div', class_='td-con').find_all('tr')[2:]
            i = 0
            while i < len(content_text_list):
                if ('季度' in content_text_list[i].text) or ('企业名称' in content_text_list[i].text):
                    del (content_text_list[i])
                else:
                    i = i + 1
            result_map_list = []
            for content_text in content_text_list:
                context = content_text.find_all('td')
                # 处罚机构
                announcement_org = '山西生态环境厅'
                # 处罚日期
                real_publish_date = format_date(each_document['publishDate'].split(' ')[0])
                # 文号
                announcement_code = ''
                # 当事人
                litigant = context[2].text
                # 违规事实
                facts = ''
                # 认定意见
                punishment_basis = ''
                # 申辩意见
                defenseOpinion = ''

                # 申辩意见反馈
                defenseResponse = ''

                # 处罚决定
                punishment_decision = context[3].text

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defenseOpinion,
                    'defenseResponse': defenseResponse,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('山西生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('山西生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('山西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue
        elif announcement_url == 'http://sthjt.shanxi.gov.cn/html/xzcfjd/20130929/53912.html':
            content_text_list = content_soup.find('div', class_='td-con').find_all('tr')[2:]
            i = 0
            while i < len(content_text_list):
                if ('序号' in content_text_list[i].text):
                    del (content_text_list[i])
                else:
                    i = i + 1
            result_map_list = []
            for content_text in content_text_list:
                context = content_text.find_all('td')
                # 处罚机构
                announcement_org = '山西生态环境厅'
                # 处罚日期
                real_publish_date = format_date(each_document['publishDate'].split(' ')[0])
                # 文号
                announcement_code = ''
                # 当事人
                litigant = context[1].text
                # 违规事实
                facts = context[2].text
                # 认定意见
                punishment_basis = '违反' + context[3].text
                # 申辩意见
                defenseOpinion = ''

                # 申辩意见反馈
                defenseResponse = ''

                # 处罚决定
                punishment_decision = context[4].text

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defenseOpinion,
                    'defenseResponse': defenseResponse,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('山西生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('山西生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('山西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue
        elif (announcement_url == 'http://sthjt.shanxi.gov.cn/html/xzcfjd/20160504/53960.html') or \
                (announcement_url == 'http://sthjt.shanxi.gov.cn/html/xzcfjd/20160504/53961.html') or \
                (announcement_url == 'http://sthjt.shanxi.gov.cn/html/xzcfjd/20160222/53931.html'):  # 此页面信息为空
            try:
                announcement_code_compiler = re.compile(r'(\S环(罚|法|责|停)[\s\S]*?号)')
                announcement_code = announcement_code_compiler.search(announcement_title).group(1).strip()
            except Exception as e:
                logger.info(e)
                announcement_code = ''
            real_publish_date = format_date(each_document['publishDate'].split(' ')[0])
            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': '',
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': '',
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': '',
                'punishmentBasement': '',
                'punishmentDecision': '',
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('山西生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('山西生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('山西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        # 处罚机构
        announcement_org = '山西生态环境厅'
        # 处罚日期
        real_publish_date = format_date(each_document['publishDate'].split(' ')[0])
        # 文号
        announcement_code_compiler = re.compile(r'(\S环(罚|法|责|停)[\s\S]*?号)')
        announcement_code = announcement_code_compiler.search(content_text).group(1).strip()

        # 当事人
        if re.search(r'(行政处罚决定书|行 政 处 罚 决 定 书)([\s\S]*?)(一、环境违法事实|我厅)', content_text):
            litigant = re.search(r'(行政处罚决定书|行 政 处 罚 决 定 书)([\s\S]*?)(一、环境违法事实|我厅)', content_text).group(2).strip()
        elif re.search(r'书([\s\S]*?)(一、环境违法事实|我厅|环境违法事实)', content_text):
            litigant = re.search(r'书([\s\S]*?)(一、环境违法事实|我厅|环境违法事实)', content_text).group(1).strip()

        # 违规事实
        facts = re.search(r'((根据|我厅)[\s\S]*?)(以上事实|你公司违反了|以上行为)', content_text).group(1).strip()

        # 认定意见
        punishment_basis = re.search(r'((你公司违反|你公司|你厂|尔厂|你单位)[\s\S]*?的规定)', content_text).group(1).strip()

        # 申辩意见
        defenseOpinion = ''

        # 申辩意见反馈
        defenseResponse = ''

        # 处罚决定
        if re.search(r'履行方式和期限([\s\S]*?)收款银行', content_text):
            punishment_decision = re.search(r'履行方式和期限([\s\S]*?)收款银行', content_text).group(1).strip()
        if re.search(r'依据([\s\S]*?)收款银行', content_text):
            punishment_decision = re.search(r'依据([\s\S]*?)收款银行', content_text).group(1).strip()

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defenseOpinion,
            'defenseResponse': defenseResponse,
            'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('山西生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('山西生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('山西生态环境厅 数据解析 ' + ' -- 修改parsed完成')


def neimenggu():
    for each_document in db.environment_data.find({'origin': '内蒙古自治区环境保护厅', 'status': {'$nin': ['ignored']}}):
        announcement_url = each_document['url']
        announcement_title = each_document['title']
        # ignored
        if announcement_url in ['http://sthjt.nmg.gov.cn/hjfw/hjjc/cfgg/201509/t20150901_1491748.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/dbdc/cfgg/201810/t20181023_1582872.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/hjjc/cfgg/201009/t20100929_42960.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/hjjc/cfgg/201009/t20100929_42962.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/hjjc/cfgg/201009/t20100929_42961.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/dbdc/cfgg/201903/t20190307_1588400.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/dbdc/cfgg/201810/t20181023_1582874.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/dbdc/cfgg/201810/t20181023_1582873.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201903/t20190318_1588796.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201903/t20190318_1588795.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201903/t20190318_1588794.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201903/t20190313_1588570.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201810/t20181023_1582859.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201810/t20181023_1582858.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201810/t20181023_1582857.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201810/t20181023_1582854.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201809/t20180918_1577944.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201804/t20180409_1560969.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201804/t20180409_1560968.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201609/t20160908_1510936.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201609/t20160908_1510935.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201512/t20151222_1499827.html', \
                                'http://sthjt.nmg.gov.cn/hjfw/xbdc/cfgg/201507/t20150721_1486152.html']:
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue

        logger.info('url to parse ' + announcement_url)
        logger.info(announcement_title)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        content_text = get_content_text(content_soup.find('div', id='zoomfont'))
        tmp_content_text = get_content_text(content_soup.find('div', class_='TRS_Editor'))
        # 如果公告在.doc附件中
        if re.search(r'\.doc', content_text):
            doc_link = re.search(r'http://[\s\S]*?\.doc', content_text)[0]
            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': doc_link}).count() == 0:
                doc_response = request_site_page(doc_link)
                if doc_response is None:
                    logger.error('doc文件下载失败 %s' % doc_link)
                    continue
                with open('./test/' + announcement_title + '.doc', 'wb') as tmp_file:
                    for chunk in doc_response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)
                # 转换为docx文件
                if not os.path.exists('./test/' + announcement_title + '.docx'):
                    shell_str = 'soffice --headless --convert-to docx ' + './test/' + announcement_title + '.doc' + ' --outdir ./test'
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()
                # 提取docx里的文字和图片
                if not os.path.exists('./test/tmp'):
                    os.makedirs('./test/tmp')
                    content_text = docx2txt.process('./test/' + announcement_title + '.docx', './test/tmp')
                img_list = []
                for filename in os.listdir('./test/tmp'):
                    img_list.append(filename)
                if len(img_list) == 0:
                    logger.error('无法提取docx文件中的图片: ' + doc_link)
                    logger.info('删除.doc&.docx文件')
                    if os.path.exists('./test/' + announcement_title + '.doc'):
                        os.remove('./test/' + announcement_title + '.doc')
                    if os.path.exists('./test/' + announcement_title + '.docx'):
                        os.remove('./test/' + announcement_title + '.docx')
                    logger.info('删除tmp文件夹')
                    if os.path.exists('./test/tmp'):
                        os.removedirs('./test/tmp')
                    continue

                img_list.sort()  # 按文件名顺序排序
                # 转换为pdf
                if not os.path.exists('./test/tmp.pdf'):
                    shell_str = 'img2pdf '
                    for num in range(len(img_list)):
                        shell_str += './test/tmp/' + img_list[num] + ' '
                    shell_str += '--imgsize 20cmx30cm  -o ./test/tmp.pdf'  # 加入图片尺寸参数，百度ocr有图片尺寸限制
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()

                result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')

                with open('./test/tmp.pdf', 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': doc_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False,
                    'if_ocr': True,
                    'ocr_result': result_text
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                content_text = result_text.replace('\n', '')

                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.pdf'):
                    os.remove('./test/tmp.pdf')
                logger.info('删除.doc&.docx文件')
                if os.path.exists('./test/' + announcement_title + '.doc'):
                    os.remove('./test/' + announcement_title + '.doc')
                if os.path.exists('./test/' + announcement_title + '.docx'):
                    os.remove('./test/' + announcement_title + '.docx')
                logger.info('删除tmp文件夹')
                if os.path.exists('./test/tmp'):
                    os.removedirs('./test/tmp')
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': doc_link})['_id']
                content_text = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': doc_link})['ocr_result']

            # 解析
            logger.info(content_text)
            # 处罚机构
            announcement_org = '内蒙古生态环境厅'
            # 处罚日期
            try:
                real_publish_date = format_date(re.findall(r'\d{4}年\d{1,2}月\d{1,2}日')[-1].strip())
            except Exception as e:
                logger.info('没有正确识别处罚日期')
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)
            # 文号
            announcement_code = re.search(r'(内环[\s\S]*?\d{4}[\s\S]*?\d+[\s\S]*?号)', content_text).group(1).strip()

            # 当事人
            litigant = re.search(r'\d+号([\s\S]*?)(一、|我[厅斤]于|环境违法事实和证据|你单位)', content_text).group(
                1).strip()

            # 违规事实
            facts = re.search(r'(事实和证据|环境违法行为|环境违法问题)([\s\S]*?)(以上事实|以上事|上述环境违法行为|上述环境问题)', content_text).group(
                2).strip()

            # 认定意见
            punishment_basis = re.search(r'((你(单位|公司)的上述行为|上述环境违法行为|上述环境问题|上述行为)[\s\S]*?(依法应当予以处罚|条|法律法规|的规定))',
                                         content_text).group(1).strip()
            # 申辩意见：
            try:
                defense_opinion = re.search(
                    r'((\d{4}年\d{1,2}月\d{1,2}日.我厅作出的[\s\S]*?听证申请权)|(我[厅斤]于\d{4}年\d{1,2}月\d{1,2}日以[\s\S]*?听证申请权))',
                    content_text).group(1).strip()
            except Exception as e:
                logger.info(e)
                defense_opinion = ''
            # 申辩意见反馈：
            try:
                defense_response = re.search(r'(在规定的期限内[\s\S]*?也未申请听证)', content_text).group(1).strip()
            except Exception as e:
                logger.info(e)
                defense_response = ''

            # 处罚决定
            if re.search(r'(依据《[\s\S]*?》[\s\S]*?)(收[到款]银行|我厅将对你|如接到通知|三、|你单位如对本决定不服)', content_text):
                punishment_decision = re.search(r'(依据《[\s\S]*?》[\s\S]*?)(收[到款]银行|我厅将对你|如接到通知|三、|你单位如对本决定不服)',
                                                content_text).group(1).strip()
            elif re.search(r'(现责令)([\s\S]*?)(收[到款]银行|我厅将对你|如接到通知|三、|你单位如对本决定不服)', content_text):
                punishment_decision = re.search(r'(现责令)([\s\S]*?)(收[到款]银行|我厅将对你|如接到通知|三、|你单位如对本决定不服)',
                                                content_text).group(2).strip()

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': announcement_org,
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': defense_opinion,
                'defenseResponse': defense_response,
                'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find(
                    {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 修改parsed完成')

        # 如果公告在.pdf附件中
        elif re.search(r'\.pdf', content_text):
            pdf_link_list = re.findall(r'http://[\s\S]*?\.pdf', content_text)
            for pdf_link in pdf_link_list:
                if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': pdf_link}).count() == 0:
                    doc_response = request_site_page(pdf_link)
                    if doc_response is None:
                        logger.error('pdf文件下载失败 %s' % pdf_link)
                        continue
                    with open('./test/tmp.pdf', 'wb') as tmp_file:
                        for chunk in doc_response.iter_content(chunk_size=1024):
                            if chunk:
                                tmp_file.write(chunk)

                    result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')
                    with open('./test/tmp.pdf', 'rb') as pdf_file:
                        pdf_content = pdf_file.read()
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': pdf_link,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'pdf',
                        'oss_file_name': announcement_title,
                        'oss_file_content': pdf_content,
                        'parsed': False,
                        'if_ocr': True,
                        'ocr_result': result_text
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    content_text = result_text.replace('\n', '')

                    logger.info('删除TMP文件')
                    if os.path.exists('./test/tmp.pdf'):
                        os.remove('./test/tmp.pdf')
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = db.parsed_data.find_one(
                        {'origin_url': announcement_url, 'oss_file_origin_url': pdf_link})['_id']
                    content_text = db.parsed_data.find_one(
                        {'origin_url': announcement_url, 'oss_file_origin_url': pdf_link})['ocr_result']

                # 解析
                logger.info(content_text)
                # 处罚机构
                announcement_org = '内蒙古生态环境厅'
                # 处罚日期
                try:
                    real_publish_date = re.findall(r'\S{4}年\S{1,2}月\S{1,2}日', content_text)[-1].strip()
                except Exception as e:
                    logger.info('没有正确识别处罚日期')
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)
                # 文号
                announcement_code = re.search(r'(内环[\s\S]*?\d{4}[\s\S]*?\d+[\s\S]*?号)', content_text).group(1).strip()
                # 当事人
                litigant = re.search(r'\d+号([\s\S]*?)(一、|我[厅斤]于|环境违法事实和证据|你单位)', content_text).group(1).strip()

                # 违规事实
                facts = re.search(r'(事实和证据|环境违法行为|环境违法问题)([\s\S]*?)(以上事实|以上事|上述环境违法行为|上述环境问题)', content_text).group(
                    2).strip()

                # 认定意见
                punishment_basis = re.search(r'((你(单位|公司)的上述行为|上述环境违法行为|上述环境问题|上述行为)[\s\S]*?(依法应当予以处罚|条|法律法规|的规定))',
                                             content_text).group(1).strip()

                # 申辩意见：
                try:
                    defense_opinion = re.search(
                        r'((\d{4}年\d{1,2}月\d{1,2}日.我厅作出的[\s\S]*?听证申请权)|(我[厅斤]于\d{4}年\d{1,2}月\d{1,2}日以[\s\S]*?听证申请权))',
                        content_text).group(1).strip()
                except Exception as e:
                    logger.info(e)
                    defense_opinion = ''
                # 申辩意见反馈：
                try:
                    defense_response = re.search(r'(在规定的期限内[\s\S]*?也未申请听证)', content_text).group(1).strip()
                except Exception as e:
                    logger.info(e)
                    defense_response = ''

                # 处罚决定
                if re.search(r'(依据《[\s\S]*?》[\s\S]*?)(收[到款]银行|我厅将对你|如接到通知|三、|你单位如对本决定不服)', content_text):
                    punishment_decision = re.search(r'(依据《[\s\S]*?》[\s\S]*?)(收[到款]银行|我厅将对你|如接到通知|三、|你单位如对本决定不服)',
                                                    content_text).group(1).strip()
                elif re.search(r'(现责令)([\s\S]*?)(收[到款]银行|我厅将对你|如接到通知|三、|你单位如对本决定不服)', content_text):
                    punishment_decision = re.search(r'(现责令)([\s\S]*?)(收[到款]银行|我厅将对你|如接到通知|三、|你单位如对本决定不服)',
                                                    content_text).group(2).strip()

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find(
                        {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 修改parsed完成')

        # 如果是.xlsx文件：
        elif re.search(r'\.xlsx', content_text):
            xlsx_link = re.search(r'http://[\s\S]*?\.xlsx', content_text)[0]
            file_response = request_site_page(xlsx_link)
            if file_response is None:
                logger.error('xlsx文件下载失败 %s' % xlsx_link)
                return
            with open('./test/tmp.xlsx', 'wb') as tmp_file:
                for chunk in file_response.iter_content(chunk_size=1024):
                    if chunk:
                        tmp_file.write(chunk)
            with open('./test/tmp.xlsx', 'rb') as xlsx_file:
                xlsx_content = xlsx_file.read()
            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': xlsx_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': xlsx_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'xlsx',
                    'oss_file_name': announcement_title,
                    'oss_file_content': xlsx_content,
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.xlsx', xlsx_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': xlsx_link})[
                    '_id']
            logger.info('存入parsed_data')

            excel_data = xlrd.open_workbook('./test/tmp.xlsx')
            sheet = excel_data.sheets()[0]
            result_map_list = []
            for i in range(sheet.nrows):
                if re.search(r'(信息公告|序号)', str(sheet.cell(i, 0).value)):
                    continue
                if sheet.ncols >= 6:
                    # 文号
                    announcement_code = ''

                    # 当事人
                    litigant = sheet.cell(i, 1).value

                    # 违规事实
                    facts = sheet.cell(i, 2).value

                    # 认定意见
                    punishment_basis = sheet.cell(i, 3).value

                    # 申辩意见
                    defenseOpinion = ''

                    # 申辩意见反馈
                    defenseResponse = ''

                    # 处罚决定
                    punishment_decision = sheet.cell(i, 4).value

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': '内蒙古生态环境厅',
                        'announcementDate': format_date(each_document['publishDate'].split(' ')[0]),
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': defenseOpinion,
                        'defenseResponse': defenseResponse,
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据已经存在')

            logger.info('删除xlsx文件')
            if os.path.exists('./test/tmp.xlsx'):
                os.remove('./test/tmp.xlsx')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 修改parsed完成')

        # 如果是限期改正/拒不执行处罚企业：
        elif re.search(r'(\d{4}年限期改正|\d{4}年拒不执行处罚企业|^关于.*的行政处罚$|东部督查中心\d{4}年.季度处罚信息|西部环保督查中心\d{4}年第.季度处罚信息)',
                       announcement_title) \
                and (not re.search(r'(一、环境违法事实和证据|二、行政处罚的依据、种类及其履行方式和期限)', tmp_content_text)):  # 判断不是完整的公告
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             content_response.text.encode(content_response.encoding).decode('utf-8'))
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url})['_id']

            content_text = get_content_text(content_soup.find(class_='TRS_Editor'))
            # 解析
            logger.info(content_text)

            # 当事人
            if re.search(r'----(.*)', announcement_title):
                litigant = re.search(r'----(.*)', announcement_title).group(1).strip()
            elif re.search(r'(企业名称：|公司名称：)([\s\S]*?)违法事实', content_text):
                litigant = re.search(r'(企业名称：|公司名称：)([\s\S]*?)违法事实', content_text).group(2).strip()
            elif re.search(r'(.*?(公司|厂|集团))', announcement_title):
                litigant = re.search(r'(.*?(公司|厂|集团))', announcement_title).group(1).strip()
            else:
                litigant = ''

            # 处罚机构
            announcement_org = '内蒙古生态环境厅'
            # 处罚日期
            try:
                publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,2}日', content_text)[-1]
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.info(e)
                real_publish_date = format_date(each_document['publishDate'].split(' ')[0])
            # 文号
            if re.search(r'(内环[\s\S]*?\d{4}.\d+号)', content_text):
                announcement_code = re.search(r'(内环[\s\S]*?\d{4}.\d+号)', content_text).group(1).strip()
            else:
                announcement_code = ''
            facts = re.search(r'((违法事实|违法行为)[\s\S]*?)(行政命令作出的依据|处罚内容|处罚依据)', content_text).group(1).strip()
            if re.search(r'((行政命令作出的依据|处罚依据)[\s\S]*?)(改正违法行为的期限|改正违法行为的具体形式|处罚内容)', content_text):
                punishment_basis = re.search(r'((行政命令作出的依据|处罚依据)[\s\S]*?)(改正违法行为的期限|改正违法行为的具体形式|处罚内容)',
                                             content_text).group(1).strip()
            else:
                punishment_basis = ''
            punishment_decision = re.search(r'((改正违法行为的期限|处罚内容|改正违法行为的具体形式：)[\s\S]*?)(命令作出机关|执行情况|下达日期)',
                                            content_text).group(1).strip()
            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': announcement_org,
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 修改parsed完成')

        # 如果是限期改正信息公告汇总
        elif re.search(
                r'(\d{4}年第.季度限期改正信息公告（.部地区汇总）|\d{4}年第.季度行政处罚信息公告（.部地区汇总）|\d{4}年度处罚信息公告（汇总）|\d{4}年处罚信息.|\d{4}年限期整改处罚公告汇总|西部环保督查中心\d{4}年\d{1,2}月限期改正信息|内蒙古西部环保督查中心\d{4}年\d{1,2}月处罚信息)',
                announcement_title):
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             content_response.text.encode(content_response.encoding).decode('utf-8'))
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url})['_id']

            if re.search(r'2013年处罚信息二', announcement_title):
                continue

            litigant_list = re.findall(r'(企业名称：|公司名称：)([\s\S]*?)违法事实', content_text)
            facts_list = re.findall(r'违法事实：([\s\S]*?)(行政命令作出的依据|处罚依据)', content_text)
            punishment_basis_list = re.findall(r'(行政命令作出的依据：|处罚依据：)([\s\S]*?)(改正违法行为的期限及具体形式|处罚内容|改正形式)', content_text)
            punishment_decision_list = re.findall(r'(改正违法行为的期限及具体形式：|处罚内容：|改正形式：)([\s\S]*?)(命令作出机关|执行情况|改正期限)',
                                                  content_text)

            # 处罚机构
            announcement_org = '内蒙古生态环境厅'
            # 处罚日期
            real_publish_date = format_date(each_document['publishDate'].split(' ')[0])
            # 文号
            announcement_code = ''
            result_map_list = []
            for i in range(len(litigant_list)):
                # 当事人
                litigant = litigant_list[i][1]
                facts = facts_list[i][0]
                punishment_basis = punishment_basis_list[i][1]
                punishment_decision = punishment_decision_list[i][1]
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 修改parsed完成')

        # 如果是html页面的表格形式
        elif re.search(r'(自治区环保专项行动领导小组召开联席会议研究制定集中检查和整治工作重点|挂牌督办案件名单|举报案件情况表)', announcement_title):
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             content_response.text.encode(content_response.encoding).decode('utf-8'))
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url})['_id']

            table_soup = content_soup.find('div', class_="TRS_Editor").find('table')
            # 处罚机构
            announcement_org = '内蒙古生态环境厅'
            # 处罚日期
            real_publish_date = format_date(each_document['publishDate'].split(' ')[0])
            announcement_code = ''
            result_map_list = []
            tr_list = table_soup.find_all('tr')
            if re.search(r'举报案件情况表', announcement_title):
                del (tr_list[0])
            for each_tr in tr_list:
                if re.search(r'(序号)', each_tr.text.replace('\n', '')):
                    continue
                td_list = each_tr.find_all('td')
                if re.search(r'举报案件情况表', announcement_title):
                    facts = td_list[3].text.strip().replace('\n', '')
                    litigant = td_list[2].text.strip().replace('\n', '')
                    punishment_decision = td_list[4].text.strip().replace('\n', '')
                else:
                    for each_td in td_list:
                        # 违规事实
                        if re.search(r'(超标排放|违反|破坏|影响|国家部际联席会议督办项目|污染|超标|配套的缄回收工程|未按|禁止|未办理|“新五小”企业|不符合|湖面萎缩)',
                                     each_td.text.replace('\n', '')):
                            facts = each_td.text.strip().replace('\n', '')

                        # 当事人
                        elif re.search(r'(公司|厂|集团|企业|骆驼山沿山区域污染综合整治|水泥二部|沙德格（盖）工业园区|炜烨热力有限公司热力项目)',
                                       each_td.text.replace('\n', '')):
                            litigant = each_td.text.strip().replace('\n', '')

                        # 处罚决定
                        elif re.search(r'(限期|以处罚|停产|不得恢复生产|清理整顿|关停取缔|停止建设|实施冬储|关停|补办|完善)',
                                       each_td.text.replace('\n', '')):
                            punishment_decision = each_td.text.strip().replace('\n', '')
                        elif re.search(r'(\d{1,2}月\d{1,2}日)', each_td.text.replace('\n', '')):
                            punishment_decision.__add__(each_td.text.strip().replace('\n', ''))
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 修改parsed完成')

        # 如果是一整篇完整的公告
        else:
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             content_response.text.encode(content_response.encoding).decode('utf-8'))
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url})['_id']
            content_text = get_content_text(content_soup.find(class_="TRS_Editor")).replace('\n', '')
            # 解析
            logger.info(content_text)
            # 处罚机构
            announcement_org = '内蒙古生态环境厅'
            # 处罚日期
            try:
                publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,2}日', content_text)[-1]
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.info(e)
                real_publish_date = format_date(each_document['publishDate'].split(' ')[0])
            # 文号
            try:
                announcement_code = re.search(r'(内环[\s\S]*?\d{4}.\d+号)', content_text).group(1).strip()
            except Exception as e:
                logger.info(e)
                announcement_code = ''

            # 当事人
            litigant = re.search(
                announcement_code.replace('[', '\[').replace(']', '\]') + r'([\s\S]*?)(我[厅斤]于|一、环境违法事实和证据|你单位)',
                content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'(事实和证据|环境违法行为|环境违法问题)([\s\S]*?)(以上事实|以上事|上述环境违法行为|上述环境问题)', content_text).group(
                2).strip()

            # 认定意见
            punishment_basis = re.search(r'((你(单位|公司)的上述行为|上述行为|上述环境违法行为|上述环境问题)[\s\S]*?(依法应当予以处罚|的规定|条|法律法规))',
                                         content_text).group(1).strip()

            # 申辩意见：
            try:
                defense_opinion = re.search(
                    r'((\d{4}年\d{1,2}月\d{1,2}日.我厅作出[\s\S]*?听证申请权)|(我[厅斤]于\d{4}年\d{1,2}月\d{1,2}日以[\s\S]*?听证申请权))',
                    content_text).group(1).strip()
            except Exception as e:
                logger.info(e)
                defense_opinion = ''
            # 申辩意见反馈：
            try:
                defense_response = re.search(r'(在规定的期限内[\s\S]*?也未申请听证)', content_text).group(1).strip()
            except Exception as e:
                logger.info(e)
                defense_response = ''

            # 处罚决定
            punishment_decision = re.search('(履行方式和期限|责令改正的依据、种类|现责令|要求如下)([\s\S]*?)(收(款|到)银行|我厅将对|如接到通知|三、)',
                                            content_text).group(2).strip()

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': announcement_org,
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': defense_opinion,
                'defenseResponse': defense_response,
                'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find(
                    {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('内蒙古生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 6.辽宁生态环境厅
def liaoning():
    for each_document in db.environment_data.find({'origin': '辽宁省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text.encode(content_response.encoding).decode('utf-8'))
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            if content_soup.find('div', class_='headInfo'):
                head_soup = content_soup.find('div', class_='headInfo')
            if content_soup.find('div', class_='TRS_PreAppend'):
                body_soup = content_soup.find('div', class_='TRS_PreAppend')
            elif content_soup.find('div', class_='zy_text'):
                body_soup = content_soup.find('div', class_='zy_text')
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']
            if content_soup.find('div', class_='headInfo'):
                head_soup = content_soup.find('div', class_='headInfo')
            if content_soup.find('div', class_='TRS_PreAppend'):
                body_soup = content_soup.find('div', class_='TRS_PreAppend')
            elif content_soup.find('div', class_='zy_text'):
                body_soup = content_soup.find('div', class_='zy_text')

        logger.info('存入parsed_data')

        # 处罚机构
        announcement_org = '辽宁生态环境厅'
        # 处罚日期
        try:
            publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', body_soup.text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.warning(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)
        # 文号
        announcement_code_compiler = re.compile(r'文号:([\s\S]*?号)')
        try:
            announcement_code = announcement_code_compiler.search(head_soup.text).group(1).strip()
        except Exception as e:
            logger.info(e)
            announcement_code = ''

        p_list = body_soup.find_all('p')
        # 网页里什么内容都没有
        if announcement_url in ['http://sthj.ln.gov.cn/xxgkml/cfjc/hjcf/201104/t20110414_41056.html']:
            announcement_code = ''
            litigant = ''
            facts = ''
            punishment_basis = ''
            defenseOpinion = ''
            defenseResponse = ''
            punishment_decision = ''
        elif announcement_url in ['http://sthj.ln.gov.cn/xxgkml/cfjc/wfsjbg/201308/t20130803_50077.html', \
                                  'http://sthj.ln.gov.cn/xxgkml/cfjc/hbzxxd/201104/t20110414_41174.html', \
                                  'http://sthj.ln.gov.cn/xxgkml/cfjc/hjcf/201104/t20110414_41058.html']:
            content_text = body_soup.text
            # 文号
            announcement_code_compiler = re.compile(r'((辽环函[\s\S]*?号)|(环函[\s\S]*?号))')
            announcement_code = announcement_code_compiler.search(content_text).group(1).strip()

            # 当事人
            litigant = re.search(r'((各市[\s\S]*?局)|(河南省环境保护局))：', content_text).group(1).strip()

            # 违规事实
            if re.search(r'(为(巩固|进一步)[\s\S]*?)(予以挂牌督办|特提出如下要求)', content_text):
                facts = re.search(r'(为(巩固|进一步)[\s\S]*?)(予以挂牌督办|特提出如下要求)', content_text).group(1).strip()
            elif re.search(r'(你局[\s\S]*?收悉)', content_text):
                facts = re.search(r'(你局[\s\S]*?收悉)', content_text).group(1).strip()
            else:
                facts = ''
            # 认定意见
            punishment_basis = ''

            # 申辩意见
            defenseOpinion = ''

            # 申辩意见反馈
            defenseResponse = ''

            # 处罚决定
            punishment_decision = re.search(r'((具体要求|特提出如下要求|函复如下)[\s\S]*?)(联 系 人|辽宁省环境保护厅|二○○七年)', content_text).group(
                1).strip()

        elif announcement_url in ['http://sthj.ln.gov.cn/xxgkml/cfjc/wfsjbg/201308/t20130803_50085.html', \
                                  'http://sthj.ln.gov.cn/xxgkml/cfjc/wfsjbg/201104/t20110414_41178.html', \
                                  'http://sthj.ln.gov.cn/xxgkml/cfjc/hjcf/201104/t20110414_41085.html']:
            content_text = body_soup.text
            # 当事人
            try:
                litigant = re.search(r'(各市[\s\S]*?(各外资银行|局))', content_text).group(1).strip()
            except Exception as e:
                logger.info(e)
                litigant = ''
            # 违规事实
            facts = re.search(r'((为有效打击|现将|在检查中也发现)[\s\S]*?)(现将有关事宜通知如下|并着重做好以下工作|鉴于以上问题)', content_text).group(
                1).strip()

            # 认定意见
            punishment_basis = ''

            # 申辩意见
            defenseOpinion = ''

            # 申辩意见反馈
            defenseResponse = ''

            # 处罚决定
            punishment_decision = re.search(r'(现将有关事宜通知如下：|并着重做好以下工作|鉴于以上问题)([\s\S]*?)(联系人|二○一○年|挂牌督办和绿色信贷)',
                                            content_text).group(2).strip()

        elif announcement_url in ['http://sthj.ln.gov.cn/xxgkml/cfjc/wfsjbg/201104/t20110414_41179.html']:
            content_text = body_soup.text
            result_map_list = []
            facts_list = re.findall(r'存在问题：([\s\S]*?)责任单位', content_text)
            litigant_list = re.findall(r'责任单位：([\s\S]*?)督办要求', content_text)
            punishment_decision_list = re.findall(r'督办要求：([\s\S]*?日)', content_text)
            for i in range(8):
                facts = facts_list[i].strip()
                litigant = litigant_list[i].strip()
                punishment_decision = punishment_decision_list[i].strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': '',
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('辽宁生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('辽宁生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('辽宁生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue
        elif announcement_url in ['http://sthj.ln.gov.cn/xxgkml/cfjc/wfsjbg/201104/t20110414_41180.html']:
            content_text = body_soup.text
            result_map_list = []
            facts_list = re.findall(
                r'((被取缔|“三无”|环保|沈阳成达牧业发展有限公司生鸡屠宰厂无环保手续|无环保审批手续|无环保手续|环保设施未经验收|2台水泥立窑生产线|制浆生产线规模不到3.4万吨|六家企业小堆浸选金项目，生产工艺落后；环保设施不完善|采用鼓风炉炼铜|采用鼓风炉炼铜|未经环保审批|废水|未履行环保批复要求和|50万吨磁选铁矿生产线未经验收|50万吨磁选铁矿生产线未经验收)[\s\S]*?。)',
                content_text)
            litigant_list = re.findall(
                r'((沈阳大|沈阳辽中|辽中|本溪市|沈阳桃|沈阳市|沈阳欣龙|沈阳成达|本溪刘云龙|本溪满族|凤城|营口市|盘锦市|开原市|调兵山|铁岭市|凌源|朝阳汇通|本溪犀牛|辽宁?|凌海市|锦州广玉|阜新市|辽中县|大连北方|大连电镀|鞍山市|抚顺市|本溪钢铁|丹东山水|丹东湘东|辽宁海鹰|义县|阜新东鑫|辽阳西洋|辽宁辽东水泥集团山河水泥)[\s\S]*?(厂|公司))',
                content_text)
            for i in range(16):
                facts = facts_list[i][0].strip()
                litigant = litigant_list[i][0].strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': '',
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': '',
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            for i in range(16, 27):
                facts = facts_list[i - 1][0].strip()
                litigant = litigant_list[i][0].strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': '',
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': '',
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            for i in range(27, 32):
                facts = facts_list[25][0].strip()
                litigant = litigant_list[i][0].strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': '',
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': '',
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            for i in range(32, 35):
                facts = facts_list[i - 6][0].strip()
                litigant = litigant_list[i][0].strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': '',
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': '',
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            for i in range(35, 44):
                facts = facts_list[i - 7][0].strip()
                litigant = litigant_list[i][0].strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': '',
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': '',
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('辽宁生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('辽宁生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('辽宁生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue


        elif len(p_list) == 0:
            content_text = body_soup.text
            # 当事人
            title_text = content_soup.find('h2', class_='zy_title').text
            litigant = re.search(r'(^.+(公司|厂))', title_text).group(1).strip()

            # 违规事实
            facts = re.search(r'基本情况：([\s\S]*?)督办要求', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 申辩意见
            defenseOpinion = ''

            # 申辩意见反馈
            defenseResponse = ''

            # 处罚决定
            punishment_decision = re.search(r'督办要求：([\s\S]*?)督办期限', content_text).group(1).strip()

        elif ('一、' in body_soup.text):
            for index, each_p in enumerate(p_list):
                # 找到当事人所在p标签
                if ('保护局' not in each_p.text) and ('领导小组' not in each_p.text):
                    continue
                # 找到开始有内容的第一个标签下标
                start_index = index
                break
            # 当事人
            litigant = p_list[start_index].text.replace('：', '')
            # 违规事实
            facts = p_list[start_index + 1].text

            # 认定意见
            punishment_basis = ''

            # 申辩意见
            defenseOpinion = ''
            # 申辩意见反馈
            defenseResponse = ''
            # 处罚决定
            punishment_decision = ''
            for index, each_p in enumerate(p_list):
                if ('一、' not in each_p.text):
                    continue
                # 找到一、标签下标
                start_index = index
                break
            for i in range(start_index, len(p_list)):
                punishment_decision += p_list[i].text
        else:
            for index, each_p in enumerate(p_list):
                # 找到当事人所在p标签
                if ('保护局' not in each_p.text) and ('领导小组' not in each_p.text):
                    continue
                # 找到开始有内容的第一个标签下标
                start_index = index
                break
            # 当事人
            litigant = p_list[start_index].text.replace('：', '')

            # 违规事实
            facts = p_list[start_index + 1].text

            # 认定意见
            punishment_basis = ''

            # 申辩意见
            defenseOpinion = ''
            # 申辩意见反馈
            defenseResponse = ''
            # 处罚决定
            punishment_decision = p_list[start_index + 2].text

            if ('附件：解除环境违法案件挂牌督办企业名单' in body_soup.text) or ('附件：解除挂牌督办企业名单' in body_soup.text):
                for index, each_p in enumerate(p_list):
                    # 找到名单起始标签
                    if (each_p.text != '解除环境违法案件挂牌督办企业名单') and (each_p.text != '解除挂牌督办企业名单'):
                        continue
                    # 找到名单的第一个标签下标
                    start_index = index
                    break
                for i in range(start_index, len(p_list)):
                    punishment_decision += p_list[i].text

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defenseOpinion,
            'defenseResponse': defenseResponse,
            'litigant': litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('辽宁生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('辽宁生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('辽宁生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 7.吉林生态环境厅
def jilin():
    for each_document in db.environment_data.find({'origin': '吉林省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            # 如果是pdf
            # 如果是图片
            if '.pdf' in announcement_url:
                doc_response = request_site_page(announcement_url)
                with open('./test/tmp.pdf', 'wb') as tmp_file:
                    for chunk in doc_response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')

                with open('./test/tmp.pdf', 'rb') as pdf_file:
                    pdf_content = pdf_file.read()

                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False,
                    'if_ocr': True,
                    'ocr_result': result_text
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                content_text = result_text

                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.pdf'):
                    os.remove('./test/tmp.pdf')
            else:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             content_response.text.encode(content_response.encoding).decode('utf-8'))
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                if content_soup.find('div', class_='TRS_Editor'):
                    content_text = content_soup.find('div', class_='TRS_Editor').text

        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']
            if '.pdf' in announcement_url:
                content_text = db.parsed_data.find_one({'origin_url': announcement_url,
                                                        'oss_file_origin_url': announcement_url})['ocr_result']
            else:
                if content_soup.find('div', class_='TRS_Editor'):
                    content_text = content_soup.find('div', class_='TRS_Editor').text

        logger.info('存入parsed_data')

        # 处罚机构
        announcement_org = '吉林生态环境厅'
        # 处罚日期
        try:
            publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.warning(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)
        # 文号
        announcement_code_compiler = re.compile(r'(吉环[\s\S]*?号)')
        announcement_code = announcement_code_compiler.search(content_text).group(1).strip()

        # 当事人
        litigant = re.search(r'号([\s\S]*?)(我厅于|我厅|一、环境违法事实)', content_text).group(1).strip()

        # 违规事实
        facts = re.search(r'((我[厅斤]于|我[厅斤])[\s\S]*?)(以上事实有|有现场勘察笔录|有调查询问笔录)', content_text).group(1).strip()

        # 认定意见
        if re.search(r'((你单位的上述行为|你单位将煤焦油)[\s\S]*?的规定)', content_text):
            punishment_basis = re.search(r'((你单位的上述行为|你单位将煤焦油)[\s\S]*?的规定)', content_text).group(1).strip()
        elif re.search(r'等证据为凭。([\s\S]*?的规定)', content_text):
            punishment_basis = re.search(r'等证据为凭。([\s\S]*?的规定)', content_text).group(1).strip()
        elif re.search(r'((你单位的上述行为|你单位将煤焦油)[\s\S]*?)我厅于', content_text):
            punishment_basis = re.search(r'((你单位的上述行为|你单位将煤焦油)[\s\S]*?)我厅于', content_text).group(1).strip()
        elif re.search(r'等证据为凭。([\s\S]*?)我厅于', content_text):
            punishment_basis = re.search(r'等证据为凭。([\s\S]*?)我厅于', content_text).group(1).strip()
        else:
            punishment_basis = ''

        # 处罚决定
        punishment_decision = re.search(r'(依据[\s\S]*?)(收缴银行|我厅将对你单位改正违法行为的情况进行监督)', content_text).group(1).strip()

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('吉林生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('吉林生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('吉林生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 8.黑龙江生态环境厅
def heilongjiang():
    for each_document in db.environment_data.find({'origin': '黑龙江省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue

        logger.info('url to parse ' + announcement_url)
        logger.info(announcement_title)

        # ignored
        if re.search(r'(关于黑龙江省鸡西市多家石墨矿环境问题调查处理情况的报告|关于开展违规建设项目查处情况信息公开的通知|关于切实加强环境监管执法依法履行环境行政权力的通知)',
                     announcement_title):
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text.encode(content_response.encoding).decode('utf-8'))
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        if content_soup.find('div', class_='txt').find('table', class_='MsoNormalTable'):  # 如果页面中是表格汇总表
            content_text = content_soup.find('div', class_='txt').find('div', class_="WordSection1").text
            content_table = content_soup.find('div', class_='txt').find('table', class_='MsoNormalTable').find_all('tr')
        elif content_soup.find('div', class_='txt').find('div', class_="WordSection1"):
            content_text = content_soup.find('div', class_='txt').find('div', class_="WordSection1").text
        else:
            content_text = content_soup.find('div', class_='txt').text

        logger.info('存入parsed_data')

        # 处罚机构
        announcement_org = '黑龙江生态环境厅'
        # 处罚日期
        try:
            publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.warning(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)
        # 文号
        announcement_code_compiler = re.compile(r'(黑环[\s\S]*?号)')
        try:
            announcement_code = announcement_code_compiler.search(announcement_title).group(1).strip()
        except Exception as e:
            logger.info(e)
            announcement_code = ''
        if announcement_url in ['http://www.hljdep.gov.cn/hbzl/zzhjwfxwztbd/2018/04/18538.html', \
                                'http://www.hljdep.gov.cn/hbzl/zzhjwfxwztbd/2017/02/14691.html', \
                                'http://www.hljdep.gov.cn/hbzl/zzhjwfxwztbd/2017/02/14694.html', \
                                'http://www.hljdep.gov.cn/hbzl/zzhjwfxwztbd/2017/02/14699.html']:
            content_table = content_soup.find('div', class_='txt').find_all('tr')
            if announcement_url in ['http://www.hljdep.gov.cn/hbzl/zzhjwfxwztbd/2018/04/18538.html']:
                # 认定意见
                punishment_basis = ''
                del (content_table[0:3])
                result_map_list = []
                for each_line in content_table:
                    td_list = each_line.find_all('td')

                    # 当事人
                    litigant = td_list[2].text.replace('\n', '')

                    # 违规事实
                    facts = (td_list[5].text + td_list[6].text).replace('\n', '')

                    # 处罚决定
                    punishment_decision = td_list[7].text.replace('\n', '')

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            elif announcement_url in ['http://www.hljdep.gov.cn/hbzl/zzhjwfxwztbd/2017/02/14691.html']:
                # 认定意见
                punishment_basis = ''
                del (content_table[0:3])
                result_map_list = []
                for each_line in content_table:
                    td_list = each_line.find_all('td')

                    # 当事人
                    litigant = td_list[1].text.replace('\n', '')

                    # 违规事实
                    facts = td_list[5].text.replace('\n', '')

                    # 处罚决定
                    punishment_decision = '罚款' + td_list[6].text.replace('\n', '') + '万元 ' + td_list[7].text.replace(
                        '\n', '')

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            elif announcement_url in ['http://www.hljdep.gov.cn/hbzl/zzhjwfxwztbd/2017/02/14694.html',
                                      'http://www.hljdep.gov.cn/hbzl/zzhjwfxwztbd/2017/02/14699.html']:
                del (content_table[0])
                result_map_list = []
                for each_line in content_table:
                    td_list = each_line.find_all('td')

                    # 当事人
                    litigant = td_list[1].text.replace('\n', '')

                    # 违规事实
                    facts = td_list[2].text.replace('\n', '')

                    # 认定意见
                    punishment_basis = td_list[3].text.replace('\n', '')

                    # 处罚决定
                    punishment_decision = td_list[5].text.replace('\n', '')

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据已经存在')

        elif content_soup.find('div', class_='txt').find('table', class_='MsoNormalTable'):  # 如果页面中是表格汇总表
            # 认定意见
            punishment_basis = ''

            # 处罚决定
            bureau_name = re.search(r'(.*保护局)', content_text).group(1).strip()
            punishment_decision = re.search(r'(经研究[\s\S]*?)按有关要求', content_text).group(1).strip().replace('你局',
                                                                                                          bureau_name)

            table_title_text = content_table[0].text
            del (content_table[0])
            result_map_list = []
            if re.search(r'区县', table_title_text):
                for each_line in content_table:
                    td_list = each_line.find_all('td')
                    # 当事人
                    litigant = td_list[2].text.replace('\n', '')

                    # 违规事实
                    facts = td_list[3].text.replace('\n', '')

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            else:
                for each_line in content_table:
                    td_list = each_line.find_all('td')
                    # 当事人
                    litigant = td_list[1].text.replace('\n', '')

                    # 违规事实
                    facts = td_list[2].text.replace('\n', '')

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)

            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据已经存在')

        elif announcement_url in ['http://www.hljdep.gov.cn/hbzl/zzhjwfxwztbd/2017/08/16342.html']:  # 如果页面中是表格汇总表
            content_table = content_soup.find('div', style='page:WordSection4;layout-grid:15.6pt').find(
                'table').find_all('tr')
            del (content_table[0])
            # 认定意见
            punishment_basis = ''

            # 处罚决定
            bureau_name = re.search(r'(.*保护局)', content_text).group(1).strip()
            punishment_decision = re.search(r'(经研究[\s\S]*?)按有关要求', content_text).group(1).strip().replace('你局',
                                                                                                          bureau_name)

            result_map_list = []
            for each_line in content_table:
                td_list = each_line.find_all('td')
                # 当事人
                litigant = td_list[1].text.replace('\n', '')

                # 违规事实
                facts = td_list[2].text.replace('\n', '')

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据已经存在')
        elif re.search(r'解除查封决定书', content_text):  # 解除查封决定书
            # 当事人
            litigant = re.search(r'解除查封决定书([\s\S]*?)：', content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'((我厅于|我厅)[\s\S]*?(因查封期限届满|进行了称重))', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'((现根据|根据)[\s\S]*?决定解除查封)', content_text).group(1).strip()

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': announcement_org,
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据已经存在')

        else:
            # 当事人
            litigant = re.search(r'([\s\S]*?)(我厅于|一、环境违法事实和证据|一、有关法律要求和你公司行为的性质)', content_text).group(1).strip()

            # 违规事实
            if re.search(r'(我厅于[\s\S]*?)以上事实', content_text):
                facts = re.search(r'(我厅于[\s\S]*?)以上事实', content_text).group(1).strip()
            elif re.search(r'(一、环境违法事实和证据|一、有关法律要求和你公司行为的性质)([\s\S]*?)(以上行为违反了|根据)', content_text):
                facts = re.search(r'(一、环境违法事实和证据|一、有关法律要求和你公司行为的性质)([\s\S]*?)(以上行为违反了|根据)', content_text).group(
                    2).strip()

            # 认定意见
            if re.search(r'((以上行为违反了|你单位上述行为违反了|上述行为违反了)[\s\S]*?([的之]规定|规定))', content_text):
                punishment_basis = re.search(r'((以上行为违反了|你单位上述行为违反了|上述行为违反了)[\s\S]*?([的之]规定|规定))', content_text).group(
                    1).strip()
            elif re.search(r'根据[\s\S]*?规定', content_text):
                punishment_basis = re.search(r'根据[\s\S]*?规定', content_text).group(0).strip()
            # 处罚决定
            if re.search(
                    r'(二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类|二、告知内容)([\s\S]*?)(收款银行|收款单位|根据《中华人民共和国行政处罚法》第三十一条、第三十二条和第四十二条的规定)',
                    content_text):
                punishment_decision = re.search(
                    r'(二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类|二、告知内容)([\s\S]*?)(收款银行|收款单位|根据《中华人民共和国行政处罚法》第三十一条、第三十二条和第四十二条的规定)',
                    content_text).group(2).replace('三、责令改正和行政处罚决定的履行方式和期限', '').replace('三、处罚决定的履行方式和期限', '')
            elif re.search(r'([依根]据[\s\S]*?)(如你对本行政强制措施不服|三、处罚决定的履行方式和期限|收款银行)', content_text):
                punishment_decision = re.search(r'([依根]据[\s\S]*?)(如你对本行政强制措施不服|三、处罚决定的履行方式和期限|收款银行)',
                                                content_text).group(1).strip()

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': announcement_org,
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 数据已经存在')

        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('黑龙江生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 浙江生态环境厅
def zhejiang():
    for each_document in db.environment_data.find({'origin': '浙江省环境保护厅', 'status': {'$nin': ['ignored']}}):
        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue

        logger.info('url to parse ' + announcement_url)
        logger.info(announcement_title)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        # 如果行政处罚决定书是在.doc附件里
        if re.search(r'(行政处罚决定书|行政处罚听证告知书)', announcement_title) and content_soup.find(id='zoom').find('a'):
            docx_link = urljoin(announcement_url, content_soup.find(id='zoom').find('a')['href'])
            link_type = docx_link.split('.')[-1]
            response = request_site_page(docx_link)
            if response is None:
                logger.error('网页请求错误')
                return
            with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        tmp_file.write(chunk)

            if not os.path.exists('./test/tmp.docx'):
                shell_str = 'soffice --headless --convert-to docx ' + \
                            './test/tmp.' + link_type + ' --outdir ./test'
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.communicate()

            with open('./test/tmp.docx', 'rb') as docx_file:
                docx_content = docx_file.read()

            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': docx_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': docx_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'docx',
                    'oss_file_name': announcement_title,
                    'oss_file_content': docx_content,
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': docx_link})[
                    '_id']
            logger.info('存入parsed_data')

            # 开始解析
            content_text = docx2txt.process('./test/tmp.docx')
            # 处罚机构
            announcement_org = '浙江生态环境厅'
            # 处罚日期
            try:
                publish_date = re.findall(r'\S{4}年\S{1,2}月\S{1,3}日', content_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.warning(e)
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)
            # 文号
            try:
                announcement_code = re.search(r'(\w环\w{1,2}.\d{4}.\d+号)', content_text).group(1).strip()
            except Exception as e:
                logger.info(e)
                announcement_code = ''

            # 当事人
            litigant = re.search(r'([\s\S]*?)(经查|我厅|根据群众举报|\d{4}年\d{1,2}月\d{1,2}日)', content_text).group(
                1).strip().replace(announcement_code, '')

            # 违规事实
            if re.search(litigant.replace('(', '\(').replace(')',
                                                             '\)') + r'([\s\S]*?)(我厅认为|以上事实，有现场检查|你公司上述行为违反|据此，你(厂|公司)的行为违反了|根据《[\s\S]*?》)',
                         content_text):
                facts = re.search(litigant.replace('(', '\(').replace(')',
                                                                      '\)') + r'([\s\S]*?)(我厅认为|以上事实，有现场检查|你公司上述行为违反了|据此，你(厂|公司)的行为违反了|根据《[\s\S]*?》)',
                                  content_text).group(1).strip()
            else:
                facts = re.search(
                    r'(\d{4}年\d{1,2}月\d{1,2}日[\s\S]*?)(我厅认为|以上事实，有现场检查|你公司上述行为违反了|据此，你(厂|公司)的行为违反了|根据《[\s\S]*?》)',
                    content_text).group(1).strip()

            # 认定意见
            punishment_basis = re.search(
                r'((上述行为违反了|你(厂|公司)的行为违反了)[\s\S]*?的规定|我厅认为[\s\S]*?(已构成违法|申请该建设项目竣工环境保护验收”的规定)|根据《[\s\S]*?》[\s\S]*?已构成违法)',
                content_text).group(1).strip()

            # 申辩意见
            if re.search(
                    r'(\d{4}年\d{1,2}月\d{1,2}日.我.向你.{1,2}送达[\s\S]*?(充分考虑了你公司提出的意见和要求|也未(进行|提出)陈述申辩|也未提出陈述、申辩意见|你公司逾期未提出听证申请|未提出听证申请))',
                    content_text):
                defense_opinion = re.search(
                    r'(\d{4}年\d{1,2}月\d{1,2}日.我.向你.{1,2}送达[\s\S]*?(充分考虑了你公司提出的意见和要求|也未(进行|提出)陈述申辩|也未提出陈述、申辩意见|你公司逾期未提出听证申请|未提出听证申请))',
                    content_text).group(1).strip()
            else:
                defense_opinion = ''
            # 申辩意见反馈
            if re.search(
                    r'(环境保护相关法律并未规定未经验收合格可投入使用的例外情形，故对你公司所作的陈述申辩意见不予采纳|我厅认为，根据法律规定，涉案建设项目应当履行建设项目环境影响评价制度[\s\S]*?依法应当追究法律责任|我局对你的厂陈述进行了复核，决定不予采纳)',
                    content_text):
                defense_response = re.search(
                    r'(环境保护相关法律并未规定未经验收合格可投入使用的例外情形，故对你公司所作的陈述申辩意见不予采纳|我厅认为，根据法律规定，涉案建设项目应当履行建设项目环境影响评价制度[\s\S]*?依法应当追究法律责任|我局对你的厂陈述进行了复核，决定不予采纳)',
                    content_text).group(1).strip()
            else:
                defense_response = ''

            # 处罚决定
            if re.search(
                    r'((根据《中华人民共和国环境影响评价法》|根据《建设项目环境保护管理条例》|根据《建设项目竣工环境保护验收管理办法》|根据《浙江省环境污染监督管理办法》|根据《中华人民共和国放射性污染防治法》|根据《中华人民共和国水污染防治法》|根据《中华人民共和国固体废物污染环境防治法》)[\s\S]*?)(\S{4}年\S{1,3}月\S{1,3}日|我厅地址)',
                    content_text):
                punishment_decision = re.search(
                    r'((根据《中华人民共和国环境影响评价法》|根据《建设项目环境保护管理条例》|根据《建设项目竣工环境保护验收管理办法》|根据《浙江省环境污染监督管理办法》|根据《中华人民共和国放射性污染防治法》|根据《中华人民共和国水污染防治法》|根据《中华人民共和国固体废物污染环境防治法》)[\s\S]*?)(\S{4}年\S{1,3}月\S{1,3}日|我厅地址)',
                    content_text).group(1)
            elif re.search(r'(根据《[\s\S]*?》[\s\S]*?实施按日连续处罚)', content_text):
                punishment_decision = re.search(r'(根据《[\s\S]*?》[\s\S]*?实施按日连续处罚)', content_text).group(1)
            else:
                punishment_decision = re.search(r'已构成违法([\s\S]*?)\S{4}年\S{1,2}月\S{1,2}日', content_text).group(1).strip()

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': announcement_org,
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': defense_opinion,
                'defenseResponse': defense_response,
                'litigant': litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('浙江生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('浙江生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('浙江生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            # 删除文件
            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.doc'):
                os.remove('./test/tmp.doc')
            if os.path.exists('./test/tmp.docx'):
                os.remove('./test/tmp.docx')
            continue

        # 附件中是.docx的复议案件表格的情况
        if re.search(r'(行政复议案件情况|行政复议情况|行政处罚结果信息公开一览表|环境违法案件查处登记表)', announcement_title) and content_soup.find(
                id='zoom').find('a'):
            docx_link = urljoin(announcement_url, content_soup.find(id='zoom').find('a')['href'])
            link_type = docx_link.split('.')[-1]
            response = request_site_page(docx_link)
            if response is None:
                logger.error('网页请求错误')
                return
            with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        tmp_file.write(chunk)

            if not os.path.exists('./test/tmp.docx'):
                shell_str = 'soffice --headless --convert-to docx ' + \
                            './test/tmp.' + link_type + ' --outdir ./test'
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.communicate()

            with open('./test/tmp.docx', 'rb') as docx_file:
                docx_content = docx_file.read()

            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': docx_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': docx_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'docx',
                    'oss_file_name': announcement_title,
                    'oss_file_content': docx_content,
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': docx_link})[
                    '_id']
            logger.info('存入parsed_data')

            # 开始解析
            document = Document('./test/tmp.docx')
            result_map_list = []
            tables = document.tables
            for table in tables:
                if (len(table.columns)) == 8:
                    for row in table.rows:
                        if '序号' in row.cells[0].text.replace('\n', '') or row.cells[2].text.replace('\n', '') == '':
                            continue
                        # 处罚机构
                        announcement_org = '浙江生态环境厅'
                        # 处罚日期
                        real_publish_date = format_date(row.cells[4].text)
                        # 文号
                        announcement_code = row.cells[1].text

                        # 当事人
                        litigant = row.cells[2].text

                        # 违规事实
                        facts = row.cells[5].text

                        # 处罚决定
                        punishment_decision = row.cells[6].text + ' ' + row.cells[7].text

                        # 处罚依据
                        punishment_basis = ''
                        result_map = {
                            'announcementTitle': announcement_title,
                            'announcementOrg': announcement_org,
                            'announcementDate': real_publish_date,
                            'announcementCode': announcement_code,
                            'facts': facts,
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': litigant,
                            'punishmentBasement': punishment_basis,
                            'punishmentDecision': punishment_decision,
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                elif (len(table.columns)) == 7:
                    for row in table.rows:
                        if '序号' in row.cells[0].text.replace('\n', '') or row.cells[2].text.replace('\n', '') == '':
                            continue
                        # 处罚机构
                        announcement_org = '浙江生态环境厅'
                        # 处罚日期
                        real_publish_date = format_date(each_document['publishDate'])
                        # 文号
                        announcement_code = row.cells[1].text

                        # 当事人
                        litigant = row.cells[2].text

                        # 违规事实
                        facts = row.cells[5].text

                        # 处罚决定
                        punishment_decision = row.cells[6].text

                        # 处罚依据
                        punishment_basis = ''
                        result_map = {
                            'announcementTitle': announcement_title,
                            'announcementOrg': announcement_org,
                            'announcementDate': real_publish_date,
                            'announcementCode': announcement_code,
                            'facts': facts,
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': litigant,
                            'punishmentBasement': punishment_basis,
                            'punishmentDecision': punishment_decision,
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                elif (len(table.columns)) == 10:
                    for row in table.rows:
                        if '序号' in row.cells[0].text.replace('\n', '') or row.cells[2].text.replace('\n', '') == '':
                            continue
                        # 处罚机构
                        announcement_org = '浙江生态环境厅'
                        # 处罚日期
                        real_publish_date = format_date(row.cells[4].text)
                        # 文号
                        announcement_code = row.cells[5].text

                        # 当事人
                        litigant = row.cells[2].text + ' 法人代表(负责人)：' + row.cells[3].text

                        # 违规事实
                        facts = row.cells[6].text

                        # 处罚决定
                        punishment_decision = row.cells[7].text + ' ' + row.cells[8].text + ' ' + row.cells[9].text

                        # 处罚依据
                        punishment_basis = ''
                        result_map = {
                            'announcementTitle': announcement_title,
                            'announcementOrg': announcement_org,
                            'announcementDate': real_publish_date,
                            'announcementCode': announcement_code,
                            'facts': facts,
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': litigant,
                            'punishmentBasement': punishment_basis,
                            'punishmentDecision': punishment_decision,
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                elif (len(table.columns)) == 14:
                    for row in table.rows:
                        if '序号' in row.cells[0].text.replace('\n', '') or row.cells[2].text.replace('\n', '') == '':
                            continue
                        # 处罚机构
                        announcement_org = '浙江生态环境厅'
                        # 处罚日期
                        try:
                            year = re.search(r'(\d{4}年)', announcement_title).group(1).strip()
                            real_publish_date = format_date(year + row.cells[10].text.replace('\n', ''))
                        except Exception as e:
                            logger.info(e)
                            real_publish_date = format_date(each_document['publishDate'])

                        # 文号
                        announcement_code = row.cells[1].text.replace('\n', '')

                        # 当事人
                        litigant = row.cells[2].text.replace('\n', '') + ' 法人代表(负责人)：' + row.cells[3].text.replace('\n',
                                                                                                                   '')

                        # 违规事实
                        facts = row.cells[4].text.replace('\n', '')

                        # 处罚决定
                        punishment_decision = row.cells[5].text.replace('\n', '') + row.cells[7].text.replace('\n',
                                                                                                              '') + \
                                              row.cells[8].text.replace('\n', '')

                        # 处罚依据
                        punishment_basis = row.cells[6].text.replace('\n', '')
                        result_map = {
                            'announcementTitle': announcement_title,
                            'announcementOrg': announcement_org,
                            'announcementDate': real_publish_date,
                            'announcementCode': announcement_code,
                            'facts': facts,
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': litigant,
                            'punishmentBasement': punishment_basis,
                            'punishmentDecision': punishment_decision,
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('浙江生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('浙江生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('浙江生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            # 删除文件
            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.doc'):
                os.remove('./test/tmp.doc')
            if os.path.exists('./test/tmp.docx'):
                os.remove('./test/tmp.docx')
            continue

        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                'content_id_name': 'zoom',
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text.encode(content_response.encoding).decode('utf-8'))
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        content_id_soup = content_soup.find('div', id='zoom')
        content_text = get_content_text(content_soup.find('div', id='zoom'))
        logger.info('存入parsed_data')

        # 人工处理
        if re.search(r'关于吊销温州市中金岭南科技环保有限公司的危险废物经营许可证的公示|国网浙江省电力公司宁波供电公司行政处罚案', announcement_title):
            logger.warning('需人工处理 ...')
            continue

        # 处罚机构
        announcement_org = '浙江生态环境厅'
        # 处罚日期
        try:
            publish_date = re.findall(r'\S{4}年\S{1,2}月\S{1,3}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.warning(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

        # 如果是html页面表格的形式
        if (content_id_soup.find('table')):
            tr_list = content_id_soup.find('table').find_all('tr')
            result_map_list = []
            for each_tr in tr_list:
                if re.search(r'(序号)', get_content_text(each_tr).strip().replace('\n', '')):
                    continue

                td_list = each_tr.find_all('td')
                # 文号
                announcement_code = get_content_text(td_list[1])

                # 当事人
                litigant = get_content_text(td_list[2]) + ' 法定代表人（负责人）:' + get_content_text(td_list[3])

                # 违规事实
                facts = get_content_text(td_list[4])

                # 处罚决定
                punishment_decision = get_content_text(td_list[5])

                # 处罚依据
                punishment_basis = get_content_text(td_list[6])

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('浙江生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('浙江生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('浙江生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(r'第.季度主要污染物严重超标国家重点监控企业处理情况', announcement_title):
            # 文号
            announcement_code = re.search(r'(\w环.{2,3}\d{4}.\d+号)', content_text).group(1).strip()

            # 当事人
            litigant = re.search(r'处理情况如下：([\s\S]*?(厂：|公司：|集团：|单位：))', content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'(厂：|公司：|集团：|单位：)([\s\S]*?要求恢复生产)', content_text).group(2).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'(经监测[\s\S]*?同意其恢复生产)', content_text).group(1)

            defense_opinion = ''
            defense_response = ''

        elif re.search(r'行政处罚强制执行申请书', announcement_title):
            # 文号
            announcement_code = ''

            # 当事人
            litigant = re.search(r'([\s\S]*?)(\d{4}年\d{1,2}月\d{1,2}日|我厅)', content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'((\d{4}年\d{1,2}月\d{1,2}日|我厅)[\s\S]*?)根据《中华人民共和国行政诉讼法》', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'(根据《中华人民共和国行政诉讼法》[\s\S]*?及相关证据材料)', content_text).group(1)

            defense_opinion = ''
            defense_response = ''

        elif re.search(r'履行行政处罚决定催告书', announcement_title):
            # 文号
            announcement_code = ''

            # 当事人
            litigant = re.search(r'([\s\S]*?)(\d{4}年\d{1,2}月\d{1,2}日|我厅)', content_text).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                        '\]') + r'([\s\S]*?)你公司在法定期限内未申请复议',
                              content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'(你公司在法定期限内未申请复议[\s\S]*?申请强制执行)', content_text).group(1)

            defense_opinion = ''
            defense_response = ''

        elif re.search(r'行政处罚听证通知书', announcement_title):
            # 文号
            announcement_code = re.search(r'(\w环.{2,3}\d{4}.\d+号)', content_text).group(1).strip()

            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                                    '\]') + r'([\s\S]*?)(根据《中华人民共和国行政处罚法》)',
                                 content_text).group(1).strip()

            # 违规事实
            facts = ''

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'(根据《中华人民共和国行政处罚法》[\s\S]*?事先告知我厅联系人)', content_text).group(1)

            defense_opinion = ''
            defense_response = ''

        elif re.search(r'准许[\s\S]*撤回行政处罚听证申请通知书', announcement_title):
            # 文号
            announcement_code = re.search(r'(\w环.{2,3}\d{4}.\d+号)', content_text).group(1).strip()

            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                                    '\]') + r'([\s\S]*?)(\d{4}年\d{1,2}月\d{1,2}日)',
                                 content_text).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                        '\]') + r'([\s\S]*?)经审查，现准许你单位撤回听证申请',
                              content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'(经审查，现准许你单位撤回听证申请)', content_text).group(1)

            defense_opinion = ''
            defense_response = ''

        elif re.search(r'行政处罚案件查处通知书', announcement_title):
            # 文号
            announcement_code = ''

            # 当事人
            litigant = re.search(r'([\s\S]*?)(经我厅\d{4}年\d{1,2}月)', content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'(经我厅\d{4}年\d{1,2}月[\s\S]*?环境违法行为)', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'(因该案中涉嫌违反建设项目[\s\S]*?)特此通知', content_text).group(1)

            defense_opinion = ''
            defense_response = ''

        # 如果是一篇完整的行政处罚决定书
        else:
            try:
                # 文号
                announcement_code = re.search(r'(\w环\w{1,2}.\d{4}.\d+号)', content_text).group(1).strip()
            except Exception as e:
                logger.info(e)
                announcement_code = ''

            # 当事人
            litigant = re.search(r'([\s\S]*?)(经查|我厅|根据群众举报|\d{4}年\d{1,2}月\d{1,2}日)', content_text).group(
                1).strip().replace(announcement_code, '')

            # 违规事实
            if re.search(litigant.replace('(', '\(').replace(')',
                                                             '\)') + r'([\s\S]*?)(我厅认为|以上事实，有现场检查|你公司上述行为违反|据此，你(厂|公司)的行为违反了|根据《[\s\S]*?》)',
                         content_text):
                facts = re.search(litigant.replace('(', '\(').replace(')',
                                                                      '\)') + r'([\s\S]*?)(我厅认为|以上事实，有现场检查|你公司上述行为违反了|据此，你(厂|公司)的行为违反了|根据《[\s\S]*?》)',
                                  content_text).group(1).strip()
            else:
                facts = re.search(
                    r'(\d{4}年\d{1,2}月\d{1,2}日[\s\S]*?)(我厅认为|以上事实，有现场检查|你公司上述行为违反了|据此，你(厂|公司)的行为违反了|根据《[\s\S]*?》)',
                    content_text).group(1).strip()

            # 认定意见
            punishment_basis = re.search(
                r'((上述行为违反了|你(厂|公司)的行为违反了)[\s\S]*?的规定|我厅认为[\s\S]*?(已构成违法|申请该建设项目竣工环境保护验收”的规定)|根据《[\s\S]*?》[\s\S]*?已构成违法)',
                content_text).group(1).strip()

            # 申辩意见
            if re.search(
                    r'(\d{4}年\d{1,2}月\d{1,2}日.我.向你.{1,2}送达[\s\S]*?(充分考虑了你公司提出的意见和要求|也未(进行|提出)陈述申辩|也未提出陈述、申辩意见|你公司逾期未提出听证申请|未提出听证申请))',
                    content_text):
                defense_opinion = re.search(
                    r'(\d{4}年\d{1,2}月\d{1,2}日.我.向你.{1,2}送达[\s\S]*?(充分考虑了你公司提出的意见和要求|也未(进行|提出)陈述申辩|也未提出陈述、申辩意见|你公司逾期未提出听证申请|未提出听证申请))',
                    content_text).group(1).strip()
            else:
                defense_opinion = ''
            # 申辩意见反馈
            if re.search(
                    r'(环境保护相关法律并未规定未经验收合格可投入使用的例外情形，故对你公司所作的陈述申辩意见不予采纳|我厅认为，根据法律规定，涉案建设项目应当履行建设项目环境影响评价制度[\s\S]*?依法应当追究法律责任|我局对你的厂陈述进行了复核，决定不予采纳)',
                    content_text):
                defense_response = re.search(
                    r'(环境保护相关法律并未规定未经验收合格可投入使用的例外情形，故对你公司所作的陈述申辩意见不予采纳|我厅认为，根据法律规定，涉案建设项目应当履行建设项目环境影响评价制度[\s\S]*?依法应当追究法律责任|我局对你的厂陈述进行了复核，决定不予采纳)',
                    content_text).group(1).strip()
            else:
                defense_response = ''

            # 处罚决定
            if re.search(
                    r'((根据《中华人民共和国环境影响评价法》|根据《建设项目环境保护管理条例》|根据《建设项目竣工环境保护验收管理办法》|根据《浙江省环境污染监督管理办法》|根据《中华人民共和国放射性污染防治法》|根据《中华人民共和国水污染防治法》|根据《中华人民共和国固体废物污染环境防治法》)[\s\S]*?)(\S{4}年\S{1,3}月\S{1,3}日|我厅地址)',
                    content_text):
                punishment_decision = re.search(
                    r'((根据《中华人民共和国环境影响评价法》|根据《建设项目环境保护管理条例》|根据《建设项目竣工环境保护验收管理办法》|根据《浙江省环境污染监督管理办法》|根据《中华人民共和国放射性污染防治法》|根据《中华人民共和国水污染防治法》|根据《中华人民共和国固体废物污染环境防治法》)[\s\S]*?)(\S{4}年\S{1,3}月\S{1,3}日|我厅地址)',
                    content_text).group(1)
            elif re.search(r'(根据《[\s\S]*?》[\s\S]*?实施按日连续处罚)', content_text):
                punishment_decision = re.search(r'(根据《[\s\S]*?》[\s\S]*?实施按日连续处罚)', content_text).group(1)
            else:
                punishment_decision = re.search(r'已构成违法([\s\S]*?)\S{4}年\S{1,2}月\S{1,2}日', content_text).group(1).strip()

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant.replace('：', ''),
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('浙江生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('浙江生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('浙江生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 安徽生态环境厅
def anhui():
    for each_document in db.environment_data.find({'origin': '安徽省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        # ignored
        if announcement_url in ['http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159094',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=158185',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=157046',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=156508',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=155793',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=155504',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=155122',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=154775',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=154774',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=154773',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=154752',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=150926',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=150922',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=150802',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=150802',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=150951',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=150956',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=100626',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=72254',
                                'http://sthjt.ah.gov.cn/pages/Sspx?NTypID=4316',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=43016',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=64154',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=64153',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=64149',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=64150',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63622',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63621',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63620',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63619',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63614',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63613',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63611',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63609',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63604',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63615',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62577',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62570',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=55075',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=55530',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=55532',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=54986',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=54726',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=54733',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=54468',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=51525',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=49914',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=49915',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=49924',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=47947',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=47853',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=47210',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=46736',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44119',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=45342',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=43017',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44058',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44112',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44112',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44113',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44115',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44116',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44117',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44120',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44121',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44122',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44124',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44123',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44003',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44004',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44114',
                                'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159068']:
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        # 如果是多篇pdf文档的
        if re.search(r'主要环境问题清单', announcement_title) and content_soup.find(class_='ContentDiv').find(
                'a'):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=156668
            pdf_link_list = []
            file_name_list = []
            a_list = content_soup.find(class_='ContentDiv').find_all('a')
            for each_a in a_list:
                pdf_link_list.append(urljoin(announcement_url, each_a['href']))
                file_name_list.append(get_content_text(each_a).replace('附件：', '').replace('.pdf', ''))
            for index, pdf_link in enumerate(pdf_link_list):
                file_name = file_name_list[index]
                response = request_site_page(pdf_link)
                if response is None:
                    logger.error('pdf文件下载失败%s' % pdf_link)
                    return
                with open('./test/tmp.pdf', 'wb') as tmp_file:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)
                with open('./test/tmp.pdf', 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')
                content_text = result_text

                if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': pdf_link}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': pdf_link,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'pdf',
                        'oss_file_name': file_name,
                        'oss_file_content': pdf_content,
                        'parsed': False,
                        'if_ocr': True,
                        'ocr_result': result_text
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + file_name + '.pdf', pdf_content)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = \
                        db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': pdf_link})[
                            '_id']
                logger.info('存入parsed_data')
                logger.info(content_text)

                # 处罚机构
                announcement_org = '安徽生态环境厅'
                # 处罚日期
                try:
                    publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
                    real_publish_date = format_date(publish_date)
                except Exception as e:
                    logger.warning(e)
                    publish_date = '20' + each_document['publishDate']
                    real_publish_date = format_date(publish_date)
                # 文号
                try:
                    announcement_code = re.search(r'((淮府办秘|环察|芜政秘)[\s\S]*?号)', content_text).group(1).strip()
                except Exception as e:
                    logger.warning(e)
                    announcement_code = ''
                # 当事人
                litigant = re.search(r'(各县、区人民政府,市政府各部门、各有关单位|省环保[厅斤]|中国石化销售有限公司芜湖石油分公司)', content_text).group(
                    1).strip()

                # 违规事实
                if re.search(r'(你公司油轮码头位我市[\s\S]*?饮用水源地一级\n保护区内)', content_text):
                    facts = re.search(r'(你公司油轮码头位我市[\s\S]*?饮用水源地一级\n保护区内)', content_text).group(1).strip()
                else:
                    facts = ''
                # 认定意见
                punishment_basis = ''

                # 处罚决定
                punishment_decision = re.search(
                    r'((整治目标\n2017年年底前,取缔市级集中式饮用水水源地|现将饮用水水源地环境问题及整\n改方案汇报如下|依据《中华人民共和国水\n污染防治法》)[\s\S]*?(并会同有关部门做好辖区内\n饮用水水源地保护工作|配合相关单位加快船舶拆解企业关停整治\n市矿|确保拆除过程中饮用水水源安全))',
                    content_text).group(1).strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据已经存在')
                # 删除pdf文档
                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.pdf'):
                    os.remove('./test/tmp.pdf')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('安徽生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        # 如果是附件中多篇doc文件的
        if re.search(r'整治方案', announcement_title) and content_soup.find(class_='ContentDiv').find('a'):
            docx_link_list = []
            file_name_list = []
            a_list = content_soup.find(class_='ContentDiv').find_all('a')
            for each_a in a_list:
                docx_link_list.append(urljoin(announcement_url, each_a['href']))
                file_name_list.append(
                    get_content_text(each_a).replace('附件：', '').replace('.doc', '').replace('.docx', ''))
            for index, docx_link in enumerate(docx_link_list):
                file_name = file_name_list[index]
                link_type = docx_link.split('.')[-1]
                response = request_site_page(docx_link)
                if response is None:
                    logger.error('docx文件下载失败%s' % docx_link)
                    return
                with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                if not os.path.exists('./test/tmp.docx'):
                    shell_str = 'soffice --headless --convert-to docx ' + \
                                './test/tmp.' + link_type + ' --outdir ./test'
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()

                with open('./test/tmp.docx', 'rb') as docx_file:
                    docx_content = docx_file.read()

                if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': docx_link}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': docx_link,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'docx',
                        'oss_file_name': file_name,
                        'oss_file_content': docx_content,
                        'parsed': False
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + file_name + '.docx', docx_content)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = \
                        db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': docx_link})[
                            '_id']
                logger.info('存入parsed_data')

                # 开始解析
                document = Document('./test/tmp.docx')
                # 获取全部文本
                content_text = ''
                for paragraph in document.paragraphs:
                    content_text += paragraph.text + '\n'

                # 需人工处理
                if re.search(r'肥西县纳污坑塘整治实施方案|新安镇鲍湾村涉污企业排放口周边', content_text):
                    logger.warning('需人工处理 ...')
                    # 删除doc、docx文档
                    logger.info('删除TMP文件')
                    if os.path.exists('./test/tmp.doc'):
                        os.remove('./test/tmp.doc')
                    if os.path.exists('./test/tmp.docx'):
                        os.remove('./test/tmp.docx')
                    continue

                # 处罚机构
                announcement_org = '安徽生态环境厅'
                # 处罚日期
                try:
                    publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,3}日', content_text)[-1].strip()
                    real_publish_date = format_date(publish_date)
                except Exception as e:
                    logger.warning(e)
                    publish_date = '20' + each_document['publishDate']
                    real_publish_date = format_date(publish_date)
                # 文号
                announcement_code = ''
                # 当事人
                litigant = re.search(r'([\s\S]*?(公司|厂|集团|垃圾堆放场))', content_text).group(1).strip()

                # 违规事实
                if re.search(r'\n存在问题\n([\s\S]*?)\n整改目标\n', content_text):
                    facts = re.search(r'\n存在问题\n([\s\S]*?)\n整改目标\n', content_text).group(1).strip()
                else:
                    facts = ''
                # 认定意见
                punishment_basis = ''

                # 处罚决定
                punishment_decision = re.search(r'((\n整改目标\n|\n总体目标\n|一、总体目标)[\s\S]*?)(\d{4}年\d{1,2}月\d{1,3}日\n)',
                                                content_text).group(1).strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据已经存在')
                # 删除doc、docx文档
                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.doc'):
                    os.remove('./test/tmp.doc')
                if os.path.exists('./test/tmp.docx'):
                    os.remove('./test/tmp.docx')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('安徽生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        # 如果附件中是多篇excel表格的 http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=156018
        if re.search(r'安徽省\d{4}年重点行业环境保护专项执法检查信息|安徽省环保厅关于取缔“十小”企业工作进展情况的报告', announcement_title) and content_soup.find(
                class_='ContentDiv').find('a'):
            # 日期
            publish_date = '20' + each_document['publishDate']
            real_publish_date = format_date(publish_date)
            # 文号
            announcement_code = ''
            # 违规事实
            facts = ''

            # 认定意见
            punishment_basis = ''

            # 申辩意见
            defenseOpinion = ''

            # 申辩意见反馈
            defenseResponse = ''

            xlsx_link_list = []
            file_name_list = []
            a_list = content_soup.find(class_='ContentDiv').find_all('a')
            for each_a in a_list:
                xlsx_link_list.append(urljoin(announcement_url, each_a['href']))
                file_name_list.append(
                    get_content_text(each_a).replace('附件：', '').replace('.xls', '').replace('.xlsx', ''))
            for index, xlsx_link in enumerate(xlsx_link_list):
                file_name = file_name_list[index]

                # 跟处罚无关的表格掠过
                if file_name in ['城镇污水处理厂信息公开数据20161108']:
                    logger.warning('xlsx has nothing to do with punishment ...' + file_name)
                    continue

                link_type = xlsx_link.split('.')[-1]
                response = request_site_page(xlsx_link)
                if response is None:
                    logger.error('xlsx文件下载失败%s' % xlsx_link)
                    return
                with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                if not os.path.exists('./test/tmp.xlsx'):
                    shell_str = 'soffice --headless --convert-to xlsx ' + \
                                './test/tmp.' + link_type + ' --outdir ./test'
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()

                with open('./test/tmp.xlsx', 'rb') as docx_file:
                    xlsx_content = docx_file.read()

                if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': xlsx_link}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': xlsx_link,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'xlsx',
                        'oss_file_name': file_name,
                        'oss_file_content': xlsx_content,
                        'parsed': False
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + file_name + '.xlsx', xlsx_content)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = \
                        db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': xlsx_link})[
                            '_id']
                logger.info('存入parsed_data')

                # 开始解析 #
                excel_data = xlrd.open_workbook('./test/tmp.xlsx')
                sheet = excel_data.sheets()[0]
                if re.search(r'安徽省环保厅关于取缔“十小”企业工作进展情况的报告', announcement_title):
                    sheet = excel_data.sheets()[1]
                result_map_list = []
                for i in range(sheet.nrows):
                    if sheet.ncols == 16:
                        if (sheet.cell(i, 0).value == '序号') or (sheet.cell(i, 15).value == ''):
                            continue
                        # 当事人
                        litigant = '企业名称：' + sheet.cell(i, 1).value + ' 省份：' + sheet.cell(i,
                                                                                          2).value + ' 地市：' + sheet.cell(
                            i, 3).value + ' 区县：' + sheet.cell(i, 4).value

                        # 处罚决定
                        punishment_decision = '企业生产状态：' + sheet.cell(i, 5).value + ' 处理处罚情况：' + sheet.cell(i, 15).value
                    elif sheet.ncols == 13:
                        if re.search(r'附件|填报单位|序号', str(sheet.cell(i, 0).value)):
                            continue
                        # 当事人
                        litigant = '企业名称：' + sheet.cell(i, 4).value + ' 市别：' + sheet.cell(i,
                                                                                          1).value + ' 行政区域：' + sheet.cell(
                            i, 2).value + ' 行业：' + sheet.cell(i, 3).value

                        # 处罚决定
                        punishment_decision = '企业简要情况及取缔情况：' + sheet.cell(i, 6).value + ' 取缔完成时间：' + str(
                            sheet.cell(i, 5).value)

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': '安徽生态环境厅',
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': defenseOpinion,
                        'defenseResponse': defenseResponse,
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
                logger.info(result_map_list)
                if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_many(result_map_list)
                    logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据已经存在')
                logger.info('删除xls文件')
                if os.path.exists('./test/tmp.xls'):
                    os.remove('./test/tmp.xls')
                logger.info('删除xlsx文件')
                if os.path.exists('./test/tmp.xlsx'):
                    os.remove('./test/tmp.xlsx')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('安徽生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        # 如果公告是在网页里
        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                'content_class_name': 'ContentDiv',
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text.encode(content_response.encoding).decode('utf-8'))
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        logger.info('存入parsed_data')
        content_text = get_content_text(content_soup.find(class_='ContentDiv'))

        # 需人工处理
        if re.search(r'行动简报', announcement_title) or announcement_url in [
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=155970',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=155527',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=154736',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=154644',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=154205',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=154639',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=153975',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=152528',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=154617',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=151814',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=150785',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=150600',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=150525',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=124328',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=83247',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62607',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62593',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62464',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62466',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62462',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62460',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62459',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=62592',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=60201',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=60202',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=60203',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=59240',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=58623',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=58622',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=59199',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=55160',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=55527',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=54520',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=54299',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=54197',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=53866',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=50984',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=49913',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=49793',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=49792',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=49790',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=49506',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=48829',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=48425',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=48438',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=48439',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=48514',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=48158',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=48052',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=48049',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=47878',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=47875',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=47628',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=47204',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=47115',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=47011',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=46867',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=45528',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=45274',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=45203',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=45204',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=45205',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=45210',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=45162',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=45154',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=43870',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44126',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44130',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44133',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44131',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=44135',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159125',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159071',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159066',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159065',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159061',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159060',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159058',
            'http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=159057']:
            logger.info('需人工处理')
            continue

        # 处罚机构
        announcement_org = '安徽生态环境厅'
        # 处罚日期
        try:
            publish_date = re.findall(r'安徽省环境保护厅[\s\S]*?(\d{4}年\d{1,2}月\d{1,3}日)', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.warning(e)
            publish_date = '20' + each_document['publishDate']
            real_publish_date = format_date(publish_date)

        # 文号
        try:
            total_text = get_content_text(content_soup.find(class_='i_div'))  # 有的时候文号会在表格头出现
            announcement_code = re.search(r'((淮府办秘|环察|芜政秘|皖环函)[\s\S]*?号)', total_text).group(1).strip()
        except Exception as e:
            logger.warning(e)
            announcement_code = ''

        # 如果是多起督办事项
        if re.search(r'等\d+起[\s\S]*?实施挂牌督办的通知', announcement_title):
            if announcement_url in ['http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=158272']:
                p_list = content_soup.find(class_='ContentDiv').find_all('p')
                p_list.reverse()
                liti_fact_list = []
                for each_p in p_list:
                    if re.search(r'一、|二、|三、|四、|五、|六、|七、|八、|九、|十、', get_content_text(each_p)):
                        liti_fact_list.append(get_content_text(each_p))
                    if re.search(r'一、', get_content_text(each_p)):
                        break
                # 认定意见
                punishment_basis = ''
                # 处罚决定
                punishment_decision = re.search(r'(根据《生态环境部办公厅关于配合做好开展打击固体废物环境违法行为专项行动的通知》[\s\S]*?落实督办要求情况进行督查)',
                                                content_text).group(1).strip()
                result_map_list = []
                for each_one in liti_fact_list:
                    # 当事人
                    litigant = get_content_text(each_one).split("：")[0].strip()

                    # 违规事实
                    facts = get_content_text(each_one).split("：")[-1].strip()
                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            elif content_soup.find(class_='LitContentXXGK').find(
                    'div'):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=155623
                if re.search(r'实施挂牌督办的通知([\s\S]*?)：', content_text):
                    agency = re.search(r'实施挂牌督办的通知([\s\S]*?)：', content_text).group(1).strip()
                else:
                    agency = re.search(r'([\s\S]*?)：', content_text).group(1).strip()
                pre_punishment_decision = re.search(r'((请你局对照督办事项|请你局落实督办事项|请你县对照整改事项)[\s\S]*?并抄送省环境监察局)',
                                                    content_text).group(1).strip().replace('你局', agency).replace('你县',
                                                                                                                 agency)
                div_list = content_soup.find(class_='LitContentXXGK').find_all('div')
                del (div_list[0])  # 删除开头
                for each_div in div_list:
                    if get_content_text(each_div) == '1':  # 删除多余的div
                        div_list.remove(each_div)
                ## 认定意见
                punishment_basis = ''
                result_map_list = []
                for each_div in div_list:
                    div_text = get_content_text(each_div)
                    # 当事人
                    litigant = re.search(r'([\s\S]*?)存在的问题及督办事项', div_text).group(1).strip().replace('\n', '').replace(
                        '附件', '')

                    # 违规事实
                    facts = re.search(r'(一、存在的问题|一、存在的环境问题)([\s\S]*?)二、督办事项', div_text).group(2).strip()

                    # 处罚决定
                    punishment_decision = pre_punishment_decision + '\n' + re.search(
                        r'二、督办事项([\s\S]*?(\d{4}年\d{1,2}月底前完成挂牌督办事项|\d{4}年底前完成挂牌督办事项))', div_text).group(1).strip()

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('安徽生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        # 如果是企业取缔情况明细表http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=156339
        elif re.search(r'企业取缔情况明细表|“十小”企业取缔名单', announcement_title):
            # 认定意见
            punishment_basis = ''
            tr_list = content_soup.find(class_='ContentDiv').find('table').find_all('tr')
            result_map_list = []
            for each_tr in tr_list:
                if re.search(r'企业取缔情况明细表|序号|备注：根据《水污染防治行动计划》', get_content_text(each_tr).replace('\n', '')):
                    continue
                td_list = each_tr.find_all('td')
                # 当事人
                litigant = '企业名称：' + get_content_text(td_list[4]) + ' 市别：' + get_content_text(
                    td_list[1]) + ' 行政区域：' + get_content_text(td_list[2]) + ' 行业：' + get_content_text(td_list[3])
                if len(td_list) == 7:
                    # 处罚决定
                    punishment_decision = '取缔情况：' + get_content_text(td_list[6]) + ' 取缔完成时间：' + get_content_text(
                        td_list[5])
                elif len(td_list) == 6:
                    # 处罚决定
                    punishment_decision = '取缔情况：' + get_content_text(td_list[5])
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': '',
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('安徽生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(r'\d{4}年第.季度(环保不达标生产|污染物超标排放)企业名单', announcement_title):
            if content_soup.find(class_='LitContentXXGK').find(
                    'p'):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=111360
                text_list = content_soup.find(class_='LitContentXXGK').find_all('p')
            elif content_soup.find(class_='LitContentXXGK').find(
                    'br'):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=123506
                text_list = content_text.split("\n")
            result_map_list = []
            # 认定意见
            punishment_basis = ''
            # 处罚决定
            punishment_decision = ''
            for each_one in text_list:
                each_text = get_content_text(each_one)
                if not re.search(r'：', each_text):
                    continue
                litigant = each_text.split("：")[0].strip()

                facts = each_text.split("：")[-1].strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('安徽生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(
                r'环境违法问题调查处理情况的函|环境污染问题查处情况的函|环境问题调查情况的函|环境信访问题查处情况的函|环境信访问题处理情况的函|受污染问题查处情况的函|污染环境问题查处情况的函|夜间噪音扰民问题查处情况的函|申请环保整改摘牌验收的复函|信访件查处情况的函|信访问题查处情况的函|水污染问题调查情况的函',
                announcement_title):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=162529/http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=156323
            # 当事人
            litigant = re.search(
                r'(关于)(\S*?)(存在|环境污染问题查处情况的函|环境问题调查情况的函|环境信访问题查处情况的函|环境信访问题处理情况的函|受污染问题查处情况的函|污染环境问题查处情况的函|夜间噪音扰民问题查处情况的函|申请环保整改摘牌验收的复函|信访件查处情况的函|信访问题查处情况的函|水污染问题调查情况的函)',
                announcement_title).group(2).strip()

            # 违规事实
            facts = re.search(
                r'((一、基本情况|一、企业基本情况|一、查处情况|一、投诉人反映的问题|一、调查情况|一、企业概况|一、未按督办要求落实整改工作)[\s\S]*?)(二、调查处理情况|二、废水处理情况|二、调查情况|包河区环保局已联合公安、城管等部门对该公司进行了断电|三、处理情况|二、查处情况|三、调查结论及处理情况|二、处理情况|二、调查及处理情况|四、关于信访问题的调查情况|三、关于信访人投诉的问题|根据检查情况，包河区环保局对该歌舞厅下达整改通知书|二、下一步工作要求|三、处理措施)',
                content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(
                r'((二、调查处理情况|二、废水处理情况|二、调查情况|包河区环保局已联合公安、城管等部门对该公司进行了断电|三、处理情况|二、查处情况|三、调查结论及处理情况|二、处理情况|二、调查及处理情况|四、关于信访问题的调查情况|三、关于信访人投诉的问题|根据检查情况，包河区环保局对该歌舞厅下达整改通知书|二、下一步工作要求|三、处理措施)[\s\S]*?)(四、反馈情况|三、反馈情况|七、反馈情况|二、反馈情况|五、反馈情况|安徽省环境保护厅)',
                content_text).group(1).strip()

        elif re.search(r'关于落实[\s\S]*?环境违法案件挂牌督办要求的通知',
                       announcement_title):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=155629
            agency = re.search(r'挂牌督办要求的通知([\s\S]*?)：', content_text).group(1).strip()
            # 当事人
            litigant = re.search(r'(关于落实)([\s\S]*?)(环境违法案件挂牌督办要求的通知)', announcement_title).group(2).strip()
            # 违规事实
            facts = re.search(agency + r'([\s\S]*?)(为落实督办要求)', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'((为落实督办要求)[\s\S]*?我厅将适时进行督查督办)', content_text).group(1).strip().replace(
                '你局', agency)
        # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=157898
        elif re.search(
                r'(关于解除|关于延长|关于同意解除对|关于申请延期解除|关于同意)[\s\S]*?(环境违法案件挂牌督办|限批|环境违法案件挂牌督办时限|环境问题挂牌督办|挂牌督办|申请延期解除挂牌督办|挂牌督办环境问题摘牌|环境违法问题摘牌|环境违法问题挂牌督办摘牌|摘牌)的(函|通知|复函)',
                announcement_title) \
                or re.search(r'(关于)[\s\S]*?(申请延期解除挂牌督办|申请解除挂牌督办|解除挂牌督办|环境违法问题挂牌督办摘牌)的(函|通知|复函)', announcement_title):
            agency = re.search(r'([\s\S]*?)(：|:)', content_text).group(1).strip()

            # 当事人
            if re.search(r'(关于解除|关于延长|关于同意解除对)([\s\S]*?)(环境违法案件挂牌督办|限批|环境违法案件挂牌督办时限|环境问题挂牌督办)的(函|通知|复函)',
                         announcement_title):
                litigant = re.search(r'(关于解除|关于延长|关于同意解除对)([\s\S]*?)(环境违法案件挂牌督办|限批|环境违法案件挂牌督办时限|环境问题挂牌督办)的(函|通知|复函)',
                                     announcement_title).group(2).strip()
            elif re.search(r'(关于同意解除对|关于申请延期解除|关于解除)([\s\S]*?)挂牌督办的(函|通知|复函)', announcement_title):
                litigant = re.search(r'(关于同意解除对|关于申请延期解除|关于解除)([\s\S]*?)挂牌督办的(函|通知|复函)', announcement_title).group(
                    2).strip()
            elif re.search(r'(关于同意)([\s\S]*?)(挂牌督办|挂牌督办环境问题摘牌|环境违法问题摘牌|环境违法问题挂牌督办摘牌)的(函|通知|复函)', announcement_title):
                litigant = re.search(r'(关于同意)([\s\S]*?)(挂牌督办|挂牌督办环境问题摘牌|环境违法问题摘牌|环境违法问题挂牌督办摘牌)的(函|通知|复函)',
                                     announcement_title).group(2).strip()
            elif re.search(r'(关于同意)([\s\S]*?)(摘牌)的(函|通知|复函)', announcement_title):
                litigant = re.search(r'(关于同意)([\s\S]*?)(摘牌)的(函|通知|复函)', announcement_title).group(2).strip()
            elif re.search(r'关于([\s\S]*?)(申请延期解除挂牌督办|申请解除挂牌督办|解除挂牌督办|环境违法问题挂牌督办摘牌)的(函|通知|复函)', announcement_title):
                litigant = re.search(r'关于([\s\S]*?)(申请延期解除挂牌督办|申请解除挂牌督办|解除挂牌督办|环境违法问题挂牌督办摘牌)的(函|通知|复函)',
                                     announcement_title).group(1).strip()
            else:
                litigant = re.search(r'关于([\s\S]*?)挂牌督办的(函|通知|复函)', announcement_title).group(1).strip()

                # 违规事实
            facts = re.search(agency.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                      '\]') + r'([\s\S]*?)(鉴此，经研究，[我省]厅同意解除|经研究，(我厅|省厅|省环保厅)同意|经研究，同意|经研究，现函复如下|经研究，我厅决定|同意苏深表面处理有限公司|经对(核查组|验收组|现场核查组|验收专家组)意见审查，(我厅同意|同意)|综上，我厅同意)',
                              content_text).group(1).strip().lstrip(':').lstrip('：')

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            if re.search(
                    r'((鉴此，经研究，[我省]厅同意解除|经研究，(我厅|省厅|省环保厅)同意|经研究，同意|经研究，现函复如下|经研究，我厅决定|综上，我厅同意)[\s\S]*?)(安徽省环境保护厅|安徽省环\n境保护厅|安徽省生态环境厅)',
                    content_text):
                punishment_decision = re.search(
                    r'((鉴此，经研究，[我省]厅同意解除|经研究，(我厅|省厅|省环保厅)同意|经研究，同意|经研究，现函复如下|经研究，我厅决定|综上，我厅同意)[\s\S]*?)(安徽省环境保护厅|安徽省环\n境保护厅|安徽省生态环境厅)',
                    content_text).group(1).strip().replace('你局', agency)
            else:
                punishment_decision = re.search(
                    r'((鉴此，经研究，[我省]厅同意解除|经研究，(我厅|省厅|省环保厅)同意|经研究，同意|经研究，现函复如下|经研究，我厅决定|同意苏深表面处理有限公司|经对(核查组|验收组|现场核查组|验收专家组)意见审查，(我厅同意|同意))[\s\S]*)',
                    content_text).group(1).strip().replace('你局', agency)
                if re.search(r'二○.{2}年.{1,2}月.{1,2}日', punishment_decision):
                    tmp_date = re.search(r'二○.{2}年.{1,2}月.{1,2}日', punishment_decision).group(0)
                    punishment_decision = punishment_decision.replace(tmp_date, '')
        elif re.search(r'实施挂牌督办的通知|关于暂停审批\S*?新、改、扩建项目环境影响评价文件的函|环境问题挂牌督办的通知|环境违法案件挂牌督办的通知|环境问题的监察通知',
                       announcement_title):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=158275  http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=155834#http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=155472
            # 当事人
            litigant = re.search(
                r'(关于对|关于|关于暂停审批)(\S*?)(环境违法案件|新、改、扩建项目环境影响评价文件的函|环境问题挂牌督办的通知|有关环境问题|环境违法案件挂牌督办的通知|环境问题的监察通知)',
                announcement_title).group(2).strip()

            # 违规事实
            facts = re.search(
                r'(一、存在的问题|一、怀远县经济开发区存在的环境问题|一、调查结论|一、存在的环境问题|一、基本情况|一、违法事实|一、环境违法问题|现将检查中发现的问题及处理意见通知如下：)([\s\S]*?)(二、督办事项|二、整改事项|二、督办要求|二、处理意见和建议|二、检查处理情况|二、处理意见|四、处理意见|三、处理情况|二、查处情况|针对存在的环境问题，提出如下处理意见，请你局：)',
                content_text).group(2).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            agency = re.search(r'([\s\S]*?)：', content_text).group(1).strip().replace('\n', '').replace(
                announcement_title, '')
            try:
                pre_punishment_decision = re.search(r'((请你局对照督办事项|请你局落实督办事项|请你县对照整改事项)[\s\S]*?并抄送省环境监察局)',
                                                    content_text).group(1).strip().replace('你局', agency).replace('你县',
                                                                                                                 agency)
            except Exception as e:
                logger.info(e)
                pre_punishment_decision = ''
            if re.search(
                    r'(二、督办事项|二、整改事项|二、督办要求|二、处理意见和建议|二、检查处理情况|二、处理意见|四、处理意见|三、处理情况|二、查处情况|针对存在的环境问题，提出如下处理意见，请你局：)([\s\S]*?((\d{4}年\d{1,2}月底前完成(挂牌督办事项|区域限批整改事项))|并抄送省环境监察局|同时抄送池州市环保局|减少对周边群众的影响|你局每月底前要将两企业整改落实情况书面报送省厅|省厅将对该企业的环境违法行为依法予以处理|解除挂牌督办申请并附相关材料上报省厅|将解除挂牌督办书面申请并附相关材料一并报省环保厅、省监察厅|确保按时完成挂牌督办事项|直至该公司环境问题彻底解决))',
                    content_text):
                punishment_decision = pre_punishment_decision + '\n' + re.search(
                    r'(二、督办事项|二、整改事项|二、督办要求|二、处理意见和建议|二、检查处理情况|二、处理意见|四、处理意见|三、处理情况|二、查处情况|针对存在的环境问题，提出如下处理意见，请你局：)([\s\S]*?((\d{4}年\d{1,2}月底前完成(挂牌督办事项|区域限批整改事项))|并抄送省环境监察局|同时抄送池州市环保局|减少对周边群众的影响|你局每月底前要将两企业整改落实情况书面报送省厅|省厅将对该企业的环境违法行为依法予以处理|解除挂牌督办申请并附相关材料上报省厅|将解除挂牌督办书面申请并附相关材料一并报省环保厅、省监察厅|确保按时完成挂牌督办事项|直至该公司环境问题彻底解决))',
                    content_text).group(2).strip().replace('你局', agency)
            else:
                punishment_decision = pre_punishment_decision + '\n' + re.search(
                    r'(二、督办事项|二、整改事项|二、督办要求|二、处理意见和建议|二、检查处理情况|二、处理意见|四、处理意见|三、处理情况|二、查处情况)([\s\S]*?)(四、反馈情况|五、反馈情况|三、反馈情况)',
                    content_text).group(2).strip()
        elif re.search(r'专项执法检查情况的通报',
                       announcement_title):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=157848

            # 当事人
            litigant = re.search(r'([\s\S]*?)：', content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'(一、总体情况[\s\S]*?)(三、处理情况)', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'((三、处理情况)[\s\S]*?)安徽省环境保护厅', content_text).group(1).strip()

        elif re.search(r'环境违法问题曝光|\d{4}年\d{1,2}月份现场执法检查情况',
                       announcement_title):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=156693
            # 当事人
            if re.search(r'([\s\S]*?)环境违法问题曝光', announcement_title):
                litigant = re.search(r'([\s\S]*?)环境违法问题曝光', announcement_title).group(1).strip()
            else:
                litigant = re.search(r'现场执法检查情况.([\s\S]*?(公司|厂|集团|单位))', announcement_title).group(1).strip()
            # 违规事实
            facts = content_text

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = ''
        elif re.search(r'关于对[\s\S]*?(环境问题整改情况的监察意见|环境问题的监察意见)',
                       announcement_title):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=64142
            # 当事人
            litigant = re.search(r'([\s\S]*?：)', content_text).group(1).strip()
            # 违规事实
            facts = re.search(litigant + r'([\s\S]*?)现对你公司环境问题提出监察意见如下', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            if re.search(r'(现对你公司环境问题提出监察意见如下[\s\S]*?)二○..年.{1,2}月.{1,2}日', content_text):
                punishment_decision = re.search(r'(现对你公司环境问题提出监察意见如下[\s\S]*?)二○..年.{1,2}月.{1,2}日', content_text).group(
                    1).strip()
            else:
                punishment_decision = re.search(r'(现对你公司环境问题提出监察意见如下[\s\S]*)', content_text).group(1).strip()

        elif re.search(r'省环境监察局追踪督导企业有效整改|省环境监察局督导违法企业彻底整改环境问题',
                       announcement_title):  # http://sthjt.ah.gov.cn/pages/ShowNews.aspx?NType=2&NewsID=63605
            # 当事人
            litigant = re.search(r'(省环保检查组在执法检查时，发现|省环境监察局在检查)([\s\S]*?)(存在以下问题|执行环保法律法规情况时)', content_text).group(
                2).strip()
            # 违规事实
            facts = re.search(r'([\s\S]*)(为确保整改成果、有力推动污染减排|\d{1,2}月\d{1,2}日，为巩固整改成果，省环境监察局要求该公司)', content_text).group(
                1).strip()

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'(省环境监察局要求该公司：[\s\S]*)', content_text).group(1).strip().replace('该公司',
                                                                                                             litigant)

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': '',
            'defenseResponse': '',
            'litigant': litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('安徽生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('安徽生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 福建生态环境厅
def fujian():
    for each_document in db.environment_data.find({'origin': '福建省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if content_soup.find(class_='TRS_Editor'):
            content_text = get_content_text(content_soup.find(class_='TRS_Editor'))
            content_class_name = 'TRS_Editor'
        else:
            content_text = get_content_text(content_soup.find(class_='xl_con1'))
            content_class_name = 'xl_con1'
        head_text = get_content_text(content_soup.find('div', id='divheadContainer'))

        # 如果公告是在网页里
        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'htm',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                'content_class_name': content_class_name,
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.htm',
                         content_response.text.encode(content_response.encoding).decode('utf-8'))
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        logger.info('存入parsed_data')

        # 需人工处理
        if announcement_url in ['http://hbt.fujian.gov.cn/zwgk/zfxxgkzl/zfxxgkml/mlflfg/201506/t20150608_108719.htm']:
            logger.info('需人工处理')
            continue

        # 开始解析
        # 处罚机构
        announcement_org = '福建生态环境厅'
        # 处罚日期
        publish_date = each_document['publishDate']
        real_publish_date = format_date(publish_date)

        # 文号
        try:
            announcement_code = re.search(r'(闽环[\s\S]*?号)', head_text).group(1).strip()
        except Exception as e:
            logger.warning(e)
            announcement_code = ''

        # 不予受理行政复议申请决定书
        if re.search(r'不予受理行政复议申请决定书', announcement_title):
            # 当事人
            litigant = re.search(r'(被申请人：[\s\S]*?)\n申请人', content_text).group(1).strip()

            # 违规事实
            facts_1 = re.search(r'(申请人：[\s\S]*?)被申请人：', content_text).group(1).strip()
            facts_2 = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                          '\]') + r'([\s\S]*?)经审查，本机关认为：',
                                content_text).group(1).strip()
            facts = facts_1 + facts_2

            # 申辩意见
            defense_opinion = ''

            # 申辩意见反馈
            defense_response = ''

            # 认定意见
            punishment_basis = re.search(r'(经审查，本机关认为：[\s\S]*?)(综上，根据《行政复议法》|因此，根据《中华人民共和国行政复议法》)', content_text).group(
                1).strip()

            # 处罚决定
            punishment_decision = re.search(r'((综上，根据《行政复议法》|因此，根据《中华人民共和国行政复议法》)[\s\S]*?)(福建省生态环境厅|福建省环境保护厅)',
                                            content_text).group(1).strip()

        # 行政复议决定书
        elif re.search(r'行政复议决定书', announcement_title):
            if re.search(r'一、申请人的申请事项和主要理由',
                         content_text):  # http://hbt.fujian.gov.cn/zwgk/zfxxgkzl/zfxxgkml/mlflfg/201505/t20150504_424303.htm
                # 当事人
                litigant = re.search(r'(被申请人：[\s\S]*?)\n申请人', content_text).group(1).strip()

                # 违规事实
                facts_1 = re.search(r'(申请人：[\s\S]*?)被申请人：', content_text).group(1).strip()
                facts_2 = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                              '\]') + r'([\s\S]*?)二、被申请人的主要答辩意见和理由',
                                    content_text).group(1).strip()
                facts = facts_1 + '\n' + facts_2

                # 申辩意见
                defense_opinion = re.search(r'(二、被申请人的主要答辩意见和理由)([\s\S]*?)(三、我厅审理后认定的事实和证据|三、我厅审理后认定的事实)',
                                            content_text).group(2).strip()

                # 申辩意见反馈
                defense_response = ''

                # 认定意见
                punishment_basis = re.search(r'(三、我厅审理后认定的事实和证据|三、我厅审理后认定的事实)([\s\S]*?)(四、复议决定)', content_text).group(
                    2).strip()

                # 处罚决定
                punishment_decision = re.search(r'(四、复议决定)([\s\S]*?)(福建省生态环境厅|福建省环境保护厅)', content_text).group(2).strip()

            else:
                # 当事人
                litigant = re.search(r'(被申请人：[\s\S]*?)\n申请人', content_text).group(1).strip()

                # 违规事实
                facts_1 = re.search(r'(申请人：[\s\S]*?)被申请人：', content_text).group(1).strip()
                facts_2 = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                              '\]') + r'([\s\S]*?)被申请人称(：|，)',
                                    content_text).group(1).strip()
                facts = facts_1 + '\n' + facts_2

                # 申辩意见
                if re.search(r'(被申请人称(：|，)[\s\S]*?)(本机关经审理查明：|本机关经审理查明事实如下：)', content_text):
                    defense_opinion = re.search(r'(被申请人称(：|，)[\s\S]*?)(本机关经审理查明：|本机关经审理查明事实如下：)', content_text).group(
                        1).strip()
                else:
                    defense_opinion = re.search(r'(被申请人称(：|，)[\s\S]*?)(经审理查明：)', content_text).group(1).strip()

                # 申辩意见反馈
                defense_response = re.search(r'((本机关经审理查明：|经审理查明：|本机关经审理查明事实如下：)[\s\S]*?)(本机关认为：|本机关认为如下：)',
                                             content_text).group(1).strip()

                # 认定意见
                if re.search(r'((本机关认为：|本机关认为如下：)[\s\S]*?)(综上，|综上所述，)', content_text):
                    punishment_basis = re.search(r'((本机关认为：|本机关认为如下：)[\s\S]*?)(综上，|综上所述，)', content_text).group(
                        1).strip()
                else:
                    punishment_basis = re.search(r'((本机关认为：|本机关认为如下：)[\s\S]*?)(根据《中华人民共和国行政复议法》)', content_text).group(
                        1).strip()

                # 处罚决定
                if re.search(r'((综上，|综上所述，)[\s\S]*?)(福建省生态环境厅|福建省环境保护厅)', content_text):
                    punishment_decision = re.search(r'((综上，|综上所述，)[\s\S]*?)(福建省生态环境厅|福建省环境保护厅)', content_text).group(
                        1).strip()
                else:
                    punishment_decision = re.search(r'((根据《中华人民共和国行政复议法》)[\s\S]*?)(福建省生态环境厅|福建省环境保护厅)',
                                                    content_text).group(1).strip()

        # 行政复议终止决定书
        elif re.search(r'行政复议终止决定书', announcement_title):
            # 当事人
            litigant = re.search(r'(被申请人：[\s\S]*?)\n申请人', content_text).group(1).strip()

            # 违规事实
            facts_1 = re.search(r'(申请人：[\s\S]*?)被申请人：', content_text).group(1).strip()
            facts_2 = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                          '\]') + r'([\s\S]*?)根据《中华人民共和国行政复议法》',
                                content_text).group(1).strip()
            facts = facts_1 + '\n' + facts_2

            # 申辩意见
            defense_opinion = ''

            # 申辩意见反馈
            defense_response = ''

            # 认定意见
            punishment_basis = ''

            # 处罚决定
            punishment_decision = re.search(r'((根据《中华人民共和国行政复议法》)[\s\S]*?)(福建省生态环境厅|福建省环境保护厅)', content_text).group(
                1).strip()

        # 行政处罚决定书
        elif re.search(r'行政处罚决定书', announcement_title):
            if re.search(r'鉴于你是[\s\S]*?(董事长、法定代表人|矿长)',
                         content_text):  # http://hbt.fujian.gov.cn/zwgk/zfxxgkzl/zfxxgkml/mlflfg/201012/t20101228_124026.htm  http://hbt.fujian.gov.cn/zwgk/zfxxgkzl/zfxxgkml/mlflfg/201012/t20101228_124029.htm
                # 当事人
                if re.search(
                        r'([\s\S]*)(你(公司|厂|集团)环境违法一案，我厅已经审查终结)', content_text):
                    litigant = re.search(r'([\s\S]*)(你(公司|厂|集团)环境违法一案，我厅已经审查终结)', content_text).group(
                        1).strip().replace(announcement_code, '')
                else:
                    litigant = re.search(r'([\s\S]*)一、环境违法事实和证据', content_text).group(1).strip().rstrip('：')
                # 违规事实
                facts = re.search(r'一、环境违法事实和证据([\s\S]*?)(你(公司|厂|集团)的上述行为，*.{0,2}违反|以上行为违反了)', content_text).group(
                    1).strip()

                defense_text = '我厅' + re.search(r'(依法于\n\d{4}年\d{1,2}月\d{1,2}日\n向你送达了[\s\S]*?)以上事实，*有如下证据为证',
                                                content_text).group(1).strip()
                if re.search(r'你收到我厅告知书后，未提出陈述申辩与听证要求', defense_text):
                    # 申辩意见
                    defense_opinion = defense_text

                    # 申辩意见反馈
                    defense_response = ''
                else:
                    # 申辩意见
                    defense_opinion = re.search(r'([\s\S]*(你提出[\s\S]*?的陈述与申辩|等申辩意见|等陈述与申辩理由。|等陈述与申辩。))',
                                                defense_text).group(1).strip()

                    # 申辩意见反馈
                    defense_response = defense_text.lstrip(defense_opinion)

                # 认定意见
                punishment_basis = re.search(r'(你(公司|厂|集团)的上述行为，*.{0,2}违反[\s\S]*?”(的|之)规定|以上行为违反了[\s\S]*?”(的|之)规定。)',
                                             content_text).group(1).strip()

                # 处罚决定
                punishment_decision = re.search(r'二、行政处罚的依据、种类([\s\S]*?)四、申请复议或者提起诉讼的途径和期限', content_text).group(
                    1).strip()
            else:
                # 当事人
                if re.search(
                        r'([\s\S]*)(你(公司|厂|集团)环境违法一案，我厅已经审查终结|福建省清流县东莹化工有限公司（以下简称公司）环境违法一案，我厅已经审查终结|邵武华孚新材料发展有限公司（以下简称公司）环境违法一案，我厅已经审查终结)',
                        content_text):
                    litigant = re.search(
                        r'([\s\S]*)(你(公司|厂|集团)环境违法一案，我厅已经审查终结|福建省清流县东莹化工有限公司（以下简称公司）环境违法一案，我厅已经审查终结|邵武华孚新材料发展有限公司（以下简称公司）环境违法一案，我厅已经审查终结)',
                        content_text).group(1).strip().replace(announcement_code, '')
                else:
                    litigant = re.search(r'([\s\S]*)一、环境违法事实和证据', content_text).group(1).strip().rstrip('：')
                # 违规事实
                facts = re.search(r'一、环境违法事实和证据([\s\S]*?)(你(公司|厂|集团)的上述行为，*.{0,2}违反|你(公司|厂|集团)以上行为违反了)',
                                  content_text).group(1).strip()

                defense_text = re.search(r'”(的|之)规定。([\s\S]*?)以上事实，*有如下证据为证', content_text).group(2).strip()
                if re.search(r'你(公司|厂|集团)收到我厅告知书后，未提出(陈述申辩和听证要求|陈述与申辩)', defense_text):
                    # 申辩意见
                    defense_opinion = defense_text

                    # 申辩意见反馈
                    defense_response = ''
                else:
                    # 申辩意见
                    defense_opinion = re.search(r'([\s\S]*(你(公司|厂|集团)提出[\s\S]*?等申辩意见|等陈述与申辩理由。|等陈述与申辩。))',
                                                defense_text).group(1).strip()

                    # 申辩意见反馈
                    defense_response = defense_text.lstrip(defense_opinion)

                # 认定意见
                punishment_basis = re.search(
                    r'(你(公司|厂|集团)的上述行为，*.{0,2}违反[\s\S]*?”(的|之)规定。|你(公司|厂|集团)以上行为违反了[\s\S]*?”(的|之)规定。)',
                    content_text).group(1).strip()

                # 处罚决定
                punishment_decision = re.search(r'二、行政处罚的依据、种类([\s\S]*?)四、申请复议或者提起诉讼的途径和期限', content_text).group(
                    1).strip()
        # 行政处罚事先（听证）告知书 http://hbt.fujian.gov.cn/zwgk/zfxxgkzl/zfxxgkml/mlflfg/201401/t20140110_406145.htm
        elif re.search(r'行政处罚事先（听证）告知书', announcement_title):
            # 当事人
            litigant = re.search(r'([\s\S]*)经调查，你(公司|厂|集团)存在以下环境违法行为：', content_text).group(1).strip().rstrip("：")

            # 违规事实
            facts = re.search(r'经调查，你(公司|厂|集团)存在以下环境违法行为：([\s\S]*?)以上事实，*有[\s\S]*?(证据为证|等证据为凭)', content_text).group(
                2).strip()

            # 申辩意见
            defense_opinion = ''

            # 申辩意见反馈
            defense_response = ''

            # 认定意见
            punishment_basis = re.search(r'(你(公司|厂|集团)的上述行为，违反了[\s\S]*?”(的|之)规定)', content_text).group(1).strip()

            # 处罚决定
            punishment_decision = re.search(r'”(的|之)规定。([\s\S]*?)我厅地址：', content_text).group(2).strip()

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('福建生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('福建生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('福建生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 江西生态环境厅
def jiangxi():
    for each_document in db.environment_data.find({'origin': '江西省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        # ignore
        if re.search(r'无下达行政处罚决定|无处罚案件|本月无', get_content_text(content_soup.find(id='inthe'))) or \
                re.search(r'无下达行政处罚决定|无处罚案件|本月无', get_content_text(content_soup.find(id='text-box'))):
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        if content_soup.find(class_='wenbendax'):
            class_name = 'wenbendax'
        else:
            class_name = 'text-box'

        file_link_list = [urljoin(announcement_url, each_a['href']) for each_a in
                          content_soup.find(class_=class_name).find_all('a')]

        for index, file_link in enumerate(file_link_list):
            # 如果附件中是rar文件
            if file_link.endswith('.rar'):
                response = request_site_page(file_link)
                if response is None:
                    logger.error('网页请求错误')
                    return
                with open('./test/tmp.rar', 'wb') as tmp_file:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                if not os.path.exists('./test/tmp'):
                    os.mkdir('./test/tmp')
                try:
                    patoolib.extract_archive('./test/tmp.rar', outdir='./test/tmp')
                except Exception as e:
                    logger.warning(e)
                    logger.info('.RAR文件解压失败')
                    logger.info('删除TMP文件')
                    if os.path.exists('./test/tmp'):
                        os.removedirs('./test/tmp')
                    if os.path.exists('./test/tmp.rar'):
                        os.remove('./test/tmp.rar')
                    continue

                doc_file_list = []
                for root, dirs, files in os.walk("./test/tmp", topdown=False):
                    for name in files:
                        doc_file_list.append(os.path.join(root, name))

                for each_doc_file in doc_file_list:
                    doc_title = re.split(r'[./]', each_doc_file)[-2]
                    if not os.path.exists('./test/tmp/' + doc_title + '.docx'):
                        shell_str = 'soffice --headless --convert-to docx ' + \
                                    './test/tmp/' + doc_title + '.doc --outdir ./test/tmp'
                        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                                   shell=True, stdout=subprocess.PIPE)
                        process.communicate()

                    with open('./test/tmp/' + doc_title + '.docx', 'rb') as docx_file:
                        docx_content = docx_file.read()

                    if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_name': doc_title}).count() == 0:
                        oss_file_map = {
                            'origin_url': announcement_url,
                            'oss_file_origin_url': file_link,
                            'origin_url_id': each_document['_id'],
                            'oss_file_type': 'docx',
                            'oss_file_name': doc_title,
                            'oss_file_content': docx_content,
                            'parsed': False
                        }
                        response = db.parsed_data.insert_one(oss_file_map)
                        file_id = response.inserted_id
                        oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                        db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    else:
                        db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                        file_id = \
                            db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': file_link})[
                                '_id']
                    logger.info('存入parsed_data')

                    # 开始解析
                    document = Document('./test/tmp.docx')
                    # 获取全部文本
                    content_text = ''
                    for paragraph in document.paragraphs:
                        content_text += paragraph.text + '\n'
                    logger.info(content_text)

                    # 处罚机构
                    announcement_org = '江西生态环境厅'
                    # 处罚日期
                    try:
                        publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,2}日', content_text)[-1]
                        real_publish_date = format_date(publish_date)
                    except Exception as e:
                        logger.info(e)
                        publish_date = each_document['publishDate']
                        real_publish_date = format_date(publish_date)

                    # 文号
                    try:
                        announcement_code = re.search(r'(赣环[\s\S]*?号)', content_text).group(1).strip()
                    except Exception as e:
                        logger.warning(e)
                        announcement_code = ''

                    # 当事人
                    litigant = re.search(
                        announcement_code.replace(r'(', r'\(').replace(r')', r'\)').replace(r'[', r'\[').replace(
                            r']',
                            r'\]') + r'([\s\S]*?)一、环境违法事实、证据和陈述申辩（听证）及采纳情况：',
                        content_text).group(1).strip()

                    # 违规事实
                    if re.search(r'以上事实，*有(我厅)*[\s\S]*?等*(证据)*为凭。', content_text):
                        facts = re.search(r'一、环境违法事实、证据和陈述申辩（听证）及采纳情况：([\s\S]*?)以上事实，*有(我厅)*[\s\S]*?等*(证据)*为凭。',
                                          content_text).group(1).strip()
                    else:
                        facts = re.search(
                            r'一、环境违法事实、证据和陈述申辩（听证）及采纳情况：([\s\S]*?)(你(公司|集团|厂|单位)(的上述违法行为违反了|上述违法行为违反了|的上述行为违反了|上述行为违反了|超标排放|未按环评批复要求)[\s\S]*?的*规定[。\n])',
                            content_text).group(1).strip()

                    # 认定意见(有的没有认定意见，最后加个try except)
                    try:
                        punishment_basis = re.search(
                            r'(你(公司|集团|厂|单位)(的上述违法行为违反了|上述违法行为违反了|的上述行为违反了|上述行为违反了|超标排放|未按环评批复要求)[\s\S]*?的*规定[。\n])',
                            content_text).group(1).strip()
                    except Exception as e:
                        logger.info(e)
                        punishment_basis = ''

                    if punishment_basis == '':
                        defense_response_text = re.search(r'(以上事实有[\s\S]*?为凭。\n)([\s\S]*?)(二、行政处罚的依据、种类及其履行方式、期限)',
                                                          content_text).group(2).strip()
                    else:
                        defense_response_text = re.search(
                            punishment_basis.replace(r'(', r'\(').replace(r')', r'\)').replace(r'[', r'\[').replace(
                                r']', r'\]') + r'([\s\S]*?)(二、行政处罚的依据、种类及其履行方式、期限)',
                            content_text).group(1).strip()

                    if re.search(r'未提出书面听证申请|未进行陈述申辩|未提[出交]听证申请|(无|没有)异议|不进行陈述、申辩|不申请听证|接受省厅的处罚决定',
                                 defense_response_text) and not re.search(
                        r'收到[\s\S]*?递交的陈述申辩材料|递交了*书面陈述申辩材料|提交书面陈述申辩材料|递交我厅书面陈述申辩材料',
                        defense_response_text):
                        defense_opinion = defense_response_text
                        defense_response = ''
                    else:
                        # 申辩意见
                        defense_opinion = re.search(
                            r'([\s\S]*?)(\n(我厅)*经审理查明，|\n我厅经审(查|理)，|\n经查询，|\n我厅案件审理委员会经审理查明，|\n\s*经审理查明，)',
                            defense_response_text).group(1).strip()

                        # 申辩意见反馈
                        defense_response = defense_response_text.lstrip(defense_opinion)

                    # 处罚决定
                    punishment_decision = re.search(r'(二、行政处罚的依据、种类及其履行方式、期限)([\s\S]*?)(限于接到本处罚决定之日起)',
                                                    content_text).group(2).strip()
                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': defense_opinion,
                        'defenseResponse': defense_response,
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    logger.info(result_map)
                    if db.announcement.find(
                            {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                        db.announcement.insert_one(result_map)
                        logger.info('江西生态环境厅 数据解析 ' + ' -- 数据导入完成')
                    else:
                        logger.info('江西生态环境厅 数据解析 ' + ' -- 数据已经存在')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('江西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp'):
                    os.removedirs('./test/tmp')
                if os.path.exists('./test/tmp.rar'):
                    os.remove('./test/tmp.rar')

            # 如果附件中是word文档
            elif file_link.endswith('.doc') or file_link.endswith('.docx'):
                link_type = file_link.split('.')[-1]
                response = request_site_page(file_link)
                if response is None:
                    logger.error('网页请求错误')
                    return
                with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                if not os._exists('./test/tmp.docx'):
                    shell_str = 'soffice --headless --convert-to docx ' + \
                                './test/tmp.' + link_type + ' --outdir ./test'
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()

                with open('./test/tmp.docx', 'rb') as docx_file:
                    docx_content = docx_file.read()

                if db.parsed_data.find(
                        {'origin_url': announcement_url, 'oss_file_origin_url': file_link}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': file_link,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'docx',
                        'oss_file_name': announcement_title,
                        'oss_file_content': docx_content,
                        'parsed': False
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = \
                        db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': file_link})[
                            '_id']
                logger.info('存入parsed_data')

                # 人工处理
                if file_link in ['http://www.jxepb.gov.cn/resource/uploadfile/file/20150401/20150401102528389.doc',
                                 'http://www.jxepb.gov.cn/resource/uploadfile/file/20141022/20141022141951132.doc',
                                 'http://www.jxepb.gov.cn/resource/uploadfile/file/20160704/20160704144513835.doc']:
                    logger.warning('需人工处理 ...')
                    logger.info('删除TMP文件')
                    if os.path.exists('./test/tmp.doc'):
                        os.remove('./test/tmp.doc')
                    if os.path.exists('./test/tmp.docx'):
                        os.remove('./test/tmp.docx')
                    continue

                # 开始解析
                document = Document('./test/tmp.docx')
                # 获取全部文本
                content_text = ''
                for paragraph in document.paragraphs:
                    content_text += paragraph.text + '\n'
                logger.info(content_text)

                # 处罚机构
                announcement_org = '江西生态环境厅'
                # 处罚日期
                try:
                    publish_date = re.findall(r'\d{4}年\d{1,2}月\d{1,2}日', content_text)[-1]
                    real_publish_date = format_date(publish_date)
                except Exception as e:
                    logger.info(e)
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)

                # 文号
                try:
                    announcement_code = re.search(r'(赣环[\s\S]*?号)', content_text).group(1).strip()
                except Exception as e:
                    logger.warning(e)
                    announcement_code = ''

                # 当事人
                litigant = re.search(
                    announcement_code.replace(r'(', r'\(').replace(r')', r'\)').replace(r'[', r'\[').replace(r']',
                                                                                                             r'\]') + r'([\s\S]*?)一、环境违法事实、证据和陈述申辩（听证）及采纳情况：',
                    content_text).group(1).strip()

                # 违规事实
                if re.search(r'以上事实，*有(我厅)*[\s\S]*?等*(证据)*为凭。', content_text):
                    facts = re.search(r'一、环境违法事实、证据和陈述申辩（听证）及采纳情况：([\s\S]*?)以上事实，*有(我厅)*[\s\S]*?等*(证据)*为凭。',
                                      content_text).group(1).strip()
                else:
                    facts = re.search(
                        r'一、环境违法事实、证据和陈述申辩（听证）及采纳情况：([\s\S]*?)(你(公司|集团|厂|单位)(的上述违法行为违反了|上述违法行为违反了|的上述行为违反了|上述行为违反了|超标排放|未按环评批复要求)[\s\S]*?的*规定[。\n])',
                        content_text).group(1).strip()

                # 认定意见(有的没有认定意见，最后加个try except)
                try:
                    punishment_basis = re.search(
                        r'(你(公司|集团|厂|单位)(的上述违法行为违反了|上述违法行为违反了|的上述行为违反了|上述行为违反了|超标排放|未按环评批复要求)[\s\S]*?的*规定[。\n])',
                        content_text).group(1).strip()
                except Exception as e:
                    logger.info(e)
                    punishment_basis = ''

                if punishment_basis == '':
                    defense_response_text = re.search(r'(以上事实有[\s\S]*?为凭。\n)([\s\S]*?)(二、行政处罚的依据、种类及其履行方式、期限)',
                                                      content_text).group(2).strip()
                else:
                    defense_response_text = re.search(
                        punishment_basis.replace(r'(', r'\(').replace(r')', r'\)').replace(r'[', r'\[').replace(r']',
                                                                                                                r'\]') + r'([\s\S]*?)(二、行政处罚的依据、种类及其履行方式、期限)',
                        content_text).group(1).strip()

                if re.search(r'未提出书面听证申请|未进行陈述申辩|未提[出交]听证申请|(无|没有)异议|不进行陈述、申辩|不申请听证|接受省厅的处罚决定',
                             defense_response_text) and not re.search(
                    r'收到[\s\S]*?递交的陈述申辩材料|递交了*书面陈述申辩材料|提交书面陈述申辩材料|递交我厅书面陈述申辩材料', defense_response_text):
                    defense_opinion = defense_response_text
                    defense_response = ''
                else:
                    # 申辩意见
                    defense_opinion = re.search(
                        r'([\s\S]*?)(\n(我厅)*经审理查明，|\n我厅经审(查|理)，|\n经查询，|\n我厅案件审理委员会经审理查明|\n\s*经审理查明，)',
                        defense_response_text).group(1).strip()

                    # 申辩意见反馈
                    defense_response = defense_response_text.lstrip(defense_opinion)

                # 处罚决定
                punishment_decision = re.search(r'(二、行政处罚的依据、种类及其履行方式、期限)([\s\S]*?)(限于接到本处罚决定之日起)',
                                                content_text).group(2).strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find(
                        {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('江西生态环境厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('江西生态环境厅 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('江西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.doc'):
                    os.remove('./test/tmp.doc')
                if os.path.exists('./test/tmp.docx'):
                    os.remove('./test/tmp.docx')
            # 如果是excel表格
            elif file_link.endswith('.xls') or file_link.endswith(
                    '.xlsx'):  # http://www.jxepb.gov.cn/ZWGK/ZTZL/wryhjjgxx/xzcf/qyhphpb/2019/c2811ebd656c4eaebae196cdcb778e1b.htm
                link_type = file_link.split('.')[-1]
                response = request_site_page(file_link)
                if response is None:
                    logger.error('网页请求错误')
                    return
                with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                if not os._exists('./test/tmp.xlsx'):
                    shell_str = 'soffice --headless --convert-to xlsx ' + \
                                './test/tmp.' + link_type + ' --outdir ./test'
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()

                with open('./test/tmp.xlsx', 'rb') as xlsx_file:
                    xlsx_content = xlsx_file.read()

                if db.parsed_data.find(
                        {'origin_url': announcement_url, 'oss_file_origin_url': file_link}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': file_link,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'xlsx',
                        'oss_file_name': announcement_title,
                        'oss_file_content': xlsx_content,
                        'parsed': False
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.xlsx', xlsx_content)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = \
                        db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': file_link})[
                            '_id']
                logger.info('存入parsed_data')

                # 开始解析
                excel_data = xlrd.open_workbook('./test/tmp.xlsx')
                result_map_list = []
                for j in range(excel_data.nsheets):
                    sheet = excel_data.sheets()[j]
                    for i in range(2, sheet.nrows):
                        if sheet.cell(i, 0).value == '':
                            break
                        if sheet.ncols == 5 and re.search(r'环境违法黑名单', announcement_title) and sheet.cell(1,
                                                                                                         3).value == '处罚依据':  # http://www.jxepb.gov.cn/ZWGK/ZTZL/wryhjjgxx/xzcf/jbzxcfjdqymd/2014/ead778816d8d459bbc9223b517a84217.htm
                            # 处罚机构
                            announcement_org = '江西生态环境厅'
                            # 处罚日期
                            publish_date = each_document['publishDate']
                            real_publish_date = format_date(publish_date)

                            # 文号
                            announcement_code = ''

                            # 当事人
                            litigant = sheet.cell(i, 1).value

                            # 违规事实
                            facts = sheet.cell(i, 2).value

                            punishment_basis = ''

                            defense_opinion = ''

                            defense_response = ''

                            # 处罚决定
                            punishment_decision = sheet.cell(i, 3).value + ' ' + sheet.cell(i, 4).value
                        elif sheet.ncols == 5 and re.search(r'环境违法黑名单',
                                                            announcement_title):  # http://www.jxepb.gov.cn/ZWGK/ZTZL/wryhjjgxx/xzcf/jbzxcfjdqymd/2016/b73df622243f4716967a215c966c9522.htm
                            # 处罚机构
                            announcement_org = '江西生态环境厅'
                            # 处罚日期
                            publish_date = each_document['publishDate']
                            real_publish_date = format_date(publish_date)

                            # 文号
                            announcement_code = sheet.cell(i, 1).value

                            # 当事人
                            litigant = sheet.cell(i, 2).value

                            # 违规事实
                            facts = sheet.cell(i, 3).value + ' ' + sheet.cell(i, 4).value

                            punishment_basis = ''

                            defense_opinion = ''

                            defense_response = ''

                            # 处罚决定
                            punishment_decision = ''

                        elif sheet.ncols == 5:
                            # 处罚机构
                            announcement_org = '江西生态环境厅'
                            # 处罚日期
                            publish_date = each_document['publishDate']
                            real_publish_date = format_date(publish_date)

                            # 文号
                            announcement_code = ''

                            # 当事人
                            litigant = sheet.cell(i, 2).value + ' 法定代表人（经营者）: ' + sheet.cell(i,
                                                                                             3).value + ' 社会统一信用代码（组织机构代码、工商注册号、税务登记证号）:' + str(
                                sheet.cell(i, 1).value)

                            # 违规事实
                            facts = sheet.cell(i, 4).value

                            punishment_basis = ''

                            defense_opinion = ''

                            defense_response = ''

                            # 处罚决定
                            punishment_decision = ''

                        elif sheet.ncols == 4:
                            # 处罚机构
                            announcement_org = '江西生态环境厅'
                            # 处罚日期
                            publish_date = each_document['publishDate']
                            real_publish_date = format_date(publish_date)

                            # 文号
                            announcement_code = sheet.cell(i, 1).value

                            # 当事人
                            litigant = sheet.cell(i, 2).value

                            # 违规事实
                            facts = sheet.cell(i, 3).value

                            punishment_basis = ''

                            defense_opinion = ''

                            defense_response = ''

                            # 处罚决定
                            punishment_decision = ''

                        elif sheet.ncols == 3:
                            # 处罚机构
                            announcement_org = '江西生态环境厅'
                            # 处罚日期
                            publish_date = each_document['publishDate']
                            real_publish_date = format_date(publish_date)

                            # 文号
                            announcement_code = ''

                            # 当事人
                            litigant = sheet.cell(i, 1).value

                            # 违规事实
                            facts = sheet.cell(i, 2).value

                            punishment_basis = ''

                            defense_opinion = ''

                            defense_response = ''

                            # 处罚决定
                            punishment_decision = ''

                        result_map = {
                            'announcementTitle': announcement_title,
                            'announcementOrg': announcement_org,
                            'announcementDate': real_publish_date,
                            'announcementCode': announcement_code,
                            'facts': facts,
                            'defenseOpinion': defense_opinion,
                            'defenseResponse': defense_response,
                            'litigant': litigant,
                            'punishmentBasement': punishment_basis,
                            'punishmentDecision': punishment_decision,
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                logger.info(result_map_list)
                # 没有具体处罚
                if result_map_list == []:
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('江西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
                    logger.info('删除TMP文件')
                    if os.path.exists('./test/tmp.xls'):
                        os.remove('./test/tmp.xls')
                    if os.path.exists('./test/tmp.xlsx'):
                        os.remove('./test/tmp.xlsx')
                    continue
                if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_many(result_map_list)
                    logger.info('江西生态环境厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('江西生态环境厅 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('江西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.xls'):
                    os.remove('./test/tmp.xls')
                if os.path.exists('./test/tmp.xlsx'):
                    os.remove('./test/tmp.xlsx')


# 山东生态环境厅
def shandong():
    for each_document in db.environment_data.find({'origin': '山东省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        content_text = get_content_text(content_soup.find(class_='zw_new'))

        # 如果公告是在网页里
        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text.encode(content_response.encoding).decode('utf-8'),
                'content_class_name': 'zw_new',
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                         content_response.text.encode(content_response.encoding).decode('utf-8'))
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']

        logger.info('存入parsed_data')

        # 需人工处理
        if announcement_url in ['http://zfc.sdein.gov.cn/xzcf/201804/t20180413_1260652.html',
                                'http://zfc.sdein.gov.cn/xzcf/201601/t20160115_821122.html']:
            logger.info('需人工处理')
            continue

        # 开始解析
        # 处罚机构
        announcement_org = '山东生态环境厅'

        # 处罚日期
        try:
            publish_date = re.findall(r'\n\d{4}年\d{1,2}月\d{1,2}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.info(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

        # 文号
        announcement_code = re.search(r'(.环.字.\d{4}.\d+号)', announcement_title).group(1).strip()

        if re.search(r'一、主要违法事实和证据', content_text):  # http://zfc.sdein.gov.cn/xzcf/201609/t20160905_821123.html
            # 当事人
            litigant = re.search(r'(被处罚单位名称：[\s\S]*?)一、主要违法事实和证据', content_text).group(1).replace(announcement_code,
                                                                                                  '').strip()

            if re.search(r'(\n以上[\s\S]*?(等证据为凭。|作为证据。|询问笔录为证。)\n|\n你公司超标排放水污染物行为有污染源自动监测数据报告、)', content_text):
                # 违规事实
                facts = re.search(
                    r'一、主要违法事实和证据\n([\s\S]*?)\n(以上[\s\S]*?(等证据为凭。|作为证据。|询问笔录为证。)\n|你公司超标排放水污染物行为有污染源自动监测数据报告、)',
                    content_text).group(1).strip()
                # 认定意见
                punishment_basis = re.search(r'(作为证据。\n|等证据为凭。\n|询问笔录为证。\n)([\s\S]*?[的之]规定[。；])', content_text).group(
                    2).strip()
            else:
                facts = re.search(r'一、主要违法事实和证据\n([\s\S]*?)\n你公司的上述行为违反了', content_text).group(1).strip()

                punishment_basis = re.search(r'\n(你公司的上述行为违反了[\s\S]*?[的之]规定。)', content_text).group(1).strip()

            # 申辩意见
            defense_response_text = re.search(
                r'(我厅已于\n*\d{4}年\d{1,2}月\d{1,2}日\n*以《行政处罚事先告知书》[\s\S]*?)二、行政处罚的依据、种类及其履行方式、期限',
                content_text).group(1).strip()

            if re.search(r'我厅审查认为，|我厅认为', defense_response_text):
                defense_opinion = re.search(r'([\s\S]*?)(我厅审查认为，|我厅认为)', defense_response_text).group(1).strip()
                # 申辩意见反馈
                defense_response = defense_response_text.lstrip(defense_opinion)
            else:
                defense_opinion = defense_response_text
                defense_response = ''

            # 处罚决定
            punishment_decision = re.search(r'二、行政处罚的依据、种类及其履行方式、期限([\s\S]*?)\n根据《行政处罚法》和《罚款决定与罚款收缴分离实施办法》',
                                            content_text).group(
                1).strip()

        else:
            # 当事人
            litigant = re.search(r'([\s\S]*?)(\n我厅于\d{4}年\d{1,2}月\d{1,2}日|\n我厅对你公司进行现场检查)',
                                 content_text).group(1).replace(announcement_code, '').strip()

            # 违规事实
            facts = re.search(r'发现你公司实施了以下违法行为：\n([\s\S]*?)\n以上事实有', content_text).group(1).strip()

            if re.search(r'\n以上事实有我厅\d{4}年\d{1,2}月\d{1,2}日《山东省环境保护厅行政处罚听证告知书》[\s\S]*?为证。\n', content_text):
                flag = 1
                envidence = re.search(r'\n以上事实有我厅\d{4}年\d{1,2}月\d{1,2}日《山东省环境保护厅行政处罚听证告知书》[\s\S]*?为证。\n',
                                      content_text).group(0).strip()
                defense_response_text = re.search(
                    r'[的之]规定。\n(我厅[\s\S]*?)\n' + envidence.replace('(', '\(').replace(')', '\)').replace('[',
                                                                                                         '\[').replace(
                        ']', '\]'), content_text).group(1).strip()
            else:
                try:
                    main_fact = re.search(r'证据为凭。\n([\s\S]*?的(行为|违法行为)，)', content_text).group(1).strip()
                    defense_response_text = re.search(
                        r'[的之]规定。\n(我厅[\s\S]*?)' + main_fact.replace('(', '\(').replace(')', '\)').replace('[',
                                                                                                           '\[').replace(
                            ']', '\]'),
                        content_text).group(1).strip()
                    flag = 0
                except Exception:
                    defense_response_text = re.search(
                        r'[的之]规定。\n(我厅[\s\S]*?)(\n依据《中华人民共和国大气污染防治法》|\n依据《中华人民共和国水污染防治法》|\n依据《中华人民共和国固体废物污染环境防治法》)',
                        content_text).group(1).strip()
                    flag = 2
            # 申辩意见
            if re.search(r'经我厅集体讨论|经听证，|我厅审查认为，', defense_response_text):
                defense_opinion = re.search(r'([\s\S]*?)(经我厅集体讨论|经听证，|我厅审查认为，)', defense_response_text).group(1).strip()
                # 申辩意见反馈
                defense_response = defense_response_text.lstrip(defense_opinion)
            else:
                defense_opinion = defense_response_text
                defense_response = ''

            # 认定意见
            punishment_basis = re.search(r'(等证据为凭。|等为证。)\n(你(公司|集团|厂|单位)[\s\S]*?，违反了[\s\S]*?[的之]规定。\n)',
                                         content_text).group(2).strip()
            if flag == 1:
                # 处罚决定
                punishment_decision = re.search(
                    envidence.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                               '\]') + r'([\s\S]*?)\n根据《行政处罚法》和《罚款决定与罚款收缴分离实施办法》',
                    content_text).group(1).strip()
            elif flag == 0:
                punishment_decision = re.search(
                    main_fact.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                               '\]') + r'(依据[\s\S]*?)\n根据《行政处罚法》和《罚款决定与罚款收缴分离实施办法》',
                    content_text).group(1).strip()
            elif flag == 2:
                punishment_decision = re.search(
                    r'((\n依据《中华人民共和国大气污染防治法》|\n依据《中华人民共和国水污染防治法》|\n依据《中华人民共和国固体废物污染环境防治法》)[\s\S]*?)\n根据《行政处罚法》和《罚款决定与罚款收缴分离实施办法》',
                    content_text).group(1).strip()
        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('山东生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('山东生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('山东生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 河南生态环境厅
def henan():
    for each_document in db.environment_data.find({'origin': '河南省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        content_text = get_content_text(content_soup.find(class_='zwcontent'))

        # 如果公告是在网页里
        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'htm',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'content_class_name': 'zwcontent',
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.htm', content_response.text)
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']
        logger.info('存入parsed_data')

        # 需人工处理
        if announcement_url in ['http://www.hnep.gov.cn/xxgk/hbywxxgk/jczf/xzcf/webinfo/2016/08/1494291219689099.htm']:
            logger.info('需人工处理')
            continue

        # 开始解析
        # 处罚机构
        announcement_org = '河南生态环境厅'

        # 处罚日期
        try:
            publish_date = re.findall(r'\n\d{4}年\d{1,2}月\d{1,2}日', content_text)[-1].strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.info(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

        # 文号
        announcement_code = re.search(r'(.环.{2,4}\d{4}.{1,2}(\d+)*号)', content_text).group(1).strip()

        if re.search(r'一、违法行为|一、违法事实和证据', content_text):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                                    '\]') + r'([\s\S]*?)(一、违法行为|一、违法事实和证据)',
                                 content_text).group(1).strip()
            # 违规事实
            facts = re.search(r'(一、违法行为|一、违法事实和证据)([\s\S]*?)(以上事实有我厅|以上事实有“调查询问笔录”)', content_text).group(2).strip()
            # 认定意见
            if re.search(
                    r'((\n上述拆除项目配套的窑炉烟气静电除尘器行为违反了|\n你公司的行为违反了)[\s\S]*?(该建设项目方可正式投入生产或者使用”之规定。|其污染物排放浓度不得超过国家和地方规定的排放标准”。\n))',
                    content_text):  # http://www.hnep.gov.cn/xxgk/hbywxxgk/jczf/xzcf/webinfo/2014/10/1494291220088026.htm
                punishment_basis = re.search(
                    r'((\n上述拆除项目配套的窑炉烟气静电除尘器行为违反了|\n你公司的行为违反了)[\s\S]*?(该建设项目方可正式投入生产或者使用”之规定。|其污染物排放浓度不得超过国家和地方规定的排放标准”。\n))',
                    content_text).group(1).strip()
            else:
                punishment_basis = re.search(
                    r'((\n你(单位|公司|集团|厂)的(上述)*行为违反了|\n上述行为违反了|\n上述3号、4号锅炉未取得环评批准文件|\n你公司调节池和水解酸化池未采取有效措施|\n上述1号锅炉未经环保部门同意擅自投入试生产违反了)[\s\S]*?([的之]规定。|已构成违法。|”规定。)\n)',
                    content_text).group(1).strip()

            defense_response_text = re.search(
                punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]') +
                r'([\s\S]*?)(以上事实有我厅《行政处罚事先（听证）告知书》|以上事实有我厅《按日连续处罚事先（听证）告知书》|\n根据你(单位|公司|集团|厂)违法行为的事实、性质、情节)',
                content_text).group(1).strip()

            if re.search(r'(我厅对你(单位|公司|集团|厂)的(申辩材料|申辩理由|陈述申辩意见)进行了(复核|审查)，)', defense_response_text):
                # 申辩意见
                defense_opinion = re.search(r'([\s\S]*?)(我厅对你(单位|公司|集团|厂)的(申辩材料|申辩理由|陈述申辩意见)进行了(复核|审查)，)',
                                            defense_response_text).group(1).strip()
                # 申辩意见反馈
                defense_response = re.search(r'((我厅对你(单位|公司|集团|厂)的(申辩材料|申辩理由|陈述申辩意见)进行了(复核|审查)，)[\s\S]*)',
                                             defense_response_text).group(1).strip()
            else:
                defense_opinion = defense_response_text
                defense_response = ''

            # 处罚决定
            if re.search(
                    r'二、行政处罚的依据、种类及其履行方式和期限([\s\S]*?)(\n根据《中华人民共和国行政处罚法》和《罚款决定与罚款收缴分离实施办法》的规定|\n你(单位|公司|集团|厂)应当自接到本处罚决定书之日起)',
                    content_text):
                punishment_decision = re.search(
                    r'二、行政处罚的依据、种类及其履行方式和期限([\s\S]*?)(\n根据《中华人民共和国行政处罚法》和《罚款决定与罚款收缴分离实施办法》的规定|\n你(单位|公司|集团|厂)应当自接到本处罚决定书之日起)',
                    content_text).group(1).strip()
            else:
                punishment_decision = re.search(r'二、行政处罚的依据、种类及其履行方式和期限([\s\S]*?)三、申请行政复议或者提起行政诉讼的途径和期限',
                                                content_text).group(1).strip()

        elif re.search(r'一、环境违法事实和证据',
                       content_text):  # http://www.hnep.gov.cn/xxgk/hbywxxgk/jczf/xzcf/webinfo/2012/12/1494291220305243.htm
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                                    '\]') + r'([\s\S]*?)(一、环境违法事实和证据)',
                                 content_text).group(1).strip()
            # 违规事实
            facts = re.search(r'(一、环境违法事实和证据)([\s\S]*?)(你(单位|公司|集团|厂)的*上述行为违反了)', content_text).group(2).strip()
            # 认定意见
            punishment_basis = re.search(r'(你(单位|公司|集团|厂)的*上述行为违反了[\s\S]*?”[的之]*规定)', content_text).group(1).strip()

            defense_response_text = re.search(
                r'(\n\d{4}年\d{1,2}月\d{1,2}日\n*，我厅向你(单位|公司|集团|厂)(邮寄|直接)*送达了《行政处罚事先（听证）告知书》[\s\S]*?)(二、行政处罚的依据、种类)',
                content_text).group(1).strip()

            if re.search(r'((我厅对你(单位|公司|集团|厂)的(申辩材料|申辩理由|陈述申辩意见|陈述和申辩)进行了(复核|审查))|经研究，我厅认为：)', defense_response_text):
                # 申辩意见
                defense_opinion = re.search(
                    r'([\s\S]*?)((我厅对你(单位|公司|集团|厂)的(申辩材料|申辩理由|陈述申辩意见|陈述和申辩)进行了(复核|审查))|经研究，我厅认为：)',
                    defense_response_text).group(1).strip()
                # 申辩意见反馈
                defense_response = re.search(
                    r'(((我厅对你(单位|公司|集团|厂)的(申辩材料|申辩理由|陈述申辩意见|陈述和申辩)进行了(复核|审查))|经研究，我厅认为：)[\s\S]*)',
                    defense_response_text).group(1).strip()
            else:
                defense_opinion = defense_response_text
                defense_response = ''

            # 处罚决定
            punishment_decision = re.search(r'二、行政处罚的依据、种类([\s\S]*?)三、处罚决定的履行方式和期限', content_text).group(1).strip()
        else:
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                                    '\]') + r'([\s\S]*?)(\n我厅)',
                                 content_text).group(1).strip()
            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]') +
                              r'([\s\S]*?)(你公司作为矿产采挖企业其行为违反了|你单位作为建材生产企业其行为违反了|你公司的行为违反了)', content_text).group(
                1).strip()
            # 认定意见
            punishment_basis = re.search(r'((你公司作为矿产采挖企业其行为违反了|你单位作为建材生产企业其行为违反了|你公司的行为违反了)[\s\S]*?已构成环境违法。)',
                                         content_text).group(1).strip()

            defense_response_text = re.search(
                r'(你(单位|公司|集团|厂)于\d{4}年\d{1,2}月\d{1,2}日签收了我厅《行政处罚事先（听证）告知书》[\s\S]*?)\n以上事实有我厅《行政处罚事先（听证）告知书》[\s\S]*?等材料为证。\n',
                content_text).group(1).strip()

            if re.search(r'(我厅对你公司的陈述申辩意见进行了复核，|我厅对你公司的申辩理由进行了审查，)', defense_response_text):
                # 申辩意见
                defense_opinion = re.search(r'([\s\S]*?)(我厅对你公司的陈述申辩意见进行了复核，|我厅对你公司的申辩理由进行了审查，)',
                                            defense_response_text).group(1).strip()
                # 申辩意见反馈
                defense_response = re.search(r'((我厅对你公司的陈述申辩意见进行了复核，|我厅对你公司的申辩理由进行了审查，)[\s\S]*)',
                                             defense_response_text).group(1).strip()
            else:
                defense_opinion = defense_response_text
                defense_response = ''

            # 处罚决定
            punishment_decision = re.search(r'等材料为证。\n([\s\S]*?)\n根据《中华人民共和国行政处罚法》和《罚款决定与罚款收缴分离实施办法》的规定',
                                            content_text).group(1).strip()

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('河南生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('河南生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('河南生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 湖南生态环境厅
def hunan():
    for each_document in db.environment_data.find({'origin': '湖南省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() == 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        # 如果公告是在网页图片里
        if content_soup.find(class_='main_con_zw').find('img') and not content_soup.find(class_='main_con_zw').find(
                'img', height='9') and not content_soup.find(class_='main_con_zw').find('a'):
            img_link_list = [urljoin(announcement_url, kk.attrs['src'])
                             for kk in content_soup.find('div', class_='main_con_zw').find_all('img')]
            for index, each_img_link in enumerate(img_link_list):
                img_response = request_site_page(each_img_link)
                with open('./test/' + str(index) + '.jpg', 'wb') as tmp_file:
                    for chunk in img_response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

            if not os.path.exists('./test/tmp.pdf'):
                shell_str = 'img2pdf '
                for index in range(len(img_link_list)):
                    shell_str += './test/' + str(index) + '.jpg '
                shell_str += '--imgsize 20cmx30cm  -o ./test/tmp.pdf'  # 加入图片尺寸参数，百度ocr有图片尺寸限制
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.communicate()

            result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')

            with open('./test/tmp.pdf', 'rb') as pdf_file:
                pdf_content = pdf_file.read()

            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False,
                    'if_ocr': True,
                    'ocr_result': result_text
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']
            logger.info('存入parsed_data')
            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.pdf'):
                os.remove('./test/tmp.pdf')
            for index in range(len(img_link_list)):
                if os.path.exists('./test/' + str(index) + '.jpg'):
                    os.remove('./test/' + str(index) + '.jpg')

            # 开始解析
            # 需要人工处理
            logger.info('需人工处理')
            continue

        # 如果公告是doc文档或pdf文档
        elif content_soup.find(class_='main_con_zw').find('a'):
            file_link = urljoin(announcement_url, content_soup.find(class_='main_con_zw').find('a')['href'])
            if file_link.endswith('.doc') or file_link.endswith('.docx'):
                link_type = file_link.split('.')[-1]
                response = request_site_page(file_link)
                if response is None:
                    logger.error('doc文件下载失败%s' % file_link)
                    return
                with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)

                # 如果doc文档里是图片，直接转成pdf
                if get_content_text(content_soup.find(class_='main_con_zw').find('a').find('font')).endswith(
                        '.PDF.doc'):  # http://sthjt.hunan.gov.cn/xxgk/zdly/jdzf/ajcc/201501/t20150114_4667113.html
                    if not os._exists('./test/tmp.pdf'):
                        shell_str = 'soffice --headless --convert-to pdf ' + \
                                    './test/tmp.' + link_type + ' --outdir ./test'
                        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                                   shell=True, stdout=subprocess.PIPE)
                        process.communicate()

                    result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')

                    with open('./test/tmp.pdf', 'rb') as pdf_file:
                        pdf_content = pdf_file.read()

                    if db.parsed_data.find(
                            {'origin_url': announcement_url, 'oss_file_origin_url': file_link}).count() == 0:
                        oss_file_map = {
                            'origin_url': announcement_url,
                            'oss_file_origin_url': file_link,
                            'origin_url_id': each_document['_id'],
                            'oss_file_type': 'pdf',
                            'oss_file_name': announcement_title,
                            'oss_file_content': pdf_content,
                            'parsed': False,
                            'if_ocr': True,
                            'ocr_result': result_text
                        }
                        response = db.parsed_data.insert_one(oss_file_map)
                        file_id = response.inserted_id
                        oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                        db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    else:
                        db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                        file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                           'oss_file_origin_url': file_link})['_id']
                    logger.info('存入parsed_data')
                    logger.info('删除TMP文件')
                    if os.path.exists('./test/tmp.pdf'):
                        os.remove('./test/tmp.pdf')
                    if os.path.exists('./test/tmp.' + link_type):
                        os.remove('./test/tmp.' + link_type)
                    # 开始解析
                    # 需要人工处理
                    logger.info('需人工处理')
                    continue

                else:
                    if not os._exists('./test/tmp.docx'):
                        shell_str = 'soffice --headless --convert-to docx ' + \
                                    './test/tmp.' + link_type + ' --outdir ./test'
                        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                                   shell=True, stdout=subprocess.PIPE)
                        process.communicate()

                    with open('./test/tmp.docx', 'rb') as docx_file:
                        docx_content = docx_file.read()

                    if db.parsed_data.find(
                            {'origin_url': announcement_url, 'oss_file_origin_url': file_link}).count() == 0:
                        oss_file_map = {
                            'origin_url': announcement_url,
                            'oss_file_origin_url': file_link,
                            'origin_url_id': each_document['_id'],
                            'oss_file_type': 'docx',
                            'oss_file_name': announcement_title,
                            'oss_file_content': docx_content,
                            'parsed': False
                        }
                        response = db.parsed_data.insert_one(oss_file_map)
                        file_id = response.inserted_id
                        oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                        db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    else:
                        db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                        file_id = \
                            db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': file_link})[
                                '_id']
                    logger.info('存入parsed_data')

                    # 获取docx文档内容
                    document = Document('./test/tmp.docx')
                    content_text = '\n'.join([each_paragraph.text.strip() for each_paragraph in document.paragraphs])
                    logger.info(content_text)

                    logger.info('删除TMP文件')
                    if os.path.exists('./test/tmp.doc'):
                        os.remove('./test/tmp.doc')
                    if os.path.exists('./test/tmp.docx'):
                        os.remove('./test/tmp.docx')

                    # 开始解析
                    # 处罚机构
                    announcement_org = '湖南生态环境厅'

                    # 处罚日期
                    try:
                        publish_date = re.findall(r'\n\w{4}年\w{1,3}月\w{1,3}日', content_text)[-1].strip()
                        real_publish_date = format_date(publish_date)
                    except Exception as e:
                        logger.info(e)
                        publish_date = each_document['publishDate']
                        real_publish_date = format_date(publish_date)

                    # 文号
                    announcement_code = re.search(r'(湘环[\s\S]*?号)', announcement_title).group(1).strip()

                    # 当事人
                    litigant = re.search(r'(行政处罚决定书\n|行政处罚事先告知书\n)([\s\S]*?)(一、环境违法事实和证据|一、环境违法事实、证据和陈述申辩（听证）情况)',
                                         content_text).group(2).strip()
                    # 违规事实
                    if re.search(r'(一、环境违法事实和证据|一、环境违法事实、证据和陈述申辩（听证）情况)([\s\S]*?)以上违法事实有', content_text):
                        facts = re.search(r'(一、环境违法事实和证据|一、环境违法事实、证据和陈述申辩（听证）情况)([\s\S]*?)以上违法事实有', content_text).group(
                            2).strip()
                    else:
                        facts = re.search(r'(一、环境违法事实和证据|一、环境违法事实、证据和陈述申辩（听证）情况)([\s\S]*?)((你(单位|公司|集团|厂))*上述行为违反了)',
                                          content_text).group(2).strip()
                    # 认定意见
                    punishment_basis = re.search(r'((你(单位|公司|集团|厂))*上述行为违反了[\s\S]*?[的之]规定。)', content_text).group(
                        1).strip()

                    if re.search(
                            punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                              '\]') +
                            r'([\s\S]*?)(以上事实，有我厅《行政处罚事先（听证）告知书》|以上事实，有我厅《行政处罚事先告知书》)', content_text):
                        defense_response_text = re.search(
                            punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                              '\]') +
                            r'([\s\S]*?)(以上事实，有我厅《行政处罚事先（听证）告知书》|以上事实，有我厅《行政处罚事先告知书》)', content_text).group(1).strip()
                    else:
                        defense_response_text = re.search(
                            punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                              '\]') +
                            r'([\s\S]*?)(二、行政处罚的依据、种类及其履行方式、期限)', content_text).group(1).replace(
                            '以上事实有现场照片、现场监察记录、调查询问笔录和现场检查（勘查）为证。', '').strip()

                    if re.search(r'(考虑到你单位积极整改|鉴于以上情况)', defense_response_text):
                        # 申辩意见
                        defense_opinion = re.search(r'([\s\S]*)(考虑到你单位积极整改|鉴于以上情况)', defense_response_text).group(
                            1).strip()
                        # 申辩意见反馈
                        defense_response = re.search(r'((考虑到你单位积极整改|鉴于以上情况)[\s\S]*)', defense_response_text).group(
                            1).strip()
                    else:
                        defense_opinion = defense_response_text
                        defense_response = ''

                    # 处罚决定
                    punishment_decision = re.search(
                        r'二、行政处罚的依据、种类及其履行方式、期限([\s\S]*?)(\n你单位应于接到本处罚决定书之日起|三、申请行政复议或者提起行政诉讼的途径和期限)',
                        content_text).group(1).strip()

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': defense_opinion,
                        'defenseResponse': defense_response,
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    logger.info(result_map)
                    if db.announcement.find(
                            {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                        db.announcement.insert_one(result_map)
                        logger.info('湖南生态环境厅 数据解析 ' + ' -- 数据导入完成')
                    else:
                        logger.info('湖南生态环境厅 数据解析 ' + ' -- 数据已经存在')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('湖南生态环境厅 数据解析 ' + ' -- 修改parsed完成')

            elif file_link.endswith('.PDF') or file_link.endswith('.pdf'):
                response = request_site_page(file_link)
                if response is None:
                    logger.error('pdf文件下载失败%s' % file_link)
                    return
                with open('./test/tmp.pdf', 'wb') as tmp_file:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)
                with open('./test/tmp.pdf', 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')
                content_text = result_text

                if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': file_link}).count() == 0:
                    oss_file_map = {
                        'origin_url': announcement_url,
                        'oss_file_origin_url': file_link,
                        'origin_url_id': each_document['_id'],
                        'oss_file_type': 'pdf',
                        'oss_file_name': announcement_title,
                        'oss_file_content': pdf_content,
                        'parsed': False,
                        'if_ocr': True,
                        'ocr_result': result_text
                    }
                    response = db.parsed_data.insert_one(oss_file_map)
                    file_id = response.inserted_id
                    oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                else:
                    db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                    file_id = \
                        db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': file_link})[
                            '_id']
                logger.info('存入parsed_data')
                logger.info(content_text)

                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.pdf'):
                    os.remove('./test/tmp.pdf')

                # 开始解析
                # 处罚机构
                announcement_org = '湖南生态环境厅'

                # 处罚日期
                try:
                    publish_date = re.findall(r'\n\w{4}年\w{1,3}月\w{1,3}日', content_text)[-1].strip()
                    real_publish_date = format_date(publish_date)
                except Exception as e:
                    logger.info(e)
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)

                # 文号
                announcement_code = re.search(r'(湘环[\s\S]*?号)', announcement_title).group(1).strip()

                # 当事人
                litigant = re.search(r'(行政处罚决定书\n)([\s\S]*?)(\n环境违法事实和证据|\n环境违法事实、证据和陈述申(辩|辨|辦)\(听证\)情况)',
                                     content_text).group(2).strip()
                # 违规事实
                facts = re.search(
                    r'(\n环境违法事实和证据|\n环境违法事实、证据和陈述申[辩辨辦]\(听证\)情况)([\s\S]*?)((你(单位|公司|集团|厂))*上述行为(违反了|违\n反了|违反\n))',
                    content_text).group(2).strip()
                if re.search(r'(以上违法事实[\s\S]*?为证。)', facts):
                    tmp = re.search(r'(以上违法事实[\s\S]*?为证。)', facts).group(1)
                    facts = facts.replace(tmp, '').strip()

                # 认定意见
                punishment_basis = re.search(r'((你(单位|公司|集团|厂))*上述行为(违反了|违\n反了|违反\n)[\s\S]*?[的之]规\n*定。*)',
                                             content_text).group(1).strip()

                defense_response_text = re.search(
                    r'((\n我[斤厅]已委托|我[斤厅]于\d{4}年\d{1,2}月\d{1,2}[日目][以将]《行政处罚事先\(听证\)告知书》)[\s\S]*?)(以\n*上事实有|以上事实,有我[斤厅]《行政处罚事先\(听证\)告知书》)',
                    content_text).group(1).strip()

                if re.search(r'(经研究,|经集体研究,)', defense_response_text):
                    # 申辩意见
                    defense_opinion = re.search(r'([\s\S]*)(经研究,|经集体研究,)', defense_response_text).group(
                        1).strip()
                    # 申辩意见反馈
                    defense_response = re.search(r'((经研究,|经集体研究,)[\s\S]*)', defense_response_text).group(
                        1).strip()
                else:
                    defense_opinion = defense_response_text
                    defense_response = ''

                # 处罚决定
                punishment_decision = re.search(
                    r'行政处罚的依据、种类及其履行方式、期限([\s\S]*?)(\n你单位应于接到本处罚决定书之日起|三*、申请行政复议或者提起行政诉讼的途径和期限)',
                    content_text).group(1).strip()

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find(
                        {'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('湖南生态环境厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('湖南生态环境厅 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('湖南生态环境厅 数据解析 ' + ' -- 修改parsed完成')
        # 如果公告在网页上
        else:
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text,
                    'content_class_name': 'main_con_zw',
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']
            logger.info('存入parsed_data')
            content_text = get_content_text(content_soup.find(class_='main_con_zw'))

            # 需人工处理
            if announcement_url in ['http://sthjt.hunan.gov.cn/xxgk/zdly/jdzf/ajcc/201509/t20150906_4667128.html']:
                logger.info('需人工处理')
                continue

            # 开始解析
            # 处罚机构
            announcement_org = '湖南生态环境厅'

            # 处罚日期
            try:
                publish_date = re.findall(r'\n\d{4}年\d{1,2}月\d{1,2}日', content_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.info(e)
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

            # 文号
            announcement_code = re.search(r'(湘环[\s\S]*?号|湘环罚告字\[2015\])', announcement_title).group(1).strip()
            if re.search(r'一、违法事实和证据',
                         content_text):  # http://sthjt.hunan.gov.cn/xxgk/zdly/jdzf/ajcc/201701/t20170119_4667157.html  http://sthjt.hunan.gov.cn/xxgk/zdly/jdzf/ajcc/201701/t20170119_4667153.html
                # 当事人
                litigant = re.search(r'行政处罚决定书\n([\s\S]*?)(一、违法事实和证据)', content_text).group(1).strip()
                # 违规事实
                facts = re.search(r'(一、违法事实和证据)([\s\S]*?)\n上述行为违反了', content_text).group(2).strip()
                # 认定意见
                punishment_basis = re.search(r'(\n上述行为违反了[\s\S]*?[的之]规定。\n)', content_text).group(1).strip()

                # 申辩意见
                defense_opinion = re.search(r'(\n我厅于\d{4}年\d{1,2}月\d{1,2}日以《行政处罚事先告知书》[\s\S]*?)(\n根据你院违法行为的事实、性质、情节)',
                                            content_text).group(1).strip()
                # 申辩意见反馈
                defense_response = re.search(r'(\n根据你院违法行为的事实、性质、情节[\s\S]*?)二、行政处罚的依据、种类及其履行方式和期限', content_text).group(
                    1).strip()

                # 处罚决定
                punishment_decision = re.search(r'二、行政处罚的依据、种类及其履行方式和期限([\s\S]*?)\n限于接到本处罚决定之日起', content_text).group(
                    1).strip()
            elif re.search(r'一、环境违法事实和证据',
                           content_text):  # http://sthjt.hunan.gov.cn/xxgk/zdly/jdzf/ajcc/201312/t20131213_4667096.html
                # 当事人
                litigant = re.search(r'行政处罚决定书\n([\s\S]*?)(一、环境违法事实和证据)', content_text).group(1).strip()
                # 违规事实
                facts = re.search(r'(一、环境违法事实和证据)([\s\S]*?)以上(违法)*事实有', content_text).group(2).strip()
                # 认定意见
                punishment_basis = re.search(r'((\n你(单位|公司|集团|厂))*上述行为违反了[\s\S]*?[的之]*规定。)', content_text).group(
                    1).strip()

                if re.search(
                        punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]') +
                        r'([\s\S]*?)(\n以上事实，有我厅《行政处罚事先（听证）告知书》)', content_text):
                    defense_response_text = re.search(
                        punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]') +
                        r'([\s\S]*?)(\n以上事实，有我厅《行政处罚事先（听证）告知书》)', content_text).group(1).strip()
                else:
                    defense_response_text = re.search(
                        punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]') +
                        r'([\s\S]*?)(二、行政处罚的依据、种类及其履行方式、期限)', content_text).group(1).strip()
                if re.search(r'(经研究|考虑到你公司)', defense_response_text):
                    # 申辩意见
                    defense_opinion = re.search(r'([\s\S]*)(经研究|考虑到你公司)', defense_response_text).group(
                        1).strip()
                    # 申辩意见反馈
                    defense_response = re.search(r'((经研究|考虑到你公司)[\s\S]*)', defense_response_text).group(
                        1).strip()
                else:
                    defense_opinion = defense_response_text
                    defense_response = ''

                # 处罚决定
                punishment_decision = re.search(r'二、行政处罚的依据、种类及其履行方式、期限([\s\S]*?)\n你单位应于接到本处罚决定书之日起',
                                                content_text).group(
                    1).strip()

            else:
                if re.search(r'行政处罚决定书\n([\s\S]*?)\n(经调查核实，发现你(单位|公司|集团|厂)存在以下环境违法行为：|我厅于\d{4}年\d{1,2}月\d{1,2}日)',
                             content_text):
                    # 当事人
                    litigant = re.search(
                        r'行政处罚决定书\n([\s\S]*?)\n(经调查核实，发现你(单位|公司|集团|厂)存在以下环境违法行为：|我厅于\d{4}年\d{1,2}月\d{1,2}日)',
                        content_text).group(1).strip()
                else:
                    litigant = re.search(
                        r'(else{}|TRS_EditorA{line-height:2;font-family:宋体;font-size:12pt;})\n([\s\S]*?)\n(\d{4}年\d{1,2}月\d{1,2}日，(我厅|湖南省环境保护厅)|你(单位|公司|集团|厂)已于\d{4}年底停产|我厅于\d{4}年\d{1,2}月\d{1,2}日)',
                        content_text).group(2).strip()
                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                            '\]') + r'([\s\S]*?)(以上事实有现场照片|以上违法事实有现场照片|该行为违反了《放射性同位素与射线装置安全和防护条例》)',
                                  content_text).group(1).strip()
                # 认定意见
                punishment_basis = re.search(
                    r'((\n你(单位|公司|集团|厂)的*上述行为违反了|该行为违反了《放射性同位素与射线装置安全和防护条例》)[\s\S]*?[的之]规定。\n*)',
                    content_text).group(1).strip()

                try:
                    # 申辩意见
                    defense_opinion = re.search(
                        punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                          '\]') + r'([\s\S]*?)(\n依据《湖南省建设项目环境保护管理办法》|\n依据《中华人民共和国水污染防治法》)',
                        content_text).group(1).strip()
                except Exception as e:
                    defense_opinion = ''
                    logger.warning(e)
                    logger.info('no defense_opinion')
                # 申辩意见反馈
                defense_response = ''

                # 处罚决定
                punishment_decision = re.search(
                    r'((\n依据《湖南省建设项目环境保护管理办法》|\n依据《中华人民共和国水污染防治法》|现依据《放射性同位素与射线装置安全和防护条例》)[\s\S]*?)(\n你单位应于接到本处罚决定书之日起|\n根据《中华人民共和国行政处罚法》第三十一条的规定)',
                    content_text).group(1).strip()
            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': announcement_org,
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': defense_opinion,
                'defenseResponse': defense_response,
                'litigant': litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('湖南生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖南生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖南生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 广东生态环境厅
def guangdong():
    for each_document in db.environment_data.find({'origin': '广东省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'html5lib')

        content_text = get_content_text(content_soup.find(class_='content'))

        # 如果公告是在网页里
        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'content_class_name': 'content',
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']
        logger.info('存入parsed_data')

        # 需人工处理
        if announcement_url in ['http://pub.gdepb.gov.cn/pub/pubcatalog/extranet_pub_document_view.jsp?docId=212434']:
            logger.info('需人工处理')
            continue

        # 开始解析
        # 处罚机构
        announcement_org = '广东生态环境厅'

        # 处罚日期
        try:
            publish_date = re.search(r'(广东省环境保护厅)\n(\d{4}年\d{1,2}月\d{1,2}日)', content_text).group(2).strip()
            real_publish_date = format_date(publish_date)
        except Exception as e:
            logger.info(e)
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

        # 文号
        try:
            announcement_code = re.search(r'\n(粤环[\s\S]*?号)\n', content_text).group(1).strip()
        except Exception as e:
            logger.info(e)
            announcement_code = ''

        # http://pub.gdepb.gov.cn/pub/pubcatalog/extranet_pub_document_view.jsp?docId=197522
        if re.search(r'限期治理决定书', announcement_title):
            # http://pub.gdepb.gov.cn/pub/pubcatalog/extranet_pub_document_view.jsp?docId=167261
            if re.search(r'一、适用限期治理之事实和证据', content_text):
                # 当事人
                litigant = re.search(
                    announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                       '\]') + r'([\s\S]*?)\n经调查核实，',
                    content_text).group(1).strip()
                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                            '\]') + r'([\s\S]*?)(该行为违反了)',
                                  content_text).group(1).strip()
                # 认定意见
                punishment_basis = re.search(r'(该行为违反了[\s\S]*?的规定。)', content_text).group(1).strip()

                defense_opinion = ''
                defense_response = ''
                # 处罚决定
                punishment_decision = re.search(r'二、适用限期治理之依据和要求([\s\S]*?)(（二）法律权利|2、法律权利)', content_text).group(
                    1).strip()

            else:
                # 当事人
                litigant = re.search(
                    announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                       '\]') + r'([\s\S]*?)：\n',
                    content_text).group(1).strip()
                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                            '\]') + r'([\s\S]*?)(经研究，)',
                                  content_text).group(1).lstrip('：').strip()
                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''
                # 处罚决定
                punishment_decision = re.search(r'(经研究，[\s\S]*?)广东省环境保护厅', content_text).group(1).strip()

        elif re.search(r'限期治理验收意见的函|延长限期治理期限的函', announcement_title):
            # 当事人
            litigant = re.search(r'(限期治理验收意见的函|延长限期治理期限的函)\n([\s\S]*?)：\n', content_text).group(2).strip()
            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                        '\]') + r'([\s\S]*?)(经研究，)',
                              content_text).group(1).lstrip('：').strip()
            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''
            # 处罚决定
            punishment_decision = re.search(r'(经研究，[\s\S]*?)广东省环境保护厅', content_text).group(1).strip()

        # http://pub.gdepb.gov.cn/pub/pubcatalog/extranet_pub_document_view.jsp?docId=195785
        elif re.search(r'督促履行义务催告书', announcement_title):
            # 当事人
            litigant = re.search(r'督促履行义务催告书\n([\s\S]*?)：\n', content_text).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                        '\]') + r'([\s\S]*?)(你公司在法定期限内未申请行政复议)',
                              content_text).group(1).lstrip('：').strip()
            # 认定意见
            punishment_basis = ''

            defense_opinion = re.search(r'(你公司在法定期限内未申请行政复议[\s\S]*?也未履行上述责令改正决定。)', content_text).group(1).strip()

            defense_response = ''

            # 处罚决定
            punishment_decision = re.search(r'\n(《中华人民共和国行政强制法》[\s\S]*?)\n你单位有权进行陈述和申辩。', content_text).group(1).strip()

        elif re.search(r'环保验收手续的通知*', announcement_title):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                                    '\]') + r'([\s\S]*?)：\n',
                                 content_text).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                        '\]') + r'([\s\S]*?)(根据《建设项目环境保护管理条例》)',
                              content_text).group(1).lstrip('：').strip()
            # 认定意见
            punishment_basis = ''

            defense_opinion = ''

            defense_response = ''

            # 处罚决定
            punishment_decision = re.search(r'(根据《建设项目环境保护管理条例》[\s\S]*?依法进行处理。)\n', content_text).group(1).strip()

        elif re.search(r'责令改正违法行为决定书|责令停产整治决定书', announcement_title):
            if re.search(r'一、环境违法事实和理由|一、环境违法事实和证据|一、调查情况及发现的环境违法事实、证据|一、环境违法事实和依据', content_text):
                # 当事人
                litigant = re.search(
                    announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                       '\]') + r'([\s\S]*?)(一、环境违法事实和理由|一、环境违法事实和证据|一、调查情况及发现的环境违法事实、证据|一、环境违法事实和依据)',
                    content_text).group(1).strip()
                # 违规事实
                facts = re.search(
                    r'(一、环境违法事实和理由|一、环境违法事实和证据|一、调查情况及发现的环境违法事实、证据|一、环境违法事实和依据)([\s\S]*?)(以上事实，有|上述事实有|《调查询问笔录》，2012年4月17日的《现场检查笔录》及粤环审〔2011〕494号文为证。)',
                    content_text).group(2).strip()
                # 认定意见
                if re.search(r'\n(你公司通过生产废水处理设施的冷却塔下方集水池与[\s\S]*?中华人民共和国水污染防治法》第二十二条第二款的规定。)\n', content_text):
                    punishment_basis = re.search(r'\n(你公司通过生产废水处理设施的冷却塔下方集水池与[\s\S]*?中华人民共和国水污染防治法》第二十二条第二款的规定。)\n',
                                                 content_text).group(1).strip()
                else:
                    punishment_basis = re.search(
                        r'((你(单位|公司|集团|厂)违反了|你(单位|公司|集团|厂)的*上述行为(分别)*违反了|你(单位|公司|集团|厂)(上述排放水污染物超过国家规定的排放标准的行为|下辖运营管理的|上述排放大气污染物|通过生产废水处理设施)|声污染防治法》第十四条第二款)[\s\S]*?规定。)',
                        content_text).group(1).strip()

                # 个别公告会产生认定意见和违规事实的乱序
                if re.search(punishment_basis, facts):
                    facts = facts.replace(punishment_basis, '').strip()
                    defense_response_text = re.search(
                        r'为证。\n([\s\S]*?)(二、行政决定的依据、种类及其履行方式、期限|二、责令改正的依据、种类及履行方式|二、责令改正的依据、种类和拒不改正的法律后果|二、责令改正的依据、种类)',
                        content_text).group(1).strip()
                else:
                    defense_response_text = re.search(
                        punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                          '\]') + r'([\s\S]*?)(二、行政决定的依据、种类及其履行方式、期限|二、责令改正的依据、种类及履行方式|二、责令改正的依据、种类和拒不改正的法律后果|二、责令改正的依据、种类)',
                        content_text).group(1).strip()

                if re.search(r'根据《中华人民共和国行政处罚法》第三十二条和《环境行政处罚办法》第四十九条的规定，我厅依法进行了复核。', defense_response_text):
                    # 申辩意见
                    defense_opinion = re.search(r'([\s\S]*)(根据《中华人民共和国行政处罚法》第三十二条和《环境行政处罚办法》第四十九条的规定，我厅依法进行了复核。)',
                                                defense_response_text).group(1).strip()

                    # 申辩意见反馈
                    defense_response = re.search(r'(根据《中华人民共和国行政处罚法》第三十二条和《环境行政处罚办法》第四十九条的规定，我厅依法进行了复核。[\s\S]*)',
                                                 defense_response_text).group(1).strip()
                else:
                    defense_opinion = defense_response_text
                    defense_response = ''
                # 处罚决定
                if re.search(r'(二、责令改正的依据、种类和拒不改正的法律后果|二、责令改正的依据、种类|二、行政决定的依据、种类及其履行方式、期限)([\s\S]*?)我厅将自送达本决定书之日起',
                             content_text):
                    punishment_decision = re.search(
                        r'(二、责令改正的依据、种类和拒不改正的法律后果|二、责令改正的依据、种类|二、行政决定的依据、种类及其履行方式、期限)([\s\S]*?)我厅将自送达本决定书之日起',
                        content_text).group(2).strip()
                else:
                    punishment_decision = re.search(
                        r'(二、责令改正的依据、种类和拒不改正的法律后果|二、责令改正的依据、种类|二、行政决定的依据、种类及其履行方式、期限)([\s\S]*?)(三、申请行政复议或者*提起行政诉讼的途径和期限|四、申请行政复议或者*提起行政诉讼的途径和期限)',
                        content_text).group(2).strip()

            else:
                # 当事人
                litigant = re.search(
                    announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                       '\]') + r'([\s\S]*?)(\n\d{4}年\d{1,2}月，我厅在审查|\n我厅在审查|\n我厅执法人员)',
                    content_text).group(1).strip()

                # 违规事实
                facts = re.search(
                    litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]').replace('\\',
                                                                                                                 '\\\\') +
                    r'([\s\S]*?)(\n上述事实有|\n以上事实，有)', content_text).group(1).strip()
                # 认定意见
                punishment_basis = re.search(r'((你(单位|公司|集团|厂)的上述行为违反了)[\s\S]*?(的规定。|需要配套建设的环境保护设施未经验收即投入使用。))',
                                             content_text).group(1).strip()

                # 申辩意见
                defense_opinion = ''

                # 申辩意见反馈
                defense_response = ''

                # 处罚决定
                punishment_decision = re.search(
                    punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                      '\]') + r'([\s\S]*?)\n你公司如对本决定不服',
                    content_text).group(1).strip()

        elif re.search(r'行政处罚决定书', announcement_title):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                                    '\]') + r'([\s\S]*?)(一、调查情况及发现的环境违法事实、证据|一、调查情况及认定的环境违法事实、证据)',
                                 content_text).group(1).strip()
            # 违规事实
            facts = re.search(
                r'(一、调查情况及发现的环境违法事实、证据|一、调查情况及认定的环境违法事实、证据)([\s\S]*?)(以上事实，有|上述事实有|\n上述违事实有|\n上述违法事实有以下证据)',
                content_text).group(2).strip()

            # 认定意见
            if re.search(
                    r'(等为证。|等证据为证。|等证据。)\n([\s\S]*?)(\n\d{4}年\d{1,2}月\d{1,2}日，我厅(分别)*向你(单位|公司|集团|厂)送达了|我厅于\d{4}年\d{1,2}月\d{1,2}日向你(单位|公司|集团|厂)送达了|我厅已于\d{4}年\d{1,2}月\d{1,2}日以留置送达方式向你(单位|公司|集团|厂)送达了|\d{4}年\d{1,2}月\d{1,2}日和(\d{4}年)*\d{1,2}月\d{1,2}日，我厅(分别)*向你(单位|公司|集团|厂)送达了)',
                    content_text):
                punishment_basis = re.search(
                    r'(等为证。|等证据为证。|等证据。)\n([\s\S]*?)(\n\d{4}年\d{1,2}月\d{1,2}日，我厅(分别)*向你(单位|公司|集团|厂)送达了|我厅于\d{4}年\d{1,2}月\d{1,2}日向你(单位|公司|集团|厂)送达了|我厅已于\d{4}年\d{1,2}月\d{1,2}日以留置送达方式向你(单位|公司|集团|厂)送达了|\d{4}年\d{1,2}月\d{1,2}日和(\d{4}年)*\d{1,2}月\d{1,2}日，我厅(分别)*向你(单位|公司|集团|厂)送达了)',
                    content_text).group(2).strip()
            else:
                punishment_basis = re.search(r'(你(单位|公司|集团|厂)的上述行为违反了[\s\S]*?的规定。)', content_text).group(1).strip()

            defense_response_text = re.search(
                r'((\n\d{4}年\d{1,2}月\d{1,2}日，我厅(分别)*向你(单位|公司|集团|厂)送达了|我厅于\d{4}年\d{1,2}月\d{1,2}日向你(单位|公司|集团|厂)送达了|我厅已于\d{4}年\d{1,2}月\d{1,2}日以留置送达方式向你(单位|公司|集团|厂)送达了|\d{4}年\d{1,2}月\d{1,2}日和(\d{4}年)*\d{1,2}月\d{1,2}日，我厅(分别)*向你(单位|公司|集团|厂)送达了)[\s\S]*?)(二、行政处罚的依据、种类及其履行方式、期限)',
                content_text).group(1).strip()

            if re.search(
                    r'根据《中华人民共和国行政处罚法》第三十二条和《环境行政处罚办法》第四十九条的规定，我厅依法进行了复核|我厅对你(单位|公司|集团|厂)的陈述申辩的事实、理由和证据进行了复核|我厅对你公司提出的陈述申辩理由进行了核实',
                    defense_response_text):
                # 申辩意见
                defense_opinion = re.search(
                    r'([\s\S]*)(根据《中华人民共和国行政处罚法》第三十二条和《环境行政处罚办法》第四十九条的规定，我厅依法进行了复核|我厅对你(单位|公司|集团|厂)的陈述申辩的事实、理由和证据进行了复核|我厅对你公司提出的陈述申辩理由进行了核实)',
                    defense_response_text).group(1).strip()

                # 申辩意见反馈
                defense_response = re.search(
                    r'((根据《中华人民共和国行政处罚法》第三十二条和《环境行政处罚办法》第四十九条的规定，我厅依法进行了复核|我厅对你(单位|公司|集团|厂)的陈述申辩的事实、理由和证据进行了复核|我厅对你公司提出的陈述申辩理由进行了核实)[\s\S]*)',
                    defense_response_text).group(1).strip()
            else:
                defense_opinion = defense_response_text
                defense_response = ''
            # 处罚决定
            punishment_decision = re.search(r'(二、行政处罚的依据、种类及其履行方式、期限)([\s\S]*?)(限你(单位|公司|集团|厂)自接到本处罚决定书*之日起)',
                                            content_text).group(2).strip()

        elif re.search(r'行政处罚情况|行政处罚信息', announcement_title):
            tr_list = content_soup.find(class_='content').find_all('tr')
            result_map_list = []
            del (tr_list[0])
            for each_tr in tr_list:
                td_list = each_tr.find_all('td')
                if len(td_list) == 5:
                    real_publish_date = format_date(get_content_text(td_list[4]))

                    announcement_code = get_content_text(td_list[0]).replace('\n', '')

                    # 当事人
                    litigant = get_content_text(td_list[1]).strip()

                    content_text = get_content_text(td_list[2])

                    # 违规事实
                    facts = re.search(r'([\s\S]*)(该(单位|公司|集团|厂)的上述行为违反了|的行为违反了《医疗废物管理条例》)', content_text).group(
                        1).strip()

                    # 认定意见
                    punishment_basis = re.search(r'((该(单位|公司|集团|厂)的上述行为违反了|违反了《医疗废物管理条例》)[\s\S]*?的规定。)',
                                                 content_text).group(1).strip()

                    defense_opinion = ''

                    defense_response = ''

                    # 处罚决定
                    punishment_decision = re.search(r'的规定。([\s\S]*)', content_text).group(1).strip()

                else:
                    real_publish_date = format_date(get_content_text(td_list[5]))

                    announcement_code = get_content_text(td_list[3]).replace('\n', '')

                    # 当事人
                    litigant = get_content_text(td_list[1]).strip()

                    content_text = get_content_text(td_list[2])

                    # 违规事实
                    facts = re.search(r'主要违法事实：([\s\S]*?)(该(单位|公司|集团|厂)的上述行为违反了)', content_text).group(1).strip()

                    # 认定意见
                    punishment_basis = re.search(r'(该(单位|公司|集团|厂)的上述行为违反了[\s\S]*?)处罚种类：', content_text).group(1).strip()

                    defense_opinion = ''

                    defense_response = ''

                    # 处罚决定
                    punishment_decision = re.search(r'(处罚依据：[\s\S]*)', content_text).group(1).strip()

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('广东生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('广东生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('广东生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue
        # http://pub.gdepb.gov.cn/pub/pubcatalog/extranet_pub_document_view.jsp?docId=133500
        elif re.search(r'行政应诉情况', announcement_title):
            tr_list = content_soup.find(class_='content').find_all('tr')
            result_map_list = []
            del (tr_list[0])
            for each_tr in tr_list:
                td_list = each_tr.find_all('td')
                if len(td_list) == 6:
                    real_publish_date = format_date(
                        re.search(r'\d{4}年\d{1,2}月\d{1,2}日', get_content_text(td_list[4])).group(0).strip())

                    announcement_code = ''

                    # 当事人
                    litigant = get_content_text(td_list[1]).strip()

                    # 违规事实
                    facts = get_content_text(td_list[0]) + '\n' + get_content_text(td_list[2])

                    # 认定意见
                    punishment_basis = ''

                    defense_opinion = ''

                    defense_response = ''

                    # 处罚决定
                    punishment_decision = get_content_text(td_list[3])

                    announcement_org = get_content_text(td_list[5])

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('广东生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('广东生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('广东生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue
        # http://pub.gdepb.gov.cn/pub/pubcatalog/extranet_pub_document_view.jsp?docId=125203
        elif re.search(r'行政复议情况|行政应诉信息|行政复议信息', announcement_title):
            tr_list = content_soup.find(class_='content').find_all('tr')
            result_map_list = []
            del (tr_list[0])
            for each_tr in tr_list:
                td_list = each_tr.find_all('td')
                if len(td_list) == 6:
                    try:
                        real_publish_date = format_date(
                            re.search(r'\d{4}年\d{1,2}月\d{1,2}日', get_content_text(td_list[4])).group(0).strip())
                    except Exception as e:
                        real_publish_date = format_date(each_document['publishDate'])
                        logger.info(e)

                    announcement_code = ''

                    # 当事人
                    litigant = get_content_text(td_list[0]).strip()

                    # 违规事实
                    facts = get_content_text(td_list[1]) + '\n' + get_content_text(td_list[2])

                    # 认定意见
                    punishment_basis = ''

                    defense_opinion = ''

                    defense_response = ''

                    # 处罚决定
                    punishment_decision = get_content_text(td_list[3])

                    announcement_org = get_content_text(td_list[5])

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('广东生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('广东生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('广东生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('广东生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('广东生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('广东生态环境厅 数据解析 ' + ' -- 修改parsed完成')


# 湖北生态环境厅
def hubei():
    for each_document in db.environment_data.find({'origin': '湖北省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        # ignored
        if re.search(r'环保法配套办法执行情况汇总表|环保法配套办法执行情况的通报|曝光：2019年1月份部分大气环境污染问题|' \
                     r'仙桃警方破获部督污染环境案 涉案金额千万元的“炼铅黑窝点”被端|环保亮剑，不避重点工程|阳新严肃问责环保不作为慢作为|'\
                     r'湖北武汉巡司河洪山区段黑臭现象依然严重|十堰查处破坏环境资源案205件|纳污坑塘排查整治工作台账和整治方案|'\
                r'环境空气自动监测站点周边喷雾作业有关情况的通报|饮用水水源地主要环境问题清单及整改方案|黄陂两电镀作坊严重污染环境|'\
                r'湖北省环境执法情况汇总表|焚烧电子垃圾牟利，3名男子当阳获刑|安陆城区垃圾处理场负责人被拘|“躲猫猫”被逮个正着|'\
                r'东西湖一工业园超标排烟被查处|鄂州一老板被拘留|沙洋侦破一起污染环境案|一经营户偷排重金属废水被刑拘|'\
                r'枝江一污水处理厂废水超标入长江 厂长被拘留|环保法配套办法执行情况|新《环保法》实施元年 我省按日连续处罚案件23件|'\
                r'远安侦破首例污染环境案件|公司法人被行拘|武汉开出最大环保罚单|黄石一公司员工被判刑|我省印发《湖北省实行网格化环境监管体系实施方案（试行）》|'\
                r'推进新《环保法》实施|多部门联合执法依法查封|黄石查处环保违法企业100家|十堰查处环境违法案51件|新《环保法》实施|'\
                r'潜江市对非法收集储存危险废物案件进行查处|湖北加大企业新《环保法》培训|十堰通报10起环境违法典型案例|环保公安齐亮剑 环境违法被制裁|'\
                r'襄阳市采取限产限排等强硬措施|湖北约谈治气不力8个地市|下发三次督办通知视若无睹|荆门一企业违规排污致8人入院|'\
                r'恩施一养鸡场万吨鸡粪直排天坑|新环保法及五个配套办法实施的有关情况|黄石市查破一起非法买卖、储存危险废物案|'\
                r'麻城市排污收费违反规定|省环保厅“向污染宣战、实施环境保护三大行动”环境违法案件曝光台|2013年—2014年5月全省环境违法案件处理情况|'\
                r'关于印发《环境违法案件挂牌督办管理办法》的通知|省环保局、省监察厅关于挂牌督办环境污染问题的通知', announcement_title):
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        # 如果是.xls文件：
        if re.search(r'全省双护双促砖瓦行业统计表', announcement_title):
            xlsx_link = urljoin(announcement_url, content_soup.find('ul', id = 'list').find('a')['href'])
            link_type = xlsx_link.split('.')[-1]
            response = request_site_page(xlsx_link)
            if response is None:
                logger.error('xlsx文件下载失败%s' % xlsx_link)
                return
            with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        tmp_file.write(chunk)

            if not os.path.exists('./test/tmp.xlsx'):
                shell_str = 'soffice --headless --convert-to xlsx ' + \
                            './test/tmp.' + link_type + ' --outdir ./test'
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.communicate()

            with open('./test/tmp.xlsx', 'rb') as xlsx_file:
                xlsx_content = xlsx_file.read()
            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': xlsx_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': xlsx_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'xlsx',
                    'oss_file_name': announcement_title,
                    'oss_file_content': xlsx_content,
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.xlsx', xlsx_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': xlsx_link})[
                    '_id']
            logger.info('存入parsed_data')

            excel_data = xlrd.open_workbook('./test/tmp.xlsx')
            sheet = excel_data.sheets()[0]
            result_map_list = []
            for i in range(2, sheet.nrows):
                if sheet.ncols == 27:
                    # 文号
                    announcement_code = ''

                    # 当事人
                    litigant = sheet.cell(i, 1).value + ' ' + sheet.cell(i, 2).value + ' ' + sheet.cell(i, 3).value + ' ' + sheet.cell(i, 4).value

                    # 违规事实
                    facts = sheet.cell(i, 20).value

                    # 认定意见
                    punishment_basis = ''

                    # 申辩意见
                    defenseOpinion = ''

                    # 申辩意见反馈
                    defenseResponse = ''

                    # 处罚决定
                    punishment_decision = sheet.cell(i, 21).value

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': '湖北生态环境厅',
                        'announcementDate': format_date(each_document['publishDate'].split(' ')[0]),
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': defenseOpinion,
                        'defenseResponse': defenseResponse,
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')

            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.xlsx'):
                os.remove('./test/tmp.xlsx')
            if os.path.exists('./test/tmp.xls'):
                os.remove('./test/tmp.xls')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        #如果是doc文档
        if re.search(r'关于取缔“十小”企业情况的公示', announcement_title):
            docx_link = urljoin(announcement_url, content_soup.find(class_ = 'display_doc').find('ul', id = 'list').find('a')['href'])
            link_type = docx_link.split('.')[-1]
            response = request_site_page(docx_link)
            if response is None:
                logger.error('doc文件下载失败%s' % docx_link)
                return
            with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        tmp_file.write(chunk)
            if not os.path.exists('./test/tmp.docx'):
                shell_str = '/usr/local/bin/soffice --headless --convert-to docx ' + \
                            './test/tmp.doc' + ' --outdir ./test'
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.communicate()

            with open('./test/tmp.docx', 'rb') as docx_file:
                docx_content = docx_file.read()
            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': docx_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': docx_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'docx',
                    'oss_file_name': announcement_title,
                    'oss_file_content': docx_content,
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url})['_id']

            document = Document('./test/tmp.docx')
            tables = document.tables
            table = tables[0]
            result_map_list = []
            for i in range(2, len(table.rows)):
                # 文号
                announcement_code = ''

                # 当事人
                litigant = table.cell(i, 3).text + ' ' + table.cell(i, 1).text

                # 违规事实
                facts = table.cell(i, 5).text + ' 取缔完成时间:' + table.cell(i, 4).text

                # 认定意见
                punishment_basis = ''

                # 申辩意见
                defenseOpinion = ''

                # 申辩意见反馈
                defenseResponse = ''

                # 处罚决定
                punishment_decision = ''

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': '湖北生态环境厅',
                    'announcementDate': format_date(each_document['publishDate'].split(' ')[0]),
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defenseOpinion,
                    'defenseResponse': defenseResponse,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')

            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.docx'):
                os.remove('./test/tmp.docx')
            if os.path.exists('./test/tmp.doc'):
                os.remove('./test/tmp.doc')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        #图片表格
        if re.search(r'环境保护部通报2015年3月份大气污染防治督查情况|关于湖北省2013年挂牌督办环境违法企业的通报', announcement_title):
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                img_link_list = [urljoin(announcement_url, kk.attrs['src'])
                                 for kk in content_soup.find('div', class_='display_article').find_all('img')]
                link_type = '.' + img_link_list[0].split('.')[-1]
                for index, each_img_link in enumerate(img_link_list):
                    img_response = request_site_page(each_img_link)
                    with open('./test/' + str(index) + link_type, 'wb') as tmp_file:
                        for chunk in img_response.iter_content(chunk_size=1024):
                            if chunk:
                                tmp_file.write(chunk)

                if not os.path.exists('./test/tmp.pdf'):
                    shell_str = 'img2pdf '
                    for index in range(len(img_link_list)):
                        shell_str += './test/' + str(index) + link_type + ' '
                    shell_str += '--imgsize 20cmx30cm  -o ./test/tmp.pdf'  # 加入图片尺寸参数，百度ocr有图片尺寸限制
                    process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                               shell=True, stdout=subprocess.PIPE)
                    process.communicate()

                result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')
                content_text = result_text

                with open('./test/tmp.pdf', 'rb') as pdf_file:
                    pdf_content = pdf_file.read()

                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False,
                    'if_ocr': True,
                    'ocr_result': result_text
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']
                content_text = db.parsed_data.find_one({'origin_url': announcement_url,
                                                    'oss_file_origin_url': announcement_url})['ocr_result']
            logger.info('存入parsed_data')
            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.pdf'):
                os.remove('./test/tmp.pdf')
            for index in range(len(img_link_list)):
                if os.path.exists('./test/' + str(index) + link_type):
                    os.remove('./test/' + str(index) + link_type)
            continue

        content_text = get_content_text(content_soup.find(class_='display_article'))

        # 如果公告是在网页里
        if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'content_class_name': 'display_article',
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']
        logger.info('存入parsed_data')

        # 需人工处理
        if announcement_url in ['http://sthjt.hubei.gov.cn/xwzx/hjxw/201805/t20180531_113536.shtml',
                                'http://sthjt.hubei.gov.cn/xxgk/xxgkml/hjjc/zdhjwf/201804/t20180402_112193.shtml',
                                'http://sthjt.hubei.gov.cn/xxgk/xxgkml/hjjc/zdhjwf/201804/t20180402_112194.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201711/t20171101_109081.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201801/t20180118_111064.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201706/t20170612_105281.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201706/t20170605_105047.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201706/t20170605_105048.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201701/t20170110_101729.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201701/t20170106_101634.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201701/t20170106_101633.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201611/t20161124_99992.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201610/t20161031_99290.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201609/t20160914_98144.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201608/t20160815_97064.shtml',
                                'http://sthjt.hubei.gov.cn:8080/pub/root8/tjgzs/hjjc/zdzhxx/201607/t20160728_96710.html',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201607/t20160720_96530.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201606/t20160606_85783.shtml',
                                'http://sthjt.hubei.gov.cn/xxgk/xxgkml/hjjc/zdhjwf/201603/t20160328_84121.shtml',
                                'http://sthjt.hubei.gov.cn/xxgk/xxgkml/hjjc/zdhjwf/201603/t20160314_83774.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201603/t20160314_83772.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201603/t20160302_83497.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201602/t20160205_83128.shtml',
                                'http://sthjt.hubei.gov.cn/xxgk/xxgkml/hjjc/zdhjwf/201602/t20160203_83090.shtml',
                                'http://sthjt.hubei.gov.cn/xxgk/xxgkml/hjjc/zdhjwf/201601/t20160119_82819.shtml',
                                'http://sthjt.hubei.gov.cn/xxgk/xxgkml/hjjc/zdhjwf/201509/t20150909_79844.shtml',
                                'http://sthjt.hubei.gov.cn:8080/pub/root8/tjgzs/hjjc/zdzhxx/201507/t20150706_78403.html',
                                'http://sthjt.hubei.gov.cn/xxgk/xxgkml/hjjc/zdhjwf/201505/t20150526_77466.shtml',
                                'http://sthjt.hubei.gov.cn/xxgk/xxgkml/hjjc/zdhjwf/201505/t20150526_77465.shtml',
                                'http://sthjt.hubei.gov.cn:8080/pub/root8/tjgzs/hjjc/zdzhxx/201502/t20150215_75460.html',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201412/t20141208_74149.shtml',
                                'http://sthjt.hubei.gov.cn/xwzx/hjxw/201409/t20140922_72331.shtml',
                                'http://sthjt.hubei.gov.cn:8080/pub/root8/tjgzs/hjjc/zdzhxx/201409/t20140905_72003.html']:
            logger.info('需人工处理')
            continue

        # 开始解析
        # 处罚机构
        announcement_org = '湖北生态环境厅'

        if re.search(r'处罚决定书', announcement_title):
            # 处罚日期
            try:
                publish_date = re.search(r'(湖北省环境保护厅)\n(\d{4}年\d{1,2}月\d{1,2}日)', content_text).group(2).strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.info(e)
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)
            # 文号
            announcement_code = re.search(r'\n(鄂环[\s\S]*?号)\n', content_text).group(1).replace('\n', '').strip()
            if re.search(r'一、调查情况及环境违法事实、证据和陈述申辩（听证）采纳情况|一、调查情况及环境违法事实、证据和陈述申辩（听证）情况' \
                         r'|一、调查情况及发现的环境违法事实、证据和陈述申辩（听证）及采纳情况|一、环境违法事实和证据',
                         content_text):
                corp = re.search(r'（([\s\S]*)）', announcement_title).group(1).strip()
                if re.search(r'(鄂环[\s\S]*?号)', corp):
                    corp = re.search(r'关于([\s\S]*?)环境违法行为', announcement_title).group(1).strip()

                # 当事人
                litigant = re.search(r'处罚决定书\n([\s\S]*?)\n(针对你(公司|单位|集团|厂)环境违法一案|你作为湖北华尔靓科技' \
                                     r'|你(公司|单位|集团|厂)环境违法案|武汉壮盛生物科技有限公司|' + corp + ')',
                                     content_text).group(1).strip()

                # 违规事实
                facts = re.search(
                    r'(一、调查情况及环境违法事实、证据和陈述申辩（听证）采纳情况|一、调查情况及环境违法事实、证据和陈述申辩（听证）情况' \
                    r'|一、调查情况及发现的环境违法事实、证据和陈述申辩（听证）及采纳情况|一、环境违法事实和证据)([\s\S]*?)(\n以上事实，*有|\n上述事实有)',
                    content_text).group(2).strip()
                # 认定意见
                punishment_basis = re.search(
                    r'(等证据为凭。|等为证。|为证。|可以认定。)\n([\s\S]*?(的规定。|应当承担(相应)*的*法律责任。|依法应当予以处罚。|依法应当予以按日连续处罚。))',
                    content_text).group(
                    2).strip()

                if re.search(
                        r'((我[厅局]于\n*\d{4}年\d{1,2}月\d{1,2}日向你(公司|单位|集团|厂)*送达了*(《湖北省环境保护厅行政处罚事先（听证）告知书》|《湖北省环境保护厅行政处罚事先告知书及听证告知书》|限期改正通知书)' \
                        r'|我厅于\d{4}年\d{1,2}月\d{1,2}日委托老河口市环境保护局向你(公司|单位|集团|厂)送达了《湖北省环境保护厅行政处罚事先（听证）告知书》' \
                        r'|我厅于\d{4}年\d{1,2}月\d{1,2}日以《行政处罚事先（听证）告知书》|我厅于\d{4}年\d{1,2}月\d{1,2}日告知你(公司|单位|集团|厂)违法事实' \
                        r'|我厅于\d{4}年\d{1,2}月\d{1,2}日责令你(公司|单位|集团|厂)立即停止违法排污行为)[\s\S]*?)(以上事实有《行政处罚事先（听证）告知书》' \
                        r'|以上事实，有我厅《湖北省环境保护厅行政处罚事先（听证）告知书》' \
                        r'|以上事实，有我厅《行政处罚事先（听证）告知书》' \
                        r'|以上事实，有我厅\d{4}年\d{1,2}月\d{1,2}日《湖北省环境保护厅行政处罚事先告知和听证告知书》' \
                        r'|以上事实，有我厅\d{4}年\d{1,2}月\d{1,2}日《责令改正违法行为决定书》'\
                        r'|以上事实，有我局\n*\d{4}年\d{1,2}月\d{1,2}日限期改正通知书)', content_text):
                    defense_response_text = re.search(
                        r'((我[厅局]于\n*\d{4}年\d{1,2}月\d{1,2}日向你(公司|单位|集团|厂)*送达了*(《湖北省环境保护厅行政处罚事先（听证）告知书》|《湖北省环境保护厅行政处罚事先告知书及听证告知书》|限期改正通知书)' \
                        r'|我厅于\d{4}年\d{1,2}月\d{1,2}日委托老河口市环境保护局向你(公司|单位|集团|厂)送达了《湖北省环境保护厅行政处罚事先（听证）告知书》' \
                        r'|我厅于\d{4}年\d{1,2}月\d{1,2}日以《行政处罚事先（听证）告知书》|我厅于\d{4}年\d{1,2}月\d{1,2}日告知你(公司|单位|集团|厂)违法事实' \
                        r'|我厅于\d{4}年\d{1,2}月\d{1,2}日责令你(公司|单位|集团|厂)立即停止违法排污行为)[\s\S]*?)(以上事实有《行政处罚事先（听证）告知书》' \
                        r'|以上事实，有我厅《湖北省环境保护厅行政处罚事先（听证）告知书》' \
                        r'|以上事实，有我厅《行政处罚事先（听证）告知书》' \
                        r'|以上事实，有我厅\d{4}年\d{1,2}月\d{1,2}日《湖北省环境保护厅行政处罚事先告知和听证告知书》' \
                        r'|以上事实，有我厅\d{4}年\d{1,2}月\d{1,2}日《责令改正违法行为决定书》'\
                        r'|以上事实，有我局\n*\d{4}年\d{1,2}月\d{1,2}日限期改正通知书)', content_text).group(1).strip()
                else:
                    defense_response_text = re.search(
                        r'((我厅于\d{4}年\d{1,2}月\d{1,2}日向你(公司|单位|集团|厂)*送达了*(《湖北省环境保护厅行政处罚事先（听证）告知书》|《湖北省环境保护厅行政处罚事先告知书及听证告知书》)' \
                        r'|我厅于\d{4}年\d{1,2}月\d{1,2}日委托老河口市环境保护局向你(公司|单位|集团|厂)送达了《湖北省环境保护厅行政处罚事先（听证）告知书》' \
                        r'|我厅于\d{4}年\d{1,2}月\d{1,2}日以《行政处罚事先（听证）告知书》)[\s\S]*?)' \
                        r'(二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类)', content_text).group(1).strip()

                if re.search(r'(根据《中华人民共和国行政处罚法》第三十二条和《环境行政处罚办法》第四十九条的规定|经研究，' \
                             r'|我厅对上述陈述意见进行了核实|因你公司听证申请超出法定期限，我厅未予受理|我厅对上述申辩意见进行了核实' \
                             r'|我厅对上述陈述和申辩意见进行了核实)', defense_response_text):
                    defense_opinion = re.search(r'([\s\S]*)(根据《中华人民共和国行政处罚法》第三十二条和《环境行政处罚办法》第四十九条的规定' \
                                                r'|经研究，|我厅对上述陈述意见进行了核实|因你公司听证申请超出法定期限，我厅未予受理' \
                                                r'|我厅对上述申辩意见进行了核实|我厅对上述陈述和申辩意见进行了核实)', defense_response_text).group(
                        1).strip()
                    defense_response = re.search(r'((根据《中华人民共和国行政处罚法》第三十二条和《环境行政处罚办法》第四十九条的规定|经研究，' \
                                                 r'|我厅对上述陈述意见进行了核实|因你公司听证申请超出法定期限，我厅未予受理' \
                                                 r'|我厅对上述申辩意见进行了核实|我厅对上述陈述和申辩意见进行了核实)[\s\S]*)',
                                                 defense_response_text).group(1).strip()
                else:
                    defense_opinion = defense_response_text
                    defense_response = ''
                # 处罚决定
                punishment_decision = re.search(r'(二、行政处罚的依据、种类|二、责令改正和行政处罚的依据、种类|二、行政处罚的依据、种类及履行方式和期限)' \
                                                r'\n([\s\S]*?)\n(三、行政处罚决定的履行方式、期限|三、行政处罚决定的履行方式和期限|三、责令改正和行政处罚决定的履行方式和期限' \
                                                r'|三、申请复议或者提起诉讼的途径和期限|四、申请复议或者提起诉讼的途径和期限)', content_text).group(
                    2).strip()
                if re.search(r'(根据《中华人民共和国行政处罚法》和《罚款决定与罚款收缴分离实施办法》的规定)', punishment_decision):
                    punishment_decision = re.search(r'([\s\S]*)(根据《中华人民共和国行政处罚法》和《罚款决定与罚款收缴分离实施办法》的规定)',
                                                    punishment_decision).group(1).strip()
            else:
                # 当事人
                litigant = re.search(r'处罚决定书\n([\s\S]*?)\n我厅于', content_text).group(1).strip()

                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                            '\]') + r'([\s\S]*)(以上事实有你公司营业执照（复印件）)',
                                  content_text).group(1).strip()
                # 认定意见
                punishment_basis = re.search(r'(等为证。)\n([\s\S]*?的规定。)', content_text).group(2).strip()

                defense_response_text = re.search(
                    r'(我厅于\d{4}年\d{1,2}月\d{1,2}日以《湖北省环境保护厅行政处罚事先（听证）告知书》[\s\S]*?)\n依据《中华人民共和国行政处罚法》',
                    content_text).group(
                    1).strip()

                if re.search(r'(\d{4}年\d{1,2}月\d{1,2}日，我厅依法组织了听证)', defense_response_text):
                    defense_opinion = re.search(r'([\s\S]*)(\d{4}年\d{1,2}月\d{1,2}日，我厅依法组织了听证)',
                                                defense_response_text).group(1).strip()
                    defense_response = re.search(r'((\d{4}年\d{1,2}月\d{1,2}日，我厅依法组织了听证)[\s\S]*)',
                                                 defense_response_text).group(1).strip()
                else:
                    defense_opinion = defense_response_text
                    defense_response = ''
                # 处罚决定
                punishment_decision = re.search(
                    r'\n(依据《中华人民共和国行政处罚法》第二十三条[\s\S]*?)(根据《中华人民共和国行政处罚法》和《罚款决定与罚款收缴分离实施办法》的规定)',
                    content_text).group(1).strip()

        elif re.search(r'督办通知', announcement_title):
            # 处罚日期
            try:
                publish_date = re.findall(r'\n(\d{4}年\d{1,3}月\d{1,3}日)', content_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.info(e)
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)
            # 文号
            announcement_code = re.search(r'(鄂环[\s\S]*?号)', announcement_title).group(1).strip()

            # 当事人
            litigant = re.search(r'督办通知\n([\s\S]*?)：\n', content_text).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]') + r'([\s\S]*?)(\n请你局|根据现场调查的情况，|现请你局)', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            # 处罚决定
            punishment_decision = re.search(r'((\n请你局|根据现场调查的情况，|现请你局)[\s\S]*?)\n(联 系 人：|联系人：)', content_text).group(1).strip()

        elif re.search(r'关于解除.+挂牌督办的公示', announcement_title):
            # 处罚日期
            try:
                publish_date = re.findall(r'\n(\d{4}年\d{1,3}月\d{1,3}日)', content_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.info(e)
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)
            # 文号
            announcement_code = ''

            # 当事人
            litigant = re.search(r'关于解除([\s\S]*?)挂牌督办的公示', announcement_title).group(1).strip()

            # 违规事实
            facts = re.search(r'挂牌督办的公示([\s\S]*?)(省环保厅认为)', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            # 处罚决定
            punishment_decision = re.search(r'((省环保厅认为)[\s\S]*?)\n为体现公开、公平、公正', content_text).group(1).strip()

        elif re.search(r'挂牌督办的公示', announcement_title):
            # 处罚日期
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

            # 文号
            announcement_code = ''

            # 当事人
            litigant = re.search(r'关于对([\s\S]*?)环境问题实施挂牌督办的公示', announcement_title).group(1).strip()

            # 违规事实
            facts = re.search(r'挂牌督办的公示([\s\S]*?)(为落实省委、省政府)', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            # 处罚决定
            punishment_decision = re.search(r'((为落实省委、省政府)[\s\S]*)', content_text).group(1).strip()

        elif re.search(r'配套办法实施典型案例', announcement_title):
            content_text = content_text.replace('湖北日报讯案例一', '案例一')
            content_text_list = content_text.split('\n案例')
            del(content_text_list[0])
            result_map_list = []
            for each_text in content_text_list:
                # 处罚日期
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

                # 文号
                announcement_code = ''

                # 当事人
                litigant = re.search(r'[：:]([\s\S]*?(公司|集团|单位|厂|局))', each_text).group(1).strip()

                # 违规事实
                facts = re.search(r'((\d{4}年\d{1,2}月|省环保厅会同黄冈市环保局|湖北省环境监察总队接到群众实名举报|潜江市环保局执法人员|2015年第1季度监督性监测结果显示)'\
                                  r'[\s\S]*?(仍然超标排放。|仍在超标排放水污染物。|不符合国家产业政策的小型电镀企业。|'\
                                  r'生产废水直接排入附近自然水体内。|举报内容属实。|大门钥匙进行扣押。|总磷浓度均超标。|应列入关停范围。|过程中违法排放水污染物。|'\
                                  r'仍在建设。|在线监控数据均显示达标。|的环境违法行为。|非法排放污染物。|色度严重超标。|监控设施参数的违法事实。|'\
                                  r'处置相关资质。|超过国家规定的排放限值。|确定该公司二氧化硫超标排放。|超过国家规定的水污染排放标准50%以上。|'\
                                  r'未采取有效防治设施违规排放大气污染物。|王店断面水质监测异常超标。|以逃避环保部门的监管。|涉嫌环境污染犯罪。|'\
                                  r'超过国家规定的排放标准10倍以上。|移送谷城县人民检察院提起公诉。|同时抄送天门市人民检察院。|废气超标排放。|'\
                                  r'仍在继续违法排放污染物。|漫流进入三湖连江。|生产废水排向长江蕲州段水体。|挥发酚超标161.8倍。|氨氮超标66.8倍。|硫磺渣渗滤液直排长江。|'\
                                  r'4.4倍和11倍。|排放废水的违法事实。|除尘设施闲置等。|移送至鄂州市公安局。|移送至潜江市公安局。|0.85倍、0.22倍。|汇入自然水体杨家新沟。|'\
                                  r'将此案移送仙桃市公安局。|运输设备进行了查封、扣押。|401倍、6.8倍。|超过国家规定的排放标准和限值。|直接排放至外界水环境。))', each_text).group(1).strip()
                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                # 处罚决定
                punishment_decision = re.search(facts.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]') + r'([\s\S]*)', each_text).group(1).strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(r'武汉市公布环保违法十大典型案例', announcement_title):
            content_text_list = content_text.split('\n案例')
            del(content_text_list[0])
            result_map_list = []
            for each_text in content_text_list:
                # 处罚日期
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

                # 文号
                announcement_code = ''

                # 当事人
                litigant = re.search(r'[：:]([\s\S]*?(公司|集团|单位|厂|局))', each_text).group(1).strip()
                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                if re.search(r'(被重罚|按日连续处罚|案)([\s\S]*?)((市环保局|市环保局决定)*责令|依照|市环保局对)', each_text):
                    # 违规事实
                    facts = re.search(r'(被重罚|按日连续处罚|案)([\s\S]*?)((市环保局|市环保局决定)*责令|依照|市环保局对)', each_text).group(2).strip()

                    # 处罚决定
                    punishment_decision = re.search(facts.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]') + r'([\s\S]*)', each_text).group(1).strip()
                else:
                    facts = re.search(r'(被重罚|按日连续处罚|案)([\s\S]*)', each_text).group(2).strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(r'恩施州立案查处违法案件24件', announcement_title):
            content_text_list = content_text.split('\n案例')
            del(content_text_list[0])
            result_map_list = []
            for each_text in content_text_list:
                # 处罚日期
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

                # 文号
                announcement_code = ''

                # 当事人
                litigant = ''
                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                # 违规事实
                facts = re.search(r'案由：([\s\S]*?)(处理结果)', each_text).group(1).strip()

                # 处罚决定
                punishment_decision = re.search(r'(处理结果[\s\S]*)', each_text).group(1).strip()

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(r'新环保法实施两月成绩单 限产停产8家查封扣押3家行政拘留4人', announcement_title):
            content_text_list = content_text.split('\n案例')
            del(content_text_list[0])
            result_map_list = []
            for each_text in content_text_list:
                # 处罚日期
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

                # 文号
                announcement_code = ''

                # 当事人
                litigant = ''
                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                # 违规事实
                facts = re.search(r'：([\s\S]*)', each_text).group(1).strip()

                # 处罚决定
                punishment_decision = ''

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(r'关于2011年环境保护部挂牌督办环境违法案件的通知|环境保护部挂牌督办10起重金属排放环境问题', announcement_title):
            # 文号
            try:
                announcement_code = re.search(r'(环办[\s\S]*号)', announcement_title).group(1).strip()
            except Exception as e:
                announcement_code = ''
            # 处罚日期
            try:
                publish_date = re.findall(r'\n(\d{4}年\d{1,3}月\d{1,3}日)', content_text)[-1].strip()
                real_publish_date = format_date(publish_date)
            except Exception as e:
                logger.info(e)
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

            content_text = re.search(r'(\n环境违法案件基本情况及督办要求[\s\S]*)', content_text).group(1).strip()

            content_text_list = content_text.split('环境违法案\n')
            result_map_list = []
            for each_text in content_text_list[1:]:
                index = content_text_list.index(each_text)

                litigant = re.search(r'\n.{1,2}、([\s\S]*)', content_text_list[index-1]).group(1).strip()

                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                # 违规事实
                facts = re.search(r'(基本情况：[\s\S]*)督办要求：', each_text).group(1).strip()

                if index == len(content_text_list) - 1:
                    # 处罚决定
                    punishment_decision = re.search(r'(督办要求：[\s\S]*)', each_text).group(1).strip()
                else:
                    # 处罚决定
                    punishment_decision = re.search(r'(督办要求：[\s\S]*?)\n.{1,2}、', each_text).group(1).strip()

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue
        elif re.search(r'尚未完成整改的突出环境问题清单|突出环境问题整改完成清单|湖北省挂牌督办（未解除）环境违法案件情况', announcement_title):
            tr_list = content_soup.find(class_ = 'display_article').find('table').find_all('tr')
            del(tr_list[0])
            result_map_list = []
            for each_tr in tr_list:
                td_list = each_tr.find_all('td')
                if re.search(r'尚未完成整改的突出环境问题清单', announcement_title):
                    # 处罚日期
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)

                    # 文号
                    announcement_code = ''

                    # 当事人
                    litigant = get_content_text(td_list[2])

                    # 违规事实
                    facts = get_content_text(td_list[4])

                    # 认定意见
                    punishment_basis = ''

                    defense_opinion = ''
                    defense_response = ''

                    # 处罚决定
                    punishment_decision = get_content_text(td_list[5]) + get_content_text(td_list[6])
                elif re.search(r'突出环境问题整改完成清单', announcement_title):
                    # 处罚日期
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)

                    # 文号
                    announcement_code = ''

                    # 当事人
                    litigant = get_content_text(td_list[2])

                    # 违规事实
                    facts = get_content_text(td_list[3])

                    # 认定意见
                    punishment_basis = ''

                    defense_opinion = ''
                    defense_response = ''

                    # 处罚决定
                    punishment_decision = '处理意见和要求:' + get_content_text(td_list[4]) + '\n整改完成情况:' + get_content_text(td_list[6])
                elif re.search(r'湖北省挂牌督办（未解除）环境违法案件情况', announcement_title):
                    # 处罚日期
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)

                    # 文号
                    announcement_code = ''

                    # 当事人
                    litigant = get_content_text(td_list[2])

                    # 违规事实
                    facts = get_content_text(td_list[3])

                    # 认定意见
                    punishment_basis = ''

                    defense_opinion = ''
                    defense_response = ''

                    # 处罚决定
                    punishment_decision = '督办要求:' + get_content_text(td_list[4]) + ' 整改进展情况:' + get_content_text(td_list[5])
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(r'核实反馈表', announcement_title):
            tables = content_soup.find(class_ = 'display_article').find_all('table')
            result_map_list = []
            for each_table in tables:
                tr_list = each_table.find_all('tr')
                del(tr_list[0])
                for each_tr in tr_list:
                    td_list = each_tr.find_all('td')
                    # 处罚日期
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)

                    announcement_org = get_content_text(td_list[5])

                    # 文号
                    announcement_code = get_content_text(td_list[6])

                    # 当事人
                    litigant = get_content_text(td_list[1])

                    # 认定意见
                    punishment_basis = ''

                    defense_opinion = ''
                    defense_response = ''

                    facts = '超标率:' + get_content_text(td_list[2]) + ' 上一季度是否严重超标:' + get_content_text(td_list[3]) + ' 本年度严重超标次数:' + get_content_text(td_list[4])

                    # 处罚决定
                    punishment_decision = '计划整改完成时间:' + get_content_text(td_list[8]) + ' 具体整改措施:' + get_content_text(td_list[9])

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': defense_opinion,
                        'defenseResponse': defense_response,
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(r'重点环境违法案件查处情况', announcement_title):
            tables = content_soup.find(class_ = 'display_article').find_all('table')
            result_map_list = []
            for each_table in tables:
                tr_list = each_table.find_all('tr')
                del(tr_list[0])
                for each_tr in tr_list:
                    td_list = each_tr.find_all('td')
                    # 处罚日期
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)

                    # 文号
                    announcement_code = ''

                    # 当事人
                    litigant = ''

                    # 认定意见
                    punishment_basis = ''

                    defense_opinion = ''
                    defense_response = ''

                    facts = '案发地:' + get_content_text(td_list[1]) + ' 简要案情:' + get_content_text(td_list[2])

                    # 处罚决定
                    punishment_decision = get_content_text(td_list[3])

                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': real_publish_date,
                        'announcementCode': announcement_code,
                        'facts': facts,
                        'defenseOpinion': defense_opinion,
                        'defenseResponse': defense_response,
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        elif re.search(r'关于对[\s\S]*的通报', announcement_title):
            if re.search(r'一、现场检查情况', content_text):
                # 处罚日期
                try:
                    publish_date = re.search(r'(湖北省环境保护厅办公室)\n(\d{4}年\d{1,2}月\d{1,2}日)', content_text).group(2).strip()
                    real_publish_date = format_date(publish_date)
                except Exception as e:
                    logger.info(e)
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)
                announcement_code = ''
                # 当事人
                litigant = re.search(r'的通报\n([\s\S]*?)：\n', content_text).group(1).strip()

                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                 '\]') + r'([\s\S]*?)(三、有关要求)', content_text).group(1).strip()
                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''
                # 处罚决定
                punishment_decision = re.search(r'三、有关要求([\s\S]*?)(湖北省环境保护厅)', content_text).group(1).strip()

            else:
                # 处罚日期
                try:
                    publish_date = re.search(r'(湖北省环境保护厅办公室)\n(\d{4}年\d{1,2}月\d{1,2}日)', content_text).group(2).strip()
                    real_publish_date = format_date(publish_date)
                except Exception as e:
                    logger.info(e)
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)
                announcement_code = ''
                # 当事人
                litigant = re.search(r'的通报\n([\s\S]*?)：\n', content_text).group(1).strip()

                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                            '\]') + r'([\s\S]*?)(鉴于以上情况)',
                                  content_text).group(1).strip()
                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''
                # 处罚决定
                punishment_decision = re.search(r'(鉴于以上情况[\s\S]*?)(湖北省环境保护厅)', content_text).group(1).strip()

        elif re.search(r'(关于)([\s\S]*?)(环境犯罪案件处理情况|环境违法案件处理情况)', announcement_title):
            # 处罚日期
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

            announcement_code = ''

            # 当事人
            litigant = re.search(r'(关于)([\s\S]*?)(环境犯罪案件处理情况|环境违法案件处理情况)', announcement_title).group(2).strip()

            if re.search(r'((该公司上述行为违反了)[\s\S]*?的规定。)', content_text):
                # 违规事实
                facts = re.search(r'([\s\S]*)(该公司上述行为违反了)', content_text).group(1).strip()

                # 认定意见
                punishment_basis = re.search(r'((该公司上述行为违反了)[\s\S]*?的规定。)', content_text).group(1).strip()

                # 处罚决定
                punishment_decision = re.search(r'的规定。([\s\S]*)', content_text).group(1).strip()
            else:
                facts = re.search(r'([\s\S]*)(根据《中华人民共和国环境保护法》)', content_text).group(1).strip()
                # 认定意见
                punishment_basis = ''

                # 处罚决定
                punishment_decision = re.search(r'(根据《中华人民共和国环境保护法》[\s\S]*)', content_text).group(1).strip()
            defense_opinion = ''
            defense_response = ''

        elif re.search(r'(环境违法案件处理情况)', announcement_title):
            # 处罚日期
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

            announcement_code = ''

            # 当事人
            litigant = re.search(r'([\s\S]*)(环境违法案件处理情况)', announcement_title).group(1).strip()

            if re.search(r'((该(公司|集团|厂|单位|作坊)上述行为违反了)[\s\S]*?”(的规定。)*)', content_text):
                # 违规事实
                facts = re.search(r'([\s\S]*)(该(公司|集团|厂|单位|作坊)上述行为违反了)', content_text).group(1).strip()

                # 认定意见
                punishment_basis = re.search(r'((该(公司|集团|厂|单位|作坊)上述行为违反了)[\s\S]*?”(的规定。)*)', content_text).group(1).strip()

                # 处罚决定
                punishment_decision = re.search(r'”(的规定。)*([\s\S]*)', content_text).group(2).strip()
            else:
                facts = re.search(r'([\s\S]*)(根据《中华人民共和国环境保护法》|依据《最高人民法院最高人民检察院关于办理环境污染刑事案件适用法律若干问题的解释》)', content_text).group(1).strip()
                # 认定意见
                punishment_basis = ''

                # 处罚决定
                punishment_decision = re.search(r'((根据《中华人民共和国环境保护法》|依据《最高人民法院最高人民检察院关于办理环境污染刑事案件适用法律若干问题的解释》)[\s\S]*)', content_text).group(1).strip()
            defense_opinion = ''
            defense_response = ''

        elif re.search(r'(环境违法案查处情况)', announcement_title):
            # 处罚日期
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

            announcement_code = ''

            # 当事人
            litigant = re.search(r'([\s\S]*)(环境违法案查处情况)', announcement_title).group(1).strip()

            # 违规事实
            facts = re.search(r'环境违法案查处情况([\s\S]*)(\n该(公司|集团|厂|单位|作坊)|该公司的*上述行为违反了)', content_text).group(1).strip()

            # 认定意见
            punishment_basis = re.search(facts.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                      '\]') + r'([\s\S]*?的*规定)', content_text).group(1).strip()

            # 处罚决定
            punishment_decision = re.search(punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                      '\]') + r'([\s\S]*)', content_text).group(1).strip()

            defense_opinion = ''
            defense_response = ''

        elif re.search(r'防治督查情况', announcement_title):
            tr_list = content_soup.find(class_='display_article').find('table').find_all('tr')
            del (tr_list[0])
            result_map_list = []
            for each_tr in tr_list:
                td_list = each_tr.find_all('td')
                # 处罚日期
                publish_date = each_document['publishDate']
                real_publish_date = format_date(publish_date)

                # 文号
                announcement_code = ''

                # 当事人
                litigant = get_content_text(td_list[3])

                # 违规事实
                facts = get_content_text(td_list[4])

                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                # 处罚决定
                punishment_decision = get_content_text(td_list[5])

                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        else:
            return

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('湖北生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('湖北生态环境厅 数据解析 ' + ' -- 修改parsed完成')

#广西生态环境厅
def guangxi():
    for each_document in db.environment_data.find({'origin': '广西壮族自治区环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        # ignored
        if re.search(r'广西壮族自治区环境保护厅行政处罚项目目录|关于《环境行政处罚案卷评查指南》（征求意见稿）意见的函|'\
                r'自治区环保厅2009年行政处罚工作新突破|申请人主动撤回行政复议申请 自治区环保厅复议调解工作起作用|'\
                r'行政公开新举措|柳州市环保局依法举行市南荣家具制造有限公司行政处罚听证会|自治区环保局前三季度行政处罚立案数取得新增长|'\
                r'维护当事人合法权利依法举行三场听证会|工作座谈会在南宁召开|2005年(8|7)月行政处罚表|(污染源自动监控)*飞行抽检情况的通报|'\
                r'砖瓦行业环保专项执法检查清理整顿情况|4名相关责任人员被问责', announcement_title) \
                or announcement_url in ['http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201206/t20120619_982180.html']:
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        content_response = request_site_page(announcement_url)
        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if content_soup.find(class_ = 'article-inner'):
            class_name = 'article-inner'
        else:
            class_name = 'article'

        #如果是网页上的表格
        if content_soup.find(class_ = class_name).find('table') and not content_soup.find(class_ = class_name).find('table').find('a'):
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text,
                    'content_class_name': 'article-inner',
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']
            logger.info('存入parsed_data')

            if re.search(r'对2007年脱硫设施非正常运行燃煤电厂进行处罚的公告', announcement_title):
                content_text = get_content_text(content_soup.find(class_='article-inner').find('table'))
                facts = re.search(r'一、核查结果([\s\S]*?)二、处罚措施及相关要求', content_text).group(1).strip()
                punishment_decision = re.search(r'二、处罚措施及相关要求([\s\S]*?)特此公告', content_text).group(1).strip()
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': '广西生态环境厅',
                    'announcementDate': format_date(each_document['publishDate']),
                    'announcementCode': '',
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': '',
                    'punishmentBasement': '',
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('广西生态环境厅 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('广西生态环境厅 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('广西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
                continue
            else:
                tr_list = content_soup.find(class_='article-inner').find('table').find_all('tr')
                result_map_list = []
                for each_tr in tr_list:
                    if re.search(r'序号', get_content_text(each_tr)):
                        continue
                    td_list = each_tr.find_all('td')
                    if len(td_list) == 5:
                        announcement_org = get_content_text(td_list[3])
                        facts = get_content_text(td_list[2])
                        litigant = get_content_text(td_list[1])
                        punishment_basis = ''
                        punishment_decision = get_content_text(td_list[4])
                    else:
                        logger.error('网页表格有未考虑到的情况')
                        return
                    result_map = {
                        'announcementTitle': announcement_title,
                        'announcementOrg': announcement_org,
                        'announcementDate': format_date(each_document['publishDate']),
                        'announcementCode': '',
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('广西生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('广西生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('广西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        # 网页公告的形式
        if not content_soup.find(class_ = class_name).find('a') or \
                (not content_soup.find(class_ = class_name).find('a')['href'].endswith('.pdf') and \
                not content_soup.find(class_= class_name).find('a')['href'].endswith('.doc') and \
                not content_soup.find(class_= class_name).find('a')['href'].endswith('.docx')):
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': content_response.text,
                    'content_class_name': 'article-inner',
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']
            logger.info('存入parsed_data')
            content_text = get_content_text(content_soup.find(class_='article-inner'))

            # 人工处理
            if announcement_url in ['http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/200907/t20090724_926491.html',
                                    'http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/200907/t20090724_926490.html',
                                    'http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/200811/t20081119_922536.html',
                                    'http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/200810/t20081009_921884.html',
                                    'http://sthjt.gxzf.gov.cn/zw/qnyw/201607/t20160728_30254.html',
                                    'http://sthjt.gxzf.gov.cn/ztzl/lsztzl/hjwfbgt/201408/t20140818_20182.html']:
                logger.warning('需人工处理 ...')
                continue
            #开始解析
            # 处罚机构
            announcement_org = '广西生态环境厅'

            # 处罚日期
            publish_date = each_document['publishDate']
            real_publish_date = format_date(publish_date)

            # 文号
            try:
                announcement_code = re.search(r'(桂环[\s\S]*?号)', get_content_text(content_soup.find(class_ = 'dg-content otr').find('table'))).group(1).strip()
            except Exception as e:
                announcement_code = ''
            if re.search(r'行政复议决定书', announcement_title):
                # 当事人
                litigant = re.search(r'(被申请人：[\s\S]*?)(柳州市丰源纸业有限责任公司|赵仰新等59户村民)', content_text).group(1).strip()

                # 违规事实
                facts1 = re.search(r'(申请人：[\s\S]*)(被申请人：)', content_text).group(1).strip()
                facts2 = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                        '\]') + r'([\s\S]*?)(二、被申请人的答辩意见和理由)', content_text).group(1).strip()
                facts = facts1 + '\n' + facts2

                defense_opinion = re.search(r'(二、被申请人的答辩意见和理由[\s\S]*?)(三、我厅认定的事实和理由|四、我厅认定的事实、理由|三、我厅认定的事实、理由和证据)', content_text).group(1).strip()
                defense_response = re.search(r'((三、我厅认定的事实和理由|四、我厅认定的事实、理由|三、我厅认定的事实、理由和证据)[\s\S]*?)(四、相关证据|五、以上事实，有以下证据为证，可以认定|上述事实有下列证据证明：)', content_text).group(1).strip()

                # 认定意见
                try:
                    punishment_basis = re.search(r'(根据双方当事人提出的理由和证据材料以及我厅的审查情况，我厅认为[\s\S]*?)(四、复议决定)', content_text).group(1).strip()
                except Exception as e:
                    punishment_basis = ''

                punishment_decision = re.search(r'(五、复议决定|六、复议决定|四、复议决定)([\s\S]*?)(\S{4}年\S{1,3}月\S{1,3}日|申请人柳州市丰源纸业有限责任公司如对本决定不服)', content_text).group(
                    2).strip()
            elif re.search(r'行政复议期限延长通知书|行政复议终止通知书', announcement_title):
                # 当事人
                litigant = re.search(r'([\s\S]*?)：\n', content_text).group(1).strip()

                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                      '\]') + r'([\s\S]*?)(根据《中华人民共和国行政复议法》)', content_text).group(1).lstrip('：').strip()

                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                punishment_decision = re.search(facts.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                     '\]') + r'([\s\S]*?)特此通知。', content_text).group(1).strip()
            elif re.search(r'行政处罚决定书', announcement_title):
                # 当事人
                litigant = re.search(r'([\s\S]*)一、环境违法事实和证据', content_text).group(1).strip()

                # 违规事实
                facts = re.search(r'一、环境违法事实和证据([\s\S]*?)(因此，你(公司|集团|厂|单位)的*上述行为违反了|以上事实，有|以上违法*事实，有)', content_text).group(1).strip()

                # 认定意见
                punishment_basis = re.search(r'(你(公司|集团|厂|单位)的*上述行为违反了[\s\S]*?(的*规定，依法应当予以处罚。|规定。))', content_text).group(1).strip()

                defense_response_text = re.search(punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)\n*(以上事实，*有我[厅局]|以上程序合法，有我[厅局]|以上事实，有下列证据为证：)', content_text).group(1).strip()
                if re.search(r'经复核，我厅认为|我厅认为：',defense_response_text):
                    defense_opinion = re.search(r'([\s\S]*)(经复核，我厅认为|我厅认为：)', defense_response_text).group(1).strip()
                    defense_response = re.search(r'((经复核，我厅认为|我厅认为：)[\s\S]*)', defense_response_text).group(1).strip()
                else:
                    defense_opinion = defense_response_text
                    defense_response = ''

                punishment_decision = re.search(r'(二、责令改正和行政处罚的依据和种类|二、责令改正和行政处罚的依据、种类|二、行政处罚的依据[和、]种类)([\s\S]*?)(四、申请行政复议或者提起行政诉讼的途径和期限|四、申请复议或者提起诉讼的途径和期限|你(公司|集团|厂|单位)如不服本处罚决定)',
                                                content_text).group(2).strip()
            elif re.search(r'行政处罚事先告知书|行政处罚听证告知书', announcement_title):
                # 当事人
                litigant = re.search(r'([\s\S]*?)：\n', content_text).group(1).strip()

                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)(该行为违反了|你(公司|集团|厂|单位)的上述行为违反了|你(公司|集团|厂|单位)以上行为违反了|违反了《中华人民共和国环境保护法》)', content_text).group(1).lstrip('：').strip()

                # 认定意见
                punishment_basis = re.search(r'((该行为违反了|你(公司|集团|厂|单位)的上述行为违反了|你(公司|集团|厂|单位)以上行为违反了|违反了《中华人民共和国环境保护法》)[\s\S]*?的*规定。)', content_text).group(1).strip()

                defense_opinion = ''
                defense_response = ''

                punishment_decision = re.search(punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)根据《中华人民共和国行政处罚法》', content_text).group(1).strip()
            elif re.search(r'环境违法问题的监察通知', announcement_title):
                # 当事人
                litigant = re.search(r'([\s\S]*?)：\n', content_text).group(1).strip()

                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)为切实履行环保部门职责，加强对企业的环境监管，现责成你局做好以下工作', content_text).group(1).lstrip('：').strip()

                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                punishment_decision = re.search(r'(为切实履行环保部门职责，加强对企业的环境监管，现责成你局做好以下工作[\s\S]*?)\n联系人', content_text).group(1).strip()
            elif re.search(r'责令改正违法行为决定书', announcement_title):
                # 当事人
                litigant = re.search(r'([\s\S]*)一、环境违法事实和证据', content_text).group(1).strip()

                # 违规事实
                facts = re.search(r'一、环境违法事实和证据([\s\S]*?)以上事实，*有', content_text).group(1).strip()

                # 认定意见
                punishment_basis = re.search(r'((你(公司|集团|厂|单位)的上述行为违反了|你(公司|集团|厂|单位)无许可证使用放射源的行为违反了)[\s\S]*?的规定。)', content_text).group(1).strip()

                defense_opinion = ''
                defense_response = ''

                punishment_decision = re.search(r'二、责令改正的依据、种类([\s\S]*?)四、申请行政复议或者提起行政诉讼的途径和期限', content_text).group(1).strip()
            elif re.search(r'行政处罚意见的复函', announcement_title):
                # 处罚日期
                try:
                    publish_date = re.findall(r'\S{4}年\S{1,3}月\S{1,3}日', content_text)[-1]
                    real_publish_date = format_date(publish_date)
                except Exception as e:
                    publish_date = each_document['publishDate']
                    real_publish_date = format_date(publish_date)

                # 文号
                try:
                    announcement_code = re.search(r'(环函[\s\S]*?号)', content_text).group(1).strip()
                except Exception as e:
                    announcement_code = ''

                # 当事人
                litigant = re.search(r'行政处罚意见的复函\n([\s\S]*?)：\n', content_text).group(1).strip()

                # 违规事实
                facts = ''

                # 认定意见
                punishment_basis = ''
                defense_opinion = ''
                defense_response = ''

                punishment_decision = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)\n主题词', content_text).group(1).strip()
            else:
                logger.error('出现未考虑到的情况')
                return

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': announcement_org,
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': defense_opinion,
                'defenseResponse': defense_response,
                'litigant': litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('广西生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('广西生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('广西生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue

        if content_soup.find(class_ = class_name).find('a')['href'].endswith('.pdf'):
            if content_soup.find(class_= class_name).find('a')['href'].startswith('http://'):
                pdf_link = content_soup.find(class_= class_name).find('a')['href']
            else:
                pdf_link = urljoin(announcement_url, content_soup.find(class_= class_name).find('a')['href'])
            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': pdf_link}).count() == 0:
                pdf_response = request_site_page(pdf_link)
                if pdf_response is None:
                    logger.error('pdf文件下载失败 %s' % pdf_link)
                    continue
                with open('./test/tmp.pdf', 'wb') as tmp_file:
                    for chunk in pdf_response.iter_content(chunk_size=1024):
                        if chunk:
                            tmp_file.write(chunk)
                result_text, ocr_flag = pdf_ocr_to_text('./test/tmp.pdf')
                with open('./test/tmp.pdf', 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': pdf_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'pdf',
                    'oss_file_name': announcement_title,
                    'oss_file_content': pdf_content,
                    'parsed': False,
                    'if_ocr': True,
                    'ocr_result': result_text
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.pdf', pdf_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                content_text = result_text

                logger.info('删除TMP文件')
                if os.path.exists('./test/tmp.pdf'):
                    os.remove('./test/tmp.pdf')
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': pdf_link})['_id']
                content_text = db.parsed_data.find_one(
                    {'origin_url': announcement_url, 'oss_file_origin_url': pdf_link})['ocr_result']
            logger.info('存入parsed_data')
        elif content_soup.find(class_ = class_name).find('a')['href'].endswith('.doc')\
                or content_soup.find(class_ = class_name).find('a')['href'].endswith('.docx'):
            docx_link = urljoin(announcement_url, content_soup.find(class_='article-inner').find('a')['href'])
            link_type = docx_link.split('.')[-1]
            response = request_site_page(docx_link)
            if response is None:
                logger.error('docx文件下载失败 %s' % docx_link)
                return
            with open('./test/tmp.' + link_type, 'wb') as tmp_file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        tmp_file.write(chunk)

            if not os.path.exists('./test/tmp.docx'):
                shell_str = 'soffice --headless --convert-to docx ' + \
                            './test/tmp.' + link_type + ' --outdir ./test'
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.communicate()

            with open('./test/tmp.docx', 'rb') as docx_file:
                docx_content = docx_file.read()

            if db.parsed_data.find({'origin_url': announcement_url, 'oss_file_origin_url': docx_link}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': docx_link,
                    'origin_url_id': each_document['_id'],
                    'oss_file_type': 'docx',
                    'oss_file_name': announcement_title,
                    'oss_file_content': docx_content,
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.docx', docx_content)
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
                file_id = \
                    db.parsed_data.find_one({'origin_url': announcement_url, 'oss_file_origin_url': docx_link})[
                        '_id']
            logger.info('存入parsed_data')
            document = Document('./test/tmp.docx')
            content_text = '\n'.join([each_paragraph.text.strip() for each_paragraph in document.paragraphs])
            logger.info('删除TMP文件')
            if os.path.exists('./test/tmp.doc'):
                os.remove('./test/tmp.doc')
            if os.path.exists('./test/tmp.docx'):
                os.remove('./test/tmp.docx')


        # 人工处理
        if announcement_url in ['http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/hjyxpj/201410/t20141022_200000785.html']:
            logger.warning('需人工处理 ...')
            continue
        #开始解析
        logger.info(content_text)

        # 处罚机构
        announcement_org = '广西生态环境厅'

        # 处罚日期
        publish_date = each_document['publishDate']
        real_publish_date = format_date(publish_date)

        # 文号
        try:
            announcement_code = re.search(r'([来桂]环[\s\S]*?号)', content_text).group(1).strip()
        except Exception as e:
            logger.info('文本中提取文号失败')
            announcement_code = get_content_text(content_soup.find(class_ = 'article-inner').find('a')).split('.')[0]

        #http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201903/t20190305_200012383.html
        if re.search(r'行政复议调解书', announcement_title):
            facts1 = re.search(r'(申请人[\s\S]*?)\n被申请人[:：]', content_text).group(1).strip()

            corp = re.search(r'申请人[:：]([\s\S]*?)\n', content_text).group(1).strip()
            # 当事人
            litigant = re.search(r'(被申请人[:：][\s\S]*?)' + corp.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]'), content_text).group(1).strip()

            # 违规事实
            facts2 = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)\n被申请人称', content_text).group(1).strip()
            facts = facts1 + facts2

            defense_opinion = re.search(r'\n(被申请人称[\s\S]*?)\n本机关经审理查明', content_text).group(1).strip()

            defense_response = re.search(r'\n(本机关经审理查明[\s\S]*?)\n鉴于申请人已落实整改', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            punishment_decision = re.search(r'\n(鉴于申请人已落实整改[\s\S]*?)\n申请人[:：]', content_text).group(1).strip()

        #http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201812/t20181207_200011620.html
        elif re.search(r'延期审理通知书', announcement_title):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)[:：]\n', content_text).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)(经审理[，,])', content_text).group(1).lstrip(':').lstrip('：').strip()

            # 认定意见
            punishment_basis = re.search(r'(经审理[，,][\s\S]*?)根据《', content_text).group(1).strip()

            defense_opinion = ''

            defense_response = ''

            punishment_decision = re.search(r'(根据《[\s\S]*?)特此通知', content_text).group(1).strip()

        #http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201811/W020181126562717236179.pdf
        elif re.search(r'撤销行政复议决定书', announcement_title):
            # 当事人
            litigant = re.search(r'撤销行政复议决定书的决定\n([\s\S]*?)[:：]\n', content_text).group(1).strip()

            # 违规事实
            facts = ''

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''

            defense_response = ''

            punishment_decision = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)(\n广西壮族自治区生态环境[厅斤])', content_text).group(1).lstrip(':').lstrip('：').strip()
        #http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201809/t20180913_200010631.html
        elif re.search(r'行政复议终止决定书', announcement_title):
            facts1 = re.search(r'(申请人[\s\S]*?)\n被申请人[:：]', content_text).group(1).strip()

            # 当事人
            litigant = re.search(r'(被申请人[:：][\s\S]*?)\n申请人', content_text).group(1).strip()

            # 违规事实
            facts2 = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                     '\]') + r'([\s\S]*?)根据.*《', content_text).group(1).strip()
            facts = facts1 + '\n' + facts2

            defense_opinion = ''

            defense_response = ''

            # 认定意见
            punishment_basis = ''

            punishment_decision = re.search(r'(根据.*《[\s\S]*?)\n复议机关', content_text).group(1).strip()

        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201811/t20181126_200011440.html
        elif re.search(r'(行政)*复议决定书', announcement_title):
            facts1 = re.search(r'(申请人[\s\S]*?)\n被申请人[:：]', content_text).group(1).strip()

            corp = re.search(r'申请人[:：]([\s\S]*?)\n', content_text).group(1).strip()
            # 当事人
            litigant = re.search(r'(被申请人[:：][\s\S]*?)' + corp.replace('(', '\(').replace(')', '\)').replace('[',
                '\[').replace(']', '\]').replace('徳', '[徳德]').replace('滨', '[滨淀]').replace('工',
                '[工エ]').replace('一', '[一ー]').replace('研', '[研硏]').replace('\(北海市海城区高朋酒楼业主\)',''), content_text).group(1).strip()

            # 违规事实
            facts2 = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                    '\]') + r'([\s\S]*?)\n(被申请人称|被申请人的答[辨辩]意见和理由)', content_text).group(1).strip()
            facts = facts1 + '\n' + facts2

            defense_opinion = re.search(r'\n((被申请人称|被申请人的答[辨辩]意见和理由)[\s\S]*?)\n(本机关经审理[查査]明|本机关认定的事实、理由|、*本厅认定的事实、理由|(三、)*我厅认定的事实[和、]理由)', content_text).group(1).strip()

            defense_response = re.search(r'\n((本机关经审理[查査]明|本机关认定的事实、理由|、*本厅认定的事实、理由|(三、)*我厅认定的事实[和、]理由)[\s\S]*?)\n(本机关认为|四、复议决定|五、复议决定)', content_text).group(1).replace('査',
                                                                                                           '查').strip()

            # 认定意见
            try:
                punishment_basis = re.search(r'\n(本机关认为[\s\S]*?)\n(综上所述|综上,)', content_text).group(1).strip()
            except Exception as e:
                punishment_basis = ''

            punishment_decision = re.search(r'\n((综上所述|综上,|四、复议决定|五、复议决定)[\s\S]*?)\n(申请人(.+)*如对本(复议)*决定不服|复议机关|(柳州静兰电镀厂|广西鸿雁食品有限公司)如对本复议决定不服)', content_text).group(1).strip()

        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201302/t20130204_991710.html
        elif re.search(r'驳回行政复议申请决定书', announcement_title):
            facts1 = re.search(r'(申请人[\s\S]*?)\n被申请人[:：]', content_text).group(1).strip()

            corp = re.search(r'申请人[:：]([\s\S]*?)\n', content_text).group(1).strip()
            # 当事人
            litigant = re.search(r'(被申请人[:：][\s\S]*?)' + corp.replace('(', '\(').replace(')', '\)').replace('[',
                    '\[').replace(']', '\]'), content_text).group(1).strip()

            # 违规事实
            facts2 = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                '\]') + r'([\s\S]*?)\n被申请人的答辨意见和理由', content_text).group(1).strip()
            facts = facts1 + facts2

            defense_opinion = re.search(r'\n(被申请人的答辨意见和理由[\s\S]*?)\n本机关认定的事实和理由', content_text).group(1).strip()

            defense_response = re.search(r'\n(本机关认定的事实和理由[\s\S]*?)\n四、复议决定', content_text).group(1).replace('査',
                                                                                                           '查').strip()

            # 认定意见
            punishment_basis = re.search(r'\n(本机关认为:[\s\S]*?)(本机关决定:)', content_text).group(1).strip()

            punishment_decision = re.search(r'(本机关决定:[\s\S]*?)\n复议机关', content_text).group(1).strip()

        #http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201804/t20180416_200008491.html
        elif re.search(r'行政处罚决定书', announcement_title) or re.search(r'自治区环境保护[厅斤]\n行政处罚决定书', content_text):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                     '\]') + r'([\s\S]*?)\n(一、)*环境违法事实和证据', content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'环境违法事实和证据([\s\S]*?)以上(违法)*事实[，,]有', content_text).group(1).strip()

            # 认定意见
            punishment_basis = re.search(r'(你*(公司|集团|厂|单位)*(的上[述远]行为|以上行为|\n违反了)[\s\S]*?(的*(\n)*规(\n)*定。|“一般辐射事故”。|的*规\n*定[,，]*(依法)*\n*应当*予以*处\n*罚。*))', content_text).group(1).strip()

            defense_response_text = re.search(punishment_basis.replace('(', '\(').replace(')', '\)').replace('[',
                '\[').replace(']', '\]') + r'([\s\S]*?)(二、)*(责令改正和)*行政处罚的依据[和、]种类', content_text).group(1).strip()

            if re.search(r'(我[厅斤]对你公司的陈述申辩意见及相关材料进行了审核|我[厅斤]认为[:：,]|经研究,)', defense_response_text):
                defense_opinion = re.search(r'([\s\S]*)(我[厅斤]对你公司的陈述申辩意见及相关材料进行了审核|我[厅斤]认为[:：,]|经研究,)', defense_response_text).group(1).strip()

                defense_response = re.search(r'((我[厅斤]对你公司的陈述申辩意见及相关材料进行了审核|我[厅斤]认为[:：,]|经研究,)[\s\S]*)', defense_response_text).group(1).strip()
            else:
                defense_opinion = defense_response_text
                defense_response = ''

            punishment_decision = re.search(r'行政处罚的依据[和、]种类([\s\S]*?)(三、)*((责令改正和)*行政处罚的履行方式和期限|责令停止生产和罚款的履行方式和期限|(责令改正和行政)*处罚决定的履行方式和期限)', content_text).group(1).strip()
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201602/t20160203_200001646.html
        elif re.search(r'按日连续处罚决定书', announcement_title) or re.search(r'壮族自治区环境保护.*[厅斤]\n按日连续处罚决定书', content_text):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
               '\]') + r'([\s\S]*?)\n(一、)*按日连续处罚的事实和证据', content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'按日连续处罚的事实和证据([\s\S]*?)以上事实[，,]有', content_text).group(1).strip()

            # 认定意见
            punishment_basis = re.search(r'\n(依据《中华人民共和国环境保护法》[\s\S]*?的行为。)',
                                         content_text).group(1).strip()

            defense_response_text = re.search(punishment_basis.replace('(', '\(').replace(')', '\)').replace('[',
               '\[').replace(']', '\]') + r'([\s\S]*?)以上事实[，,]有', content_text).group(1).strip()

            if re.search(r'(我[厅斤]对你公司的陈述申辩意见及相关材料进行了审核|我[厅斤]认为[:：,])', defense_response_text):
                defense_opinion = re.search(r'([\s\S]*)(我[厅斤]对你公司的陈述申辩意见及相关材料进行了审核|我[厅斤]认为[:：,])',
                                            defense_response_text).group(1).strip()

                defense_response = re.search(r'((我[厅斤]对你公司的陈述申辩意见及相关材料进行了审核|我[厅斤]认为[:：,])[\s\S]*)',
                                             defense_response_text).group(1).strip()
            else:
                defense_opinion = defense_response_text
                defense_response = ''

            punishment_decision = re.search(r'实施按日连续处罚的依据和总额([\s\S]*?)实施按日连续处罚的履行方式和期限', content_text).group(
                1).strip()
        #http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201804/t20180416_200008488.html
        elif re.search(r'限制生产决定书', announcement_title):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                     '\]') + r'([\s\S]*?)\n(一、)*环境违法事实和证据', content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'环境违法事实和证据([\s\S]*?)以上事实[，,]有', content_text).group(1).strip()

            # 认定意见
            punishment_basis = re.search(r'(你(公司|集团|厂|单位)(的上述(两个违法)*|以上)行为[\s\S]*?的(\n)*规定。)', content_text).group(1).strip()

            defense_response_text = re.search(punishment_basis.replace('(', '\(').replace(')', '\)').replace('[',
                '\[').replace(']', '\]') + r'([\s\S]*?)(二、)*(责令限制生产的依据、期限和幅度|责令限制生产的依据和期限)', content_text).group(1).strip()

            if re.search(r'(我[厅斤]对你公司的陈述申辩意见及相关材料进行了审核)', defense_response_text):
                defense_opinion = re.search(r'([\s\S]*)我[厅斤]对你公司的陈述申辩意见及相关材料进行了审核', defense_response_text).group(1).strip()

                defense_response = re.search(r'(我[厅斤]对你公司的陈述申辩意见及相关材料进行了审核[\s\S]*)', defense_response_text).group(1).strip()
            else:
                defense_opinion = defense_response_text
                defense_response = ''

            punishment_decision = re.search(r'(责令限制生产的依据、期限和幅度|责令限制生产的依据和期限)([\s\S]*?)(三、)*责令限制生产的履行方式和解除', content_text).group(2).strip()
        #http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201804/t20180412_200008436.html
        elif re.search(r'责令改正违法行为决定书', announcement_title):
            if re.search(r'环境违法事实和证据', content_text):
                # 当事人
                litigant = re.search(r'责令改正违法行为决定书\n([\s\S]*?)\n(一、)*环境违法事实和证据', content_text).group(1).strip()

                # 违规事实
                facts = re.search(r'环境违法事实和证据([\s\S]*?)(以上(违法)*事实[,，]*有)', content_text).group(1).strip()

                # 认定意见
                punishment_basis = re.search(r'(你(公司|集团|厂|单位|院)(的这些行为|的*上[述远]行为)[\s\S]*?的规定。*)', content_text).group(1).strip()

                defense_opinion = re.search(punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                  '\]') + r'([\s\S]*?)(二、)*责令改正的依据、种类', content_text).group(1).strip()
                defense_response = ''

                punishment_decision = re.search(r'责令改正的依据、种类\n([\s\S]*?)(\n四、申请行政复议或者提起行政诉讼的途径和期限)',
                    content_text).group(1).strip()
            else:
                # 当事人
                litigant = re.search(
                    announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                       '\]') + r'([\s\S]*?)\n(\d{4}年\d{1,2}月\d{1,2}日|自治区环境监察总队|我厅和防城港市环境保护局|我厅于\d{4}年\d{1,2}月\d{1,2}日)',
                    content_text).group(1).strip()

                # 违规事实
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                            '\]') + r'([\s\S]*?)(以上违法事实[，,]*有|以上事实[，,]有)',
                                  content_text).group(1).strip()

                # 认定意见
                if re.search(r'(你(公司|集团|厂|单位)(的*(上述)*行为|利用2号水塘塘堤)[\s\S]*?的规定。)', content_text):
                    punishment_basis = re.search(r'(你(公司|集团|厂|单位)(的*(上述)*行为|利用2号水塘塘堤)[\s\S]*?的规定。)',
                                                 content_text).group(1).strip()
                else:
                    punishment_basis = re.search(r'等证据为证。\n([\s\S]*?的规定。)', content_text).group(1).strip()
                defense_opinion = ''
                defense_response = ''

                punishment_decision = re.search(
                    punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                      '\]') + r'([\s\S]*?)(\n你(公司|集团|厂|单位)如不服本决定|如你(公司|集团|厂|单位)对本决定不服|\n我厅对你(公司|集团|厂|单位)改正违法行为的情况进行监督)',
                    content_text).group(1).strip()
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201701/t20170110_200003288.html
        elif re.search(r'听证通知书', announcement_title):
            if re.search(r'一、环境违法事实和证据', content_text):
                # 当事人
                litigant = re.search(r'行政处罚决定书\n([\s\S]*?)一、环境违法事实和证据', content_text).group(1).strip()

                # 违规事实
                facts = re.search(r'一、环境违法事实和证据([\s\S]*?)以上违法事实，有', content_text).group(1).strip()

                # 认定意见
                punishment_basis = re.search(r'(违反了[\s\S]*?的规定，应予处罚。)', content_text).group(1).strip()

                defense_opinion = re.search(punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                        '\]') + r'([\s\S]*?)以上程序，有我厅', content_text).group(1).strip()
                defense_response = ''

                punishment_decision = re.search(r'二、责令改正和行政处罚的依据、种类([\s\S]*?)三、处罚决定的履行方式和期限',
                    content_text).group(1).strip()
            else:
                # 当事人
                litigant = re.search(
                    announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                                       '\]') + r'([\s\S]*?)\n根据《中华人民共和国行政处罚法》',
                    content_text).group(1).strip()

                # 违规事实
                facts = ''

                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                punishment_decision = re.search(
                    litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                                                                                              '\]') + r'([\s\S]*?)(\n广西壮族自治区环境保护厅|\n广西壮族自治区环境保护斤办公室)',
                    content_text).group(1).strip()

        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201512/t20151222_200001549.html
        elif re.search(r'行政复议答复通知书', announcement_title):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                       '\]') + r'([\s\S]*?)\n(南宁市盛燕塑料制品有限责任公司|南宁南箱彩印包装有限公司|王永芬\(南宁市柳发机瓦厂\)请求撤销|甘振中同志向我厅提出行政复议|南宁辽大养殖有限公司请求撤销|刘贻勇请求撤销)', content_text).group(1).strip()

            # 违规事实
            facts = re.search(r'\n((南宁市盛燕塑料制品有限责任公司|南宁南箱彩印包装有限公司|王永芬\(南宁市柳发机瓦厂\)请求撤销|甘振中同志向我厅提出行政复议|南宁辽大养殖有限公司请求撤销|刘贻勇请求撤销)[\s\S]*?(本机关已)*(依法|决定)\n*受理。)', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            punishment_decision = re.search(facts.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                  '\]') + r'([\s\S]*?)(\n复议机关)', content_text).group(1).strip()
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201512/t20151222_200001548.html
        elif re.search(r'行政复议受理通知书|行政复议终止通知书', announcement_title):
            # 当事人
            litigant = re.search(r'[\(（]([\s\S]*?)[\)）]', announcement_title).group(1).strip()
            # litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
            #      '\]') + r'([\s\S]*?)\n([你您](公司)*请求[撤撒]销|南宁市多宝路酒业有限公司请求撤销|你请求责令|你们请求确认|李艳萍、曾广元等52人请求确认|广西福基混凝土有限公司不服)', content_text).group(1).rstrip(':').strip()

            # 违规事实
            # facts = re.search(r'\n(([你您](公司)*请求[撤撒]销|南宁市多宝路酒业有限公司请求撤销|你请求责令|你们请求确认|李艳萍、曾广元等52人请求确认|广西福基混凝土有限公司不服)[\s\S]*?)(经\n*审[查査]|根据《中华人民共和国行)', content_text).group(1).strip()
            if re.search(r'(你不服北海市环境保护局于[\s\S]*?)(经\n*审[查査]|根据《中华人民共和国行|现将复议申请书副本送达)', content_text):
                facts = re.search(r'(你不服北海市环境保护局于[\s\S]*?)(经\n*审[查査]|根据《中华人民共和国行|现将复议申请书副本送达)', content_text).group(1).strip()
            else:
                facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                 '\]').replace('研', '[研硏]') + r'([\s\S]*?)(经\n*审[查査]|根据《中华人民共和国行|现将复议申请书副本[送迭]达)', content_text).group(1).strip()
            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            punishment_decision = re.search(facts.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                '\]') + r'([\s\S]*?)(\n复议机关|\n特此通知。*)', content_text).group(1).strip()
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201512/t20151208_200001526.html
        elif re.search(r'行政复议期限延长通知书', announcement_title):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                   '\]') + r'([\s\S]*?)\n(你(公司)*请求[撤撒]销|刘贻勇请求撤销)', content_text).group(1).rstrip(':').strip()

            # 违规事实
            facts = re.search(r'\n((你(公司)*请求[撤撒]销|刘贻勇请求撤销)[\s\S]*?)(\n《中华人民共和国行政复议法》)', content_text).group(
                1).strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            punishment_decision = re.search(facts.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                '\]') + r'([\s\S]*?)(\n复议机关)', content_text).group(1).strip()
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201511/t20151130_200001510.html
        elif re.search(r'环境保护厅办公室转发环境保护部关于按日连续处罚计罚日数有关适用问题的函', announcement_title):
            # 当事人
            litigant = re.search(r'有关适用问题的函\n([\s\S]*?)[:：]\n', content_text).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
            '\]') + r'([\s\S]*?)\n广西壮族自治区环境保护斤办公室', content_text).group(1).lstrip(':').lstrip('：').strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            punishment_decision = ''
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201509/t20150929_200001399.html
        elif re.search(r'行政处罚(使用|适用|有关)问题的复*函|行政处罚相关事项的复函|行政复议情况的函', announcement_title):
            # 当事人
            litigant = re.search(r'(适用问题的函|有关问题的函|关事项的复函|行政复议[崝情]况的[图函])\n([\s\S]*?)\n(你局|《关于家禽宰杀建设项目行政处罚有关问题的请示》|你单位)', content_text).group(2).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                 '\]') + r'([\s\S]*?)(经研究[,\n]|经查,)', content_text).group(1).lstrip(':').lstrip('：').strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            punishment_decision = re.search(r'((经研究[,\n]|经查,)[\s\S]*?)\n(广*西壮族自治区环境保护厅|此函。|\(信息是否公开)', content_text).group(1).lstrip(':').lstrip('：').strip()
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201412/t20141211_200000840.html
        elif re.search(r'行政处罚(听证|事先)[告通]知书', announcement_title):
            # 当事人
            litigant = re.search(announcement_code.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                '\]') + r'([\s\S]*?)\n(经调[查査][,，]|我[厅斤]经现场核[查査]和调[查査]|经我[厅斤]调[查査])', content_text).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                      '\]') + r'([\s\S]*?)(你(公司|集团|厂|单位)的*上述\n*行\n*为|以上行为|2012年4月16日我厅向你公司下达了|依据《中华人民共和国环境影响评价法》第三十一条|违反了《中\n华人民共和国水污染防治法》第二十一条第二款:|违反了《建设项目环境保护管理条例》|违反了《中华人民共和国环境影响评价法》)', content_text).group(1).strip()

            # 认定意见
            if re.search(r'((你(公司|集团|厂|单位)的*上述\n*行\n*为|以上行为|你单位不正常使用水污染物处理设施|违反了《建设项目环境保护管理条例》|违反了《中华人民共和国环境影响评价法》)[\s\S]*?的*(有关)*规定。)', content_text):
                punishment_basis = re.search(r'((你(公司|集团|厂|单位)的*上述\n*行\n*为|以上行为|你单位不正常使用水污染物处理设施|违反了《建设项目环境保护管理条例》|违反了《中华人民共和国环境影响评价法》)[\s\S]*?的*(有关)*规定。)', content_text).group(1).strip()
            else:
                logger.info("提取认定意见失败")
                punishment_basis = ''
            defense_opinion = ''
            defense_response = ''

            if punishment_basis == '':
                punishment_decision = re.search(
                    facts.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']', '\]') + r'([\s\S]*?)\n根据《中华人民共和国行政处罚法》',
                    content_text).group(1).strip()
            else:
                punishment_decision = re.search(punishment_basis.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                      '\]') + r'([\s\S]*?)\n根据《中华人民共和国行政处罚法》', content_text).group(1).strip()
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201409/t20140903_200000729.html
        elif re.search(r'责令停产整治通知书', announcement_title):
            # 当事人
            litigant = re.search(r'[\(（]([\s\S]*?)[\)）]', announcement_title).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                '\]') + r'([\s\S]*?)之规定', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            punishment_decision = re.search(r'(现责令[\s\S]*?)\n广西壮族自治区环境保护厅', content_text).group(1).strip()
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201203/t20120314_982051.html
        elif re.search(r'委托检查行政处罚决定执行情况的函', announcement_title):
            # 当事人
            litigant = re.search(r'行政处罚决定执行情况的函\n([\s\S]*?)\n中信大锰田东新材料有限公司', content_text).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                '\]') + r'([\s\S]*?)\n该公司已按期缴纳罚款', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            punishment_decision = re.search(facts.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                '\]') + r'([\s\S]*?)\n\(信息是否公开:主动公开\)', content_text).group(1).strip()
        # http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201107/t20110728_981805.html
        elif re.search(r'关于同意\S*分期缴纳行政处罚罚款的决定', announcement_title):
            # 当事人
            litigant = re.search(r'关于同意([\s\S]*?)分期缴纳行政处罚罚款的决定', announcement_title).group(1).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                '\]') + r'([\s\S]*?)\n我厅认为:', content_text).group(1).strip()

            # 认定意见
            punishment_basis = re.search(r'(我厅认为:[\s\S]*?)根据《中华人民共和国行政处罚法》', content_text).group(1).strip()

            defense_opinion = ''
            defense_response = ''

            punishment_decision = re.search(r'(根据《中华人民共和国行政处罚法》[\s\S]*?逾期未缴纳完毕,我厅将\n*依法申请法院强制执行。)',
                                            content_text).group(1).strip()
        elif re.search(r'关于对([\s\S]*?)违规编制环评文件事件的通报', announcement_title):
            # 当事人
            litigant = re.search(r'关于对([\s\S]*?)违规编制环评文件事件的通报', announcement_title).group(1).strip()

            # 违规事实
            facts = re.search(r'环评文件的通报\n([\s\S]*?)\n为规范我区环评市场', content_text).group(1).strip()

            # 认定意见
            punishment_basis = ''

            defense_opinion = ''
            defense_response = ''

            punishment_decision = re.search(r'(为规范我区环评市场[\s\S]*?)广西壮族自治区环境保护厅', content_text).group(1).strip()
        else:
            logger.error('出现未考虑到的情况')
            return

        result_map = {
            'announcementTitle': announcement_title,
            'announcementOrg': announcement_org,
            'announcementDate': real_publish_date,
            'announcementCode': announcement_code,
            'facts': facts,
            'defenseOpinion': defense_opinion,
            'defenseResponse': defense_response,
            'litigant': litigant,
            'punishmentBasement': punishment_basis,
            'punishmentDecision': punishment_decision,
            'type': '行政处罚决定',
            'oss_file_id': file_id,
            'status': 'not checked'
        }
        logger.info(result_map)
        if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
            db.announcement.insert_one(result_map)
            logger.info('广西生态环境厅 数据解析 ' + ' -- 数据导入完成')
        else:
            logger.info('广西生态环境厅 数据解析 ' + ' -- 数据已经存在')
        db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
        logger.info('广西生态环境厅 数据解析 ' + ' -- 修改parsed完成')

#海南生态环境厅
def hainan():
    for each_document in db.environment_data.find({'origin': '海南省环境保护厅', 'status': {'$nin': ['ignored']}}):

        announcement_url = each_document['url']
        announcement_title = each_document['title']

        # 判断是否解析过
        if db.environment_data.find({'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
                {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue
        logger.info(announcement_title)
        logger.info('url to parse ' + announcement_url)

        # ignored
        if announcement_url in ['http://sthjt.gxzf.gov.cn/xxgkml/ztfl/hjglywxx/xzcfhfy/201206/t20120619_982180.html']:
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        content_response = request_site_page(announcement_url)
        if content_response.status_code == 404:
            logger.warning('url has nothing to do with punishment ...')
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'ignored'}})
            logger.info('update environment data success')
            continue

        if content_response is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_response.encoding = content_response.apparent_encoding
        content_soup = BeautifulSoup(content_response.text, 'lxml')

        if db.parsed_data.find(
                {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
            oss_file_map = {
                'origin_url': announcement_url,
                'oss_file_origin_url': announcement_url,
                'origin_url_id': each_document['_id'],
                'oss_file_type': 'html',
                'oss_file_name': announcement_title,
                'oss_file_content': content_response.text,
                'content_class_name': 'cen-main',
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html', content_response.text)
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
        else:
            db.environment_data.update_one({'_id': each_document['_id']}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                               'oss_file_origin_url': announcement_url})['_id']
        logger.info('存入parsed_data')
        content_text = get_content_text(content_soup.find(id = 'zoom'))
        head_text = get_content_text(content_soup.find(class_ = 'headInfo'))

        # 开始解析
        # 处罚机构
        announcement_org = '海南生态环境厅'

        # 文号
        try:
            announcement_code = re.search(r'(琼环[\s\S]*?号)', head_text).group(1).strip()
        except Exception as e:
            announcement_code = ''
            logger.info('提取文号失败')
        if content_soup.find(id = 'zoom').find('table'):
            # 处罚日期
            real_publish_date = format_date(each_document['publishDate'])
            result_map_list = []
            tr_list = content_soup.find(id = 'zoom').find('table').find_all('tr')
            for each_tr in tr_list:
                if re.search(r'公示表|序号', get_content_text(each_tr)):
                    continue
                td_list = each_tr.find_all('td')
                # 当事人
                litigant = '辖区: ' + get_content_text(td_list[1]) + ' 当事人： ' + get_content_text(td_list[2])

                # 违规事实
                facts = get_content_text(td_list[3])
                # 认定意见
                punishment_basis = ''

                defense_opinion = ''
                defense_response = ''

                punishment_decision = ''
                result_map = {
                    'announcementTitle': announcement_title,
                    'announcementOrg': announcement_org,
                    'announcementDate': real_publish_date,
                    'announcementCode': announcement_code,
                    'facts': facts,
                    'defenseOpinion': defense_opinion,
                    'defenseResponse': defense_response,
                    'litigant': litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            logger.info(result_map_list)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_many(result_map_list)
                logger.info('海南生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('海南生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('海南生态环境厅 数据解析 ' + ' -- 修改parsed完成')
            continue
        else:
            # 处罚日期
            publish_date = re.findall(r'\S{4}年\S{1,3}月\S{1,3}日', content_text)[-1]
            real_publish_date = format_date(publish_date)

            # 当事人
            litigant = re.search(r'(行政处罚决定书|责令改正违法行为决定书)\n([\s\S]*?)(我厅于*\d{4}年\d{1,2}月\d{1,2}日)', content_text).group(2).strip()

            # 违规事实
            facts = re.search(litigant.replace('(', '\(').replace(')', '\)').replace('[', '\[').replace(']',
                      '\]') + r'([\s\S]*?)(以上事实,*有)', content_text).group(1).strip()
            # 认定意见
            try:
                punishment_basis = re.search(r'(((你的)*上述行为违反了|你(公司|单位|集团|厂)的*上述行为违反了)[\s\S]*?的规定。)', content_text).group(1).strip()
            except Exception as e:
                punishment_basis = ''
                logger.info('提取认定意见失败')

            try:
                defense_response_text = re.search(r'(我厅于*\d{4}年\d{1,2}月\d{1,2}日以《行政处罚告知书》[\s\S]*?)\n依据', content_text).group(1).strip()
                if re.search(r'我厅对你公司申辩的理由进行讨论研究，认为', defense_response_text):
                    defense_opinion = re.search(r'([\s\S]*)我厅对你公司申辩的理由进行讨论研究，认为', defense_response_text).group(1).strip()
                    defense_response = re.search(r'(我厅对你公司申辩的理由进行讨论研究，认为[\s\S]*)', defense_response_text).group(1).strip()
                else:
                    defense_opinion = defense_response_text
                    defense_response = ''
            except Exception as e:
                defense_opinion = ''
                defense_response = ''

            punishment_decision = re.search(r'((依据《建设项目环境保护管理条例》|依据《中华人民共和国行政处罚法》)[\s\S]*?)(限于接到本处罚决定之日起15日内|你(公司|单位|集团|厂)如对本决定不服)',
                content_text).group(1).strip()

            result_map = {
                'announcementTitle': announcement_title,
                'announcementOrg': announcement_org,
                'announcementDate': real_publish_date,
                'announcementCode': announcement_code,
                'facts': facts,
                'defenseOpinion': defense_opinion,
                'defenseResponse': defense_response,
                'litigant': litigant,
                'punishmentBasement': punishment_basis,
                'punishmentDecision': punishment_decision,
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
            logger.info(result_map)
            if db.announcement.find({'announcementTitle': announcement_title, 'oss_file_id': file_id}).count() == 0:
                db.announcement.insert_one(result_map)
                logger.info('海南生态环境厅 数据解析 ' + ' -- 数据导入完成')
            else:
                logger.info('海南生态环境厅 数据解析 ' + ' -- 数据已经存在')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('海南生态环境厅 数据解析 ' + ' -- 修改parsed完成')

if __name__ == "__main__":
    # #生态环境部
    # mee_parse()
    # #北京生态环境局
    # beijing()
    # #天津生态环境局
    # tianjin()
    # #河北生态环境厅
    # hebei()
    # #山西生态环境厅
    # shanxi()
    # # 内蒙古生态环境厅
    # neimenggu()
    # #辽宁生态环境厅
    # liaoning()
    # #吉林生态环境厅
    # jilin()
    # # 黑龙江生态环境厅
    # heilongjiang()
    # # 浙江生态环境厅
    # zhejiang()
    # # 安徽生态环境厅
    # anhui()
    # #福建生态环境厅
    # fujian()
    # # 江西生态环境厅
    # jiangxi()
    # # 山东生态环境厅
    # shandong()
    # # 河南生态环境厅
    # henan()
    # #湖南生态环境厅
    # hunan()
    # # 广东生态环境厅
    # guangdong()
    # # 湖北生态环境厅
    # hubei()
    # #广西生态环境厅
    # guangxi()
    #海南生态环境厅
    hainan()
