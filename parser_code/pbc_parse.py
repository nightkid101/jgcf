import jsbeautifier
import js2py
import time
import re
import os
import sys
import subprocess
from urllib.parse import urljoin

import requests
import pymongo
from bs4 import BeautifulSoup as bs
from pymongo import MongoClient
from docx import Document
from xlrd import open_workbook, xldate_as_tuple
import pdfplumber
from init import logger_init, config_init
from utility import table_to_list, request_site_page, genHeader, format_date, get_content_text, format_text, \
    get_chinese_proportion
from pdf2html import pdf_ocr_to_table
from oss_utils import init_ali_oss, oss_add_file

logger = logger_init('人行及分支机构 数据解析')
config = config_init()
if config['mongodb']['dev_mongo'] == '1':
    db = MongoClient(config['mongodb']['ali_mongodb_url'], username=config['mongodb']['ali_mongodb_username'],
                     password=config['mongodb']['ali_mongodb_password'],
                     port=int(config['mongodb']['ali_mongodb_port']))[config['mongodb']['ali_mongodb_name']]
else:
    db = MongoClient(
        host=config['mongodb']['mongodb_host'],
        port=int(config['mongodb']['mongodb_port']),
        username=None if config['mongodb']['mongodb_username'] == '' else config['mongodb']['mongodb_username'],
        password=None if config['mongodb']['mongodb_password'] == '' else config['mongodb']['mongodb_password'])[
        config['mongodb']['mongodb_db_name']]

ali_bucket = init_ali_oss()


def prepare(url):
    # get real html content
    # 利用session保存cookie信息，第一次请求会设置cookie类似
    # {'wzwsconfirm': 'ab3039756ba3ee041f7e68f634d28882', 'wzwsvtime': '1488938461'}，
    # 与js解析得到的cookie合起来才能通过验证
    r = requests.session()
    content = r.get(url, timeout=(10, 30)).content
    time.sleep(3)
    # 获取页面脚本内容
    re_script = re.search(r'<script type="text/javascript">(?P<script>.*)</script>', content.decode('utf-8'),
                          flags=re.DOTALL)
    # 用点匹配所有字符，用(?P<name>...)获取：https://docs.python.org/3/howto/regex.html#regex-howto
    script = re_script.group('script')
    script = script.replace('\r\n', '')
    # 在美化之前，去掉\r\n之类的字符才有更好的效果
    res = jsbeautifier.beautify(script)
    # 美化并一定程度解析js代码：https://github.com/beautify-web/js-beautify
    js_code_list = res.split('function')
    var_ = js_code_list[0]
    var_list = var_.split('\n')
    template_js = var_list[3]  # 依顺序获取，亦可用正则
    template_py = js2py.eval_js(template_js)
    # 将所有全局变量插入第一个函数变为局部变量并计算
    function1_js = 'function' + js_code_list[1]
    position = function1_js.index('{') + 1
    function1_js = function1_js[:position] + var_ + function1_js[position:]
    function1_py = js2py.eval_js(function1_js)
    cookie1 = function1_py(str(template_py))  # 结果类似'NA=='
    # 保存得到的第一个cookie
    cookies = {'wzwstemplate': cookie1}
    # 对第三个函数做类似操作
    function3_js = 'function' + js_code_list[3]
    position = function3_js.index('{') + 1
    function3_js = function3_js[:position] + var_ + function3_js[position:]
    function3_py = js2py.eval_js(function3_js)
    middle_var = function3_py()
    cookie2 = function1_py(middle_var)
    cookies['wzwschallenge'] = cookie2
    # 关于js代码中的document.cookie参见 https://developer.mozilla.org/zh-CN/docs/Web/API/Document/cookie
    dynamic_url = js2py.eval_js(var_list[0])

    # 利用新的cookie对提供的动态网址进行访问即是我们要达到的内容页面了
    r.cookies.update(cookies)
    rep3 = r.get(urljoin(url, dynamic_url), timeout=(10, 30))
    time.sleep(3)
    if not rep3:
        logger.error('网页请求错误 %s' % url)
        return
    return rep3


# 人民银行
def pbc_parse(url, data_id):
    logger.info('人民银行解析 Url to parse: %s' % url)
    r = prepare(url)
    if r is None:
        logger.error('网页请求错误')
        return
    content_soup = bs(r.text.encode(r.encoding).decode('utf-8'), 'lxml')
    table_content = content_soup.find(id='zwgk_pre').find_all('table')[-1]
    title = content_soup.find(id='zwgk_pre').find_all('table')[0].text

    if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
        oss_file_map = {
            'origin_url': url,
            'oss_file_origin_url': url,
            'origin_url_id': data_id,
            'oss_file_type': 'html',
            'oss_file_name': title,
            'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
            'parsed': False
        }
        response = db.parsed_data.insert_one(oss_file_map)
        file_id = response.inserted_id
        oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html', r.text.encode(r.encoding).decode('utf-8'))
        db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
    else:
        db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
        file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

    each_result_list = []
    table_list = table_to_list(table_content)
    if db.announcement.find({'oss_file_id': file_id}).count() > 0:
        logger.warning('人民银行解析 -- 数据已经存在')
        return
    for each_row in table_list[1:]:
        if len(each_row) == 6:
            result_map = {
                'announcementTitle': title,
                'announcementOrg': '人民银行',
                'announcementDate': format_date(each_row[5]),
                'announcementCode': each_row[1],
                'facts': each_row[2],
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': each_row[0],
                'punishmentBasement': '',
                'punishmentDecision': '我行对' + each_row[0] + '作出以下处罚：' + each_row[3],
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
        else:
            result_map = {
                'announcementTitle': title,
                'announcementOrg': '人民银行',
                'announcementDate': format_date(each_row[6]),
                'announcementCode': each_row[2],
                'facts': each_row[3],
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': each_row[1],
                'punishmentBasement': '',
                'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            }
        each_result_list.append(result_map)

    if len(each_result_list) > 0:
        logger.info('人民银行解析 -- 一共有%d条数据' % len(each_result_list))
        db.announcement.insert_many(each_result_list)
        logger.info('人民银行解析 -- 数据导入完成')
        db.parsed_data.update_one({'origin_url_id': data_id}, {'$set': {'parsed': True}})
        logger.info('人民银行解析 -- 修改parsed完成')
    else:
        logger.warning('人民银行解析 -- 解析未能完成')


# pdf
def pdf_to_json(file_link, file_name, origin_url, origin_data_id, city_info):
    pdf_link = file_link
    pdf_name = file_name.replace('.pdf', '')

    response = prepare(pdf_link)
    if response is None:
        logger.error('网页请求错误')
        return

    with open('./test/tmp.pdf', 'wb') as tmp_file:
        for chunk in response.iter_content(chunk_size=1024):
            if chunk:
                tmp_file.write(chunk)

    with open('./test/tmp.pdf', 'rb') as pdf_file:
        pdf_content = pdf_file.read()

    pdf = pdfplumber.open('./test/tmp.pdf')
    pdf_text = ''
    for i in range(len(pdf.pages)):
        if pdf.pages[i].extract_text() is not None:
            pdf_text += pdf.pages[i].extract_text()
    pdf_text = pdf_text.strip()
    if get_chinese_proportion(pdf_text)[0] >= 0.2:
        if db.parsed_data.find({'origin_url': origin_url, 'oss_file_origin_url': pdf_link}).count() == 0:
            oss_file_map = {
                'origin_url': origin_url,
                'oss_file_origin_url': pdf_link,
                'origin_url_id': origin_data_id,
                'oss_file_type': 'pdf',
                'oss_file_name': pdf_name,
                'oss_file_content': pdf_content,
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + pdf_name + '.pdf', pdf_content)
            db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
        else:
            db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': origin_url, 'oss_file_origin_url': pdf_link})['_id']

        if db.announcement.find({'oss_file_id': file_id}).count() == 0:
            result_map_list = []
            announcement_count = 0
            with pdfplumber.open('./test/tmp.pdf') as pdf:
                page_count = len(pdf.pages)
                for pdf_page in range(page_count):
                    table_info = pdf.pages[pdf_page].extract_tables()
                    announcement_count += len(table_info)
                    for i in range(len(table_info)):
                        for each_row in table_info[i]:
                            if len(each_row) >= 7 and '序' not in each_row[0] and each_row[6] != '' and \
                                    '企业名称' not in each_row[0] and '日期' not in each_row[6] and \
                                    '......' not in each_row[6] and '……' not in each_row[6]:
                                publish_date = format_date(each_row[6].replace('\n', ''))
                                document_code = each_row[2].replace('\n', '').replace(' ', '').replace('\t',
                                                                                                       '').replace(
                                    '\n', '').strip()
                                result_map = {
                                    'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' + document_code + '）',
                                    'announcementOrg': '人民银行' + city_info,
                                    'announcementDate': publish_date,
                                    'announcementCode': document_code,
                                    'facts': each_row[3].replace('\n', ''),
                                    'defenseOpinion': '',
                                    'defenseResponse': '',
                                    'litigant': each_row[1].replace('\n', ''),
                                    'punishmentBasement': '',
                                    'punishmentDecision': '我行对' + each_row[1].replace('\n', '') +
                                                          '作出以下处罚：' + each_row[4].replace('\n', ''),
                                    'type': '行政处罚决定',
                                    'oss_file_id': file_id,
                                    'status': 'not checked'
                                }
                                result_map_list.append(result_map)
            if len(result_map_list) > 0:
                logger.info('人民银行' + city_info + '解析 -- 一共有%d条数据' % len(result_map_list))
                db.announcement.insert_many(result_map_list)
                logger.info('人民银行' + city_info + '解析 -- 数据导入完成')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('人民银行' + city_info + '解析 -- 修改parsed完成')
            else:
                logger.warning('人民银行' + city_info + '解析 -- 解析未能完成')
        else:
            logger.info('人民银行' + city_info + '解析 -- 数据已经存在')
    else:
        result_text, ocr_flag = pdf_ocr_to_table('./test/tmp.pdf')
        if db.parsed_data.find({'origin_url': origin_url, 'oss_file_origin_url': pdf_link}).count() == 0:
            oss_file_map = {
                'origin_url': origin_url,
                'oss_file_origin_url': pdf_link,
                'origin_url_id': origin_data_id,
                'oss_file_type': 'pdf',
                'oss_file_name': pdf_name,
                'oss_file_content': pdf_content,
                'parsed': False,
                'if_ocr': True,
                'ocr_result': '\n'.join(result_text)
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + pdf_name + '.pdf', pdf_content)
            db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
        else:
            db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
            if ocr_flag:
                db.parsed_data.update_one({'origin_url': origin_url,
                                           'oss_file_origin_url': pdf_link},
                                          {'$set': {'if_ocr': True, 'ocr_result': '\n'.join(result_text)}})
            file_id = db.parsed_data.find_one({'origin_url': origin_url, 'oss_file_origin_url': pdf_link})['_id']

        if db.announcement.find({'oss_file_id': file_id}).count() == 0:
            result_map_list = []
            for each_row in result_text:
                if str(each_row).startswith('序号') or str(each_row).startswith('企业名称') or \
                        str(each_row).startswith('序企业名称') or str(each_row).startswith('号企业名称') or \
                        str(each_row).startswith('行政处罚违法行为') or str(each_row).startswith('违法行为类型'):
                    continue
                try:
                    litigant = re.search(r'^\d*(.*?)(\(潮银\)罚字|涪银罚|抚银罚字|黔南银罚字|遵银综?罚字)',
                                         each_row).group(1).strip()
                    document_code = re.search(litigant.replace(r'[', r'\[').replace(r']', r'\]').
                                              replace(r'(', r'\(').replace(r')', r'\)') + r'(.*?号)',
                                              each_row).group(1).strip()
                    facts = re.search(document_code.replace(r'[', r'\[').replace(r']', r'\]').
                                      replace(r'(', r'\(').replace(r')', r'\)') + r'(.*?)(对该单位处以|警告|处以罚款|根据|'
                                                                                  r'2016年7月1日至2017年6月30日期间)',
                                      each_row).group(1).strip()
                    punishment_decision = re.search(facts.replace(r'[', r'\[').replace(r']', r'\]').
                                                    replace(r'(', r'\(').replace(r')', r'\)') +
                                                    r'(.*?)(中国人民银行)',
                                                    each_row).group(1).strip()
                    publish_date = re.search(r'(\d{4}年\d+月\d+日)(无)?$', each_row).group(1).strip()
                    if format_date(publish_date) != 'Bad date':
                        publish_date = format_date(publish_date)
                    result_map = {
                        'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                             document_code.replace('\n', '') + '）',
                        'announcementOrg': '人民银行' + city_info,
                        'announcementDate': publish_date,
                        'announcementCode': document_code,
                        'facts': facts,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': litigant,
                        'punishmentBasement': '',
                        'punishmentDecision': '我行对' + litigant + '作出以下处罚：' + punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
                except Exception as e:
                    logger.warning(e)
                    result_map = {
                        'announcementTitle': each_row,
                        'announcementOrg': '人民银行' + city_info,
                        'announcementDate': '',
                        'announcementCode': '',
                        'facts': '',
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': '',
                        'punishmentBasement': '',
                        'punishmentDecision': '',
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
                    continue
            if len(result_map_list) > 0:
                logger.info('人民银行' + city_info + '解析 -- 一共有%d条数据' % len(result_map_list))
                db.announcement.insert_many(result_map_list)
                logger.info('人民银行' + city_info + '解析 -- 数据导入完成')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('人民银行' + city_info + '解析 -- 修改parsed完成')
            else:
                logger.warning('人民银行' + city_info + '解析 -- 解析未能完成')
        else:
            logger.info('人民银行' + city_info + '解析 -- 数据已经存在')

    logger.info('删除TMP文件')
    os.remove('./test/tmp.pdf')


# docx or doc file
def docx_to_json(file_link, file_name, origin_url, origin_data_id, city_info):
    doc_link = file_link
    doc_name = file_name.replace('.docx', '').replace('.doc', '')

    link_type = doc_link.split('.')[-1]
    response = prepare(doc_link)
    if response is None:
        logger.error('网页请求错误')
        return

    with open('./test/tmp.' + link_type, 'wb') as tmp_file:
        for chunk in response.iter_content(chunk_size=1024):
            if chunk:
                tmp_file.write(chunk)

    if not os.path.exists('./test/tmp.docx'):
        shell_str = '/usr/local/bin/soffice --headless --convert-to docx ' + \
                    './test/tmp.' + link_type + ' --outdir ./test'
        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                   shell=True, stdout=subprocess.PIPE)
        process.communicate()

    with open('./test/tmp.docx', 'rb') as docx_file:
        docx_content = docx_file.read()

    if db.parsed_data.find({'origin_url': origin_url, 'oss_file_origin_url': doc_link}).count() == 0:
        oss_file_map = {
            'origin_url': origin_url,
            'oss_file_origin_url': doc_link,
            'origin_url_id': origin_data_id,
            'oss_file_type': 'docx',
            'oss_file_name': doc_name,
            'oss_file_content': docx_content,
            'parsed': False
        }
        response = db.parsed_data.insert_one(oss_file_map)
        file_id = response.inserted_id
        oss_add_file(ali_bucket, str(file_id) + '/' + doc_name + '.docx', docx_content)
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
    else:
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
        file_id = db.parsed_data.find_one({'origin_url': origin_url, 'oss_file_origin_url': doc_link})['_id']

    if db.announcement.find({'oss_file_id': file_id}).count() == 0:
        document = Document('./test/tmp.docx')
        result_map_list = []
        announcement_count = 0
        tables = document.tables
        for table in tables:
            announcement_count += len(table.rows)
            for row in table.rows:
                if '序号' in row.cells[0].text:
                    continue
                row_list = sorted(set([each_cell.text for each_cell in row.cells]),
                                  key=[each_cell.text for each_cell in row.cells].index)
                row_list = [format_text(kk) for kk in row_list]
                if len(row_list) == 8 and row_list[0] != '序号' and row_list[6] != '' and \
                        row_list[1] != '' and '作出行政处' not in row_list[6] and \
                        row_list[1] != row_list[6] and '业务部门' not in row_list[0]:
                    if doc_link == \
                            'http://ningbo.pbc.gov.cn/ningbo/127076/127098/127105/3523416/2018041913585811693.doc':
                        publish_date = '2018年4月19日'
                    else:
                        publish_date = format_date(row_list[6].replace('\n', ''))
                    if publish_date == '':
                        continue
                    result_map = {
                        'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                             row_list[2].replace('\n', '') + '）',
                        'announcementOrg': '人民银行' + city_info,
                        'announcementDate': publish_date,
                        'announcementCode': row_list[2].replace('\n', ''),
                        'facts': row_list[3],
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': row_list[1].replace('\n', ''),
                        'punishmentBasement': '',
                        'punishmentDecision': '我行对' + row_list[1].replace('\n', '') + '作出以下处罚：' + row_list[4],
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
                else:
                    if len(row_list) == 7 and row_list[0] != '序号' and row_list[6] != '' and \
                            row_list[1] != '' and '作出行政处' not in row_list[6] and \
                            row_list[1] != row_list[6] and '业务部门' not in row_list[0]:

                        publish_date = format_date(row_list[6].replace('\n', ''))
                        if publish_date == '':
                            continue
                        result_map = {
                            'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                                 row_list[2].replace('\n', '') + '）',
                            'announcementOrg': '人民银行' + city_info,
                            'announcementDate': publish_date,
                            'announcementCode': row_list[2].replace('\n', ''),
                            'facts': row_list[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': row_list[1].replace('\n', ''),
                            'punishmentBasement': '',
                            'punishmentDecision': '我行对' + row_list[1].replace('\n', '') + '作出以下处罚：' + row_list[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                    else:
                        if doc_link == \
                                'http://ningbo.pbc.gov.cn/ningbo/127076/127098/127105/3523416/2018041913585811693.doc' \
                                and len(row_list) == 7 and row_list[0] != '序号' and \
                                row_list[1] != '' and '作出行政处' not in row_list[6] and \
                                row_list[1] != row_list[6] and '业务部门' not in row_list[0]:
                            result_map_list.append({
                                'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                                     row_list[2].replace('\n', '') + '）',
                                'announcementOrg': '人民银行' + city_info,
                                'announcementDate': '2018年4月19日',
                                'announcementCode': row_list[2].replace('\n', ''),
                                'facts': row_list[3],
                                'defenseOpinion': '',
                                'defenseResponse': '',
                                'litigant': row_list[1].replace('\n', ''),
                                'punishmentBasement': '',
                                'punishmentDecision': '我行对' + row_list[1].replace('\n', '') + '作出以下处罚：' + row_list[4],
                                'type': '行政处罚决定',
                                'oss_file_id': file_id,
                                'status': 'not checked'
                            })
        if len(result_map_list) == 0:
            if not os.path.exists('./test/tmp.pdf'):
                shell_str = '/usr/local/bin/soffice --headless --convert-to pdf ' + \
                            './test/tmp.' + link_type + ' --outdir ./test'
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.communicate()

            with pdfplumber.open('./test/tmp.pdf') as pdf:
                page_count = len(pdf.pages)
                for pdf_page in range(page_count):
                    table_info = pdf.pages[pdf_page].extract_tables()
                    announcement_count += len(table_info)
                    for i in range(len(table_info)):
                        for each_row in table_info[i]:
                            if len(each_row) >= 7 and each_row[0] and '序' not in each_row[0] and \
                                    each_row[6] != '' and \
                                    '企业名称' not in each_row[0] and '日期' not in each_row[6] and \
                                    '......' not in each_row[6] and '……' not in each_row[6]:
                                publish_date = format_date(each_row[6].replace('\n', ''))
                                document_code = each_row[2].replace('\n', '').replace(' ', '').replace(
                                    '\t', '').replace('\n', '').strip()
                                result_map = {
                                    'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                                         document_code + '）',
                                    'announcementOrg': '人民银行' + city_info,
                                    'announcementDate': publish_date,
                                    'announcementCode': document_code,
                                    'facts': each_row[3].replace('\n', ''),
                                    'defenseOpinion': '',
                                    'defenseResponse': '',
                                    'litigant': each_row[1].replace('\n', ''),
                                    'punishmentBasement': '',
                                    'punishmentDecision': '我行对' + each_row[1].replace('\n', '') +
                                                          '作出以下处罚：' + each_row[4].replace('\n', ''),
                                    'type': '行政处罚决定',
                                    'oss_file_id': file_id,
                                    'status': 'not checked'
                                }
                                result_map_list.append(result_map)

        if origin_url == 'http://lanzhou.pbc.gov.cn/lanzhou/117067/117091/117098/3180806/index.html':
            result_map_list.append({
                'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                     '兰银罚字〔2016〕第3号' + '）',
                'announcementOrg': '人民银行' + city_info,
                'announcementDate': '2016年11月4日',
                'announcementCode': '兰银罚字〔2016〕第3号',
                'facts': '你行于2016年10月8日至2016年10月16日存款准备金账户透支，欠缴法定存款准备金1,696,016.59元。',
                'defenseOpinion': '',
                'defenseResponse': '',
                'litigant': '被处罚单位：甘肃西固金城村镇银行股份有限公司\n地址：兰州市西固区山丹街138号\n法定代表人（主要负责人）：高晓民',
                'punishmentBasement': '',
                'punishmentDecision': '根据《中华人民共和国中国人民银行法》第三十二条第一款第一项和《中华人民共和国商业银行法》第'
                                      '七十七条第三项规定以及《中国人民银行关于加强存款准备金管理的通知》（银发〔2004〕302号）第四条'
                                      '第一款规定：“……对于未按照中国人民银行规定交存存款准备金的金融机构，依法减轻处罚的，'
                                      '对其缴存存款准备金不足部分按每日万分之六的比例处以罚款。”， 处以人民币金额1017.61元罚款。',
                'type': '行政处罚决定',
                'oss_file_id': file_id,
                'status': 'not checked'
            })
        if len(result_map_list) > 0:
            logger.info('人民银行' + city_info + '解析 -- 一共有%d条数据' % len(result_map_list))
            db.announcement.insert_many(result_map_list)
            logger.info('人民银行' + city_info + '解析 -- 数据导入完成')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('人民银行' + city_info + '解析 -- 修改parsed完成')
        else:
            logger.warning('人民银行' + city_info + '解析 -- 解析未能完成')
    else:
        logger.info('人民银行' + city_info + '解析 -- 数据已经存在')

    logger.info('删除TMP文件')
    if os.path.exists('./test/tmp.doc'):
        os.remove('./test/tmp.doc')
    if os.path.exists('./test/tmp.docx'):
        os.remove('./test/tmp.docx')
    if os.path.exists('./test/tmp.pdf'):
        os.remove('./test/tmp.pdf')


# xls or xlsx file
def xlsx_to_json(file_link, file_name, origin_url, origin_data_id, city_info):
    xlsx_link = file_link
    xlsx_name = file_name.replace('.xlsx', '').replace('.xls', '')

    if 'pbc.gov.cn' not in xlsx_link:
        logger.warning('url has nothing to do with punishment ...')
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'ignored'}})
        logger.info('update pbc data success')
        return

    link_type = xlsx_link.split('.')[-1]
    response = prepare(xlsx_link)
    if response is None:
        logger.error('网页请求错误')
        return

    with open('./test/tmp.' + link_type, 'wb') as tmp_file:
        for chunk in response.iter_content(chunk_size=1024):
            if chunk:
                tmp_file.write(chunk)

    with open('./test/tmp.' + link_type, 'rb') as xlsx_file:
        xlsx_content = xlsx_file.read()

    if db.parsed_data.find({'origin_url': origin_url, 'oss_file_origin_url': xlsx_link}).count() == 0:
        oss_file_map = {
            'origin_url': origin_url,
            'oss_file_origin_url': xlsx_link,
            'origin_url_id': origin_data_id,
            'oss_file_type': link_type,
            'oss_file_name': xlsx_name,
            'oss_file_content': xlsx_content,
            'parsed': False
        }
        response = db.parsed_data.insert_one(oss_file_map)
        file_id = response.inserted_id
        oss_add_file(ali_bucket, str(file_id) + '/' + xlsx_name + '.' + link_type, xlsx_content)
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
    else:
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
        file_id = db.parsed_data.find_one({'origin_url': origin_url, 'oss_file_origin_url': xlsx_link})['_id']

    if db.announcement.find({'oss_file_id': file_id}).count() == 0:
        if city_info == '西安分行':
            excel_data = open_workbook('./test/tmp.' + link_type)
            sheet = excel_data.sheets()[0]
            result_map_list = []
            if sheet.ncols == 19:
                for i in range(sheet.nrows):
                    if sheet.cell(i, 0).value != '企业名称' and sheet.cell(i, 0).value != '' and \
                            sheet.cell(i, 15).value != '':
                        if sheet.cell(i, 15).value == '2018/411':
                            publish_date = '2018年4月11日'
                        else:
                            if sheet.cell(i, 15).ctype == 3 or sheet.cell(i, 15).ctype == 2:
                                publish_date = xldate_as_tuple(sheet.cell_value(i, 15), excel_data.datemode)
                                publish_date = str(publish_date[0]) + '年' + str(publish_date[1]) + '月' + str(
                                    publish_date[2]) + '日'
                            else:
                                publish_date = format_date(str(sheet.cell(i, 15).value))

                        prefix_list = ['企业名称：', '统一社会信用代码：', '组织机构代码：', '工商登记码：',
                                       '法定代表人姓名：', '法定代表人身份证号：', '税务登记号：']
                        litigant_list = [str(sheet.cell(i, 0).value), str(sheet.cell(i, 1).value),
                                         str(sheet.cell(i, 2).value), str(sheet.cell(i, 3).value),
                                         str(sheet.cell(i, 4).value), str(sheet.cell(i, 5).value),
                                         str(sheet.cell(i, 6).value)]
                        litigant = ''
                        for index, each_litigant in enumerate(litigant_list):
                            if each_litigant != '':
                                litigant += prefix_list[index] + each_litigant + '\n'

                        document_code = str(int(sheet.cell(i, 7).value)) if sheet.cell(i, 7).ctype == 2 else str(
                            sheet.cell(i, 7).value)

                        result_map = {
                            'announcementTitle': '中国人民银行西安分行行政处罚信息公示表（' + document_code + '）',
                            'announcementOrg': '人民银行西安分行',
                            'announcementDate': publish_date,
                            'announcementCode': document_code,
                            'facts': sheet.cell(i, 12).value,
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': litigant,
                            'punishmentBasement': '',
                            'punishmentDecision': '我行根据' + sheet.cell(i, 13).value + '作出以下处罚：' + sheet.cell(i,
                                                                                                            14).value,
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
            else:
                if sheet.ncols == 18:
                    for i in range(sheet.nrows):
                        if sheet.cell(i, 0).value != '企业名称' and sheet.cell(i, 0).value != '' and \
                                sheet.cell(i, 14).value != '':
                            if sheet.cell(i, 14).value == '2018/411':
                                publish_date = '2018年4月11日'
                            else:
                                if sheet.cell(i, 14).ctype == 3:
                                    publish_date = xldate_as_tuple(sheet.cell_value(i, 14), excel_data.datemode)
                                    publish_date = str(publish_date[0]) + '年' + str(publish_date[1]) + '月' + str(
                                        publish_date[2]) + '日'
                                else:
                                    if '/' in str(sheet.cell(i, 14).value):
                                        publish_date = str(sheet.cell(i, 14).value).split('/')
                                        publish_date = str(publish_date[0]) + '年' + str(publish_date[1]) + '月' + str(
                                            publish_date[2]) + '日'
                                    else:
                                        if '.' in str(sheet.cell(i, 14).value):
                                            publish_date = str(sheet.cell(i, 14).value).split('.')
                                            if len(publish_date) == 3:
                                                publish_date = str(publish_date[0]) + '年' + str(publish_date[1]) + '月' \
                                                               + str(publish_date[2]) + '日'
                                            else:
                                                publish_date = str(sheet.cell(i, 14).value)[:4] + '年' + str(
                                                    int(str(sheet.cell(i, 6).value)[4:6])) + '月' + str(
                                                    int(str(sheet.cell(i, 6).value)[6:8])) + '日'
                                        else:
                                            if '-' in str(sheet.cell(i, 14).value):
                                                publish_date = str(sheet.cell(i, 14).value).split('-')
                                                publish_date = str(publish_date[0]) + '年' + str(
                                                    publish_date[1]) + '月' + str(publish_date[2]) + '日'
                                            else:
                                                if '年' in str(sheet.cell(i, 14).value):
                                                    publish_date = str(sheet.cell(i, 14).value)
                                                else:
                                                    publish_date = str(sheet.cell(i, 14).value)[:4] + '年' + str(
                                                        int(str(sheet.cell(i, 6).value)[4:6])) + '月' + str(
                                                        int(str(sheet.cell(i, 6).value)[6:8])) + '日'

                            prefix_list = ['企业名称：', '统一社会信用代码：', '组织机构代码：', '工商登记码：', '法定代表人姓名：',
                                           '税务登记号：']

                            litigant_list = [str(sheet.cell(i, 0).value), str(sheet.cell(i, 1).value),
                                             str(sheet.cell(i, 2).value), str(sheet.cell(i, 3).value),
                                             str(sheet.cell(i, 4).value), str(sheet.cell(i, 5).value),
                                             str(sheet.cell(i, 6).value)]

                            litigant = ''
                            for index, each_litigant in enumerate(litigant_list):
                                if each_litigant != '':
                                    litigant += prefix_list[index] + each_litigant + '\n'

                            document_code = str(int(sheet.cell(i, 6).value)) if sheet.cell(i, 6).ctype == 2 else str(
                                sheet.cell(i, 6).value)

                            result_map = {
                                'announcementTitle': '中国人民银行西安分行行政处罚信息公示表（' + document_code + '）',
                                'announcementOrg': '人民银行西安分行',
                                'announcementDate': publish_date,
                                'announcementCode': document_code,
                                'facts': sheet.cell(i, 11).value,
                                'defenseOpinion': '',
                                'defenseResponse': '',
                                'litigant': litigant,
                                'punishmentBasement': '',
                                'punishmentDecision': '我行根据' + sheet.cell(i, 12).value + '作出以下处罚：'
                                                      + sheet.cell(i, 13).value,
                                'type': '行政处罚决定',
                                'oss_file_id': file_id,
                                'status': 'not checked'
                            }
                            result_map_list.append(result_map)
                else:
                    if sheet.ncols == 21:
                        for i in range(sheet.nrows):
                            if sheet.cell(i, 0).value != '法人及其他组织名称' and sheet.cell(i, 0).value != '' and \
                                    sheet.cell(i, 15).value != '':
                                if sheet.cell(i, 15).value == '2018/411':
                                    publish_date = '2018年4月11日'
                                else:
                                    if sheet.cell(i, 15).ctype == 3 or sheet.cell(i, 15).ctype == 2:
                                        publish_date = xldate_as_tuple(sheet.cell_value(i, 15), excel_data.datemode)
                                        publish_date = str(publish_date[0]) + '年' + str(publish_date[1]) + '月' + str(
                                            publish_date[2]) + '日'
                                    else:
                                        publish_date = format_date(str(sheet.cell(i, 15).value))

                                prefix_list = ['企业名称：', '统一社会信用代码：', '部门登记号：', '组织机构代码：',
                                               '税务登记号：',
                                               '法定代表人姓名：', '法定代表人身份证号：']
                                litigant_list = [str(sheet.cell(i, 0).value), str(sheet.cell(i, 1).value),
                                                 str(sheet.cell(i, 2).value), str(sheet.cell(i, 3).value),
                                                 str(sheet.cell(i, 4).value), str(sheet.cell(i, 5).value),
                                                 str(sheet.cell(i, 6).value)]
                                litigant = ''
                                for index, each_litigant in enumerate(litigant_list):
                                    if each_litigant != '':
                                        litigant += prefix_list[index] + each_litigant + '\n'

                                document_code = str(int(sheet.cell(i, 7).value)) if sheet.cell(i,
                                                                                               7).ctype == 2 else str(
                                    sheet.cell(i, 7).value)

                                result_map = {
                                    'announcementTitle': '中国人民银行西安分行行政处罚信息公示表（' + document_code + '）',
                                    'announcementOrg': '人民银行西安分行',
                                    'announcementDate': publish_date,
                                    'announcementCode': document_code,
                                    'facts': sheet.cell(i, 12).value,
                                    'defenseOpinion': '',
                                    'defenseResponse': '',
                                    'litigant': litigant,
                                    'punishmentBasement': '',
                                    'punishmentDecision': '我行根据' + sheet.cell(i, 13).value + '作出以下处罚：' +
                                                          sheet.cell(i, 14).value,
                                    'type': '行政处罚决定',
                                    'oss_file_id': file_id,
                                    'status': 'not checked'
                                }
                                result_map_list.append(result_map)
                    else:
                        if sheet.ncols == 24:
                            for i in range(sheet.nrows):
                                if sheet.cell(i, 0).value != '法人及其他组织名称' and sheet.cell(i, 0).value != '' and \
                                        sheet.cell(i, 16).value != '':
                                    if sheet.cell(i, 16).value == '2018/411':
                                        publish_date = '2018年4月11日'
                                    else:
                                        if sheet.cell(i, 16).ctype == 3 or sheet.cell(i, 16).ctype == 2:
                                            try:
                                                publish_date = \
                                                    xldate_as_tuple(sheet.cell_value(i, 16), excel_data.datemode)
                                                publish_date = str(publish_date[0]) + '年' + str(
                                                    publish_date[1]) + '月' + str(
                                                    publish_date[2]) + '日'
                                            except Exception as e:
                                                logger.warning(e)
                                                publish_date = str(int(sheet.cell_value(i, 16)))[:4] + '年' + str(
                                                    int(sheet.cell_value(i, 16)))[4:6] + '月' + str(
                                                    int(sheet.cell_value(i, 16)))[6:8] + '日'
                                                publish_date = format_date(publish_date)
                                        else:
                                            publish_date = format_date(str(sheet.cell(i, 16).value))

                                    prefix_list = ['企业名称：', '统一社会信用代码：', '部门登记号：', '组织机构代码：',
                                                   '税务登记号：',
                                                   '法定代表人姓名：', '法定代表人身份证号：']
                                    litigant_list = [str(sheet.cell(i, 0).value), str(sheet.cell(i, 1).value),
                                                     str(sheet.cell(i, 2).value), str(sheet.cell(i, 3).value),
                                                     str(sheet.cell(i, 4).value), str(sheet.cell(i, 5).value),
                                                     str(sheet.cell(i, 6).value)]
                                    litigant = ''
                                    for index, each_litigant in enumerate(litigant_list):
                                        if each_litigant != '':
                                            litigant += prefix_list[index] + each_litigant + '\n'
                                    litigant = litigant.strip()

                                    document_code = str(int(sheet.cell(i, 7).value)) if \
                                        sheet.cell(i, 7).ctype == 2 else str(sheet.cell(i, 7).value)

                                    result_map = {
                                        'announcementTitle': '中国人民银行西安分行行政处罚信息公示表（' + document_code + '）',
                                        'announcementOrg': '人民银行西安分行',
                                        'announcementDate': publish_date,
                                        'announcementCode': document_code,
                                        'facts': sheet.cell(i, 9).value,
                                        'defenseOpinion': '',
                                        'defenseResponse': '',
                                        'litigant': litigant,
                                        'punishmentBasement': '',
                                        'punishmentDecision': '我行根据' + sheet.cell(i, 10).value + '作出以下处罚：' +
                                                              sheet.cell(i, 12).value,
                                        'type': '行政处罚决定',
                                        'oss_file_id': file_id,
                                        'status': 'not checked'
                                    }
                                    result_map_list.append(result_map)
            if len(result_map_list) > 0:
                logger.info('人民银行' + city_info + '解析 -- 一共有%d条数据' % len(result_map_list))
                db.announcement.insert_many(result_map_list)
                logger.info('人民银行' + city_info + '解析 -- 数据导入完成')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('人民银行' + city_info + '解析 -- 修改parsed完成')
            else:
                logger.warning('人民银行' + city_info + '解析 -- 解析未能完成')
        else:
            excel_data = open_workbook('./test/tmp.' + link_type)
            if origin_url in \
                    ['http://jinan.pbc.gov.cn/jinan/120967/120985/120994/3447690/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/3417662/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/3417657/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/3387814/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/3337132/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/3264550/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/2994420/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/2992505/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/3585576/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/3564465/index.html',
                     'http://jinan.pbc.gov.cn/jinan/120967/120985/120994/3564460/index.html',
                     'http://qingdao.pbc.gov.cn/qingdao/resource/cms/2017/05/中国人民银行即墨市支行行政处罚信息公示表（2017年4月21日）.xls',
                     'http://qingdao.pbc.gov.cn/qingdao/resource/cms/2017/04/中国人民银行即墨市支行行政处罚信息公示表（2017年4月11日）.xls',
                     'http://qingdao.pbc.gov.cn/qingdao/resource/cms/2017/04/中国人民银行即墨市支行行政处罚信息公示表（2017年3月29日）.xls',
                     'http://qingdao.pbc.gov.cn/qingdao/resource/cms/2016/08/中国人民银行即墨市支行行政处罚信息公示表（2016年5月9日）.xls',
                     'http://qingdao.pbc.gov.cn/qingdao/resource/cms/2016/08/中国人民银行即墨市支行行政处罚信息公示表（2016年4月29日）.xls',
                     'http://qingdao.pbc.gov.cn/qingdao/resource/cms/2016/08/中国人民银行即墨市支行行政处罚信息公示表（2016年4月13日）.xls',
                     'http://qingdao.pbc.gov.cn/qingdao/resource/cms/2016/08/中国人民银行即墨市支行行政处罚信息公示表（2016年3月28日）.xls',
                     'http://qingdao.pbc.gov.cn/qingdao/resource/cms/2016/08/中国人民银行即墨市支行行政处罚信息公示表（2016年2月14日）.xls',
                     'http://qingdao.pbc.gov.cn/qingdao/resource/cms/2016/08/中国人民银行即墨市支行行政处罚信息公示表（2016年1月4日）.xls'
                     ]:
                sheet = excel_data.sheets()[1]
            else:
                if origin_url in ['http://taiyuan.pbc.gov.cn/taiyuan/133960/133981/133988/3740523/index.html']:
                    sheet = excel_data.sheets()[2]
                else:
                    sheet = excel_data.sheets()[0]
            result_map_list = []
            for i in range(sheet.nrows):
                if '许可机关' in str(sheet.cell(i, 6).value):
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city_info + '解析 -- 修改parsed完成')
                    return
                if sheet.cell(i, 0).value != '序号' and sheet.cell(i, 6).value != '' and sheet.cell(i, 6).value != '　' \
                        and '报送时间' not in str(sheet.cell(i, 6).value):
                    if sheet.cell(i, 6).ctype == 3 or sheet.cell(i, 6).ctype == 2:
                        try:
                            publish_date = xldate_as_tuple(sheet.cell_value(i, 6), excel_data.datemode)
                            publish_date = str(publish_date[0]) + '年' + str(publish_date[1]) + '月' + str(
                                publish_date[2]) + '日'
                        except Exception as e:
                            publish_date = str(sheet.cell_value(i, 6)).split('.')[0]
                            publish_date = publish_date[:4] + '年' + publish_date[4:6] + '月' + publish_date[6:8] + '日'

                    else:
                        publish_date = format_date(str(sheet.cell(i, 6).value))

                    result_map = {
                        'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                             sheet.cell(i, 2).value.replace('\n', '') + '）',
                        'announcementOrg': '人民银行' + city_info,
                        'announcementDate': publish_date,
                        'announcementCode': sheet.cell(i, 2).value.replace('\n', ''),
                        'facts': sheet.cell(i, 1).value + sheet.cell(i, 3).value,
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': sheet.cell(i, 1).value,
                        'punishmentBasement': '',
                        'punishmentDecision': '我行对' + sheet.cell(i, 1).value + '作出以下处罚：' + sheet.cell(i, 4).value,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            if len(result_map_list) > 0:
                logger.info('人民银行' + city_info + '解析 -- 一共有%d条数据' % len(result_map_list))
                db.announcement.insert_many(result_map_list)
                logger.info('人民银行' + city_info + '解析 -- 数据导入完成')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('人民银行' + city_info + '解析 -- 修改parsed完成')
            else:
                logger.warning('人民银行' + city_info + '解析 -- 解析未能完成')
    else:
        logger.info('人民银行' + city_info + '解析 -- 数据已经存在')

    logger.info('删除TMP文件')
    if os.path.exists('./test/tmp.xls'):
        os.remove('./test/tmp.xls')

    if os.path.exists('./test/tmp.xlsx'):
        os.remove('./test/tmp.xlsx')


# et file
def et_to_json(file_link, file_name, origin_url, origin_data_id, city_info):
    et_link = file_link
    et_name = file_name.replace('.ett', '').replace('.et', '')

    link_type = et_link.split('.')[-1]
    response = prepare(et_link)

    with open('./test/tmp.' + link_type, 'wb') as tmp_file:
        for chunk in response.iter_content(chunk_size=1024):
            if chunk:
                tmp_file.write(chunk)

    if not os.path.exists('./test/tmp.xlsx'):
        shell_str = '/usr/local/bin/soffice --headless --convert-to xlsx ' + \
                    './test/tmp.' + link_type + ' --outdir ./test/'
        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                   shell=True, stdout=subprocess.PIPE)
        process.communicate()

    with open('./test/tmp.xlsx', 'rb') as xlsx_file:
        xlsx_content = xlsx_file.read()

    if db.parsed_data.find({'origin_url': origin_url, 'oss_file_origin_url': et_link}).count() == 0:
        oss_file_map = {
            'origin_url': origin_url,
            'oss_file_origin_url': et_link,
            'origin_url_id': origin_data_id,
            'oss_file_type': 'xlsx',
            'oss_file_name': et_name,
            'oss_file_content': xlsx_content,
            'parsed': False
        }
        response = db.parsed_data.insert_one(oss_file_map)
        file_id = response.inserted_id
        oss_add_file(ali_bucket, str(file_id) + '/' + et_name + '.xlsx', xlsx_content)
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
    else:
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
        file_id = db.parsed_data.find_one({'origin_url': origin_url, 'oss_file_origin_url': et_link})['_id']

    if db.announcement.find({'oss_file_id': file_id}).count() == 0:
        excel_data = open_workbook('./test/tmp.xlsx')
        sheet = excel_data.sheets()[0]
        result_map_list = []
        for i in range(sheet.nrows):
            if sheet.cell(i, 0).value != '序号' and sheet.cell(i, 6).value != '' and sheet.cell(i, 0).value != '' and \
                    '样  表' not in str(sheet.cell(i, 6).value) and '王晓兵' not in str(sheet.cell(i, 6).value) and \
                    '办公室意见' not in str(sheet.cell(i, 6).value) and '按规定公示' not in str(sheet.cell(i, 6).value) and \
                    '负责人签字' not in str(sheet.cell(i, 6).value) and '彭艳玲' not in str(sheet.cell(i, 6).value) and \
                    '郭涛' not in str(sheet.cell(i, 6).value):
                if sheet.cell(i, 6).ctype == 3:
                    publish_date = xldate_as_tuple(sheet.cell_value(i, 6), excel_data.datemode)
                    publish_date = str(publish_date[0]) + '年' + str(publish_date[1]) + '月' + str(
                        publish_date[2]) + '日'
                else:
                    publish_date = format_date(sheet.cell(i, 6).value)
                result_map = {
                    'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                         sheet.cell(i, 2).value.replace('\n', '') + '）',
                    'announcementOrg': '人民银行' + city_info,
                    'announcementDate': publish_date,
                    'announcementCode': sheet.cell(i, 2).value.replace('\n', ''),
                    'facts': sheet.cell(i, 3).value,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': sheet.cell(i, 1).value,
                    'punishmentBasement': '',
                    'punishmentDecision': '我行对' + sheet.cell(i, 1).value + '作出以下处罚：' + sheet.cell(i, 4).value,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
        if len(result_map_list) > 0:
            logger.info('人民银行' + city_info + '解析 -- 一共有%d条数据' % len(result_map_list))
            db.announcement.insert_many(result_map_list)
            logger.info('人民银行' + city_info + '解析 -- 数据导入完成')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('人民银行' + city_info + '解析 -- 修改parsed完成')
        else:
            logger.warning('人民银行' + city_info + '解析 -- 解析未能完成')
    else:
        logger.info('人民银行' + city_info + '解析 -- 数据已经存在')

    logger.info('删除TMP文件')
    if os.path.exists('./test/tmp.et'):
        os.remove('./test/tmp.et')
    if os.path.exists('./test/tmp.ett'):
        os.remove('./test/tmp.ett')
    if os.path.exists('./test/tmp.xlsx'):
        os.remove('./test/tmp.xlsx')


# wps file
def wps_to_json(file_link, file_name, origin_url, origin_data_id, city_info):
    wps_link = file_link
    wps_name = file_name.replace('.wps', '')

    link_type = wps_link.split('.')[-1]
    response = prepare(wps_link)

    with open('./test/tmp.' + link_type, 'wb') as tmp_file:
        for chunk in response.iter_content(chunk_size=1024):
            if chunk:
                tmp_file.write(chunk)

    if not os.path.exists('./test/tmp.docx'):
        shell_str = '/usr/local/bin/soffice --headless --convert-to docx ' + \
                    './test/tmp.' + link_type + ' --outdir ./test/'
        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                   shell=True, stdout=subprocess.PIPE)
        process.communicate()

    with open('./test/tmp.docx', 'rb') as docx_file:
        docx_content = docx_file.read()

    if db.parsed_data.find({'origin_url': origin_url, 'oss_file_origin_url': wps_link}).count() == 0:
        oss_file_map = {
            'origin_url': origin_url,
            'oss_file_origin_url': wps_link,
            'origin_url_id': origin_data_id,
            'oss_file_type': 'docx',
            'oss_file_name': wps_name,
            'oss_file_content': docx_content,
            'parsed': False
        }
        response = db.parsed_data.insert_one(oss_file_map)
        file_id = response.inserted_id
        oss_add_file(ali_bucket, str(file_id) + '/' + wps_name + '.docx', docx_content)
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
    else:
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
        file_id = db.parsed_data.find_one({'origin_url': origin_url, 'oss_file_origin_url': wps_link})['_id']

    if db.announcement.find({'oss_file_id': file_id}).count() == 0:
        document = Document('./test/tmp.docx')
        result_map_list = []
        tables = document.tables
        for table in tables:
            for row in table.rows:
                row_list = sorted(set([each_cell.text for each_cell in row.cells]),
                                  key=[each_cell.text for each_cell in row.cells].index)
                if len(row_list) == 8 and row_list[0] != '序号' and row_list[6] != '' and \
                        '作出行政处罚' not in row_list[6] and '同意' not in row_list[6] and \
                        '郑智彬' not in row_list[6]:
                    publish_date = format_date(row_list[6].replace('\n', ''))
                    if publish_date == '':
                        continue
                    result_map = {
                        'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                             row_list[2].replace('\n', '') + '）',
                        'announcementOrg': '人民银行' + city_info,
                        'announcementDate': publish_date,
                        'announcementCode': row_list[2].replace('\n', ''),
                        'facts': row_list[3],
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': row_list[1].replace('\n', ''),
                        'punishmentBasement': '',
                        'punishmentDecision': '我行对' + row_list[1].replace('\n', '') + row_list[3] + '作出以下处罚：' +
                                              row_list[4],
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
        if len(result_map_list) == 0:
            if not os.path.exists('./test/tmp.pdf'):
                shell_str = '/usr/local/bin/soffice --headless --convert-to pdf ' + \
                            './test/tmp.' + link_type + ' --outdir ./test'
                process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                           shell=True, stdout=subprocess.PIPE)
                process.communicate()

            with pdfplumber.open('./test/tmp.pdf') as pdf:
                page_count = len(pdf.pages)
                for pdf_page in range(page_count):
                    table_info = pdf.pages[pdf_page].extract_tables()
                    for i in range(len(table_info)):
                        for each_row in table_info[i]:
                            if len(each_row) >= 7 and '序' not in each_row[0] and each_row[6] != '' and \
                                    '企业名称' not in each_row[0] and '日期' not in each_row[6] and \
                                    '......' not in each_row[6] and '……' not in each_row[6]:
                                publish_date = format_date(each_row[6].replace('\n', ''))
                                document_code = each_row[2].replace('\n', '').replace(' ', '').replace(
                                    '\t', '').replace('\n', '').strip()
                                result_map = {
                                    'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                                         document_code + '）',
                                    'announcementOrg': '人民银行' + city_info,
                                    'announcementDate': publish_date,
                                    'announcementCode': document_code,
                                    'facts': each_row[3].replace('\n',
                                                                 ''),
                                    'defenseOpinion': '',
                                    'defenseResponse': '',
                                    'litigant': each_row[1].replace('\n', ''),
                                    'punishmentBasement': '',
                                    'punishmentDecision': '我行对' + each_row[1].replace('\n', '') +
                                                          '作出以下处罚：' + each_row[4].replace('\n', ''),
                                    'type': '行政处罚决定',
                                    'oss_file_id': file_id,
                                    'status': 'not checked'
                                }
                                result_map_list.append(result_map)
        if len(result_map_list) > 0:
            logger.info('人民银行' + city_info + '解析 -- 一共有%d条数据' % len(result_map_list))
            db.announcement.insert_many(result_map_list)
            logger.info('人民银行' + city_info + '解析 -- 数据导入完成')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('人民银行' + city_info + '解析 -- 修改parsed完成')
        else:
            logger.warning('人民银行' + city_info + '解析 -- 解析未能完成')
    else:
        logger.info('人民银行' + city_info + '解析 -- 数据已经存在')

    logger.info('删除TMP文件')
    if os.path.exists('./test/tmp.wps'):
        os.remove('./test/tmp.wps')
    if os.path.exists('./test/tmp.docx'):
        os.remove('./test/tmp.docx')
    if os.path.exists('./test/tmp.pdf'):
        os.remove('./test/tmp.pdf')


# tif file
def tif_to_json(file_link, file_name, origin_url, origin_data_id, city_info):
    tif_link = file_link
    link_type = file_link.split('.')[-1]
    file_name = file_name.replace('.' + link_type, '')

    response = prepare(tif_link)

    with open('./test/tmp.' + link_type, 'wb') as tmp_file:
        for chunk in response.iter_content(chunk_size=1024):
            if chunk:
                tmp_file.write(chunk)

    if not os.path.exists('./test/tmp.pdf'):
        shell_str = 'img2pdf ./test/tmp.' + link_type + ' -o ./test/tmp.pdf'
        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                   shell=True, stdout=subprocess.PIPE)
        process.communicate()

    result_text, ocr_flag = pdf_ocr_to_table('./test/tmp.pdf')

    with open('./test/tmp.pdf', 'rb') as pdf_file:
        pdf_content = pdf_file.read()

    if db.parsed_data.find({'origin_url': origin_url, 'oss_file_origin_url': tif_link}).count() == 0:
        oss_file_map = {
            'origin_url': origin_url,
            'oss_file_origin_url': tif_link,
            'origin_url_id': origin_data_id,
            'oss_file_type': 'pdf',
            'oss_file_name': file_name,
            'oss_file_content': pdf_content,
            'parsed': False,
            'if_ocr': True,
            'ocr_result': '\n'.join(result_text)
        }
        response = db.parsed_data.insert_one(oss_file_map)
        file_id = response.inserted_id
        oss_add_file(ali_bucket, str(file_id) + '/' + file_name + '.pdf', pdf_content)
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
    else:
        db.pbc_data.update_one({'_id': origin_data_id}, {'$set': {'status': 'parsed'}})
        if ocr_flag:
            db.parsed_data.update_one({'origin_url': origin_url,
                                       'oss_file_origin_url': tif_link},
                                      {'$set': {'if_ocr': True, 'ocr_result': '\n'.join(result_text)}})
        file_id = db.parsed_data.find_one({'origin_url': origin_url, 'oss_file_origin_url': tif_link})['_id']

    if db.announcement.find({'oss_file_id': file_id}).count() == 0:
        result_map_list = []
        for each_row in result_text:
            if str(each_row).startswith('序号') or str(each_row).startswith('企业名称') or \
                    str(each_row).startswith('序企业名称') or str(each_row).startswith('号企业名称'):
                continue
            try:
                litigant = re.search(r'^\d*(.*?)(银罚|沪钅艮罚|沪银|沪钅银|户银罚|沪钜灵罚|沪灵罚|沪针是罚|沪针灵罚|沪罚|沪是罚|'
                                     r'七银罚字|佳银罚字)',
                                     each_row).group(1).strip()
                document_code = re.search(litigant.replace(r'[', r'\[').replace(r']', r'\]').
                                          replace(r'(', r'\(').replace(r')', r'\)') + r'(.*?号)',
                                          each_row).group(1).strip()
                facts = re.search(document_code.replace(r'[', r'\[').replace(r']', r'\]').
                                  replace(r'(', r'\(').replace(r')', r'\)') + r'(.*?)(罚款|款[\d.,]*元|警告)',
                                  each_row).group(1).strip()
                punishment_decision = re.search(facts.replace(r'[', r'\[').replace(r']', r'\]').
                                                replace(r'(', r'\(').replace(r')', r'\)') +
                                                r'(.*?)(中国民银|中国人民银|中国人飞民|中国银|国人民银行|中国月民银行|中国丿民银行|'
                                                r'中国人人民银|国民银行|中国1人民银行|中压1?人民银行|中人民银行|中日1?人民银行|'
                                                r'中国民民银行|中国人国银行|中国人银行|中国人号银行|中国人国号银行|中国人是银行|'
                                                r'中国人只银行|中国人民号银行|中国人瓦银行|中国飞民银行|国人民银亍|中国人目民?银行|'
                                                r'中国1人民银|人民银行上海分行|中国口人民银彳|中国1人民银|民银行|国人气民银行|'
                                                r'国气民银行|、民银行|国、民银行|国丿民银行|中国丿银行|中国丿人民银行|中国飞民银|'
                                                r'国月行民银行|中国月民银|国民银上海分行|中国、民银行|中国气民银行|中国己银行)',
                                                each_row).group(1).strip()
                publish_date = re.search(r'([\d-]*)$', each_row).group(1).strip()
                if format_date(publish_date) != 'Bad date':
                    publish_date = format_date(publish_date)
                result_map = {
                    'announcementTitle': '中国人民银行' + city_info + '行政处罚信息公示表（' +
                                         document_code.replace('\n', '') + '）',
                    'announcementOrg': '人民银行' + city_info,
                    'announcementDate': publish_date,
                    'announcementCode': document_code,
                    'facts': facts,
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': litigant,
                    'punishmentBasement': '',
                    'punishmentDecision': '我行对' + litigant + '作出以下处罚：' + punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
            except Exception as e:
                result_map = {
                    'announcementTitle': each_row,
                    'announcementOrg': '人民银行' + city_info,
                    'announcementDate': '',
                    'announcementCode': '',
                    'facts': '',
                    'defenseOpinion': '',
                    'defenseResponse': '',
                    'litigant': '',
                    'punishmentBasement': '',
                    'punishmentDecision': '',
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                result_map_list.append(result_map)
                continue
        if len(result_map_list) > 0:
            logger.info('人民银行' + city_info + '解析 -- 一共有%d条数据' % len(result_map_list))
            db.announcement.insert_many(result_map_list)
            logger.info('人民银行' + city_info + '解析 -- 数据导入完成')
            db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
            logger.info('人民银行' + city_info + '解析 -- 修改parsed完成')
        else:
            logger.warning('人民银行' + city_info + '解析 -- 解析未能完成')
    else:
        logger.info('人民银行' + city_info + '解析 -- 数据已经存在')

    logger.info('删除TMP文件')
    if os.path.exists('./test/tmp.tif'):
        os.remove('./test/tmp.tif')
    if os.path.exists('./test/tmp.bmp'):
        os.remove('./test/tmp.bmp')
    if os.path.exists('./test/tmp.jpg'):
        os.remove('./test/tmp.jpg')
    if os.path.exists('./test/tmp.pdf'):
        os.remove('./test/tmp.pdf')

    return


def parse_links(links_list, url, data_id, city_info):
    for each_link in links_list:
        file_name = each_link.text.strip()
        if '行政许可信息公示表' in file_name or '行政处罚项目表' in file_name:
            continue
        if each_link.attrs.get('href', '') == '':
            continue
        real_url = urljoin(url, each_link.attrs['href'])
        file_type = real_url.split('.')[-1].lower()
        if file_type in ['doc', 'docx']:
            docx_to_json(real_url, file_name, url, data_id, city_info)
        elif file_type in ['xls', 'xlsx']:
            xlsx_to_json(real_url, file_name, url, data_id, city_info)
        elif file_type == 'pdf':
            pdf_to_json(real_url, file_name, url, data_id, city_info)
        elif file_type == 'wps':
            wps_to_json(real_url, file_name, url, data_id, city_info)
        elif file_type in ['et', 'ett']:
            et_to_json(real_url, file_name, url, data_id, city_info)
        elif file_type in ['tif', 'jpg', 'bmp']:
            tif_to_json(real_url, file_name, url, data_id, city_info)
        else:
            logger.warning(real_url + ' Not parsed!!')


# 地方人民银行
def local_pbc_parse(url, city, data_id):
    logger.info('人民银行' + city + ' Url to parse: %s' % url)
    r = prepare(url)
    if r is None:
        logger.error('网页请求错误')
        return
    if r.status_code == 404:
        logger.warning('url is 404 ...')
        db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'ignored'}})
        logger.info('update pbc data success')
        return

    r.encoding = r.apparent_encoding
    content_soup = bs(r.text, 'lxml')
    if content_soup.find(id='content_right'):
        content_text = get_content_text(content_soup.find(id='content_right'))
        content_text_first_line = content_text.split('\n')[1]
    else:
        if content_soup.find(id='zwgk_pre'):
            content_text = get_content_text(content_soup.find(id='zwgk_pre'))
            content_text_first_line = content_text.split('\n')[1]
        else:
            content_text = ''
            content_text_first_line = ''

    if '行政许可信息公示表' in content_text_first_line or '中国人民银行行政处罚文书' in content_text_first_line \
            or '中国人民银行行政处罚程序规定' in content_text_first_line or '行政处罚目录' in content_text_first_line \
            or '违反存款准备金管理情况处罚' in content_text_first_line or '行政处罚-金融市场业务处罚' in content_text_first_line \
            or '意见告知书格式' in content_text_first_line or '对代理国库、国库经收处的检查及处罚' in content_text_first_line \
            or ('反洗钱行政处罚' in content_text_first_line and '反洗钱行政处罚信息公示表' not in content_text_first_line) \
            or '行政处罚流程图' in content_text_first_line or '行政处罚项目' in content_text_first_line \
            or '行政处罚事项目录表' in content_text_first_line or '行政处罚依据一览' in content_text_first_line \
            or '行政处罚事项情况表' in content_text_first_line or '金融违法行为处罚办法' in content_text_first_line \
            or '行政处罚实施细则' in content_text_first_line or '中华人民共和国行政处罚法' in content_text_first_line \
            or '人民币银行结算账户行政处罚' in content_text_first_line \
            or '出票人签发空头支票或签发与其预留签章不符支票的行为' in content_text_first_line \
            or '行政处罚程序实施细则' in content_text_first_line \
            or '行政处罚操作流程' in content_text_first_line \
            or '行政处罚细则' in content_text_first_line \
            or '行政处罚权力运行流程图' in content_text_first_line \
            or '行政处罚程序' in content_text_first_line \
            or '行政处罚职权目录表' in content_text_first_line \
            or '行政许可信息' in content_text_first_line \
            or '行政许可公示信息' in content_text_first_line:
        if url != 'http://shanghai.pbc.gov.cn/fzhshanghai/113577/114832/114918/3196486/index.html':
            logger.warning('url has nothing to do with punishment ...')
            db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'ignored'}})
            logger.info('update pbc data success')
            return

    if city == '上海总部' or city == '上海分行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='content_right').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
            result_map_list = []
            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                for each_row in table_list:
                    if len(each_row) > 4 and each_row[0] != '序号' and each_row[6] != '' \
                            and each_row[0] != '序\n号':
                        publish_date = format_date(each_row[6])
                        result_map = {
                            'announcementTitle': '中国人民银行上海分行行政处罚信息公示表（' + each_row[2] + '）',
                            'announcementOrg': '人民银行上海分行',
                            'announcementDate': publish_date,
                            'announcementCode': each_row[2],
                            'facts': each_row[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1],
                            'punishmentBasement': '',
                            'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '天津分行':
        if '.doc' in url:
            docx_to_json(url, url.split('/')[-1], url, data_id, city)
        else:
            if '行政许可信息公示表' in content_text or '行政处罚项目表' in content_text:
                logger.warning('url has nothing to do with punishment ...')
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                logger.info('update pbc data success')
                return
            links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
            if len(links_list) > 0:
                parse_links(links_list, url, data_id, city)

    if city == '沈阳分行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='content_right').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
                result_map_list = []
                for each_row in table_list[1:]:
                    publish_date = format_date(each_row[6])
                    result_map = {
                        'announcementTitle': '中国人民银行沈阳分行行政处罚信息公示表（' + each_row[2] + '）',
                        'announcementOrg': '人民银行沈阳分行',
                        'announcementDate': publish_date,
                        'announcementCode': each_row[2],
                        'facts': each_row[3],
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': each_row[1],
                        'punishmentBasement': '',
                        'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '南京分行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='content_right').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
                result_map_list = []
                for each_row in table_list[1:]:
                    if each_row[6] != '':
                        publish_date = format_date(each_row[6])
                        result_map = {
                            'announcementTitle': '中国人民银行南京分行行政处罚信息公示表（' + each_row[2] + '）',
                            'announcementOrg': '人民银行南京分行',
                            'announcementDate': publish_date,
                            'announcementCode': each_row[2],
                            'facts': each_row[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1],
                            'punishmentBasement': '',
                            'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '济南分行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='content_right').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
                result_map_list = []
                for each_row in table_list[1:]:
                    publish_date = format_date(each_row[6])
                    result_map = {
                        'announcementTitle': '中国人民银行济南分行行政处罚信息公示表（' + each_row[2] + '）',
                        'announcementOrg': '人民银行济南分行',
                        'announcementDate': publish_date,
                        'announcementCode': each_row[2],
                        'facts': each_row[3],
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': each_row[1],
                        'punishmentBasement': '',
                        'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '武汉分行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '广州分行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '成都分行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '西安分行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '营业管理部（北京）':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='content_right').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
                result_map_list = []
                for each_row in table_list:
                    if '序号' in each_row[0]:
                        continue
                    if len(each_row) != 6:
                        publish_date = format_date(each_row[6])
                        result_map = {
                            'announcementTitle': '中国人民银行营业管理部（北京）行政处罚信息公示表（' + each_row[2] + '）',
                            'announcementOrg': '人民银行营业管理部（北京）',
                            'announcementDate': publish_date,
                            'announcementCode': each_row[2],
                            'facts': each_row[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1],
                            'punishmentBasement': '',
                            'punishmentDecision': '我部对' + each_row[1] + '作出以下处罚：' + each_row[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                    else:
                        publish_date = re.search(r'(\d{4}年\d+月\d+日)', title).group(1).strip()
                        result_map = {
                            'announcementTitle': '中国人民银行营业管理部（北京）行政处罚信息公示表（' + each_row[2] + '）',
                            'announcementOrg': '人民银行营业管理部（北京）',
                            'announcementDate': publish_date,
                            'announcementCode': each_row[2],
                            'facts': each_row[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1],
                            'punishmentBasement': '',
                            'punishmentDecision': '我部对' + each_row[1] + '作出以下处罚：' + each_row[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '重庆营业管理部':
        links_list = content_soup.find(id='zwgk_pre').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '石家庄中心支行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '太原中心支行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '呼和浩特中心支行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '长春中心支行':
        links_list = content_soup.find(id='zwgk_pre').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '哈尔滨中心支行':
        links_list = content_soup.find(id='zwgk_pre').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '杭州中心支行':
        links_list = content_soup.find(id='zwgk_pre').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '福州中心支行':
        links_list = content_soup.find(id='zwgk_pre').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0 and url != 'http://fuzhou.pbc.gov.cn/fuzhou/126805/126823/126830/2382883/index.html' \
                and url != 'http://fuzhou.pbc.gov.cn/fuzhou/126805/126823/126830/2926899/index.html':
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='zwgk_pre').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
                result_map_list = []
                for each_row in table_list:
                    if len(each_row) > 4 and each_row[0] != '序号' and each_row[6] != '':
                        publish_date = format_date(each_row[6])
                        result_map = {
                            'announcementTitle': '中国人民银行福州中心支行行政处罚信息公示表（' + each_row[2] + '）',
                            'announcementOrg': '人民银行福州中心支行',
                            'announcementDate': publish_date,
                            'announcementCode': each_row[2],
                            'facts': each_row[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1],
                            'punishmentBasement': '',
                            'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '合肥中心支行':
        links_list = content_soup.find(id='zwgk_pre').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0 and url != 'http://hefei.pbc.gov.cn/hefei/122364/122382/122389/3186858/index.html':
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='zwgk_pre').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
                result_map_list = []
                for each_row in table_list:
                    if len(each_row) > 4 and each_row[0] != '序号' and each_row[6] != '' \
                            and each_row[0] != '序\n号':
                        publish_date = format_date(each_row[6])
                        result_map = {
                            'announcementTitle': '中国人民银行合肥中心支行行政处罚信息公示表（' + each_row[2] + '）',
                            'announcementOrg': '人民银行合肥中心支行',
                            'announcementDate': publish_date,
                            'announcementCode': each_row[2],
                            'facts': each_row[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1],
                            'punishmentBasement': '',
                            'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '郑州中心支行':
        links_list = content_soup.find(id='zwgk_pre').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '长沙中心支行':
        links_list = content_soup.find(id='zwgk_pre').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '南昌中心支行':
        if '中国人民银行南昌中心支行有权对金融机构' in content_text:
            logger.warning('url has nothing to do with punishment ...')
            db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'ignored'}})
            logger.info('update pbc data success')
            return
        title = content_soup.find(id='zwgk_pre').find(class_='buleB16').text.strip()
        if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
            oss_file_map = {
                'origin_url': url,
                'oss_file_origin_url': url,
                'origin_url_id': data_id,
                'oss_file_type': 'html',
                'oss_file_name': title,
                'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                'parsed': False
            }
            response = db.parsed_data.insert_one(oss_file_map)
            file_id = response.inserted_id
            oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                         r.text.encode(r.encoding).decode('utf-8'))
            db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
        else:
            db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

        if db.announcement.find({'oss_file_id': file_id}).count() == 0:
            table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
            result_map_list = []
            for each_row in table_list:
                if len(each_row) > 4 and each_row[0] != '序号' and each_row[6] != '' \
                        and each_row[0] != '序\n号':
                    publish_date = format_date(each_row[6])
                    result_map = {
                        'announcementTitle': '中国人民银行南昌中心支行行政处罚信息公示表（' + each_row[2] + '）',
                        'announcementOrg': '人民银行南昌中心支行',
                        'announcementDate': publish_date,
                        'announcementCode': each_row[2],
                        'facts': each_row[3],
                        'defenseOpinion': '',
                        'defenseResponse': '',
                        'litigant': each_row[1],
                        'punishmentBasement': '',
                        'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    result_map_list.append(result_map)
            if len(result_map_list) > 0:
                logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                db.announcement.insert_many(result_map_list)
                logger.info('人民银行' + city + '解析 -- 数据导入完成')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('人民银行' + city + '解析 -- 修改parsed完成')
            else:
                logger.warning('人民银行' + city + '解析 -- 解析未能完成')
        else:
            logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '南宁中心支行':
        if '中国人民银行南宁中心支行有权对金融机构' in content_text:
            logger.warning('url has nothing to do with punishment ...')
            db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'ignored'}})
            logger.info('update pbc data success')
            return
        links_list = content_soup.find(id='zwgk_pre').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '海口中心支行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '昆明中心支行':
        if '云南省空头支票处罚情况' in content_text:
            logger.warning('url has nothing to do with punishment ...')
            db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'ignored'}})
            logger.info('update pbc data success')
            return
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='content_right').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
                result_map_list = []
                for each_row in table_list:
                    if len(each_row) > 4 and each_row[0] != '序号' and each_row[6] != '' \
                            and each_row[0] != '序\n号':
                        if each_row[6] == '人民银行楚雄州中心支行':
                            publish_date = '2017年12月12日'
                        else:
                            publish_date = format_date(each_row[6])
                        result_map = {
                            'announcementTitle': '中国人民银行昆明中心支行行政处罚信息公示表（' + each_row[2] + '）',
                            'announcementOrg': '人民银行昆明中心支行',
                            'announcementDate': publish_date,
                            'announcementCode': each_row[2],
                            'facts': each_row[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1],
                            'punishmentBasement': '',
                            'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '贵阳中心支行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '拉萨中心支行':
        if '中国人民银行拉萨中心支行有权对金融机构' in content_text:
            logger.warning('url has nothing to do with punishment ...')
            db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'ignored'}})
            logger.info('update pbc data success')
            return
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '兰州中心支行':
        if '.doc' in url:
            return docx_to_json(url, url.split('/')[-1], url, data_id, city)
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '西宁中心支行':
        if '.doc' in url:
            return docx_to_json(url, url.split('/')[-1], url, data_id, city)
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='content_right').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
                result_map_list = []
                for each_row in table_list:
                    if len(each_row) > 4 and each_row[0] != '序号' and each_row[6] != '' \
                            and each_row[0] != '序\n号':
                        publish_date = format_date(each_row[6])
                        result_map = {
                            'announcementTitle': '中国人民银行西宁中心支行行政处罚信息公示表（' + each_row[2] + '）',
                            'announcementOrg': '人民银行西宁中心支行',
                            'announcementDate': publish_date,
                            'announcementCode': each_row[2],
                            'facts': each_row[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1],
                            'punishmentBasement': '',
                            'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '银川中心支行':
        if '.doc' in url:
            return docx_to_json(url, url.split('/')[-1], url, data_id, city)
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '乌鲁木齐中心支行':
        if '.doc' in url:
            return docx_to_json(url, url.split('/')[-1], url, data_id, city)
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0 and url != 'http://wulumuqi.pbc.gov.cn/wulumuqi/121755/121777/121784/3397197/index.html':
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='content_right').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                result_map_list = []
                for each_table in content_soup.find(class_='hei14jj').find_all('table'):
                    table_list = table_to_list(each_table)
                    for each_row in table_list:
                        if len(each_row) == 8 and each_row[0] != '序号' and each_row[6] != '' \
                                and each_row[0] != '序\n号':
                            publish_date = format_date(each_row[6]).replace('\xa0', '')
                            result_map = {
                                'announcementTitle': '中国人民银行乌鲁木齐中心支行行政处罚信息公示表（' + each_row[2] + '）',
                                'announcementOrg': '人民银行乌鲁木齐中心支行',
                                'announcementDate': publish_date,
                                'announcementCode': each_row[2],
                                'facts': each_row[3],
                                'defenseOpinion': '',
                                'defenseResponse': '',
                                'litigant': each_row[1],
                                'punishmentBasement': '',
                                'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                                'type': '行政处罚决定',
                                'oss_file_id': file_id,
                                'status': 'not checked'
                            }
                            result_map_list.append(result_map)
                        else:
                            if len(each_row) == 19 and each_row[0] != '序号' and each_row[13] != '' \
                                    and each_row[0] != '序\n号':
                                publish_date = re.search(r'填报日期：(\d{4}年\d+月\d+日)', content_text).group(1).strip()
                                litigant = ''
                                if each_row[1] != '':
                                    litigant += '企业名称：' + str(each_row[1]) + '\n'
                                if each_row[2] != '':
                                    litigant += '社会信用代码（组织机构代码）：' + str(each_row[2]) + '\n'
                                if each_row[3] != '':
                                    litigant += '法人代表姓名：' + str(each_row[3]) + '\n'
                                if each_row[4] != '':
                                    litigant += '法人代表证件号：' + str(each_row[5]) + '\n'
                                result_map = {
                                    'announcementTitle': '中国人民银行乌鲁木齐中心支行行政处罚信息公示表（' + each_row[2] + '）',
                                    'announcementOrg': '人民银行乌鲁木齐中心支行',
                                    'announcementDate': publish_date,
                                    'announcementCode': each_row[6],
                                    'facts': each_row[13] + '，' + each_row[12],
                                    'defenseOpinion': '',
                                    'defenseResponse': '',
                                    'litigant': litigant,
                                    'punishmentBasement': each_row[11],
                                    'punishmentDecision': '处罚金额' + str(each_row[15]),
                                    'type': '行政处罚决定',
                                    'oss_file_id': file_id,
                                    'status': 'not checked'
                                }
                                result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '深圳市中心支行':
        if '.doc' in url:
            return docx_to_json(url, url.split('/')[-1], url, data_id, city)
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '大连市中心支行':
        if '.doc' in url:
            return docx_to_json(url, url.split('/')[-1], data_id, url, city)
        r = prepare(url)
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0 and url != 'http://dalian.pbc.gov.cn/dalian/123812/123830/123837/3550166/index.html':
            parse_links(links_list, url, data_id, city)
        else:
            title = content_soup.find(id='content_right').find(class_='buleB16').text.strip()
            if db.parsed_data.find({'origin_url': url, 'oss_file_origin_url': url}).count() == 0:
                oss_file_map = {
                    'origin_url': url,
                    'oss_file_origin_url': url,
                    'origin_url_id': data_id,
                    'oss_file_type': 'html',
                    'oss_file_name': title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                response = db.parsed_data.insert_one(oss_file_map)
                file_id = response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
            else:
                db.pbc_data.update_one({'_id': data_id}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': url, 'oss_file_origin_url': url})['_id']

            if db.announcement.find({'oss_file_id': file_id}).count() == 0:
                table_list = table_to_list(content_soup.find(class_='hei14jj').find_all('table')[0])
                result_map_list = []
                for each_row in table_list:
                    if len(each_row) > 4 and each_row[0] != '序号' and each_row[6] != '' \
                            and each_row[0] != '序\n号':
                        publish_date = format_date(each_row[6])
                        result_map = {
                            'announcementTitle': '中国人民银行大连市中心支行行政处罚信息公示表（' + each_row[2] + '）',
                            'announcementOrg': '人民银行大连市中心支行',
                            'announcementDate': publish_date,
                            'announcementCode': each_row[2],
                            'facts': each_row[3],
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1],
                            'punishmentBasement': '',
                            'punishmentDecision': '我行对' + each_row[1] + '作出以下处罚：' + each_row[4],
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                        result_map_list.append(result_map)
                if len(result_map_list) > 0:
                    logger.info('人民银行' + city + '解析 -- 一共有%d条数据' % len(result_map_list))
                    db.announcement.insert_many(result_map_list)
                    logger.info('人民银行' + city + '解析 -- 数据导入完成')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('人民银行' + city + '解析 -- 修改parsed完成')
                else:
                    logger.warning('人民银行' + city + '解析 -- 解析未能完成')
            else:
                logger.info('人民银行' + city + '解析 -- 数据已经存在')

    if city == '青岛市中心支行':
        if '.pdf' in url:
            return pdf_to_json(url, url.split('/')[-1], url, data_id, city)
        if '.xls' in url:
            return xlsx_to_json(url, url.split('/')[-1], url, data_id, city)
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '宁波市中心支行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)

    if city == '厦门市中心支行':
        links_list = content_soup.find(id='content_right').find(class_='hei14jj').find_all('a')
        if len(links_list) > 0:
            parse_links(links_list, url, data_id, city)


def parse():
    db.parsed_data.create_index([('origin_url', pymongo.HASHED)])
    db.parsed_data.create_index([('oss_file_origin_url', pymongo.HASHED)])
    db.announcement.create_index([('oss_file_id', pymongo.HASHED)])
    # 中国人民银行
    for each_data in db.pbc_data.find({'origin': '人民银行', 'status': {'$nin': ['ignored', 'parsed']}}):
        pbc_parse(each_data['url'], each_data['_id'])

    # 地方人民银行
    city_list = [
        '上海总部',
        '上海分行',
        '天津分行',
        '沈阳分行',
        '南京分行',
        '济南分行',
        '武汉分行',
        '广州分行',
        '成都分行',
        '西安分行',
        '营业管理部（北京）',
        '重庆营业管理部',
        '石家庄中心支行',
        '太原中心支行',
        '呼和浩特中心支行',
        '长春中心支行',
        '哈尔滨中心支行',
        '杭州中心支行',
        '福州中心支行',
        '合肥中心支行',
        '郑州中心支行',
        '长沙中心支行',
        '南昌中心支行',
        '南宁中心支行',
        '海口中心支行',
        '昆明中心支行',
        '贵阳中心支行',
        '拉萨中心支行',
        '兰州中心支行',
        '西宁中心支行',
        '银川中心支行',
        '乌鲁木齐中心支行',
        '深圳市中心支行',
        '大连市中心支行',
        '青岛市中心支行',
        '宁波市中心支行',
        '厦门市中心支行'
    ]
    for each_city in city_list:
        for each_data in db.pbc_data.find({'origin': '人民银行' + each_city,
                                           'status': {'$nin': ['ignored']}},
                                          no_cursor_timeout=True):
            if db.parsed_data.find(
                    {'origin_url_id': each_data['_id'], 'parsed': False}).count() > 0 or \
                    (db.parsed_data.find({'origin_url': each_data['url'], 'parsed': True}).count() == 0 and
                     each_data['status'] == 'not parsed'):
                try:
                    local_pbc_parse(each_data['url'], each_city, each_data['_id'])
                except Exception as e:
                    exc_type, exc_obj, exc_tb = sys.exc_info()
                    logger.error('出错行数：%s' % str(exc_tb.tb_lineno))
                    logger.info('解析出现问题！！')
                    logger.warning(e)

                    logger.info('删除TMP文件')
                    if os.path.exists('./test/tmp.xls'):
                        os.remove('./test/tmp.xls')

                    if os.path.exists('./test/tmp.xlsx'):
                        os.remove('./test/tmp.xlsx')

                    if os.path.exists('./test/tmp.doc'):
                        os.remove('./test/tmp.doc')

                    if os.path.exists('./test/tmp.docx'):
                        os.remove('./test/tmp.docx')

                    if os.path.exists('./test/tmp.pdf'):
                        os.remove('./test/tmp.pdf')

                    if os.path.exists('./test/tmp.wps'):
                        os.remove('./test/tmp.wps')

                    if os.path.exists('./test/tmp.et'):
                        os.remove('./test/tmp.et')

                    if os.path.exists('./test/tmp.tif'):
                        os.remove('./test/tmp.tif')

                    if os.path.exists('./test/tmp.jpg'):
                        os.remove('./test/tmp.jpg')

                    if os.path.exists('./test/tmp.bmp'):
                        os.remove('./test/tmp.bmp')

                    if os.path.exists('./test/tmp/'):
                        for each_img in os.listdir('./test/tmp'):
                            os.remove('./test/tmp/' + each_img)
                        os.rmdir('./test/tmp')

                    continue


if __name__ == "__main__":
    parse()
