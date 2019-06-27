import re
import os
import shutil
import subprocess
import patoolib
import docx
from urllib.parse import urljoin

from init import config_init
from oss_utils import oss_add_file, init_ali_oss
from utility import cn2dig, get_year, request_site_page, get_content_text, table_to_list, remove_strip, format_date
from bs4 import BeautifulSoup as bs

ali_bucket = init_ali_oss()
config = config_init()


def shaanxi_circ(db, logger):
    for each_circ_data in db.circ_data.find({'origin': '陕西保监局', 'status': {'$nin': ['ignored']}}):
        announcement_url = each_circ_data['url']
        announcement_title = each_circ_data['title']

        if db.circ_data.find(
                {'url': announcement_url, 'status': 'parsed'}).count() == 1 and db.parsed_data.find(
            {'origin_url': announcement_url, 'parsed': True}).count() >= 1:
            continue

        logger.info('陕西保监局 ' + 'Url to parse: %s' % announcement_url)

        r = request_site_page(announcement_url)
        if r is None:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_soup = bs(r.content, 'lxml') if r else bs('', 'lxml')

        table_content = content_soup.find(id='tab_content')
        if not table_content:
            logger.error('网页请求错误 %s' % announcement_url)
            continue
        content_text = get_content_text(table_content.find_all('tr')[3])
        if content_text == '':
            continue
        title = table_content.find_all('tr')[0].text.strip()

        if '行政处罚公布' in title or '行政处罚汇总' in title:
            rar_link = [urljoin(announcement_url, each_link.attrs['href'])
                        for each_link in content_soup.find_all('a') if '.rar' in each_link.attrs['href']][0]
            response = request_site_page(rar_link)
            with open('./test/tmp.rar', 'wb') as out_file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        out_file.write(chunk)

            if not os.path.exists('./test/tmp'):
                os.mkdir('./test/tmp')
            patoolib.extract_archive('./test/tmp.rar', outdir='./test/tmp')

            doc_file_list = []
            for root, dirs, files in os.walk("./test/tmp", topdown=False):
                for name in files:
                    doc_file_list.append(os.path.join(root, name))

            if title == '2008年行政处罚公布':
                for each_doc_file in doc_file_list:
                    doc_title = re.split(r'[./]', each_doc_file)[-2]
                    if not os.path.exists('./test/tmp/' + doc_title + '.docx'):
                        shell_str = config['soffice']['soffice_path'] + ' --headless --convert-to docx ' + \
                                    each_doc_file + ' --outdir ./test/tmp'
                        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                                   shell=True, stdout=subprocess.PIPE)
                        process.wait()

                    with open('./test/tmp/' + doc_title + '.docx', 'rb') as docx_file:
                        docx_content = docx_file.read()

                    if db.parsed_data.find(
                            {'origin_url': announcement_url, 'oss_file_name': doc_title}).count() == 0:
                        oss_file_map = {
                            'origin_url': announcement_url,
                            'oss_file_origin_url': rar_link,
                            'origin_url_id': each_circ_data['_id'],
                            'oss_file_type': 'docx',
                            'oss_file_name': doc_title,
                            'oss_file_content': docx_content,
                            'parsed': False
                        }
                        insert_response = db.parsed_data.insert_one(oss_file_map)
                        file_id = insert_response.inserted_id
                        oss_add_file(ali_bucket, str(file_id) + '/' + doc_title + '.docx', docx_content)
                        db.circ_data.update_one({'_id': each_circ_data['_id']}, {'$set': {'status': 'parsed'}})
                    else:
                        db.circ_data.update_one({'_id': each_circ_data['_id']}, {'$set': {'status': 'parsed'}})
                        file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                           'oss_file_name': doc_title})['_id']

                    doc = docx.Document('./test/tmp/' + doc_title + '.docx')
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    content_text = '\n'.join(full_text)

                    document_code_compiler = re.compile(r'(陕保监罚\n?.*?\d+\n?\d+\n?.*?\n?\d+\n?号)')
                    document_code = document_code_compiler.search(content_text).group(1).strip()

                    real_title = '中国保监会陕西监管局行政处罚决定书（' + document_code + '）'

                    litigant_compiler = re.compile(
                        document_code.replace(r'[', r'\[').replace(r']', r'\]') +
                        r'\n([\s\S]*?)\n' + r'(经查|经检查|依据.*?的有关规定|抽查|'
                                            r'经抽查|.*?存在以下(问题|违法行为)|'
                                            r'.*?认定事实|.*?存在.*?(行为|业务数据不真实)|'
                                            r'你公司于.*?期间|'
                                            r'经调查)')
                    litigant = litigant_compiler.search(content_text).group(1).strip()

                    truth_text_str = r'((经查|一、|二、|三、|经检查|.*?存在以下(问题|违法行为)|.*?认定事实|.*?存在.*?(行为|业务数据不真实)|你公司于.*?期间)' \
                                     r'([\s\S]*?))' \
                                     r'(依\n?据)'
                    truth_compiler = re.compile(truth_text_str)
                    truth_list = truth_compiler.findall(content_text)
                    if len(truth_list) > 0:
                        truth = '\n'.join([kk[0].strip() for kk in truth_list])
                    else:
                        truth_text_str = litigant + r'([\s\S]*?)' \
                                                    r'((，|，)?(综合)?(依据|鉴于|根据))'
                        truth_compiler = re.compile(truth_text_str)
                        truth = truth_compiler.search(content_text).group(1).strip()

                    if '申辩' in content_text:
                        defense_text_str = r'((针对.*?行为.*?申辩意见|(当事人)?[^，。,；\n]*?(未)?提出(了)?陈述申辩(意见)?|' \
                                           r'[^，。,；\n]*?向我局(报送|递交|提出)[^，。,；\n]*?|本案在审理过程中.*?提出陈述申辩|' \
                                           r'[^，。,；\n]*?在(申辩材料|陈述申辩|陈述申辩意见|申辩意见|陈述材料|申辩书)中称|[^，。,；\n]*?在听证阶段提出|' \
                                           r'[^，。,；\n]*?在法定期限内(未)?提出(了)?(听证要求|陈述申辩|陈述申辩及听证要求)|' \
                                           r'在法定期限内，当事人未提出|[^，。,；\n]*?提出了?陈述、申辩|[^，。,；\n]*?提出以下陈述申辩意见|' \
                                           r'[^，。,；\n]*?放弃陈述、申辩权利|[^，。,；\n]*?提出以下陈述(申辩理由|意见)|' \
                                           r'[^，。,；\n]*?陈述申辩材料中辩称|[^，。,；\n]*?陈述材料中提出|[^，。,；\n]*?提出以下申辩意见|' \
                                           r'[^，。,；\n]*?申辩中提出)' \
                                           r'([\s\S]*?))' \
                                           r'(因此，我局决定|' \
                                           r'我局经复核(认为|决定)|' \
                                           r'本案现已审理终结|' \
                                           r'我局经复查[^，。,；\n]*?情况|' \
                                           r'我局[^，。,；\n]*?认真复核|' \
                                           r'经研究，对[^，。,；\n]*?予以采纳。|' \
                                           r'我局认为.*?申辩理由|' \
                                           r'依据.*?我局认为.*?的申辩理由|' \
                                           r'经研究，我局认为.*?申辩意见|' \
                                           r'经我局审核，决定|' \
                                           r'我局认为，上述违法行为事实清楚、证据确凿、法律法规适当|' \
                                           r'我局对陈述申辩意见进行了复核|' \
                                           r'经我局审核|' \
                                           r'针对[^，。,；\n]*?的(陈述)?申辩意见，我局进行了核实|' \
                                           r'经查，我局认为|' \
                                           r'依据现场检查及听证情况|' \
                                           r'我局认为|我局经核查|对此，我局认为|我局经审理认为|我局在处罚幅度裁量时)'
                        defense_compiler = re.compile(defense_text_str, re.MULTILINE)
                        defense_list = defense_compiler.findall(content_text)
                        if len(defense_list) != 0:
                            defense = defense_list[-1][0].strip()
                            defense_response_str = defense.replace(r'[', r'\[').replace(r']', r'\]') \
                                                   + r'(([\s\S]*?)' + r'(本案现已审理终结。|不符合.*?情形。|根据.*?依法可从轻或者减轻行政处罚。|' \
                                                                      r'对[^，。,；\n]*?申辩意见(不予|予以)采纳|因此.*?申辩理由.*?成立。|' \
                                                                      r'我局认为.*?申辩(理由|意见).*?符合.*?第.*?条.*?的条件.(予以采纳。)?|' \
                                                                      r'不予采纳其陈述申辩理由。|维持原处罚决定。|不予采纳。|' \
                                                                      r'不予采纳.*?(陈述)?申辩(意见|理由)?。|' \
                                                                      r'我局认定你公司行为构成销售误导。|你公司的申辩不成立。))'
                            defense_response_compiler = re.compile(defense_response_str, re.MULTILINE)
                            if defense_response_compiler.search(content_text):
                                defense_response = defense_response_compiler.search(content_text).group(
                                    1).strip()
                            else:
                                if '未' in defense:
                                    defense_response = ''
                        else:
                            defense_text_str = '([^。；\n]*?向.*?公告送达了《行政处罚事先告知书》.*?提出陈述申辩。|' \
                                               '我局依法于2012年5月25日对你公司送达了《行政处罚事先告知书》，你公司在规定的时间内未提出陈述和申辩意见，也未要求举行听证。)'
                            defense_compiler = re.compile(defense_text_str, re.MULTILINE)
                            defense = defense_compiler.search(content_text).group(1).strip()
                            defense_response = ''
                    else:
                        defense = defense_response = ''

                    punishment_decision_text_str = r'(((依\n?据|根据)[^。；]*?第?[^。；]*?条\n?.*?(规定)?.?' \
                                                   r'(我局|我机关)?(决定|责令|给予|于.*?向.*?发出|对.*?作出|拟对你)|' \
                                                   r'(我局)?决定.*?(作出|给予)(如下|以下)(行政)?处罚：)' \
                                                   r'([\s\S]*?))' \
                                                   r'(请在本处罚决定书送达之日|当事人应当在接到本处罚决定书之日|如不服本处罚决定|' \
                                                   r'请(在)?接到本处罚决定书之日|如不服从本处罚决定|.*?如对本处罚决定不服|' \
                                                   r'请在接到本处罚决定书之日)'
                    punishment_decision_compiler = re.compile(punishment_decision_text_str)
                    punishment_decision_list = punishment_decision_compiler.findall(content_text)
                    if len(punishment_decision_list) > 0:
                        punishment_decision = '\n'.join([kk[0].strip() for kk in punishment_decision_list])

                    punishment_basis_str_list = [
                        r'((上\n?述|以上|该)(事实|实施)?行为)?违反[^。；]*?第[^，。,；]*?条[\s\S]*?',
                    ]
                    punishment_basis_str = '|'.join(punishment_basis_str_list)
                    punishment_basis_compiler = re.compile(r'[。\n；，]' + '(' + punishment_basis_str + ')' +
                                                           '.(\n?依\n?据|\n?根据|\n?鉴于|\n?上述事实有现场检查确认书)',
                                                           re.MULTILINE)
                    punishment_basis_list = punishment_basis_compiler.findall(content_text)
                    punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

                    publish_date_text = re.search(
                        punishment_decision.replace(r'(', r'\(').replace(r')', r'\)').replace(r'[', r'\[').replace(
                            r']', r'\]').
                        replace(r'*', r'\*') + r'([\s\S]*?)$', content_text).group(1).replace('\n', '')
                    if re.search(r'.{4}年.{1,2}月.{1,3}日', publish_date_text):
                        publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)[-1].replace(' ', '')
                        m = re.match(
                            "([0-9零一二两三四五六七八九十〇○ＯOΟО]+年)?([0-9一二两三四五六七八九十]+)月?([0-9一二两三四五六七八九十]+)[号日]?",
                            publish_date)
                        real_publish_date = get_year(m.group(1)) + str(cn2dig(m.group(2))) + '月' + str(
                            cn2dig(m.group(3))) + '日'
                    else:
                        publish_date_text = table_content.find_all('tr')[1].text
                        publish_date = re.findall(r'\d{4}-\d{1,2}-\d{1,2}', publish_date_text)[-1]
                        real_publish_date = publish_date.split('-')[0] + '年' + str(
                            int(publish_date.split('-')[1])) + '月' + str(
                            int(publish_date.split('-')[2])) + '日'

                    result_map = {
                        'announcementTitle': real_title,
                        'announcementOrg': '陕西保监局',
                        'announcementDate': real_publish_date,
                        'announcementCode': document_code,
                        'facts': truth,
                        'defenseOpinion': defense,
                        'defenseResponse': defense_response,
                        'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    logger.info(result_map)
                    if db.announcement.find({'announcementTitle': real_title, 'oss_file_id': file_id}).count() == 0:
                        db.announcement.insert_one(result_map)
                        logger.info('陕西保监局 数据解析 ' + ' -- 数据导入完成')
                    else:
                        logger.info('陕西保监局 数据解析 ' + ' -- 数据已经存在')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('陕西保监局 数据解析 ' + ' -- 修改parsed完成')
            else:
                for each_doc_file in doc_file_list:
                    doc_title = re.split('[./]', each_doc_file)[-2]
                    if not os.path.exists('./test/tmp/' + doc_title + '.docx'):
                        shell_str = config['soffice']['soffice_path'] + ' --headless --convert-to docx ' + \
                                    each_doc_file.replace(' ', '" "') + ' --outdir ./test/tmp'
                        process = subprocess.Popen(shell_str.replace(r'(', r'\(').replace(r')', r'\)'),
                                                   shell=True, stdout=subprocess.PIPE)
                        process.wait()

                    with open('./test/tmp/' + doc_title + '.docx', 'rb') as docx_file:
                        docx_content = docx_file.read()

                    if db.parsed_data.find(
                            {'origin_url': announcement_url, 'oss_file_name': doc_title}).count() == 0:
                        oss_file_map = {
                            'origin_url': announcement_url,
                            'oss_file_origin_url': rar_link,
                            'origin_url_id': each_circ_data['_id'],
                            'oss_file_type': 'docx',
                            'oss_file_name': doc_title,
                            'oss_file_content': docx_content,
                            'parsed': False
                        }
                        insert_response = db.parsed_data.insert_one(oss_file_map)
                        file_id = insert_response.inserted_id
                        oss_add_file(ali_bucket, str(file_id) + '/' + doc_title + '.docx', docx_content)
                        db.circ_data.update_one({'_id': each_circ_data['_id']}, {'$set': {'status': 'parsed'}})
                    else:
                        db.circ_data.update_one({'_id': each_circ_data['_id']}, {'$set': {'status': 'parsed'}})
                        file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                           'oss_file_name': doc_title})['_id']

                    doc = docx.Document('./test/tmp/' + doc_title + '.docx')
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    content_text = '\n'.join(full_text)

                    document_code_compiler = re.compile(r'(陕银?保监罚\n?.*?\d+\n?\d+\n?.*?\n?\d+\n?.*?号)')
                    document_code = document_code_compiler.search(content_text).group(1).strip()

                    real_title = '中国银保监会陕西监管局行政处罚决定书（' + document_code + '）'

                    litigant_compiler = re.compile(
                        document_code.replace(r'[', r'\[').replace(r']', r'\]') +
                        r'\n([\s\S]*?)\n' + r'(经查|依据.*?的有关规定|抽查|'
                                            r'经抽查|经检查|.*?存在以下(问题|违法行为)|'
                                            r'.*?认定事实|.*?存在.*?(行为|业务数据不真实)|'
                                            r'你公司于.*?期间|'
                                            r'经调查)')
                    litigant = litigant_compiler.search(content_text).group(1).strip()

                    truth_text_str = r'((经查|一、|二、|三、|经检查|.*?存在以下(问题|违法行为)|.*?认定事实|.*?存在.*?(行为|业务数据不真实)|你公司于.*?期间)' \
                                     r'([\s\S]*?))' \
                                     r'(依据)'
                    truth_compiler = re.compile(truth_text_str)
                    truth_list = truth_compiler.findall(content_text)
                    if len(truth_list) > 0:
                        truth = '\n'.join([kk[0].strip() for kk in truth_list])
                    else:
                        truth_text_str = litigant + r'([\s\S]*?)' \
                                                    r'((，|，)?(综合)?(依据|鉴于|根据))'
                        truth_compiler = re.compile(truth_text_str)
                        truth = truth_compiler.search(content_text).group(1).strip()

                    if '申辩' in content_text:
                        defense_text_str = r'((针对.*?行为.*?申辩意见|(当事人)?[^，。,；\n]*?(未)?提出(了)?陈述申辩(意见)?|' \
                                           r'[^，。,；\n]*?向我局(报送|递交|提出)[^，。,；\n]*?|本案在审理过程中.*?提出陈述申辩|' \
                                           r'[^，。,；\n]*?在(申辩材料|陈述申辩|陈述申辩意见|申辩意见|陈述材料|申辩书)中称|[^，。,；\n]*?在听证阶段提出|' \
                                           r'[^，。,；\n]*?在法定期限内(未)?提出(了)?(听证要求|陈述申辩|陈述申辩及听证要求)|' \
                                           r'在法定期限内，当事人未提出|[^，。,；\n]*?提出了?陈述、申辩|[^，。,；\n]*?提出以下陈述申辩意见|' \
                                           r'[^，。,；\n]*?放弃陈述、申辩权利|[^，。,；\n]*?提出以下陈述(申辩理由|意见)|' \
                                           r'[^，。,；\n]*?陈述申辩材料中辩称|[^，。,；\n]*?陈述材料中提出|[^，。,；\n]*?提出以下申辩意见|' \
                                           r'[^，。,；\n]*?申辩中提出)' \
                                           r'([\s\S]*?))' \
                                           r'(因此，我局决定|' \
                                           r'我局经复核(认为|决定)|' \
                                           r'本案现已审理终结|' \
                                           r'我局经复查[^，。,；\n]*?情况|' \
                                           r'我局[^，。,；\n]*?认真复核|' \
                                           r'经研究，对[^，。,；\n]*?予以采纳。|' \
                                           r'我局认为.*?申辩理由|' \
                                           r'依据.*?我局认为.*?的申辩理由|' \
                                           r'经研究，我局认为.*?申辩意见|' \
                                           r'经我局审核，决定|' \
                                           r'我局认为，上述违法行为事实清楚、证据确凿、法律法规适当|' \
                                           r'我局对陈述申辩意见进行了复核|' \
                                           r'经我局审核|' \
                                           r'针对[^，。,；\n]*?的(陈述)?申辩意见，我局进行了核实|' \
                                           r'经查，我局认为|' \
                                           r'依据现场检查及听证情况|' \
                                           r'我局认为|我局经核查|对此，我局认为|我局经审理认为|我局在处罚幅度裁量时)'
                        defense_compiler = re.compile(defense_text_str, re.MULTILINE)
                        defense_list = defense_compiler.findall(content_text)
                        if len(defense_list) != 0:
                            defense = defense_list[-1][0]
                            defense_response_str = defense.replace(r'[', r'\[').replace(r']', r'\]') \
                                                   + r'(([\s\S]*?)' + r'(本案现已审理终结。|不符合.*?情形。|根据.*?依法可从轻或者减轻行政处罚。|' \
                                                                      r'对[^，。,；\n]*?申辩意见(不予|予以)采纳|因此.*?申辩理由.*?成立。|' \
                                                                      r'我局认为.*?申辩(理由|意见).*?符合.*?第.*?条.*?的条件.(予以采纳。)?|' \
                                                                      r'不予采纳其陈述申辩理由。|维持原处罚决定。|不予采纳。|' \
                                                                      r'不予采纳.*?(陈述)?申辩(意见|理由)?。|' \
                                                                      r'我局认定你公司行为构成销售误导。|你公司的申辩不成立。))'
                            defense_response_compiler = re.compile(defense_response_str, re.MULTILINE)
                            if defense_response_compiler.search(content_text):
                                defense_response = defense_response_compiler.search(content_text).group(
                                    1).strip()
                            else:
                                if '未' in defense:
                                    defense_response = ''
                        else:
                            defense_text_str = '([^。；\n]*?向.*?公告送达了《行政处罚事先告知书》.*?提出陈述申辩。|' \
                                               '我局依法于2012年5月25日对你公司送达了《行政处罚事先告知书》，你公司在规定的时间内未提出陈述和申辩意见，也未要求举行听证。)'
                            defense_compiler = re.compile(defense_text_str, re.MULTILINE)
                            defense = defense_compiler.search(content_text).group(1).strip()
                            defense_response = ''
                    else:
                        defense = defense_response = ''

                    punishment_decision_text_str = r'(((依据|根据)[^。；]*?第?[^。；]*?条\n?.*?(规定)?.?(我局|我机关)?(决定|责令|给予|于.*?向.*?发出|对.*?作出|拟对)|' \
                                                   r'(我局)?决定.*?(作出|给予)(如下|以下)(行政)?处罚：)' \
                                                   r'([\s\S]*?))' \
                                                   r'(请在本处罚决定书送达之日|当事人应当在接到本处罚决定书之日|如不服本处罚决定|' \
                                                   r'请(在)?接到本处罚决定书之日|如不服从本处罚决定|.*?如对本处罚决定不服|' \
                                                   r'请在接到本处罚决定书之日)'
                    punishment_decision_compiler = re.compile(punishment_decision_text_str)
                    punishment_decision_list = punishment_decision_compiler.findall(content_text)
                    if len(punishment_decision_list) > 0:
                        punishment_decision = '\n'.join([kk[0].strip() for kk in punishment_decision_list])

                    punishment_basis_str_list = [
                        r'((上述|以上|该)(事实|实施)?行为)?违反[^。；]*?第[^，。,；\n]*?条[\s\S]*?',
                    ]
                    punishment_basis_str = '|'.join(punishment_basis_str_list)
                    punishment_basis_compiler = re.compile('[。\n；，]' + '(' + punishment_basis_str + ')' +
                                                           '.(\n?依据|\n?根据|\n?鉴于|\n?上述事实有现场检查确认书)', re.MULTILINE)
                    punishment_basis_list = punishment_basis_compiler.findall(content_text)
                    punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

                    publish_date_text = re.search(
                        punishment_decision.replace(r'(', r'\(').replace(r')', r'\)').replace(r'[', r'\[').replace(
                            r']', r'\]').
                        replace(r'*', r'\*') + r'([\s\S]*?)$', content_text).group(1).replace('\n', '')
                    if re.search(r'.{4}年.{1,2}月.{1,3}日', publish_date_text):
                        publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)[-1].replace(' ', '')
                        m = re.match(
                            "([0-9零一二两三四五六七八九十〇○ＯOΟО]+年)?([0-9一二两三四五六七八九十]+)月?([0-9一二两三四五六七八九十]+)[号日]?",
                            publish_date)
                        real_publish_date = get_year(m.group(1)) + str(cn2dig(m.group(2))) + '月' + str(
                            cn2dig(m.group(3))) + '日'
                    else:
                        publish_date_text = table_content.find_all('tr')[1].text
                        publish_date = re.findall(r'\d{4}-\d{1,2}-\d{1,2}', publish_date_text)[-1]
                        real_publish_date = publish_date.split('-')[0] + '年' + str(
                            int(publish_date.split('-')[1])) + '月' + str(
                            int(publish_date.split('-')[2])) + '日'

                    result_map = {
                        'announcementTitle': real_title,
                        'announcementOrg': '陕西银保监局',
                        'announcementDate': real_publish_date,
                        'announcementCode': document_code,
                        'facts': truth,
                        'defenseOpinion': defense,
                        'defenseResponse': defense_response,
                        'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                        'punishmentBasement': punishment_basis,
                        'punishmentDecision': punishment_decision,
                        'type': '行政处罚决定',
                        'oss_file_id': file_id,
                        'status': 'not checked'
                    }
                    logger.info(result_map)
                    if db.announcement.find({'announcementTitle': real_title, 'oss_file_id': file_id}).count() == 0:
                        db.announcement.insert_one(result_map)
                        logger.info('陕西保监局 数据解析 ' + ' -- 数据导入完成')
                    else:
                        logger.info('陕西保监局 数据解析 ' + ' -- 数据已经存在')
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('陕西保监局 数据解析 ' + ' -- 修改parsed完成')

            logger.info('删除Tmp文件')
            os.remove('./test/tmp.rar')
            shutil.rmtree('./test/tmp', ignore_errors=True)
        else:
            if db.parsed_data.find(
                    {'origin_url': announcement_url, 'oss_file_origin_url': announcement_url}).count() == 0:
                oss_file_map = {
                    'origin_url': announcement_url,
                    'oss_file_origin_url': announcement_url,
                    'origin_url_id': each_circ_data['_id'],
                    'oss_file_type': 'html',
                    'oss_file_name': announcement_title,
                    'oss_file_content': r.text.encode(r.encoding).decode('utf-8'),
                    'parsed': False
                }
                insert_response = db.parsed_data.insert_one(oss_file_map)
                file_id = insert_response.inserted_id
                oss_add_file(ali_bucket, str(file_id) + '/' + announcement_title + '.html',
                             r.text.encode(r.encoding).decode('utf-8'))
                db.circ_data.update_one({'_id': each_circ_data['_id']}, {'$set': {'status': 'parsed'}})
            else:
                db.circ_data.update_one({'_id': each_circ_data['_id']}, {'$set': {'status': 'parsed'}})
                file_id = db.parsed_data.find_one({'origin_url': announcement_url,
                                                   'oss_file_origin_url': announcement_url})['_id']

            if '行政处罚信息主动公开事项' in announcement_title:
                result_map_list = []
                table_value_list = table_to_list(table_content.find(class_='ke-zeroborder'))
                announcement_code = ''
                new_result_map = {}
                for index, each_origin_row in enumerate(table_value_list):
                    each_row = []
                    [each_row.append(i) for i in each_origin_row if i in each_origin_row]
                    if '文号' in each_row[0] or '行政处罚信息主动公开事项' in each_row[0] or '处罚决定书' in each_row[0]:
                        continue
                    if remove_strip(each_row[0]).strip() != announcement_code:
                        if new_result_map != {}:
                            punishment_basis_compiler = re.compile(r'(。|\n|；|^)' + r'(([^\n。；]*?)违反.*?第.*?条.*?规定)' +
                                                                   '.(\n?依据|\n?根据|\n?鉴于|\n?上述事实有现场检查确认书)', re.MULTILINE)
                            punishment_basis_list = punishment_basis_compiler.findall(
                                new_result_map['punishmentDecision'])
                            punishment_basis = '；'.join([kk[1].strip() for kk in punishment_basis_list])
                            new_result_map['punishmentBasement'] = punishment_basis
                            logger.info(new_result_map)
                            if db.announcement.find({'announcementTitle': new_result_map['announcementTitle'],
                                                     'oss_file_id': new_result_map['oss_file_id']}).count() == 0:
                                db.announcement.insert_one(new_result_map)
                                logger.info('陕西保监局 数据解析 ' + ' -- 数据导入完成')
                            else:
                                logger.info('陕西保监局 数据解析 ' + ' -- 数据已经存在')
                            result_map_list.append(new_result_map)
                        announcement_code = remove_strip(each_row[0]).strip()
                        this_punishment_decision = each_row[1].strip() + '：' + \
                                                   each_row[3].strip() + ' ' + each_row[4].strip()
                        new_result_map = {
                            'announcementTitle': '陕西银保监局行政处罚信息主动公开事项（' + announcement_code + '）',
                            'announcementOrg': '陕西银保监局',
                            'announcementDate': format_date(each_row[-1].strip()),
                            'announcementCode': announcement_code,
                            'facts': each_row[1].strip() + '：' + each_row[2].strip(),
                            'defenseOpinion': '',
                            'defenseResponse': '',
                            'litigant': each_row[1].strip(),
                            'punishmentBasement': '',
                            'punishmentDecision': this_punishment_decision,
                            'type': '行政处罚决定',
                            'oss_file_id': file_id,
                            'status': 'not checked'
                        }
                    else:
                        if each_row[1].strip() not in new_result_map['litigant']:
                            new_result_map['litigant'] += '，' + each_row[1].strip()
                        new_result_map['facts'] += '\n' + each_row[1].strip() + '：' + each_row[2].strip()
                        this_punishment_decision = each_row[1].strip() + '：' + \
                                                   each_row[3].strip() + ' ' + each_row[4].strip()
                        new_result_map['punishmentDecision'] += '\n' + this_punishment_decision
                    if index == len(table_value_list) - 1:
                        punishment_basis_compiler = re.compile(r'(。|\n|；|^)' + r'(([^\n。；]*?)违反.*?第.*?条.*?规定)' +
                                                               '.(\n?依据|\n?根据|\n?鉴于|\n?上述事实有现场检查确认书)', re.MULTILINE)
                        punishment_basis_list = punishment_basis_compiler.findall(new_result_map['punishmentDecision'])
                        punishment_basis = '；'.join([kk[1].strip() for kk in punishment_basis_list])
                        new_result_map['punishmentBasement'] = punishment_basis
                        logger.info(new_result_map)
                        if db.announcement.find({'announcementTitle': new_result_map['announcementTitle'],
                                                 'oss_file_id': new_result_map['oss_file_id']}).count() == 0:
                            db.announcement.insert_one(new_result_map)
                            logger.info('陕西保监局 数据解析 ' + ' -- 数据导入完成')
                        else:
                            logger.info('陕西保监局 数据解析 ' + ' -- 数据已经存在')
                        result_map_list.append(new_result_map)
                if len(result_map_list) > 0:
                    db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                    logger.info('陕西保监局 数据解析 ' + '一共有%d条数据' % len(result_map_list))
                    logger.info('陕西保监局 数据解析 ' + ' -- 修改parsed完成')
                else:
                    logger.info('陕西保监局 数据解析 ' + ' -- 没有数据')
                # for index, each_row in enumerate(table_value_list):
                #     if '信息主动公开事项' not in each_row[0] and '处罚决定文号' not in each_row[0]:

            else:
                document_code_compiler = re.compile(r'(陕银?保监罚.*?\d{4}.*?\d+号|陕银?保监罚\d{4}.\d+.号)')
                if document_code_compiler.search(content_text):
                    document_code = document_code_compiler.search(content_text).group(1).strip()
                    litigant_compiler = re.compile(
                        document_code.replace(r'[', r'\[').replace(r']', r'\]') +
                        r'\n([\s\S]*?)\n' + r'(经查|经检查|依据.*?的有关规定|抽查|'
                                            r'经抽查|.*?存在以下(问题|违法行为)|'
                                            r'.*?认定事实|.*?存在.*?(行为|业务数据不真实)|'
                                            r'你公司于2009年期间)')
                    litigant = litigant_compiler.search(content_text).group(1).strip()
                else:
                    if document_code_compiler.search(title):
                        document_code = document_code_compiler.search(title).group(1).strip()
                    else:
                        document_code = ''
                    litigant_compiler = re.compile(r'^([\s\S]*?)\n' + r'(经查|依据.*?的有关规定|抽查|经抽查|'
                                                                      r'经检查|.*?存在以下(问题|违法行为)|'
                                                                      r'.*?认定事实)')
                    litigant = litigant_compiler.search(content_text).group(1).strip()

                truth_text_str = r'((经查|一、|二、|三、|经检查|.*?存在以下(问题|违法行为)|.*?认定事实|.*?存在.*?(行为|业务数据不真实)|你公司于2009年期间)' \
                                 r'([\s\S]*?))' \
                                 r'((我局认为，)?(上述|以上).*?(事实|行为|事实)(,|，)?.*?有.*?等证据(在案)?证明(,|，|。)(足以认定。)?|' \
                                 r'(我局认为，|综上，)?(上述|以上).*?(行为|问题|事实).*?违反.*?第.*条.*?(的规定)?|' \
                                 r'.*?作为.*?时任.*?对.*?负有直接责任|' \
                                 r'依据.*?第.*?条的规定|' \
                                 r'你作为.*?对.*?负有直接责任)'
                truth_compiler = re.compile(truth_text_str)
                truth_list = truth_compiler.findall(content_text)
                if len(truth_list) > 0:
                    truth = '\n'.join([kk[0].strip() for kk in truth_list])

                if '申辩' in content_text:
                    defense_text_str = r'((针对.*?行为.*?申辩意见|(当事人)?[^，。,；\n]*?(未)?提出(了)?陈述申辩(意见)?|' \
                                       r'[^，。,；\n]*?向我局(报送|递交|提出)[^，。,；\n]*?|本案在审理过程中.*?提出陈述申辩|' \
                                       r'[^，。,；\n]*?在(申辩材料|陈述申辩|陈述申辩意见|申辩意见|陈述材料|申辩书)中称|[^，。,；\n]*?在听证阶段提出|' \
                                       r'[^，。,；\n]*?在法定期限内(未)?提出(了)?(听证要求|陈述申辩|陈述申辩及听证要求)|' \
                                       r'在法定期限内，当事人未提出|[^，。,；\n]*?提出了?陈述、申辩|[^，。,；\n]*?提出以下陈述申辩意见|' \
                                       r'[^，。,；\n]*?放弃陈述、申辩权利|[^，。,；\n]*?提出以下陈述(申辩理由|意见)|' \
                                       r'[^，。,；\n]*?陈述申辩材料中辩称|[^，。,；\n]*?陈述材料中提出|[^，。,；\n]*?提出以下申辩意见|' \
                                       r'[^，。,；\n]*?申辩中提出)' \
                                       r'([\s\S]*?))' \
                                       r'(因此，我局决定|' \
                                       r'我局经复核(认为|决定)|' \
                                       r'本案现已审理终结|' \
                                       r'我局经复查[^，。,；\n]*?情况|' \
                                       r'我局[^，。,；\n]*?认真复核|' \
                                       r'经研究，对[^，。,；\n]*?予以采纳。|' \
                                       r'我局认为.*?申辩理由|' \
                                       r'依据.*?我局认为.*?的申辩理由|' \
                                       r'经研究，我局认为.*?申辩意见|' \
                                       r'经我局审核，决定|' \
                                       r'我局认为，上述违法行为事实清楚、证据确凿、法律法规适当|' \
                                       r'我局对陈述申辩意见进行了复核|' \
                                       r'经我局审核|' \
                                       r'针对[^，。,；\n]*?的(陈述)?申辩意见，我局进行了核实|' \
                                       r'经查，我局认为|' \
                                       r'依据现场检查及听证情况|' \
                                       r'我局认为|我局经核查|对此，我局认为|我局经审理认为|我局在处罚幅度裁量时)'
                    defense_compiler = re.compile(defense_text_str, re.MULTILINE)
                    defense_list = defense_compiler.findall(content_text)
                    if len(defense_list) != 0:
                        defense = defense_list[-1][0].strip()
                        defense_response_str = defense.replace(r'[', r'\[').replace(r']', r'\]') \
                                               + r'(([\s\S]*?)' + r'(本案现已审理终结。|不符合.*?情形。|根据.*?依法可从轻或者减轻行政处罚。|' \
                                                                  r'对[^，。,；\n]*?申辩意见(不予|予以)采纳|因此.*?申辩理由.*?成立。|' \
                                                                  r'我局认为.*?申辩(理由|意见).*?符合.*?第.*?条.*?的条件.(予以采纳。)?|' \
                                                                  r'不予采纳其陈述申辩理由。|维持原处罚决定。|不予采纳。|' \
                                                                  r'不予采纳.*?(陈述)?申辩(意见|理由)?。|' \
                                                                  r'我局认定你公司行为构成销售误导。|你公司的申辩不成立。))'
                        defense_response_compiler = re.compile(defense_response_str, re.MULTILINE)
                        if defense_response_compiler.search(content_text):
                            defense_response = defense_response_compiler.search(content_text).group(1).strip()
                        else:
                            if '未' in defense:
                                defense_response = ''
                    else:
                        defense_text_str = '([^。；\n]*?向.*?公告送达了《行政处罚事先告知书》.*?提出陈述申辩。|' \
                                           '我局依法于2012年5月25日对你公司送达了《行政处罚事先告知书》，你公司在规定的时间内未提出陈述和申辩意见，也未要求举行听证。)'
                        defense_compiler = re.compile(defense_text_str, re.MULTILINE)
                        defense = defense_compiler.search(content_text).group(1).strip()
                        defense_response = ''
                else:
                    defense = defense_response = ''

                punishment_decision_text_str = r'(((依据|根据).*?第?.*?条.*?(规定)?.?(我局)?(决定|责令|给予|于.*?向.*?发出|对.*?作出|拟对你)|' \
                                               r'(我局)?决定.*?作出(如下|以下)(行政)?处罚：)' \
                                               r'([\s\S]*?))' \
                                               r'(请在本处罚决定书送达之日|当事人应当在接到本处罚决定书之日|如不服本处罚决定|' \
                                               r'请(在)?接到本处罚决定书之日|如不服从本处罚决定|.*?如对本处罚决定不服)'
                punishment_decision_compiler = re.compile(punishment_decision_text_str)
                punishment_decision_list = punishment_decision_compiler.findall(content_text)
                if len(punishment_decision_list) > 0:
                    punishment_decision = '\n'.join([kk[0].strip() for kk in punishment_decision_list])

                punishment_basis_str_list = [
                    r'([^\n。；]*?)(问题|行为|事项|情况|事实)([^\n。；\s]*?)违反.*?\n?.*?第.*?条?\n?.*?((的|之|等)(相关)?规定)?',
                    r'百倍保险代理有限公司聘任不具有任职资格的人员，违反了《保险法》第121条的规定',
                    r'我局认为，直接业务虚挂中介业务套取费用，违反了《保险法》第116条第10款的规定',
                    r'我局认为，在高管任职资格审核中提交伪造的毕业证书违反了《保险专业代理机构监管规定》第24条的规定',
                    r'我局认为，未严格按照监管要求对人身保险新型产品投保人进行回访，违反了《人身保险新型产品信息披露办法》第10条的规定',
                    r'我局认为，未取得保险兼业代理资格代理车险业务，违反了《保险法》第119条的规定',
                    r'经查，中国人民财产保险股份有限公司西安市雁塔支公司于2006年3至4月和2007年2至4月期间，'
                    r'存在违规使用费率优惠因子，向不具备代理资格的机构支付代理手续费，出具阴阳发票、保单，及提供虚假的报表、资料等，'
                    r'违反了《保险法》第107条、第109条、第134条、第122条的相关规定。\n你作为雁塔支公司高级管理人员，对上述违法行为负有直接责任',
                    r'经查，2014年，中国人民人寿保险股份有限公司陕西省分公司委托未取得合法保险销售资格的江西雅林环保科技有限公司上饶分公司销售航意险并向其支付服务费，'
                    r'涉及保费104.47万元。上述行为违反了《保险法》（2009年修订，下同）第116条第8项的规定。\n程叶卿作为中国人民人寿保险股份有限公司陕西省分公司时任战略客户部经理，对上述行为负有直接责任'
                ]
                punishment_basis_str = '|'.join(punishment_basis_str_list)
                punishment_basis_compiler = re.compile(r'[。\n；]' + '(' + punishment_basis_str + ')' +
                                                       '.(\n?依据|\n?根据|\n?鉴于|\n?上述事实有现场检查确认书)', re.MULTILINE)
                punishment_basis_list = punishment_basis_compiler.findall(content_text)
                punishment_basis = '；'.join([kk[0].strip() for kk in punishment_basis_list])

                publish_date_text = re.search(r'\n(.*?)$', content_text).group(1).replace('\n', '')
                if re.search(r'.{4}年.{1,2}月.{1,3}日', publish_date_text):
                    publish_date = re.findall('.{4}年.{1,2}月.{1,3}日', publish_date_text)[-1].replace(' ', '')
                    m = re.match("([0-9零一二两三四五六七八九十〇○ＯOΟО]+年)?([0-9一二两三四五六七八九十]+)月?([0-9一二两三四五六七八九十]+)[号日]?",
                                 publish_date)
                    real_publish_date = get_year(m.group(1)) + str(cn2dig(m.group(2))) + '月' + str(
                        cn2dig(m.group(3))) + '日'
                else:
                    publish_date_text = table_content.find_all('tr')[1].text
                    publish_date = re.findall(r'\d{4}-\d{1,2}-\d{1,2}', publish_date_text)[-1]
                    real_publish_date = publish_date.split('-')[0] + '年' + str(
                        int(publish_date.split('-')[1])) + '月' + str(
                        int(publish_date.split('-')[2])) + '日'

                result_map = {
                    'announcementTitle': title,
                    'announcementOrg': '陕西银保监局',
                    'announcementDate': real_publish_date,
                    'announcementCode': document_code,
                    'facts': truth,
                    'defenseOpinion': defense,
                    'defenseResponse': defense_response,
                    'litigant': litigant[:-1] if litigant[-1] == '：' else litigant,
                    'punishmentBasement': punishment_basis,
                    'punishmentDecision': punishment_decision,
                    'type': '行政处罚决定',
                    'oss_file_id': file_id,
                    'status': 'not checked'
                }
                logger.info(result_map)
                if db.announcement.find({'announcementTitle': title, 'oss_file_id': file_id}).count() == 0:
                    db.announcement.insert_one(result_map)
                    logger.info('陕西保监局 数据解析 ' + ' -- 数据导入完成')
                else:
                    logger.info('陕西保监局 数据解析 ' + ' -- 数据已经存在')
                db.parsed_data.update_one({'_id': file_id}, {'$set': {'parsed': True}})
                logger.info('陕西保监局 数据解析 ' + ' -- 修改parsed完成')
